#!/usr/bin/env python3
"""Create English manuscript for Pain (IASP) format.

Pain journal format:
- Unstructured abstract (single paragraph, no headings)
- American English
- Figures/tables embedded in text
- Vancouver numbered references [1], [2]
- Double-spaced, Times New Roman 12pt
- Title page with word counts
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import json

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# --- Page setup: A4, wide margins, double-spaced ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 2.0

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=5.5):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        parts = caption.split('. ', 1)
        if len(parts) == 2:
            run_b = cap.add_run(parts[0] + '. ')
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_n = cap.add_run(parts[1])
            run_n.font.size = Pt(10)
        else:
            run_b = cap.add_run(caption)
            run_b.font.size = Pt(10)

# ============================================================
# TITLE PAGE (Page 1)
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    '\u201cJapanese Patient\u201d? \u2014 A Patient.\n'
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures\n'
    'Challenges the Stereotype of a Stoic Monolith'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()  # blank line

# Authors placeholder
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[Author names and affiliations to be inserted]')
run.font.size = Pt(12)

doc.add_paragraph()

# Corresponding author
add_bold_paragraph('Corresponding author: ',
    '[Name], Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. '
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

# Keywords (3-6 for Pain)
add_bold_paragraph('Keywords: ',
    'cultural pain behavior; ecological study; gaman; neuropathic pain; regional variation; prescribing patterns')

doc.add_paragraph()

# Manuscript info (Pain format)
add_bold_paragraph('Number of text pages: ', '[to be determined after formatting]')
add_bold_paragraph('Number of figures: ', '7 (6 main + 1 supplementary)')
add_bold_paragraph('Number of tables: ', '2')

doc.add_page_break()

# ============================================================
# ABSTRACT (new page)
# ============================================================
add_heading_text('Abstract', level=1)

# Pain uses unstructured abstract (single paragraph, no subheadings)
doc.add_paragraph(
    'Cross-cultural studies consistently characterize Japanese people as more stoic toward pain '
    'than Western populations, yet whether this cultural pain endurance varies regionally within Japan '
    'remains unexplored. This ecological study used Japan\u2019s National Database (NDB) Open Data '
    '(population-complete insurance claims, April 2023\u2013March 2024) to map pain-related prescribing '
    'across 47 prefectures and nine regional blocks. In Phase 1, an inpatient analgesic-per-surgery index '
    'varied 1.97-fold (Gifu 25.20 to Kagoshima 49.75; Kruskal\u2013Wallis P < 0.001). '
    'The perioperative setting neutralizes healthcare access as a confounder, since all patients are hospitalized. '
    'Contrary to the popular perception of northern stoicism, the Tohoku region ranked seventh of nine '
    '(mean 39.97 vs national 35.78; P = 0.031; Cohen\u2019s d = 0.87), prescribing more, not fewer, analgesics. '
    f'In Phase 2, outpatient neuropathic pain prescribing (a chronic postsurgical pain proxy) '
    f'showed Tohoku with the highest unadjusted index '
    f'(mean {reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs {reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}), but this excess became nonsignificant '
    f'after adjustment for confounding diseases (diabetes, postherpetic neuralgia, depression, anxiety; '
    f'P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}), '
    'with diabetes drug prescribing as the strongest confounder (r = 0.87). '
    'These findings demonstrate that Japan\u2019s pain culture is not monolithic: '
    'nearly twofold within-country variation in analgesic prescribing challenges the assumption '
    'that \u201cJapanese\u201d constitutes a uniform category for pain behavior. '
    'The monolithic characterization of any national population\u2019s pain behavior carries '
    'a risk of clinical harm\u2014whenever cultural stereotypes narrow the range of treatment options '
    'considered for a patient. Individualized pain assessment, informed by objective nociception monitoring '
    'rather than culturally stereotyped assumptions, should be the standard of perioperative care.')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('Introduction', level=1)

# Para 1: Cross-cultural pain literature (foundation)
doc.add_paragraph(
    'Pain is a universal experience, yet its expression, tolerance, and management are profoundly shaped by culture [1,2]. '
    'Since Zborowski\u2019s seminal observation that ethnic groups in the United States differ markedly in pain behavior [3], '
    'a large body of literature has established that cultural norms influence whether individuals report pain, '
    'seek treatment, or endure in silence [4]. These cross-cultural differences have practical implications: '
    'they affect postoperative pain assessment, analgesic prescribing, and ultimately patient outcomes.')

# Para 2: Japanese stoicism specifically (gaman)
doc.add_paragraph(
    'Among the world\u2019s cultures, Japan is consistently characterized as having a particularly stoic orientation '
    'toward pain. Hobara reported that Japanese respondents rated pain behaviors (e.g., moaning, crying) as '
    'significantly less appropriate than Euro-American respondents [5]. '
    'Feng et al. demonstrated that Japanese participants were far less willing to trade time to avoid '
    'pain/discomfort than British participants using the EQ-5D: the decrement for moving from moderate to '
    'extreme pain was 0.65 in the United Kingdom versus only 0.15 in Japan [6]. '
    'This culturally mediated pain endurance, encapsulated by the Japanese concept of '
    'gaman (\u6211\u6162, stoic endurance), is widely recognized as a core cultural value [7]. '
    'Such stoicism may lead to underreporting of pain and underutilization of analgesics at the population level. '
    'Internationally, this stereotype carries clinical risk: clinicians unfamiliar with Japanese patients '
    'may assume that stoic presentation indicates lower analgesic need.')

# Para 3: Within-country regional variation (gap)
doc.add_paragraph(
    'Most studies of cultural pain behavior compare between nations. '
    'However, substantial cultural heterogeneity exists within countries. '
    'Cohen and Nisbett demonstrated that the \u201cculture of honor\u201d in the southern United States '
    'produces measurably different behavioral and physiological responses to provocation compared with '
    'the northern states [8]\u2014a within-country cultural difference with real-world consequences. '
    'In Japan, regional cultural identities are strong: the Tohoku region (northeast) is traditionally '
    'perceived as stoic and resilient, while other regions carry their own cultural characteristics [9]; '
    'even the Japanese word for throbbing pain (uzuku) shows distinct regional usage patterns [10]. '
    'A nationwide survey by Pfizer Japan found that the proportion of chronic pain patients '
    'reporting that they \u201cendure pain\u201d ranged from 48.7% to 81.6% across prefectures [11], '
    'suggesting within-country variation. Yet whether such attitudinal differences translate into '
    'measurable differences in pain-related healthcare utilization remains unexplored.')

# Para 4: Japan as natural experiment + NDB
doc.add_paragraph(
    'Japan provides an ideal setting for studying regional variation in pain behavior. '
    'Its universal health insurance system eliminates financial access barriers, '
    'standardized drug pricing removes market-driven prescribing variation, '
    'and the National Database of Health Insurance Claims (NDB) captures virtually all '
    'reimbursed healthcare utilization (\u22482.1 billion claims annually) [12]. '
    'Furthermore, recent studies have documented significant regional variation in '
    'other pain-related metrics: Wakaizumi et al. found up to 1.6-fold differences in '
    'high-impact chronic pain prevalence across prefectures [13], '
    'and Matsuoka et al. reported 4-fold regional variation in cancer opioid prescribing [14]. '
    'However, no study has examined perioperative and chronic pain-related prescribing '
    'at the national prefecture level.')

doc.add_paragraph(
    'The perioperative setting offers a distinct methodological advantage over community-based pain surveys. '
    'Because all surgical patients are already admitted to hospital, healthcare access\u2014'
    'a major confounding factor in outpatient studies of analgesic behavior\u2014'
    'is effectively neutralized. Every patient, regardless of geographic remoteness, '
    'socioeconomic status, or health-seeking tendency, is under direct clinical observation '
    'and has equal opportunity to receive analgesics. '
    'Regional differences in perioperative prescribing therefore more closely reflect '
    'differences in clinical practice and pain expression than differences in access to care.')

# Para 5: Objectives (exploratory)
doc.add_paragraph(
    'The purpose of this exploratory study was threefold. First (Phase 1), to map regional variation '
    'in acute perioperative analgesic prescribing across Japan\u2019s 47 prefectures and nine regional blocks. '
    'Second (Phase 2), to examine regional variation in outpatient neuropathic pain prescribing '
    'as a chronic postsurgical pain (CPSP) proxy, after adjustment for confounding diseases '
    '(diabetic neuropathy, postherpetic neuralgia, depression, anxiety). '
    'Third, to integrate Phase 1 and Phase 2 findings to explore whether acute perioperative '
    'analgesic use correlates with subsequent chronic pain-related prescribing at the population level.')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

add_heading_text('Study design and reporting', level=2)
doc.add_paragraph(
    'This ecological study analyzed prefecture-level aggregate data from the NDB Open Data. '
    'The study is reported following the Strengthening the Reporting of Observational Studies in Epidemiology '
    '(STROBE) guidelines [15] and the REporting of studies Conducted using Observational '
    'Routinely-collected health Data (RECORD) extension [16]. '
    'As only publicly available aggregate data were used, ethical approval was not required.')

add_heading_text('Data source', level=2)
doc.add_paragraph(
    'The 10th edition of the NDB Open Data (Ministry of Health, Labour and Welfare [MHLW]) was used, '
    'covering healthcare claims from April 2023 to March 2024 [12]. The NDB captures claims from all insurers '
    'within Japan\u2019s universal health coverage system, encompassing approximately 125 million insured individuals. '
    'Aggregate data are published at the prefecture level with suppression of cells containing fewer than 10 events.')

add_heading_text('Regional classification', level=2)
doc.add_paragraph(
    'Prefectures were grouped into nine standard regional blocks: '
    'Hokkaido (1), Tohoku (6: Aomori, Iwate, Miyagi, Akita, Yamagata, Fukushima), '
    'Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), Kinki (6), '
    'Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8). '
    'This classification follows the standard administrative grouping used by the MHLW.')

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'Inpatient prescription data were extracted for three analgesic drug classes: '
    'Class 114 (antipyretic analgesics and anti-inflammatory drugs [NSAIDs/acetaminophen]), '
    'Class 811 (opium alkaloid narcotics [morphine, oxycodone, codeine]), and '
    'Class 821 (synthetic narcotics [fentanyl, pethidine, tapentadol]). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section. '
    'The analgesic-per-surgery index was calculated for each prefecture as: '
    'total inpatient analgesic quantity (units)/total inpatient surgical procedure count.')

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing as CPSP proxy', level=2)
doc.add_paragraph(
    'Five classes of outpatient oral neuropathic pain medications were extracted by prefecture: '
    'pregabalin (78 formulations), mirogabalin (8 formulations), duloxetine (33 formulations), '
    'tramadol (3 formulations), and neurotropin (1 formulation). '
    'The neuropathic pain prescribing-per-surgery index was calculated as: '
    'total outpatient neuropathic pain drug quantity/total inpatient surgical procedure count. '
    'This index serves as a proxy for CPSP, recognizing that these drugs are also '
    'prescribed for other neuropathic conditions.')

add_heading_text('Confounder disease proxies', level=2)
doc.add_paragraph(
    'Since neuropathic pain drugs are prescribed for conditions other than CPSP, '
    'four confounder disease proxies were extracted from the outpatient prescription data:')
doc.add_paragraph('Oral hypoglycemic agents (261 formulations): proxy for diabetic neuropathy', style='List Bullet')
doc.add_paragraph('Herpes zoster antivirals (47 formulations: valacyclovir, acyclovir, famciclovir, amenamevir): '
                  'proxy for postherpetic neuralgia', style='List Bullet')
doc.add_paragraph('Antidepressants excluding duloxetine (128 formulations: SSRIs, SNRIs, NaSSAs, tricyclics): '
                  'proxy for depression', style='List Bullet')
doc.add_paragraph('Anxiolytics (112 formulations: benzodiazepines, tandospirone, hydroxyzine): '
                  'proxy for anxiety disorders', style='List Bullet')
doc.add_paragraph(
    'Each confounder proxy was expressed per surgery to maintain comparability with the neuropathic pain index.')

add_heading_text('Outpatient nerve block procedures', level=2)
doc.add_paragraph(
    'Outpatient nerve block procedure counts were extracted from the anesthesia section of the NDB Open Data '
    'as an independent CPSP proxy. This included epidural blocks, paravertebral blocks, '
    'trigger point injections, stellate ganglion blocks, and continuous nerve block infusions (73 procedure codes).')

add_heading_text('Statistical analysis', level=2)
doc.add_paragraph(
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test across '
    'nine regional blocks, followed by post hoc pairwise Mann\u2013Whitney U tests with Bonferroni correction. '
    'For Phase 2, five regression models were fitted:')
doc.add_paragraph('Model 1: Unadjusted comparison of neuropathic pain prescribing across regions', style='List Bullet')
doc.add_paragraph('Model 2: Neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + regional block', style='List Bullet')
doc.add_paragraph('Model 3: Core neuropathic drugs (pregabalin + mirogabalin only) ~ same confounders', style='List Bullet')
doc.add_paragraph('Model 4: Nerve blocks ~ same confounders', style='List Bullet')
doc.add_paragraph('Model 5: Neuropathic pain ~ acute analgesic index + confounders (integrated model)', style='List Bullet')

doc.add_paragraph(
    'The adjusted CPSP index was derived as the residual from regressing neuropathic pain prescribing '
    'on the four confounder proxies, '
    'representing the \u201cunexplained\u201d neuropathic pain prescribing after accounting for confounding diseases. '
    'Pearson and Spearman correlations assessed relationships between indices. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11, matplotlib 3.8).')

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima), a 1.97-fold difference '
    '(Kruskal\u2013Wallis P < 0.001 across nine regions).')

doc.add_paragraph(
    'Substantial regional clustering was observed. Tokai and Kinki (western Japan) had the lowest indices, '
    'while Kyushu-Okinawa and Hokkaido had the highest (Table 1). '
    'A notable finding was that Tohoku, culturally perceived as Japan\u2019s most stoic region, '
    'ranked seventh of nine regions with a mean index of 39.97 (SD 3.53), '
    'significantly above the non-Tohoku mean of 35.17 (Mann\u2013Whitney U = 190, P = 0.031; Cohen\u2019s d = 0.87). '
    'All six Tohoku prefectures ranked in the upper half nationally. '
    'This pattern was consistent across all drug classes: NSAIDs (P = 0.044), '
    'opioid alkaloids (P = 0.003), and synthetic opioids (P = 0.001).')

# Table 1
add_bold_paragraph('Table 1. ', 'Phase 1: Inpatient analgesic-per-surgery index by regional block')
table1 = doc.add_table(rows=11, cols=5)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Regional block', 'n', 'Mean (SD)', 'Range', 'Rank']):
    table1.rows[0].cells[i].text = h
data1 = [
    ['Tokai', '4', '27.62 (2.18)', '25.20\u201330.07', '1'],
    ['Kinki', '6', '30.02 (2.03)', '27.92\u201332.33', '2'],
    ['Kanto', '7', '33.00 (2.09)', '29.82\u201334.78', '3'],
    ['Hokuriku-Koshinetsu', '6', '35.38 (3.54)', '31.18\u201340.08', '4'],
    ['Chugoku', '5', '35.73 (4.18)', '31.01\u201340.17', '5'],
    ['Shikoku', '4', '36.33 (3.27)', '33.49\u201341.02', '6'],
    ['Tohoku', '6', '39.97 (3.53)', '35.18\u201344.51', '7'],
    ['Kyushu-Okinawa', '8', '42.26 (4.33)', '35.82\u201349.75', '8'],
    ['Hokkaido', '1', '46.12 (\u2014)', '46.12', '9'],
    ['National', '47', '35.78 (5.56)', '25.20\u201349.75', '\u2014'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
# Bold Tohoku row
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
doc.add_paragraph(
    'Nationally, outpatient neuropathic pain drug prescriptions totaled 2,289,549,163 units, '
    'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    'duloxetine (15.3%), and tramadol (4.9%). '
    f'Tohoku had a markedly higher index ({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\u2019s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    'with Iwate (566.7), Aomori (519.3), and Akita (461.1) occupying the top three positions nationally (Fig. 1).')

add_figure(OUTPUT_DIR + 'fig1_neuropathic_unadjusted.png',
    'Fig. 1. Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
    'Bars represent total neuropathic pain prescriptions (pregabalin + mirogabalin + duloxetine + tramadol + neurotropin) '
    'divided by inpatient surgical procedure count. Tohoku prefectures (red bars with red borders) '
    'cluster at the high end. Dashed line = national mean.')

doc.add_page_break()

add_heading_text('Confounder analysis', level=2)
doc.add_paragraph(
    'Neuropathic pain prescribing showed strong correlations with the confounder disease proxies. '
    'Diabetes drug prescribing was the strongest correlate (r = 0.87, P < 0.001), '
    'followed by anxiolytics (r = 0.75, P < 0.001), antidepressants (r = 0.46, P = 0.001), '
    'and herpes antivirals (r = 0.19, P = 0.19). '
    'These confounders collectively explained 80.4% of the variance in neuropathic pain prescribing '
    f'(R\u00b2 = {reg["model2_adjusted"]["R2"]:.3f} in Model 2; Fig. 2).')

add_figure(OUTPUT_DIR + 'fig2_confounder_correlations.png',
    'Fig. 2. Correlation between outpatient neuropathic pain drug prescribing and confounder disease proxies. '
    'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
    'Diabetes drugs show the strongest correlation (r = 0.87).')

doc.add_page_break()

add_heading_text('Confounder-adjusted analysis', level=2)
doc.add_paragraph(
    'After adjusting for all four confounder proxies, '
    f'the Tohoku effect was attenuated and became nonsignificant in Model 2 '
    f'(\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, P = {reg["model2_adjusted"]["tohoku_p"]:.3f}). '
    'This attenuation was consistent across all model specifications: '
    f'Model 3 (core neuropathic drugs: \u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}), '
    f'Model 4 (nerve blocks: P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}), and '
    f'Model 5 (integrated: \u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}; Table 2).')

doc.add_paragraph(
    'The adjusted CPSP index (residuals after removing confounder effects) showed a dramatically '
    'different pattern from the unadjusted data (Fig. 3). '
    f'The Tohoku mean shifted from markedly positive (unadjusted) to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}). '
    'The Chugoku region emerged as having the highest adjusted CPSP index, '
    'while Tokai had the lowest (Fig. 4).')

add_figure(OUTPUT_DIR + 'fig3_adjusted_cpsp_index.png',
    'Fig. 3. Confounder-adjusted CPSP index by prefecture. '
    'Residuals from regressing neuropathic pain prescribing on diabetes drugs, herpes antivirals, '
    'antidepressants, and anxiolytics. Positive values indicate higher neuropathic pain prescribing '
    'than expected given confounding disease prevalence. Tohoku prefectures (red borders) are dispersed '
    'across the distribution after adjustment.')

doc.add_page_break()

# Table 2
add_bold_paragraph('Table 2. ', 'Regression models for neuropathic pain prescribing: effect of confounder adjustment')
table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Dependent variable', 'Tohoku effect', 'P value', 'Interpretation']):
    table2.rows[0].cells[i].text = h
t2_data = [
    ['1 (unadjusted)', 'Neuropathic/surgery', f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}', f'{reg["model1_unadjusted"]["p_value"]:.4f}', 'Significant ***'],
    ['2 (all confounders)', 'Neuropathic/surgery', f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['3 (core drugs)', 'PGB + MGB/surgery', f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['4 (nerve blocks)', 'Blocks/surgery', f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['5 (integrated)', 'Neuropathic/surgery', f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['Adj. CPSP index', 'Residuals', f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}', f'{reg["adjusted_cpsp_test"]["p_value"]:.3f}', 'Nonsignificant'],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_figure(OUTPUT_DIR + 'fig4_region_unadj_vs_adj.png',
    'Fig. 4. Regional comparison of neuropathic pain prescribing: (a) unadjusted and (b) after confounder adjustment. '
    'Tohoku (red border) shifts from the highest region to mid-range after adjustment. '
    'Error bars represent standard deviation.')

doc.add_page_break()

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
doc.add_paragraph(
    f'Phase 1 acute perioperative analgesic prescribing correlated positively with '
    f'Phase 2 unadjusted neuropathic pain prescribing (r = 0.38, P = 0.008; Fig. 5a). '
    f'After confounder adjustment, this correlation was attenuated to borderline significance '
    f'(r = 0.29, P = 0.052; Fig. 5b). '
    f'In the integrated model (Model 5), the Phase 1 acute pain index was a significant predictor '
    f'of neuropathic pain prescribing even after adjusting for confounders '
    f'(\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, P = {reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, P = {reg["model5_integrated"]["tohoku_p"]:.3f}).')

add_figure(OUTPUT_DIR + 'fig5_phase1_vs_phase2.png',
    'Fig. 5. Integration of Phase 1 (acute perioperative analgesic prescribing) and Phase 2 '
    '(outpatient neuropathic pain prescribing as CPSP proxy). '
    '(a) Unadjusted: positive correlation (r = 0.38, P = 0.008). '
    '(b) Confounder-adjusted: attenuated correlation (r = 0.29, P = 0.052). '
    'Tohoku prefectures (red borders) cluster in the upper-right quadrant.')

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

doc.add_paragraph(
    'This study is the first to map perioperative and chronic pain-related prescribing across '
    'all 47 prefectures of Japan, leveraging freely available NDB Open Data. '
    'Three principal findings emerged from this exploratory analysis.')

add_heading_text('Substantial regional variation exists within a stoic culture', level=2)
doc.add_paragraph(
    'Despite Japan\u2019s well-documented cultural stoicism toward pain [5,6], '
    'we found a 1.97-fold variation in acute perioperative analgesic prescribing across prefectures '
    'and significant differences across regional blocks. '
    'This parallels findings from other domains: Cohen and Nisbett\u2019s work on the \u201cculture of honor\u201d '
    'in the United States demonstrated that even within a single nation, regional cultural norms '
    'produce measurably different behavioral outcomes [8]. '
    'Our finding suggests that Japan\u2019s pain culture is not monolithic; '
    'regional demographics, healthcare infrastructure, and local clinical practices '
    'generate heterogeneity beneath the surface of a nationally shared cultural norm.')

add_heading_text('Clinical implications: the danger of monolithic cultural stereotypes', level=2)
doc.add_paragraph(
    'The 1.97-fold within-Japan variation documented here has direct clinical relevance '
    'beyond the domestic context. A large body of evidence demonstrates that ethnic and cultural '
    'stereotypes influence clinician pain assessment and analgesic prescribing. '
    'Anderson et al. showed that racial and ethnic minorities in the United States consistently '
    'receive less adequate pain management across acute, chronic, cancer, and palliative settings [17]. '
    'Campbell and Edwards identified that clinician expectations about a patient\u2019s cultural '
    'pain behavior can lead to systematic under- or over-treatment [18]. '
    'Rogger et al. emphasized that cultural framing affects not only patient reporting but also '
    'how clinicians interpret and respond to pain cues [2].')

doc.add_paragraph(
    'This risk is compounded by the widespread perception of Japan as a culturally and ethnically '
    'homogeneous society. The nihonjinron discourse (theories of Japanese uniqueness) '
    'has long promoted the notion that Japanese people constitute a uniform population '
    'sharing a single set of behavioral norms [19]. '
    'Yet this \u201chegemony of homogeneity,\u201d as Befu termed it, is an ideological construct '
    'rather than an empirical fact [19]. '
    'Burgess showed that this \u201cillusion\u201d of a homogeneous Japan '
    'has tangible consequences for social policy and public perception [20]. '
    'In the clinical context, the combination of two stereotypes\u2014\u201cJapanese are stoic\u201d '
    'and \u201cJapanese are homogeneous\u201d\u2014creates a doubly misleading assumption: '
    'that all Japanese patients will tolerate pain equally and require less analgesia. '
    'Our finding of 1.97-fold within-Japan variation in analgesic prescribing '
    'directly refutes this assumption.')

doc.add_paragraph(
    'Put simply, there is no such entity as \u201cthe Japanese patient\u201d '
    'whose pain behavior can be predicted from nationality alone\u2014'
    'there are only individual patients from 47 diverse prefectures, '
    'each with distinct demographic profiles, clinical environments, '
    'and pain-related prescribing cultures.')

doc.add_paragraph(
    'For Japanese patients treated outside Japan, clinicians may thus be operating under '
    'a dual misapprehension: that Japanese culture is uniformly stoic, '
    'and that this stoicism translates into lower analgesic need. '
    'Our data challenge both premises. '
    'Pain-related prescribing varies nearly twofold across prefectures, '
    'and regions stereotyped as particularly stoic do not in fact use fewer analgesics. '
    'If within-country variation of this magnitude exists, applying a national-level '
    'stereotype to individual clinical decisions is neither evidence-based nor safe.')

doc.add_paragraph(
    'We therefore argue that the monolithic characterization of any national population\u2019s '
    'pain behavior carries a risk of clinical harm\u2014and this argument extends well beyond Japan. '
    'Whenever clinicians allow cultural stereotypes to narrow the range of treatment options '
    'considered for a patient, the patient is placed at a therapeutic disadvantage. '
    'A \u201cstoic Japanese\u201d label may lead to under-prescribing of analgesics; '
    'conversely, stereotypes about other populations may lead to over-prescribing, '
    'under-recognition of certain symptoms, or premature closure in differential diagnosis. '
    'Our data\u2014showing nearly twofold within-country variation beneath a single cultural label\u2014'
    'serve as a concrete case study of a universal principle: '
    'cultural generalizations, however convenient, are a poor basis for individual clinical decisions.')

doc.add_paragraph(
    'The aspiration behind this study is that fewer patients, regardless of nationality or ethnicity, '
    'should suffer the disadvantage of having their treatment options constrained by cultural stereotypes. '
    'For Japanese patients treated outside Japan, clinicians should be aware that '
    'pain endurance norms are heterogeneous and that stoic presentation '
    'does not reliably indicate lower analgesic need. '
    'But the same logic applies to every patient from every culture: '
    'the assumption of cultural homogeneity\u2014however convenient\u2014is '
    'no more scientifically defensible in pain medicine than in any other domain. '
    'Individualized pain assessment, rather than culturally stereotyped assumptions, '
    'remains the cornerstone of equitable perioperative care.')

add_heading_text('Pain as an individual experience and the need for objective nociception monitoring', level=2)
doc.add_paragraph(
    'Fundamentally, pain is a personal experience. The International Association for the Study of Pain '
    'defines pain as \u201can unpleasant sensory and emotional experience associated with, or resembling that '
    'associated with, actual or potential tissue damage\u201d\u2014a definition that is inherently subjective [21]. '
    'Our ecological data illustrate this principle at the population level: '
    'even within a single nation sharing the same language, insurance system, and broad cultural heritage, '
    'pain-related prescribing varies nearly twofold. '
    'This heterogeneity likely reflects not only differences in disease burden and clinical practice '
    'but also the irreducible individuality of pain perception and expression.')

doc.add_paragraph(
    'The clinical consequence of this individuality is that no cultural label\u2014whether '
    '\u201cstoic Japanese\u201d or any other\u2014can substitute for direct measurement of a patient\u2019s '
    'nociceptive state. Yet the tools to perform such measurement remain underdeveloped. '
    'The balance between nociception and anti-nociception (NANB) is a physiological quantity '
    'that, in principle, could be monitored objectively during and after surgery, '
    'bypassing the need for patient self-report or clinician inference altogether. '
    'Onishi et al. demonstrated that normalized pulse volume (NPV), '
    'derived from standard pulse oximetry, predicts the recovery of spontaneous respiration '
    'with significantly less interindividual variability (coefficient of variation 36.3%) '
    'than opioid effect-site concentration (62.4%), '
    'suggesting that NPV may serve as a culture-independent, objective indicator of NANB [22].')

doc.add_paragraph(
    'The development and clinical implementation of such objective nociception monitoring systems '
    'would fundamentally change the landscape of perioperative pain management. '
    'If NANB could be quantified at the bedside\u2014much as bispectral index quantifies hypnotic depth\u2014'
    'analgesic titration would no longer depend on subjective pain scores susceptible to cultural bias, '
    'nor on population-level stereotypes about which patients \u201ctolerate\u201d pain. '
    'Until such systems are clinically available, the findings of the present study serve as a reminder '
    'that cultural generalizations are a poor proxy for individual nociceptive reality.')

add_heading_text('Confounding diseases explain apparent regional variation in CPSP proxies', level=2)
doc.add_paragraph(
    'The most methodologically important finding is that the dramatic regional variation in '
    f'neuropathic pain prescribing (unadjusted d = {reg["model1_unadjusted"]["cohens_d"]:.2f} for Tohoku vs rest) was largely explained '
    'by confounding disease proxies. '
    'Diabetes drug prescribing alone correlated at r = 0.87 with neuropathic pain prescribing, '
    'reflecting the known high prevalence of diabetic neuropathy requiring gabapentinoids. '
    'After adjustment, the Tohoku effect was attenuated by 62% and became nonsignificant.')

doc.add_paragraph(
    'This has important implications for ecological pain research. '
    'Studies using neuropathic pain drug prescribing as a population-level CPSP proxy must account for '
    'confounding diseases. Without such adjustment, regional differences in diabetes prevalence '
    'could be misinterpreted as differences in CPSP. '
    'The within-database confounder adjustment demonstrated here\u2014using disease-specific drug proxies '
    'from the same data source\u2014provides a replicable framework.')

add_heading_text('A population-level acute\u2013chronic pain continuum', level=2)
doc.add_paragraph(
    'The positive correlation between Phase 1 (acute) and Phase 2 (chronic, adjusted) indices '
    '(r = 0.29, P = 0.052) suggests a modest link between regional acute pain management intensity '
    'and subsequent chronic pain-related prescribing. While ecological correlations cannot establish '
    'causation, this finding is consistent with the individual-level literature suggesting that '
    'the intensity of acute postoperative pain is a risk factor for CPSP [23]. '
    'Prefectures with relatively low acute analgesic use but high adjusted CPSP indices '
    'may warrant investigation for potential under-treatment of acute pain leading to chronification.')

add_heading_text('Strengths and limitations', level=2)
doc.add_paragraph(
    'Strengths of this study include the use of population-complete data covering all insurance-reimbursed healthcare '
    'in Japan, the novel integration of acute and chronic pain proxies, the transparent confounder-adjustment methodology, '
    'and the exploratory rather than confirmatory design that allows hypothesis generation. '
    'A further strength specific to the perioperative focus is that healthcare access does not confound '
    'the analgesic prescribing data: all patients in Phase 1 are inpatients by definition, '
    'eliminating the access-to-care heterogeneity that limits community-based pain studies.')

doc.add_paragraph('The main limitations are inherent to the ecological design:')
doc.add_paragraph('The unit of analysis is the prefecture, not the individual patient. '
                  'Ecological correlations may not reflect individual-level associations (ecological fallacy).',
                  style='List Bullet')
doc.add_paragraph('NDB Open Data do not contain diagnosis codes, '
                  'so CPSP cannot be directly identified. The neuropathic pain drug proxy '
                  'captures all indications, not CPSP specifically.', style='List Bullet')
doc.add_paragraph('Drug prescribing proxies may not capture disease prevalence accurately '
                  '(e.g., fibromyalgia has no specific drug proxy and shares pregabalin as first-line treatment).',
                  style='List Bullet')
doc.add_paragraph('The cross-sectional design cannot distinguish temporal sequences '
                  '(surgery \u2192 acute pain \u2192 CPSP).', style='List Bullet')
doc.add_paragraph('Unmeasured confounders (age distribution, surgical case mix, prescribing culture) '
                  'may contribute to residual regional variation.', style='List Bullet')

add_heading_text('Implications and future directions', level=2)
doc.add_paragraph(
    'This exploratory study demonstrates that Japan\u2019s NDB Open Data can serve as a hypothesis-generating platform '
    'for population-level pain research. Future studies using the NDB sampling dataset (Level 3 access) '
    'could enable individual-level longitudinal tracking from surgery to new neuropathic pain prescriptions, '
    'providing a direct CPSP measure. Procedure-specific analyses (e.g., total knee arthroplasty, mastectomy) '
    'would reduce surgical case-mix confounding. Linking prescribing data with patient-reported pain outcomes '
    'would clarify whether regional prescribing differences reflect differences in pain experience, '
    'pain expression, or clinical practice.')

doc.add_paragraph(
    'From an international perspective, our findings underscore the need for culturally nuanced '
    'but individually tailored pain management. Prospective studies examining how cultural '
    'stereotypes influence analgesic prescribing for Japanese patients in multicultural clinical '
    'settings would directly test whether the \u201cstoic Japanese\u201d label translates into therapeutic disadvantage. '
    'Such research is increasingly urgent as global migration and medical tourism expose more '
    'patients to healthcare systems unfamiliar with their cultural background.')

# ============================================================
# CONCLUSION
# ============================================================
add_heading_text('Conclusion', level=1)
doc.add_paragraph(
    'Despite Japan\u2019s culturally ingrained norm of pain endurance (gaman), '
    'perioperative and chronic pain-related prescribing varies up to 1.97-fold across prefectures. '
    'Confounding diseases, particularly diabetes, substantially modify the apparent regional pattern '
    'of neuropathic pain prescribing. '
    'These findings demonstrate that Japan\u2019s pain culture is not monolithic. '
    'Treating \u201cJapanese\u201d as a uniform category for pain behavior risks inadequate analgesia '
    'for Japanese patients treated abroad\u2014and the same principle applies to every cultural label '
    'applied to any patient population. '
    'Individualized pain assessment should replace culturally stereotyped assumptions '
    'to ensure equitable perioperative care across all clinical settings.')

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgments', level=1)
doc.add_paragraph(
    'The authors thank the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. [Additional acknowledgments to be inserted as appropriate.]')

# ============================================================
# CONFLICT OF INTEREST STATEMENT (Pain requires this in Acknowledgments)
# ============================================================
doc.add_paragraph()
add_heading_text('Conflict of interest statement', level=1)
doc.add_paragraph(
    'The authors have no conflicts of interest to declare.')

# ============================================================
# DATA AVAILABILITY STATEMENT (required by Pain)
# ============================================================
doc.add_paragraph()
add_heading_text('Data availability statement', level=1)
doc.add_paragraph(
    'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    'Labour and Welfare website (https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    'Analysis code is available at https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan.')

doc.add_page_break()

# ============================================================
# FIGURES (remaining)
# ============================================================
add_figure(OUTPUT_DIR + 'fig6_model_comparison_table.png',
    'Fig. 6. Summary of regression models: the Tohoku effect on neuropathic pain prescribing '
    'before and after confounder adjustment. Model 1 (unadjusted) shows a highly significant excess; '
    'all adjusted models show nonsignificant results, indicating confounders explain the regional difference.')

doc.add_page_break()
add_figure(OUTPUT_DIR + 'sfig1_heatmap.png',
    'Supplementary Fig. 1. Z-score heatmap of all indices by prefecture. '
    'Each row represents a variable; each column represents a prefecture, '
    'sorted by neuropathic pain prescribing. Red = above average; blue = below average. '
    'Tohoku prefectures are marked with red vertical lines. '
    'The adjusted CPSP index shows a different pattern from the raw neuropathic index.')

# ============================================================
# REFERENCES - Updated for new narrative
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)
refs = [
    # [1] Callister 2003 - Cultural influences on pain
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract. 2003;15:207\u2013211.',

    # [2] Rogger 2023 - Cultural framing and acute pain
    'Rogger R, Bello C, Romero CS, Urman RD, Luedi MM, Filipovic MG. '
    'Cultural framing and the impact on acute pain and pain services. '
    'Curr Pain Headache Rep. 2023;27:429\u2013436.',

    # [3] Zborowski 1969 - People in Pain
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass; 1969.',

    # [4] Okolo 2024 - Cultural variability review
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
    'a review of cross-cultural studies. Int J Sci Res Arch. 2024;11:2550\u20132556.',

    # [5] Hobara 2005 - Japanese vs Euro-American pain behavior beliefs
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain. 2005;9:389\u2013393.',

    # [6] Feng 2017 - Japan vs Europe EQ-5D
    'Feng Y, Herdman M, van Nooten F, Cleeland C, Parkin D, Ikeda S, Igarashi A, Devlin NJ. '
    'An exploration of differences between Japan and two European countries in the self-reporting '
    'and valuation of pain and discomfort on the EQ-5D. Qual Life Res. 2017;26:2067\u20132078.',

    # [7] Hayashi 2022 - PCS variation by country
    'Hayashi K, Ikemoto T, Shiro Y, Arai YC, Marcuzzi A, Costa D, Wrigley PJ. '
    'A systematic review of the variation in Pain Catastrophizing Scale reference scores '
    'based on language version and country in patients with chronic primary (non-specific) pain. '
    'Pain Ther. 2022;11:751\u2013780.',

    # [8] Cohen & Nisbett 1996 - Culture of honor, within-country variation
    'Cohen D, Nisbett RE, Bowdle BF, Schwarz N. '
    'Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d '
    'J Pers Soc Psychol. 1996;70:945\u2013960.',

    # [9] Kumagai 2020 - Media reproducing Tohoku stereotypes
    'Kumagai S. Media representations reproducing images of Tohoku: '
    'the Tohoku reconstruction corner in \u201cSecret Kenmin SHOW.\u201d '
    'Kotoba. 2020;41:21\u201338. [in Japanese]',

    # [10] Takeda & Yarimizu 2016 - Regional differences in pain expression "uzuku"
    'Takeda K, Yarimizu K. Regional differences in the pain expression uzuku. '
    'NINJAL Research Papers. 2016;10:85\u2013107. [in Japanese]',

    # [11] Pfizer 2017 - Prefecture pain tolerance survey
    'Pfizer Japan Inc. 47-prefecture survey on chronic pain. 2017. '
    'Available from: https://www.pfizer.co.jp/pfizer/company/press/2017. '
    'Accessed February 1, 2025.',

    # [12] NDB Open Data
    'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. 2024. '
    'Available from: https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
    'Accessed January 15, 2025.',

    # [13] Wakaizumi 2024 - Geographic variation in HICP in Japan
    'Wakaizumi K, Tanaka C, Shinohara Y, Wu Y, Takaoka S, Kawate M, Oka H, Matsudaira K. '
    'Geographical variation in high-impact chronic pain and psychological associations at the regional level: '
    'a multilevel analysis of a large-scale internet-based cross-sectional survey. '
    'Front Public Health. 2024;12:1482177.',

    # [14] Matsuoka 2025 - Regional opioid prescribing for cancer pain
    'Matsuoka Y, et al. Population-based claims study of regional and hospital function differences '
    'in opioid prescribing for cancer patients who died in hospital in Japan. '
    'Jpn J Clin Oncol. 2025;hyaf149.',

    # [15] STROBE
    'von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
    'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
    'guidelines for reporting observational studies. '
    'Lancet. 2007;370:1453\u20131457.',

    # [16] RECORD
    'Benchimol EI, Smeeth L, Guttmann A, Harber K, Moher D, Petersen I, Sorensen HT, '
    'von Elm E, Langan SM. '
    'The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med. 2015;12:e1001885.',

    # [17] Anderson 2009 - Racial/ethnic disparities in pain: unequal care
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
    'causes and consequences of unequal care. J Pain. 2009;10:1187\u20131204.',

    # [18] Campbell & Edwards 2012 - Ethnic differences in pain management
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag. 2012;2:219\u2013230.',

    # [19] Befu 2001 - Hegemony of Homogeneity
    'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
    'Melbourne: Trans Pacific Press; 2001.',

    # [20] Burgess 2010 - Illusion of homogeneous Japan
    'Burgess C. The \u201cillusion\u201d of homogeneous Japan and national character: '
    'discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. '
    'Asia Pac J. 2010;8(9):1\u201322.',

    # [21] IASP definition of pain
    'Raja SN, Carr DB, Cohen M, Finnerup NB, Flor H, Gibson S, Keefe FJ, Mogil JS, '
    'Ringkamp M, Sluka KA, Song XJ, Stevens B, Sullivan MD, Tutelman PR, Ushida T, Vader K. '
    'The revised International Association for the Study of Pain definition of pain: '
    'concepts, challenges, and compromises. Pain. 2020;161:1976\u20131982.',

    # [22] Onishi et al. 2024 - NPV as objective NANB monitor
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
    'and quantification of nociception anti-nociception balance compared to opioid effect site concentration: '
    'a prospective, observational study. F1000Research. 2024;13:233.',

    # [23] Kehlet 2006 - Persistent postsurgical pain
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet. 2006;367:1618\u20131625.',

]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# Save
outpath = OUTPUT_DIR + 'Pain_manuscript_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
