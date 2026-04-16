#!/usr/bin/env python3
"""Create Summary file for PAIN Reports submission (<=25 words)."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

h = doc.add_heading('Summary', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

summary_text = (
    'Perioperative analgesic prescribing varies 1.97-fold across Japan\u2019s 47 prefectures, '
    'challenging the stereotype of uniform Japanese stoicism toward pain.'
)

doc.add_paragraph(summary_text)

wc = len(summary_text.split())
print(f'Summary word count: {wc} (limit: 25)')

outpath = OUTPUT_DIR + 'PainReports_summary.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
