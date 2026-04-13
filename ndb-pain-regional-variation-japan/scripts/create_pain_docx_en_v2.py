#!/usr/bin/env python3
"""Create English manuscript for Pain (IASP) format — v2 (Editorial Manager revision).

Changes from v1:
- Introduction reduced to <=500 words
- Discussion (incl. conclusion) reduced to <=1500 words
- Page numbers added
- References: alphabetical by first author, bracketed [n], all authors, DOIs
- Tables moved to END of manuscript (after references)
- Figures NOT embedded — separate files (TIFF)
- Figure legends listed at end of manuscript (after references, before tables)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

# ============================================================
# REFERENCES — alphabetical by first author surname
# ============================================================
# Each entry: (key, full_citation)
# key = first-author surname for sorting
# After sorting, numbers are assigned; in-text we use those numbers.
refs_raw = [
    ('Anderson', 'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
     'causes and consequences of unequal care. J Pain 2009;10:1187\u20131204. '
     'doi:10.1016/j.jpain.2009.03.015'),

    ('Befu', 'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
     'Melbourne: Trans Pacific Press, 2001.'),

    ('Benchimol', 'Benchimol EI, Smeeth L, Guttmann A, Harron K, Moher D, Petersen I, '
     'Sorensen HT, von Elm E, Langan SM. '
     'The REporting of studies Conducted using Observational Routinely-collected health Data '
     '(RECORD) statement. PLoS Med 2015;12:e1001885. '
     'doi:10.1371/journal.pmed.1001885'),

    ('Burgess', 'Burgess C. The \u201cillusion\u201d of homogeneous Japan and national character: '
     'discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. '
     'Asia Pac J 2010;8(9):1\u201322.'),

    ('Callister', 'Callister LC. Cultural influences on pain perceptions and behaviors. '
     'Home Health Care Manag Pract 2003;15:207\u2013211. '
     'doi:10.1177/1084822302250687'),

    ('Campbell', 'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
     'Pain Manag 2012;2:219\u2013230. '
     'doi:10.2217/pmt.12.7'),

    ('Cohen', 'Cohen D, Nisbett RE, Bowdle BF, Schwarz N. '
     'Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d '
     'J Pers Soc Psychol 1996;70:945\u2013960. '
     'doi:10.1037/0022-3514.70.5.945'),

    ('Feng', 'Feng Y, Herdman M, van Nooten F, Cleeland C, Parkin D, Ikeda S, Igarashi A, Devlin NJ. '
     'An exploration of differences between Japan and two European countries in the self-reporting '
     'and valuation of pain and discomfort on the EQ-5D. Qual Life Res 2017;26:2067\u20132078. '
     'doi:10.1007/s11136-017-1541-4'),

    ('Hayashi', 'Hayashi K, Ikemoto T, Shiro Y, Arai YC, Marcuzzi A, Costa D, Wrigley PJ. '
     'A systematic review of the variation in Pain Catastrophizing Scale reference scores '
     'based on language version and country in patients with chronic primary (non-specific) pain. '
     'Pain Ther 2022;11:751\u2013780. '
     'doi:10.1007/s40122-022-00402-3'),

    ('Hobara', 'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
     'between Japanese and Euro-Americans. Eur J Pain 2005;9:389\u2013393. '
     'doi:10.1016/j.ejpain.2004.09.006'),

    ('Kehlet', 'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
     'Lancet 2006;367:1618\u20131625. '
     'doi:10.1016/S0140-6736(06)68700-X'),

    ('Kumagai', 'Kumagai S. Media representations reproducing images of Tohoku: '
     'the Tohoku reconstruction corner in \u201cSecret Kenmin SHOW.\u201d '
     'Kotoba 2020;41:21\u201338. [in Japanese]'),

    ('Matsuoka', 'Matsuoka Y, Morishima T, Sato A, Ogawa T, Miyashiro I. '
     'Population-based claims study of regional and hospital function differences '
     'in opioid prescribing for cancer patients who died in hospital in Japan. '
     'Jpn J Clin Oncol 2025;55:hyaf149. '
     'doi:10.1093/jjco/hyaf149'),

    ('MHLW', 'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. 2024. '
     'Available: https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
     'Accessed January 15, 2025.'),

    ('Okolo', 'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
     'a review of cross-cultural studies. Int J Sci Res Arch 2024;11:2550\u20132556. '
     'doi:10.30574/ijsra.2024.11.1.0263'),

    ('Onishi', 'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
     'and quantification of nociception anti-nociception balance compared to opioid effect site concentration: '
     'a prospective, observational study. F1000Research 2024;13:233. '
     'doi:10.12688/f1000research.146054.2'),

    ('Pfizer', 'Pfizer Japan Inc. 47-prefecture survey on chronic pain. 2017. '
     'Available: https://www.pfizer.co.jp/pfizer/company/press/2017. '
     'Accessed February 1, 2025.'),

    ('Raja', 'Raja SN, Carr DB, Cohen M, Finnerup NB, Flor H, Gibson S, Keefe FJ, Mogil JS, '
     'Ringkamp M, Sluka KA, Song XJ, Stevens B, Sullivan MD, Tutelman PR, Ushida T, Vader K. '
     'The revised International Association for the Study of Pain definition of pain: '
     'concepts, challenges, and compromises. Pain 2020;161:1976\u20131982. '
     'doi:10.1097/j.pain.0000000000001939'),

    ('Rogger', 'Rogger R, Bello C, Romero CS, Urman RD, Luedi MM, Filipovic MG. '
     'Cultural framing and the impact on acute pain and pain services. '
     'Curr Pain Headache Rep 2023;27:429\u2013436. '
     'doi:10.1007/s11916-023-01130-1'),

    ('Takeda', 'Takeda K, Yarimizu K. Regional differences in the pain expression uzuku. '
     'NINJAL Research Papers 2016;10:85\u2013107. [in Japanese]'),

    ('von Elm', 'von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
     'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
     'guidelines for reporting observational studies. '
     'Lancet 2007;370:1453\u20131457. '
     'doi:10.1016/S0140-6736(07)61602-X'),

    ('Wakaizumi', 'Wakaizumi K, Tanaka C, Shinohara Y, Wu Y, Takaoka S, Kawate M, Oka H, Matsudaira K. '
     'Geographical variation in high-impact chronic pain and psychological associations at the regional level: '
     'a multilevel analysis of a large-scale internet-based cross-sectional survey. '
     'Front Public Health 2024;12:1482177. '
     'doi:10.3389/fpubh.2024.1482177'),

    ('Zborowski', 'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.'),
]

# Sort alphabetically by first-author surname
refs_sorted = sorted(refs_raw, key=lambda x: x[0].lower())
# Build citation-number lookup: key -> number
ref_num = {}
for i, (key, _) in enumerate(refs_sorted, 1):
    ref_num[key] = i

# Convenience function for inline citation
def cite(*keys):
    """Return bracketed citation string, e.g. '[3,12]'."""
    nums = sorted(ref_num[k] for k in keys)
    return '[' + ','.join(str(n) for n in nums) + ']'

# ============================================================
# Document creation
# ============================================================
doc = Document()

# --- Page setup: A4, margins, double-spaced ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# --- Page numbers (footer, centered) ---
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)


def set_table_borders(table):
    """Set full borders on a DOCX table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


# Word-count helper
def wc(text):
    return len(text.split())


# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    '\u201cJapanese Patient\u201d? \u2014 A Patient.\n'
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures\n'
    'Challenges the Stereotype of a Stoic Monolith'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Tatsuki Onishi')
run.font.size = Pt(12)

doc.add_paragraph()

add_bold_paragraph('Corresponding author: ',
    'Tatsuki Onishi, Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. '
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

add_bold_paragraph('Keywords: ',
    'cultural pain behavior; ecological study; gaman; neuropathic pain; regional variation; prescribing patterns')

doc.add_paragraph()

add_bold_paragraph('Number of text pages: ', '[to be determined after formatting]')
add_bold_paragraph('Number of figures: ', '5 (+ 1 supplementary)')
add_bold_paragraph('Number of tables: ', '3')

doc.add_page_break()

# ============================================================
# ABSTRACT (unstructured, single paragraph)
# ============================================================
add_heading_text('Abstract', level=1)

abstract_text = (
    f'Cross-cultural studies characterize Japanese people as stoic toward pain, '
    f'yet whether this endurance varies regionally within Japan remains unexplored. '
    f'This ecological study used Japan\u2019s National Database (NDB) Open Data '
    f'(population-complete insurance claims, April 2023\u2013March 2024) to map pain-related prescribing '
    f'across 47 prefectures and nine regional blocks. '
    f'In Phase 1, an inpatient analgesic-per-surgery index '
    f'varied 1.97-fold (Gifu 25.20 to Kagoshima 49.75; Kruskal\u2013Wallis P < 0.001). '
    f'The perioperative setting neutralizes healthcare access as a confounder. '
    f'Contrary to popular perception, Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014'
    f'ranked seventh of nine (mean 39.97 vs national 35.78; P = 0.031; Cohen\u2019s d = 0.87), '
    f'prescribing more, not fewer, analgesics. '
    f'In Phase 2, outpatient neuropathic pain prescribing (a chronic postsurgical pain proxy) '
    f'showed Tohoku with the highest unadjusted index '
    f'(mean {reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs {reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}), but this excess became nonsignificant '
    f'after adjustment for confounding diseases '
    f'(P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}), '
    f'with diabetes drug prescribing as the strongest confounder (r = 0.87). '
    f'Nearly twofold within-country variation challenges the assumption '
    f'that \u201cJapanese\u201d constitutes a uniform category for pain behavior. '
    f'Culturally stereotyped assumptions risk narrowing treatment options; '
    f'individualized pain assessment should be the standard of perioperative care.'
)
doc.add_paragraph(abstract_text)
print(f'Abstract word count: {wc(abstract_text)}')

doc.add_page_break()

# ============================================================
# INTRODUCTION (target: <=500 words)
# ============================================================
add_heading_text('Introduction', level=1)

intro_parts = []

intro1 = (
    f'Pain is a universal experience, yet its expression and management are profoundly shaped '
    f'by culture {cite("Callister","Rogger")}. Since Zborowski\u2019s observation that ethnic groups '
    f'differ in pain behavior {cite("Zborowski")}, a large literature has established that cultural norms '
    f'influence pain reporting, treatment-seeking, and analgesic use {cite("Okolo")}. '
    f'Japan is consistently characterized as stoic: Hobara found that Japanese respondents rated pain '
    f'behaviors as less appropriate than Euro-Americans {cite("Hobara")}, and Feng et al. showed that '
    f'Japanese participants were far less willing to trade time to avoid pain on the EQ-5D {cite("Feng")}. '
    f'This culturally mediated endurance, encapsulated by the concept of gaman (\u6211\u6162), '
    f'carries clinical risk when clinicians assume stoic presentation indicates lower analgesic need.'
)
doc.add_paragraph(intro1)
intro_parts.append(intro1)

intro2 = (
    f'Most studies compare pain behavior between nations. However, substantial heterogeneity exists '
    f'within countries: Cohen and Nisbett showed that the \u201cculture of honor\u201d in the US South '
    f'produces measurably different behavioral responses from the North {cite("Cohen")}. '
    f'In Japan, regional identities are strong\u2014the Tohoku region is traditionally perceived as '
    f'stoic {cite("Kumagai")}, and even the word for throbbing pain (uzuku) shows regional usage '
    f'patterns {cite("Takeda")}. A Pfizer Japan survey found the proportion of chronic pain patients '
    f'\u201cenduring pain\u201d ranged from 48.7% to 81.6% across prefectures {cite("Pfizer")}. '
    f'Yet whether such differences translate into measurable healthcare utilization differences remains unknown.'
)
doc.add_paragraph(intro2)
intro_parts.append(intro2)

intro3 = (
    f'Japan\u2019s universal health insurance, standardized drug pricing, and the NDB\u2014capturing '
    f'virtually all reimbursed utilization {cite("MHLW")}\u2014provide an ideal setting. '
    f'Recent work has documented up to 1.6-fold regional variation in chronic pain prevalence '
    f'{cite("Wakaizumi")} and 4-fold variation in cancer opioid prescribing {cite("Matsuoka")}. '
    f'The perioperative setting offers a further advantage: because all patients are hospitalized, '
    f'healthcare access\u2014a major confounder in outpatient studies\u2014is neutralized.'
)
doc.add_paragraph(intro3)
intro_parts.append(intro3)

intro4 = (
    f'This exploratory study had three objectives: (1) map regional variation in acute perioperative '
    f'analgesic prescribing across 47 prefectures; (2) examine outpatient neuropathic pain prescribing '
    f'as a chronic postsurgical pain (CPSP) proxy after confounder adjustment; and (3) integrate '
    f'acute and chronic findings at the population level.'
)
doc.add_paragraph(intro4)
intro_parts.append(intro4)

intro_total = sum(wc(p) for p in intro_parts)
print(f'Introduction word count: {intro_total}')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

add_heading_text('Study design and reporting', level=2)
doc.add_paragraph(
    f'This ecological study analyzed prefecture-level aggregate data from the NDB Open Data. '
    f'It is reported following STROBE {cite("von Elm")} and the RECORD extension {cite("Benchimol")}. '
    f'As only publicly available aggregate data were used, ethical approval was not required.'
)

add_heading_text('Data source', level=2)
doc.add_paragraph(
    f'The 10th edition of the NDB Open Data (April 2023\u2013March 2024) was used {cite("MHLW")}. '
    f'The NDB captures claims from all insurers within Japan\u2019s universal coverage system, '
    f'encompassing approximately 125 million insured individuals. '
    f'Aggregate data are published at the prefecture level with suppression of cells <10 events.'
)

add_heading_text('Regional classification', level=2)
doc.add_paragraph(
    'Prefectures were grouped into nine standard regional blocks: '
    'Hokkaido (1), Tohoku (6: Aomori, Iwate, Miyagi, Akita, Yamagata, Fukushima), '
    'Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), Kinki (6), '
    'Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8).'
)

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'Inpatient prescription data were extracted for three analgesic drug classes: '
    'Class 114 (antipyretic analgesics/NSAIDs/acetaminophen), '
    'Class 811 (opium alkaloid narcotics), and Class 821 (synthetic narcotics). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section. '
    'The analgesic-per-surgery index was calculated for each prefecture as: '
    'total inpatient analgesic quantity (units)/total inpatient surgical procedure count.'
)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing as CPSP proxy', level=2)
doc.add_paragraph(
    'Five classes of outpatient oral neuropathic pain medications were extracted: '
    'pregabalin (78 formulations), mirogabalin (8), duloxetine (33), '
    'tramadol (3), and neurotropin (1). '
    'The neuropathic pain prescribing-per-surgery index was calculated as: '
    'total outpatient neuropathic pain drug quantity/total inpatient surgical procedure count.'
)

add_heading_text('Confounder disease proxies', level=2)
doc.add_paragraph(
    'Four confounder disease proxies were extracted from outpatient data:')
doc.add_paragraph('Oral hypoglycemic agents (261 formulations): proxy for diabetic neuropathy',
                  style='List Bullet')
doc.add_paragraph('Herpes zoster antivirals (47 formulations): proxy for postherpetic neuralgia',
                  style='List Bullet')
doc.add_paragraph('Antidepressants excluding duloxetine (128 formulations): proxy for depression',
                  style='List Bullet')
doc.add_paragraph('Anxiolytics (112 formulations): proxy for anxiety disorders',
                  style='List Bullet')
doc.add_paragraph(
    'Each proxy was expressed per surgery. Outpatient nerve block procedure counts '
    '(73 codes) served as an independent CPSP proxy.'
)

add_heading_text('Statistical analysis', level=2)
doc.add_paragraph(
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test, '
    'followed by post hoc Mann\u2013Whitney U tests with Bonferroni correction. '
    'For Phase 2, five regression models were fitted:')
doc.add_paragraph('Model 1: Unadjusted regional comparison', style='List Bullet')
doc.add_paragraph('Model 2: Neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + region',
                  style='List Bullet')
doc.add_paragraph('Model 3: Core neuropathic drugs (pregabalin + mirogabalin) ~ same confounders',
                  style='List Bullet')
doc.add_paragraph('Model 4: Nerve blocks ~ same confounders', style='List Bullet')
doc.add_paragraph('Model 5: Neuropathic pain ~ acute analgesic index + confounders (integrated)',
                  style='List Bullet')
doc.add_paragraph(
    'The adjusted CPSP index was derived as residuals from regressing neuropathic pain prescribing '
    'on the four confounder proxies. '
    'Pearson and Spearman correlations assessed inter-index relationships. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11, matplotlib 3.8).'
)

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima)\u2014a 1.97-fold difference '
    '(Kruskal\u2013Wallis P < 0.001 across nine regions).'
)

doc.add_paragraph(
    'Substantial regional clustering was observed. Tokai and Kinki (western Japan) had the lowest '
    'indices, while Kyushu-Okinawa and Hokkaido had the highest (Table 1). '
    'Tohoku, culturally perceived as Japan\u2019s most stoic region, '
    'ranked seventh of nine with a mean index of 39.97 (SD 3.53), '
    'significantly above the non-Tohoku mean of 35.17 '
    '(Mann\u2013Whitney U = 190, P = 0.031; Cohen\u2019s d = 0.87). '
    'All six Tohoku prefectures ranked in the upper half nationally. '
    'This pattern was consistent across drug classes: NSAIDs (P = 0.044), '
    'opioid alkaloids (P = 0.003), and synthetic opioids (P = 0.001).'
)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
doc.add_paragraph(
    f'Nationally, outpatient neuropathic pain drug prescriptions totaled 2,289,549,163 units, '
    f'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    f'duloxetine (15.3%), and tramadol (4.9%). '
    f'Tohoku had a markedly higher index ({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\u2019s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    f'with Iwate (566.7), Aomori (519.3), and Akita (461.1) occupying the top three nationally (Fig. 1).'
)

add_heading_text('Confounder analysis', level=2)
doc.add_paragraph(
    f'Neuropathic pain prescribing showed strong correlations with confounder proxies. '
    f'Diabetes drug prescribing was the strongest (r = 0.87, P < 0.001), '
    f'followed by anxiolytics (r = 0.75), antidepressants (r = 0.46), '
    f'and herpes antivirals (r = 0.19). '
    f'These collectively explained 80.4% of variance '
    f'(R\u00b2 = {reg["model2_adjusted"]["R2"]:.3f} in Model 2; Fig. 2).'
)

add_heading_text('Confounder-adjusted analysis', level=2)
doc.add_paragraph(
    f'After adjustment for all four confounders, '
    f'the Tohoku effect was attenuated and became nonsignificant in Model 2 '
    f'(\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, P = {reg["model2_adjusted"]["tohoku_p"]:.3f}). '
    f'This was consistent across specifications: '
    f'Model 3 (\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}), '
    f'Model 4 (P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}), and '
    f'Model 5 (\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}; Table 2).'
)

doc.add_paragraph(
    f'The adjusted CPSP index showed a dramatically different pattern from the unadjusted data (Fig. 3). '
    f'The Tohoku mean shifted from markedly positive to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}). '
    f'Chugoku emerged as having the highest adjusted index, while Tokai had the lowest (Fig. 4).'
)

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
doc.add_paragraph(
    f'Acute perioperative prescribing correlated positively with unadjusted neuropathic pain prescribing '
    f'(r = 0.38, P = 0.008; Fig. 5a). After confounder adjustment, this was attenuated '
    f'(r = 0.29, P = 0.052; Fig. 5b). '
    f'In Model 5, acute pain index remained a significant predictor '
    f'(\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, P = {reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, P = {reg["model5_integrated"]["tohoku_p"]:.3f}). '
    f'A comprehensive Z-score heatmap confirmed the heterogeneous multi-variable pattern (Supplementary Fig. 1). '
    f'After adjustment, the Tohoku effect was attenuated by 62% and became nonsignificant (Table 3).'
)

doc.add_page_break()

# ============================================================
# DISCUSSION (target: <=1500 words INCLUDING conclusion)
# ============================================================
add_heading_text('Discussion', level=1)

disc_parts = []

d1 = (
    f'This study is the first to map perioperative and chronic pain-related prescribing across '
    f'all 47 prefectures of Japan using NDB Open Data. Three principal findings emerged.'
)
doc.add_paragraph(d1)
disc_parts.append(d1)

add_heading_text('Regional variation within a stoic culture', level=2)
d2 = (
    f'Despite Japan\u2019s well-documented cultural stoicism {cite("Hobara","Feng")}, '
    f'we found a 1.97-fold variation in acute perioperative prescribing and significant differences '
    f'across regional blocks. This parallels Cohen and Nisbett\u2019s finding that '
    f'within-country cultural norms produce measurably different behavioral outcomes {cite("Cohen")}. '
    f'Japan\u2019s pain culture is not monolithic; regional demographics, healthcare infrastructure, '
    f'and local clinical practices generate heterogeneity beneath a nationally shared norm.'
)
doc.add_paragraph(d2)
disc_parts.append(d2)

add_heading_text('Clinical implications of monolithic stereotypes', level=2)
d3 = (
    f'The within-Japan variation has direct clinical relevance. Anderson et al. showed that '
    f'ethnic minorities in the United States consistently receive less adequate pain management '
    f'{cite("Anderson")}. Campbell and Edwards identified that clinician expectations about '
    f'cultural pain behavior lead to systematic under- or over-treatment {cite("Campbell")}. '
    f'Rogger et al. emphasized that cultural framing affects both reporting and management '
    f'{cite("Rogger")}. If a 1.97-fold difference exists within Japan, the international '
    f'label \u201cJapanese patient = stoic\u201d is an oversimplification that risks inadequate analgesia '
    f'for Japanese patients treated abroad.'
)
doc.add_paragraph(d3)
disc_parts.append(d3)

d4 = (
    f'The homogeneity myth reinforces this risk. Befu\u2019s analysis of nihonjinron ideology '
    f'{cite("Befu")} and Burgess\u2019s critique of the \u201cillusion\u201d of homogeneous Japan '
    f'{cite("Burgess")} reveal that internal diversity is systematically underappreciated. '
    f'Our data provide empirical support: there is no monolithic \u201cJapanese patient\u201d\u2014'
    f'there are 47 prefectures with distinct prescribing patterns. '
    f'This principle extends to every cultural label applied to any patient population '
    f'{cite("Raja")}; pain is an individual experience that defies categorical treatment.'
)
doc.add_paragraph(d4)
disc_parts.append(d4)

add_heading_text('Confounders explain Tohoku\u2019s apparent excess', level=2)
d5 = (
    f'The dramatic regional variation in neuropathic pain prescribing '
    f'(unadjusted d = {reg["model1_unadjusted"]["cohens_d"]:.2f}) was largely explained '
    f'by confounding diseases. Diabetes prescribing alone correlated at r = 0.87 with neuropathic '
    f'pain prescribing, reflecting high prevalence of diabetic neuropathy requiring gabapentinoids. '
    f'This has important implications for ecological pain research: studies using neuropathic pain '
    f'drug prescribing as a CPSP proxy must account for confounding diseases. '
    f'The within-database adjustment demonstrated here provides a replicable framework.'
)
doc.add_paragraph(d5)
disc_parts.append(d5)

add_heading_text('A population-level acute\u2013chronic pain continuum', level=2)
d6 = (
    f'The positive correlation between Phase 1 and adjusted Phase 2 indices '
    f'(r = 0.29, P = 0.052) suggests a modest link between acute management intensity '
    f'and chronic pain-related prescribing, consistent with individual-level evidence that '
    f'acute pain intensity is a risk factor for CPSP {cite("Kehlet")}. '
    f'Objective nociception monitoring may help standardize perioperative assessment '
    f'regardless of cultural background {cite("Onishi")}.'
)
doc.add_paragraph(d6)
disc_parts.append(d6)

add_heading_text('Strengths and limitations', level=2)
d7 = (
    f'Strengths include population-complete data covering all insurance-reimbursed healthcare in Japan, '
    f'the integration of acute and chronic pain proxies, transparent confounder adjustment, and '
    f'the perioperative focus that neutralizes access-to-care heterogeneity. '
    f'Limitations include the ecological design (prefecture-level, not individual-level analysis; '
    f'ecological fallacy cannot be excluded), absence of diagnosis codes in NDB Open Data, '
    f'inability to capture fibromyalgia separately, cross-sectional design precluding temporal '
    f'sequence inference, and unmeasured confounders such as age distribution, surgical case mix, '
    f'and prescribing culture.'
)
doc.add_paragraph(d7)
disc_parts.append(d7)

add_heading_text('Conclusion', level=2)
d8 = (
    f'Despite Japan\u2019s norm of pain endurance, perioperative and chronic pain-related prescribing '
    f'varies up to 1.97-fold across prefectures. Confounding diseases substantially modify '
    f'the apparent regional pattern. These findings demonstrate that Japan\u2019s pain culture is not '
    f'monolithic. Treating \u201cJapanese\u201d as a uniform category for pain behavior risks '
    f'inadequate analgesia\u2014and the same principle applies to every cultural label '
    f'applied to any patient population. Individualized pain assessment should replace '
    f'culturally stereotyped assumptions to ensure equitable perioperative care.'
)
doc.add_paragraph(d8)
disc_parts.append(d8)

disc_total = sum(wc(p) for p in disc_parts)
print(f'Discussion word count (incl. conclusion): {disc_total}')

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgments', level=1)
doc.add_paragraph(
    'The authors thank the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. '
    'Parts of data processing and manuscript preparation were assisted by generative AI (Claude, Anthropic). '
    'The authors take full responsibility for the accuracy and content of the manuscript.')

doc.add_paragraph()
add_heading_text('Conflict of interest statement', level=1)
doc.add_paragraph('The authors have no conflicts of interest to declare.')

doc.add_paragraph()
add_heading_text('Data availability statement', level=1)
doc.add_paragraph(
    f'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    f'Labour and Welfare website (https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    f'Analysis code is available at https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan.'
)

# ============================================================
# REFERENCES (alphabetical, bracketed)
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)
for i, (key, ref_text) in enumerate(refs_sorted, 1):
    doc.add_paragraph(f'[{i}] {ref_text}')

# ============================================================
# FIGURE LEGENDS (after references, before tables)
# ============================================================
doc.add_page_break()
add_heading_text('Figure legends', level=1)

figure_legends = [
    ('Fig. 1.', 'Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
     'Bars represent total neuropathic pain prescriptions (pregabalin + mirogabalin + duloxetine + '
     'tramadol + neurotropin) divided by inpatient surgical procedure count. '
     'Tohoku prefectures (red bars with red borders) cluster at the high end. Dashed line = national mean.'),

    ('Fig. 2.', 'Correlation between outpatient neuropathic pain drug prescribing and confounder disease proxies. '
     'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
     'Diabetes drugs show the strongest correlation (r = 0.87).'),

    ('Fig. 3.', 'Confounder-adjusted CPSP index by prefecture. '
     'Residuals from regressing neuropathic pain prescribing on diabetes drugs, herpes antivirals, '
     'antidepressants, and anxiolytics. Positive values indicate higher neuropathic pain prescribing '
     'than expected given confounding disease prevalence. Tohoku prefectures (red borders) are dispersed '
     'across the distribution after adjustment.'),

    ('Fig. 4.', 'Regional comparison of neuropathic pain prescribing: (a) unadjusted and (b) after '
     'confounder adjustment. Tohoku (red border) shifts from the highest region to mid-range after '
     'adjustment. Error bars represent standard deviation.'),

    ('Fig. 5.', 'Integration of Phase 1 (acute perioperative analgesic prescribing) and Phase 2 '
     '(outpatient neuropathic pain prescribing as CPSP proxy). '
     '(a) Unadjusted: positive correlation (r = 0.38, P = 0.008). '
     '(b) Confounder-adjusted: attenuated correlation (r = 0.29, P = 0.052). '
     'Tohoku prefectures (red borders) cluster in the upper-right quadrant.'),

    ('Supplementary Fig. 1.', 'Z-score heatmap of all indices by prefecture. '
     'Each row represents a variable; each column represents a prefecture, '
     'sorted by neuropathic pain prescribing. Red = above average; blue = below average. '
     'Tohoku prefectures are marked with red vertical lines. '
     'The adjusted CPSP index shows a different pattern from the raw neuropathic index.'),
]

for label, legend in figure_legends:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    p.add_run(legend)

# ============================================================
# TABLES (at the END, after references and figure legends)
# ============================================================
doc.add_page_break()
add_heading_text('Tables', level=1)

# --- Table 1 ---
add_bold_paragraph('Table 1. ', 'Phase 1: Inpatient analgesic-per-surgery index by regional block.')
doc.add_paragraph()
table1 = doc.add_table(rows=11, cols=5)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table1)
for i, h in enumerate(['Regional block', 'n', 'Mean (SD)', 'Range', 'Rank']):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data1 = [
    ['Tokai', '4', '27.62 (2.18)', '25.20\u201330.07', '1'],
    ['Kinki', '6', '30.02 (2.03)', '27.92\u201332.33', '2'],
    ['Kanto', '7', '33.00 (2.09)', '29.82\u201334.78', '3'],
    ['Hokuriku-Koshinetsu', '6', '35.38 (3.54)', '31.18\u201340.08', '4'],
    ['Chugoku', '5', '35.73 (4.18)', '31.01\u201340.17', '5'],
    ['Shikoku', '4', '36.33 (3.27)', '33.49\u201341.02', '6'],
    ['Tohoku', '6', '39.97 (3.53)', '35.18\u201344.51', '7'],
    ['Kyushu-Okinawa', '8', '42.26 (4.33)', '35.82\u201349.75', '8'],
    ['Hokkaido', '1', '46.12 (\u2014)', '46.12', '9'],
    ['National', '47', '35.78 (5.56)', '25.20\u201349.75', '\u2014'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
# Bold Tohoku row
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# --- Table 2 ---
add_bold_paragraph('Table 2. ', 'Regression models for neuropathic pain prescribing: '
    'effect of confounder adjustment.')
doc.add_paragraph()
table2 = doc.add_table(rows=7, cols=5)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table2)
for i, h in enumerate(['Model', 'Dependent variable', 'Tohoku effect', 'P value', 'Interpretation']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
t2_data = [
    ['1 (unadjusted)', 'Neuropathic/surgery',
     f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}',
     f'{reg["model1_unadjusted"]["p_value"]:.4f}', 'Significant ***'],
    ['2 (all confounders)', 'Neuropathic/surgery',
     f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['3 (core drugs)', 'PGB + MGB/surgery',
     f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['4 (nerve blocks)', 'Blocks/surgery',
     f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['5 (integrated)', 'Neuropathic/surgery',
     f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['Adj. CPSP index', 'Residuals',
     f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}',
     f'{reg["adjusted_cpsp_test"]["p_value"]:.3f}', 'Nonsignificant'],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# --- Table 3 ---
add_bold_paragraph('Table 3. ', 'Summary of regression models: the Tohoku effect on neuropathic pain '
    'prescribing before and after confounder adjustment.')
doc.add_paragraph()
table3 = doc.add_table(rows=7, cols=5)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table3)
for i, h in enumerate(['Model', 'Dependent variable', 'Tohoku coeff. / effect size', 'P value', 'Result']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
t3_data = [
    ['Model 1', 'Neuropathic pain drugs / surgery (unadjusted)',
     f'd = {reg["model1_unadjusted"]["cohens_d"]:.3f}',
     f'{reg["model1_unadjusted"]["p_value"]:.4f}', '***'],
    ['Model 2', 'Neuropathic pain drugs / surgery (fully adjusted)',
     f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 3', 'Core neuropathic drugs (PGB+MGB) (fully adjusted)',
     f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 4', 'Nerve blocks / surgery (fully adjusted)',
     f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 5', 'Neuropathic pain drugs (acute + confounder adj.)',
     f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.4f}', 'ns'],
    ['Adj CPSP', 'Confounder-removed residual',
     f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.3f}',
     f'{reg["adjusted_cpsp_test"]["p_value"]:.4f}', 'ns'],
]
for r, row_data in enumerate(t3_data):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val
for cell in table3.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
p_note = doc.add_paragraph('*** p < 0.001; ns = not significant')
p_note.runs[0].font.size = Pt(9)
p_note.runs[0].font.italic = True

# ============================================================
# SAVE
# ============================================================
outpath = OUTPUT_DIR + 'Pain_manuscript_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')

# Print reference number mapping for verification
print('\nReference number mapping (alphabetical):')
for i, (key, ref_text) in enumerate(refs_sorted, 1):
    print(f'  [{i}] {key}: {ref_text[:60]}...')
