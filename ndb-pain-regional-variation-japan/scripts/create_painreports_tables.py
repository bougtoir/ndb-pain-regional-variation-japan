#!/usr/bin/env python3
"""Create separate table DOCX files for PAIN Reports submission.
PAIN Reports requires tables as separate attachments, NOT embedded in manuscript.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import csv

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

# Load data
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

rows = []
with open(OUTPUT_DIR + 'cpsp_integrated_results.csv', 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

# Region mapping
REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道','東北','関東','北陸・甲信越','東海','近畿','中国','四国','九州・沖縄']

import numpy as np
from collections import defaultdict

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

def make_table_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.5
    return doc


# ============================================================
# TABLE 1: Regional summary of Phase 1
# ============================================================
doc1 = make_table_doc()
h = doc1.add_heading('Table 1', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
p = doc1.add_paragraph(
    'Phase 1: Regional summary of inpatient analgesic prescribing per surgery across nine regional blocks.')
p.runs[0].font.italic = True

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

table1 = doc1.add_table(rows=1 + len(REGION_ORDER) + 1, cols=5)
set_table_borders(table1)
headers = ['Region', 'N prefectures', 'Mean', 'SD', 'Range']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = table1.rows[idx + 1]
    row.cells[0].text = REGION_EN[reg_name]
    row.cells[1].text = str(len(vals))
    row.cells[2].text = f'{np.mean(vals):.2f}'
    row.cells[3].text = f'{np.std(vals):.2f}'
    row.cells[4].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# National row
all_vals = [r['acute_analgesic_per_surgery'] for r in rows]
nat_row = table1.rows[len(REGION_ORDER) + 1]
nat_row.cells[0].text = 'National'
nat_row.cells[1].text = str(len(rows))
nat_row.cells[2].text = f'{np.mean(all_vals):.2f}'
nat_row.cells[3].text = f'{np.std(all_vals):.2f}'
nat_row.cells[4].text = f'{min(all_vals):.2f}\u2013{max(all_vals):.2f}'
for cell in nat_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

doc1.add_paragraph()
note = doc1.add_paragraph(
    'Values represent the analgesic-per-surgery index: total inpatient analgesic prescription units '
    'divided by total inpatient surgical procedure count for each prefecture. '
    'Kruskal\u2013Wallis test across nine regions: P < 0.001.')
note.runs[0].font.size = Pt(8)
note.runs[0].font.italic = True

outpath1 = OUTPUT_DIR + 'PainReports_Table1.docx'
doc1.save(outpath1)
print(f'Saved: {outpath1}')


# ============================================================
# TABLE 2: Regression models
# ============================================================
doc2 = make_table_doc()
h = doc2.add_heading('Table 2', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
p = doc2.add_paragraph(
    'Phase 2: Regression models for outpatient neuropathic pain prescribing with Tohoku indicator and confounder adjustment.')
p.runs[0].font.italic = True

table2 = doc2.add_table(rows=7, cols=5)
set_table_borders(table2)
t2_headers = ['Model', 'Dependent variable', 'Tohoku coefficient / effect size', 'P value', 'Significance']
for i, h in enumerate(t2_headers):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t2_data = [
    ['Model 1', 'Unadjusted t-test',
     f'd = {reg["model1_unadjusted"]["cohens_d"]:.3f}',
     f'{reg["model1_unadjusted"]["p_value"]:.2e}', '***'],
    ['Model 2', 'Neuropathic pain drugs / surgery (fully adjusted)',
     f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 3', 'Core neuropathic drugs (PGB+MGB) (fully adjusted)',
     f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 4', 'Nerve blocks / surgery (fully adjusted)',
     f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.4f}', 'ns'],
    ['Model 5', 'Neuropathic pain drugs (acute + confounder adj.)',
     f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.4f}', 'ns'],
    ['Adj CPSP', 'Confounder-removed residual',
     f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.3f}',
     f'{reg["adjusted_cpsp_test"]["p_value"]:.4f}', 'ns'],
]
for r_idx, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r_idx + 1].cells[c].text = val
        for paragraph in table2.rows[r_idx + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Bold Model 1 row
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc2.add_paragraph()
note = doc2.add_paragraph('*** P < 0.001; ns = not significant. '
    'Models 2\u20135: multiple linear regression with Tohoku indicator (binary) and confounder disease proxies. '
    'Adj CPSP: adjusted CPSP index derived as residuals from regressing neuropathic pain prescribing on four confounder proxies.')
note.runs[0].font.size = Pt(8)
note.runs[0].font.italic = True

outpath2 = OUTPUT_DIR + 'PainReports_Table2.docx'
doc2.save(outpath2)
print(f'Saved: {outpath2}')


# ============================================================
# TABLE 3: Confounder adjustment effect on Tohoku
# ============================================================
doc3 = make_table_doc()
h = doc3.add_heading('Table 3', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
p = doc3.add_paragraph(
    'Effect of confounder adjustment on Tohoku regional indicators.')
p.runs[0].font.italic = True

table3 = doc3.add_table(rows=4, cols=5)
set_table_borders(table3)
t3_headers = ['Metric', 'Unadjusted', 'Confounder-adjusted', 'Change', 'Interpretation']
for i, h in enumerate(t3_headers):
    table3.rows[0].cells[i].text = h
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

# Calculate attenuation
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

t3_data = [
    ["Cohen's d (Tohoku vs rest)",
     f'{unadj_d:.3f} (P < 0.001)',
     f'{adj_d:.3f} (P = {reg["adjusted_cpsp_test"]["p_value"]:.3f})',
     f'{attenuation:.0f}% attenuation',
     'Large \u2192 Small effect'],
    ['Tohoku mean index',
     f'{reg["model1_unadjusted"]["tohoku_mean"]:.1f}',
     f'{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} (residual)',
     '\u2014',
     'Excess largely explained by confounders'],
    ['Non-Tohoku mean index',
     f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}',
     f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f} (residual)',
     '\u2014',
     'Reference group'],
]

for r_idx, row_data in enumerate(t3_data):
    for c, val in enumerate(row_data):
        table3.rows[r_idx + 1].cells[c].text = val
        for paragraph in table3.rows[r_idx + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc3.add_paragraph()
note = doc3.add_paragraph(
    'Confounders: oral hypoglycemic agents (diabetes proxy), herpes zoster antivirals, '
    'antidepressants (excluding duloxetine), and anxiolytics. '
    'Adjustment reduced the Tohoku effect by 62% and rendered it nonsignificant.')
note.runs[0].font.size = Pt(8)
note.runs[0].font.italic = True

outpath3 = OUTPUT_DIR + 'PainReports_Table3.docx'
doc3.save(outpath3)
print(f'Saved: {outpath3}')

print('\nAll three table files created as separate DOCX.')
