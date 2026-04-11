#!/usr/bin/env python3
"""Create STROBE checklist for the Pain manuscript as English DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# Page setup - A4 landscape
for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(9)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'STROBE Statement \u2014 Checklist of items that should be included '
    'in reports of cross-sectional studies')
run.bold = True
run.font.size = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    '\u201cJapanese Patient\u201d? \u2014 A Patient.: Regional Heterogeneity '
    'in Pain-Related Prescribing Challenges the Stereotype of a Stoic Monolith')
run2.italic = True
run2.font.size = Pt(10)

doc.add_paragraph()

# STROBE items for cross-sectional studies
# (Item number, Recommendation, Section in manuscript, How addressed)
# Section headers use (section_name, None, None, None)
strobe_items = [
    # --- Title and abstract ---
    ('Title and abstract', None, None, None),
    ('1',
     '(a) Indicate the study\u2019s design with a commonly used term in the title or the abstract',
     'Title page and Abstract',
     'Title includes \u201cRegional Heterogeneity in Pain-Related Prescribing\u201d; '
     'Abstract states \u201cecological cross-sectional study using NDB Open Data\u201d'),
    ('1',
     '(b) Provide in the abstract an informative and balanced summary of what was done and what was found',
     'Abstract',
     'Unstructured abstract covers objective, methods (NDB Open Data, 47 prefectures, '
     '5 regression models), key results (1.97-fold variation, Tohoku paradox, confounder attenuation), '
     'and conclusion'),

    # --- Introduction ---
    ('Introduction', None, None, None),
    ('2',
     'Explain the scientific background and rationale for the investigation being reported',
     'Introduction, paragraphs 1\u20134',
     'Cultural influences on pain perception [1\u20134]; Japanese stoicism evidence [5\u20137]; '
     'within-country variation rationale [8\u201311]; NDB as data source [12]'),
    ('3',
     'State specific objectives, including any prespecified hypotheses',
     'Introduction, final paragraph',
     'Three objectives stated: (1) map acute perioperative prescribing, '
     '(2) construct confounder-adjusted CPSP proxy, (3) integrate acute and chronic phases'),

    # --- Methods ---
    ('Methods', None, None, None),
    ('4',
     'Present key elements of study design early in the paper',
     'Methods, paragraph 1',
     '\u201cEcological cross-sectional study\u201d using NDB Open Data 10th edition '
     '(April 2023\u2013March 2024)'),
    ('5',
     'Describe the setting, locations, and relevant dates, including periods of recruitment, '
     'exposure, follow-up, and data collection',
     'Methods, \u201cData source\u201d',
     'NDB Open Data 10th edition; April 2023\u2013March 2024; all 47 prefectures of Japan; '
     'population-complete insurance claims data'),
    ('6',
     '(a) Cross-sectional study \u2014 Give the eligibility criteria, and the sources and methods '
     'of selection of participants',
     'Methods, \u201cData source\u201d',
     'All insurance-reimbursed prescriptions and surgical procedures in NDB; '
     'prefecture-level aggregation (n = 47 units)'),
    ('7',
     'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
     'Methods, \u201cPhase 1 index\u201d through \u201cStatistical analysis\u201d',
     'Outcome: neuropathic pain drug prescribing per surgery. '
     'Exposure: prefecture/region (is_Tohoku). '
     'Confounders: diabetes drugs, herpes zoster drugs, antidepressants, anxiolytics per surgery. '
     'Five regression models specified'),
    ('8',
     'For each variable of interest, give sources of data and details of methods of assessment (measurement)',
     'Methods, \u201cPhase 1 index\u201d and \u201cPhase 2 index\u201d',
     'Phase 1: inpatient analgesic prescriptions (drug classification codes specified) \u00f7 surgical procedures. '
     'Phase 2: outpatient neuropathic pain drugs (pregabalin, mirogabalin, duloxetine, tramadol, neurotropin) '
     '\u00f7 surgeries. Confounder proxies: disease-specific drug prescriptions per surgery'),
    ('9',
     'Describe any efforts to address potential sources of bias',
     'Methods, \u201cStatistical analysis\u201d',
     'Multiple regression models adjusting for confounding diseases; '
     'residual-based adjusted CPSP index; sensitivity analyses with core neuropathic drugs only (Model 3) '
     'and nerve blocks (Model 4)'),
    ('10',
     'Explain how the study size was arrived at',
     'Methods / Results',
     'Population-complete: all insurance-reimbursed healthcare in Japan during study period. '
     'N = 47 prefectures (ecological units). 7,903,515 surgical procedures; '
     '274,579,851 analgesic prescription units'),
    ('11',
     'Explain how quantitative variables were handled in the analyses',
     'Methods, \u201cStatistical analysis\u201d',
     'Per-surgery indices (continuous); is_Tohoku (binary); '
     'Pearson/Spearman correlations; OLS regression; Mann\u2013Whitney U; '
     'Cohen\u2019s d with pooled SD; adjusted CPSP index as regression residuals'),
    ('12',
     '(a) Describe all statistical methods, including those used to control for confounding',
     'Methods, \u201cStatistical analysis\u201d',
     'Five OLS regression models specified; Mann\u2013Whitney U for Tohoku vs rest; '
     'Kruskal\u2013Wallis across 9 regions; Pearson/Spearman correlations; '
     'Cohen\u2019s d effect size; residual-based confounder adjustment'),
    ('12',
     '(b) Describe any methods used to examine subgroups and interactions',
     'Methods',
     'Regional block comparisons (9 blocks); '
     'Tohoku vs non-Tohoku subgroup analysis; drug class-specific analyses (NSAIDs, opioids)'),
    ('12',
     '(c) Explain how missing data were addressed',
     'Methods',
     'NDB Open Data is population-complete with no missing prefectures. '
     'Cell suppression (<1,000 cases) noted as limitation'),
    ('12',
     '(d) Cross-sectional study \u2014 If applicable, describe analytical methods taking account '
     'of sampling strategy',
     'N/A',
     'Census data (population-complete), not a sample; no sampling weights needed'),
    ('12',
     '(e) Describe any sensitivity analyses',
     'Methods, Models 3\u20135',
     'Model 3 (core neuropathic drugs only), Model 4 (nerve blocks), '
     'Model 5 (integrated with Phase 1 acute index)'),

    # --- Results ---
    ('Results', None, None, None),
    ('13',
     '(a) Report numbers of individuals at each stage of study',
     'Results, paragraph 1',
     '7,903,515 inpatient surgical procedures; 274,579,851 analgesic prescription units; '
     '47 prefectures'),
    ('13',
     '(b) Give reasons for non-participation at each stage',
     'N/A',
     'Population-complete ecological data; no individual-level participation/exclusion'),
    ('13',
     '(c) Consider use of a flow diagram',
     'N/A',
     'Not applicable for ecological study using aggregate open data'),
    ('14',
     '(a) Give characteristics of study participants and information on exposures and potential confounders',
     'Results, Table 1; Results paragraphs',
     'Table 1: regional block characteristics (n, mean, SD, range). '
     'Confounder correlations reported (diabetes r = 0.87, anxiolytics r = 0.75, etc.)'),
    ('14',
     '(b) Indicate number of participants with missing data for each variable of interest',
     'N/A',
     'No missing data \u2014 NDB provides complete prefecture-level aggregates'),
    ('15',
     'Report numbers of outcome events or summary measures',
     'Results',
     'National mean analgesic index 35.78 (SD 5.56); range 25.20\u201349.75; '
     'Tohoku mean neuropathic index reported with SD; all model coefficients and P values in Table 2'),
    ('16',
     '(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision. '
     'Make clear which confounders were adjusted for and why',
     'Results; Table 2; Table 3',
     'Unadjusted: Cohen\u2019s d = 2.81 (P < 0.001). '
     'Adjusted models 2\u20135: \u03b2 coefficients with P values. '
     'Adjusted CPSP index: d = 0.44 (P = 0.268). '
     'Confounders specified: diabetes, herpes zoster, antidepressants, anxiolytics'),
    ('16',
     '(b) Report category boundaries when continuous variables were categorized',
     'N/A',
     'Continuous variables not categorized; binary Tohoku grouping defined by geographic region'),
    ('16',
     '(c) If relevant, consider translating estimates of relative risk into absolute risk '
     'for a meaningful time period',
     'N/A',
     'Ecological prescribing indices; not relative/absolute risk estimates'),
    ('17',
     'Report other analyses done \u2014 e.g., analyses of subgroups and interactions, '
     'and sensitivity analyses',
     'Results',
     'Drug class-specific analyses (NSAIDs P = 0.044, opioid alkaloids P = 0.003, '
     'synthetic opioids P = 0.001); Phase 1\u2013Phase 2 integration '
     '(r = 0.38 unadjusted, r = 0.29 adjusted); Supplementary Fig. 1 heatmap'),

    # --- Discussion ---
    ('Discussion', None, None, None),
    ('18',
     'Summarise key results with reference to study objectives',
     'Discussion, paragraphs 1\u20132',
     'Three principal findings: (1) 1.97-fold within-Japan variation, '
     '(2) confounders explain Tohoku excess, (3) modest acute\u2013chronic continuum'),
    ('19',
     'Discuss limitations of the study, taking into account sources of potential bias or imprecision',
     'Discussion, \u201cStrengths and limitations\u201d',
     'Ecological fallacy; no diagnosis codes in NDB; drug proxy limitations; '
     'cross-sectional design; unmeasured confounders; cell suppression'),
    ('20',
     'Give a cautious overall interpretation of results considering objectives, limitations, '
     'multiplicity of analyses, results from similar studies, and other relevant evidence',
     'Discussion',
     'Exploratory/hypothesis-generating framing throughout; comparison with Wakaizumi [13] '
     'and Matsuoka [14]; cultural stereotype critique grounded in Anderson [17], Campbell [18], '
     'Befu [19], Burgess [20]'),
    ('21',
     'Discuss the generalisability (external validity) of the study results',
     'Discussion, \u201cClinical implications\u201d and \u201cImplications and future directions\u201d',
     'Generalizability to international clinical settings discussed; '
     'argument extended beyond Japan to all cultural stereotypes; '
     'future individual-level studies recommended'),

    # --- Other information ---
    ('Other information', None, None, None),
    ('22',
     'Give the source of funding and the role of the funders for the present study and, '
     'if applicable, for the original study on which the present article is based',
     'Acknowledgments / Conflict of interest',
     'No external funding declared; no conflicts of interest; '
     'NDB Open Data is publicly available from MHLW'),
]

# Create table with header row
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)

# Header row
headers = ['Item No.', 'Recommendation', 'Reported on page / section',
           'How addressed in manuscript']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '4472C4')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
for item in strobe_items:
    row_cells = table.add_row().cells
    if item[1] is None:
        # Section header row - merge all columns
        row_cells[0].merge(row_cells[3])
        row_cells[0].text = item[0]
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        row_cells[0]._tc.get_or_add_tcPr().append(shading)
    else:
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    if len(row.cells) > 1:
        row.cells[1].width = Cm(10.0)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(10.7)

# Footer reference
doc.add_paragraph()
footer = doc.add_paragraph(
    'Reference: von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
    'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
    'guidelines for reporting observational studies. Lancet. 2007;370:1453\u20131457.')
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.italic = True

outpath = OUTPUT_DIR + 'STROBE_checklist_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
