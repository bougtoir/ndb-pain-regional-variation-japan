#!/usr/bin/env python3
"""Create Japanese manuscript for PAIN Reports format.

Key differences from Pain:
- Research Article: max 4,000 words (body only, excl. abstract/refs/legends)
- Summary: <=25 words (separate file)
- Abstract: <=250 words, structured (Introduction, Methods, Results, Conclusions)
- References: alphabetical by first author, bracketed [n] (same as PAIN)
- Tables: uploaded as SEPARATE attachments, NOT embedded in manuscript
- Figures: TIFF, separate files, legends on separate manuscript page after references
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

# ============================================================
# REFERENCES — alphabetical by first-author surname (PAIN Reports = PAIN style)
# ============================================================
refs_raw = [
    ('Anderson', 'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
     'causes and consequences of unequal care. J Pain 2009;10:1187\u20131204. '
     'doi:10.1016/j.jpain.2009.03.015'),

    ('Befu', 'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
     'Melbourne: Trans Pacific Press, 2001.'),

    ('Benchimol', 'Benchimol EI, Smeeth L, Guttmann A, Harron K, Moher D, Petersen I, '
     'Sorensen HT, von Elm E, Langan SM. '
     'The REporting of studies Conducted using Observational Routinely-collected health Data '
     '(RECORD) statement. PLoS Med 2015;12:e1001885. '
     'doi:10.1371/journal.pmed.1001885'),

    ('Burgess', 'Burgess C. The \u201cillusion\u201d of homogeneous Japan and national character: '
     'discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. '
     'Asia Pac J 2010;8(9):1\u201322.'),

    ('Callister', 'Callister LC. Cultural influences on pain perceptions and behaviors. '
     'Home Health Care Manag Pract 2003;15:207\u2013211. '
     'doi:10.1177/1084822302250687'),

    ('Campbell', 'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
     'Pain Manag 2012;2:219\u2013230. '
     'doi:10.2217/pmt.12.7'),

    ('Cohen', 'Cohen D, Nisbett RE, Bowdle BF, Schwarz N. '
     'Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d '
     'J Pers Soc Psychol 1996;70:945\u2013960. '
     'doi:10.1037/0022-3514.70.5.945'),

    ('Feng', 'Feng Y, Herdman M, van Nooten F, Cleeland C, Parkin D, Ikeda S, Igarashi A, Devlin NJ. '
     'An exploration of differences between Japan and two European countries in the self-reporting '
     'and valuation of pain and discomfort on the EQ-5D. Qual Life Res 2017;26:2067\u20132078. '
     'doi:10.1007/s11136-017-1541-4'),

    ('Hobara', 'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
     'between Japanese and Euro-Americans. Eur J Pain 2005;9:389\u2013393. '
     'doi:10.1016/j.ejpain.2004.09.006'),

    ('Kehlet', 'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
     'Lancet 2006;367:1618\u20131625. '
     'doi:10.1016/S0140-6736(06)68700-X'),

    ('Kumagai', 'Kumagai S. Media representations reproducing images of Tohoku: '
     'the Tohoku reconstruction corner in \u201cSecret Kenmin SHOW.\u201d '
     'Kotoba 2020;41:21\u201338. [in Japanese]'),

    ('Matsuoka', 'Matsuoka Y, Morishima T, Sato A, Ogawa T, Miyashiro I. '
     'Population-based claims study of regional and hospital function differences '
     'in opioid prescribing for cancer patients who died in hospital in Japan. '
     'Jpn J Clin Oncol 2025;55:hyaf149. '
     'doi:10.1093/jjco/hyaf149'),

    ('MHLW', 'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. 2024. '
     'Available: https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
     'Accessed January 15, 2025.'),

    ('Okolo', 'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
     'a review of cross-cultural studies. Int J Sci Res Arch 2024;11:2550\u20132556. '
     'doi:10.30574/ijsra.2024.11.1.0263'),

    ('Onishi', 'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
     'and quantification of nociception anti-nociception balance compared to opioid effect site concentration: '
     'a prospective, observational study. F1000Research 2024;13:233. '
     'doi:10.12688/f1000research.146054.2'),

    ('Pfizer', 'Pfizer Japan Inc. 47-prefecture survey on chronic pain. 2017. '
     'Available: https://www.pfizer.co.jp/pfizer/company/press/2017. '
     'Accessed February 1, 2025.'),

    ('Raja', 'Raja SN, Carr DB, Cohen M, Finnerup NB, Flor H, Gibson S, Keefe FJ, Mogil JS, '
     'Ringkamp M, Sluka KA, Song XJ, Stevens B, Sullivan MD, Tutelman PR, Ushida T, Vader K. '
     'The revised International Association for the Study of Pain definition of pain: '
     'concepts, challenges, and compromises. Pain 2020;161:1976\u20131982. '
     'doi:10.1097/j.pain.0000000000001939'),

    ('Rogger', 'Rogger R, Bello C, Romero CS, Urman RD, Luedi MM, Filipovic MG. '
     'Cultural framing and the impact on acute pain and pain services. '
     'Curr Pain Headache Rep 2023;27:429\u2013436. '
     'doi:10.1007/s11916-023-01130-1'),

    ('Takeda', 'Takeda K, Yarimizu K. Regional differences in the pain expression uzuku. '
     'NINJAL Research Papers 2016;10:85\u2013107. [in Japanese]'),

    ('von Elm', 'von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
     'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
     'guidelines for reporting observational studies. '
     'Lancet 2007;370:1453\u20131457. '
     'doi:10.1016/S0140-6736(07)61602-X'),

    ('Wakaizumi', 'Wakaizumi K, Tanaka C, Shinohara Y, Wu Y, Takaoka S, Kawate M, Oka H, Matsudaira K. '
     'Geographical variation in high-impact chronic pain and psychological associations at the regional level: '
     'a multilevel analysis of a large-scale internet-based cross-sectional survey. '
     'Front Public Health 2024;12:1482177. '
     'doi:10.3389/fpubh.2024.1482177'),

    ('Zborowski', 'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.'),
]

# Sort alphabetically by first-author surname
refs_sorted = sorted(refs_raw, key=lambda x: x[0].lower())
ref_num = {}
for i, (key, _) in enumerate(refs_sorted, 1):
    ref_num[key] = i

def cite(*keys):
    nums = sorted(ref_num[k] for k in keys)
    return '[' + ','.join(str(n) for n in nums) + ']'


# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# Page numbers (footer)
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
title_ja = doc.add_paragraph()
title_ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_ja.add_run(
    '\u201c日本人患者\u201dか？\u2014\u2014ただの患者だ。\n'
    '47都道府県における疼痛関連処方の地域間異質性は\n'
    '「我慢強い単一民族」というステレオタイプに疑義を呈する'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

title_en = doc.add_paragraph()
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_en = title_en.add_run(
    '\u201cJapanese Patient\u201d? \u2014 A Patient.\n'
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures\n'
    'Challenges the Stereotype of a Stoic Monolith'
)
run_en.font.size = Pt(12)
run_en.italic = True

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[著者名・所属を挿入]')
run.font.size = Pt(12)

doc.add_paragraph()

add_bold_paragraph('責任著者：',
    '[氏名]、[所属]麻酔科学教室、[住所]、[郵便番号]、日本　'
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

add_bold_paragraph('キーワード：',
    '文化的疼痛行動; 生態学的研究; 我慢; 神経障害性疼痛; 地域差; 処方パターン')

doc.add_paragraph()

add_bold_paragraph('論文種別：', 'Research Article')
add_bold_paragraph('投稿先：', 'PAIN Reports')
add_bold_paragraph('推奨セクション：', '"Pain around the world" または "Acute and Perioperative"')
add_bold_paragraph('図数：', '5（＋補足1）— 別ファイルとしてアップロード')
add_bold_paragraph('表数：', '3 — 別ファイルとしてアップロード（原稿内に埋め込まない）')

doc.add_page_break()

# ============================================================
# ABSTRACT (structured: Introduction, Methods, Results, Conclusions)
# ============================================================
add_heading_text('抄録', level=1)

abstract_intro = (
    '異文化研究は日本人が痛みに対し忍耐強いことを一貫して示してきたが、'
    'この文化的忍耐が日本国内で地域的に異なるかどうかは未解明であった。'
)

abstract_methods = (
    '本生態学的研究では、NDBオープンデータ'
    '（集団完全保険請求、2023年4月〜2024年3月）を用い、'
    '47都道府県・9地域ブロックの疼痛関連処方を検討した。'
    'Phase 1は急性期周術期鎮痛薬処方、'
    'Phase 2は交絡疾患調整済みの外来神経障害性疼痛薬処方'
    '（術後遷延性疼痛プロキシ）を分析した。'
)

abstract_results = (
    f'Phase 1では、鎮痛薬/手術指標が都道府県間で1.97倍の差を示した'
    f'（岐阜25.20〜鹿児島49.75; Kruskal\u2013Wallis P < 0.001）。'
    f'伝統的に我慢強いとされる東北は鎮痛薬使用がむしろ多かった'
    f'（平均39.97 vs 全国35.78; Cohen\u2019s d = 0.87）。'
    f'Phase 2では、東北の神経障害性疼痛薬処方は未調整で最高値'
    f'（d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）だったが、'
    f'交絡疾患調整後は有意でなくなった'
    f'（P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}）。'
    f'糖尿病薬処方が最強の交絡因子（r = 0.87）であった。'
)

abstract_conclusions = (
    '約2倍の日本国内変動は、「日本人」を疼痛行動の均質なカテゴリーとする仮定に疑義を呈する。'
    '文化的ステレオタイプに基づく仮定ではなく、個別化された疼痛評価が周術期ケアの標準であるべきである。'
)

p = doc.add_paragraph()
r = p.add_run('Introduction: ')
r.bold = True
p.add_run(abstract_intro)

p = doc.add_paragraph()
r = p.add_run('Methods: ')
r.bold = True
p.add_run(abstract_methods)

p = doc.add_paragraph()
r = p.add_run('Results: ')
r.bold = True
p.add_run(abstract_results)

p = doc.add_paragraph()
r = p.add_run('Conclusions: ')
r.bold = True
p.add_run(abstract_conclusions)

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('緒言', level=1)

# Para 1: Cross-cultural pain literature
doc.add_paragraph(
    f'疼痛は普遍的な体験であるが、その表現・忍耐・管理は文化によって深く形成される {cite("Callister","Rogger")}。'
    f'Zborowskiが米国内の民族集団間で疼痛行動が著しく異なることを観察した先駆的研究以来 {cite("Zborowski")}、'
    f'文化的規範が個人の疼痛報告・受療行動・忍耐に影響を及ぼすことを示す膨大な文献が蓄積されてきた {cite("Okolo")}。'
    f'日本は痛みに対して特に忍耐強い傾向を持つと一貫して特徴づけられている。'
    f'Hobaraは、日本人回答者が疼痛行動をヨーロッパ系アメリカ人よりも有意に不適切と評価したことを報告し {cite("Hobara")}、'
    f'Fengらは、EQ-5Dを用いて日本人参加者が痛み回避のための時間犠牲意向が著しく低いことを示した {cite("Feng")}。'
    f'この文化的に媒介された忍耐は「我慢」（\u6211\u6162）の概念に集約される。'
    f'国際的には、このステレオタイプは臨床的リスクを伴う：'
    f'忍耐的な態度を鎮痛薬の必要性が低いと誤解する可能性がある。'
)

# Para 2: Within-country variation
doc.add_paragraph(
    f'文化的疼痛行動に関するほとんどの研究は国家間比較である。'
    f'しかし国内にも相当な文化的異質性が存在する。'
    f'CohenとNisbettは、米国南部の「名誉の文化」が北部とは測定可能に異なる行動を生むことを実証した {cite("Cohen")}。'
    f'日本でも地域文化のアイデンティティは強い：'
    f'東北地方は忍耐強いと認識され {cite("Kumagai")}、'
    f'痛みの擬態語「うずく」ですら地域差がある {cite("Takeda")}。'
    f'ファイザー日本の全国調査では、痛みを「我慢する」と回答した割合が'
    f'都道府県間で48.7%〜81.6%と大きく異なっていた {cite("Pfizer")}。'
    f'しかし、態度の差が疼痛関連医療利用の測定可能な差異として反映されるかは未解明である。'
)

# Para 3: Japan as natural experiment
doc.add_paragraph(
    f'日本は疼痛行動の地域差を研究する上で理想的な環境を提供する。'
    f'国民皆保険制度がアクセスの経済的障壁を排除し、'
    f'NDBがほぼ全ての保険請求医療を収録する {cite("MHLW")}。'
    f'近年の研究は有意な地域差を報告している：'
    f'Wakaizumiらは高インパクト慢性疼痛有病率に最大1.6倍の差を認め {cite("Wakaizumi")}、'
    f'Matsuokaらはがん疼痛オピオイド処方に4倍の地域差を報告した {cite("Matsuoka")}。'
    f'しかし、都道府県レベルでの周術期鎮痛薬処方を検討した研究は存在しない。'
    f'周術期は手術患者が全員入院中であるため、'
    f'医療アクセスが交絡因子とならない方法論上の利点がある。'
)

# Para 4: Objectives
doc.add_paragraph(
    '本探索的研究の目的は3つである。'
    '第一（Phase 1）、47都道府県の急性期周術期鎮痛薬処方の地域差をマッピングすること。'
    '第二（Phase 2）、交絡疾患を調整した上で外来神経障害性疼痛薬処方の地域差を検討すること。'
    '第三、Phase 1とPhase 2の統合により、急性期鎮痛薬使用と慢性疼痛関連処方の関連を探索すること。'
)

# ============================================================
# METHODS
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
doc.add_paragraph(
    f'本研究はNDBオープンデータの都道府県別集計データを分析した生態学的研究である。'
    f'STROBE声明 {cite("von Elm")} およびRECORD拡張 {cite("Benchimol")} に準拠して報告する。'
    f'公開された集計データのみを使用するため、倫理審査は不要である。'
)

add_heading_text('データソース', level=2)
doc.add_paragraph(
    f'NDBオープンデータ第10回（厚生労働省公表）を使用し、'
    f'2023年4月〜2024年3月の診療報酬請求を対象とした {cite("MHLW")}。'
    f'NDBは全保険者からの請求を収録し、約1億2,500万人の被保険者を網羅する。'
)

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県を9つの標準地域ブロックに分類した：'
    '北海道（1）、東北（6）、関東（7）、北陸・甲信越（6）、東海（4）、近畿（6）、'
    '中国（5）、四国（4）、九州・沖縄（8）。'
)

add_heading_text('Phase 1：急性期周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから3つの鎮痛薬分類を抽出した：'
    '薬効分類114（NSAIDs・アセトアミノフェン）、'
    '811（あへんアルカロイド系麻薬）、821（合成麻薬）。'
    '入院手術件数はKコードから抽出し、'
    '各都道府県の鎮痛薬/手術指標＝入院鎮痛薬処方量÷入院手術件数として算出した。'
)

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（CPSPプロキシ）', level=2)
doc.add_paragraph(
    '5種類の外来経口神経障害性疼痛薬を都道府県別に抽出した：'
    'プレガバリン、ミロガバリン、デュロキセチン、トラマドール、ノイロトロピン。'
    '神経障害性疼痛薬/手術指標＝外来処方量÷入院手術件数として算出した。'
)

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '以下の4つの交絡疾患プロキシを外来処方データから抽出した：'
    '経口血糖降下薬（糖尿病性神経障害プロキシ）、'
    '帯状疱疹抗ウイルス薬（帯状疱疹後神経痛プロキシ）、'
    'デュロキセチンを除く抗うつ薬（うつ病プロキシ）、'
    '抗不安薬（不安障害プロキシ）。'
)

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    'Phase 1の地域差はKruskal\u2013Wallis検定で9地域ブロック間を比較した。'
    'Phase 2では5つの回帰モデルを適合した：'
    '未調整比較（モデル1）、全交絡因子調整（モデル2）、'
    'コア神経障害性薬のみ（モデル3）、神経ブロック（モデル4）、'
    '急性期指標＋交絡因子の統合モデル（モデル5）。'
    '調整済CPSP指標は交絡プロキシへの回帰残差として導出した。'
    '全解析はPython 3.11で実施した。'
)

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1：急性期周術期鎮痛薬処方の地域差', level=2)
doc.add_paragraph(
    '7,903,515件の入院手術と274,579,851単位の鎮痛薬処方を分析した。'
    '全国平均の鎮痛薬/手術指標は35.78（SD 5.56）で、'
    '岐阜（25.20）から鹿児島（49.75）まで1.97倍の差があった'
    '（9地域間Kruskal\u2013Wallis P < 0.001）。'
    '東海・近畿が最低、九州・沖縄・北海道が最高であった（Table 1）。'
)

doc.add_paragraph(
    '日本で最も我慢強いとされる東北は9地域中7位であった'
    '（平均39.97 vs 東北以外35.17; Cohen\u2019s d = 0.87; P = 0.031）。'
    '東北6県全てが全国上位半分に位置し、'
    'このパターンは全薬効分類で一貫していた：NSAIDs（P = 0.044）、'
    'あへんアルカロイド（P = 0.003）、合成麻薬（P = 0.001）。'
)

add_heading_text('Phase 2：外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    f'全国の外来神経障害性疼痛薬処方量は合計2,289,549,163単位であった。'
    f'東北は未調整で指標が突出して高かった'
    f'（平均{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; d = {reg["model1_unadjusted"]["cohens_d"]:.2f}; Fig. 1）。'
)

add_heading_text('交絡因子分析と調整後の結果', level=2)
doc.add_paragraph(
    f'神経障害性疼痛薬処方は交絡疾患プロキシと強い相関を示した：'
    f'糖尿病薬（r = 0.87）、抗不安薬（r = 0.75）、抗うつ薬（r = 0.46）（Fig. 2）。'
    f'交絡因子は分散の80.4%を説明した（R\u00b2 = {reg["model2_adjusted"]["R2"]:.3f}）。'
    f'4つの交絡プロキシで調整後、東北効果は減弱し有意でなくなった'
    f'（\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model2_adjusted"]["tohoku_p"]:.3f}; Table 2）。'
    f'この減弱は全モデルで一貫していた（Table 2）。'
)

doc.add_paragraph(
    f'調整済CPSP指標では、東北の過剰は緩やかで非有意となった'
    f'（{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs '
    f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}; '
    f'P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; Fig. 3, Fig. 4; Table 3）。'
)

add_heading_text('Phase 1\u2013Phase 2 統合', level=2)
doc.add_paragraph(
    f'Phase 1の急性期鎮痛薬/手術はPhase 2の未調整指標と正の相関を示した'
    f'（r = 0.38, P = 0.008; Fig. 5a）。'
    f'交絡調整後は減弱した（r = 0.29, P = 0.052; Fig. 5b）。'
    f'統合モデルでは急性期指標は有意な予測因子であったが'
    f'（\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P = {reg["model5_integrated"]["acute_pain_p"]:.3f}）、'
    f'東北効果は有意でなかった'
    f'（P = {reg["model5_integrated"]["tohoku_p"]:.3f}）。'
)

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    '本研究は、NDBオープンデータを用いて日本の47都道府県全てにおける'
    '周術期および慢性疼痛関連処方をマッピングした初の研究である。'
    '主要な知見は3つである。'
)

add_heading_text('忍耐強い文化の中にも相当な地域差が存在する', level=2)
doc.add_paragraph(
    f'日本の痛みに対する文化的忍耐は十分に文献化されているにもかかわらず {cite("Hobara","Feng")}、'
    f'急性期周術期鎮痛薬処方には都道府県間で1.97倍の差があり、'
    f'地域ブロック間にも有意差が認められた。'
    f'CohenとNisbettによる米国の「名誉の文化」研究と類似し {cite("Cohen")}、'
    f'単一国家内でも地域文化規範が測定可能に異なる帰結を生む。'
    f'本研究の知見は日本の疼痛文化が一枚岩でないことを示唆する。'
)

add_heading_text('臨床的含意：一枚岩的な文化的ステレオタイプの危険性', level=2)
doc.add_paragraph(
    f'1.97倍の国内変動は、国内にとどまらない臨床的意義を有する。'
    f'Andersonらは人種的・民族的マイノリティの疼痛管理格差を示し {cite("Anderson")}、'
    f'CampbellとEdwardsは文化的期待が系統的な過少治療につながりうることを指摘した {cite("Campbell")}。'
    f'日本人論の言説は、日本人が均一な行動規範を共有するという概念を広めてきたが {cite("Befu")}、'
    f'この「均質性の覇権」は経験的事実ではなくイデオロギー的構築物であり {cite("Befu")}、'
    f'Burgessはこの「幻想」が社会政策に実質的影響を及ぼすことを論証した {cite("Burgess")}。'
)

doc.add_paragraph(
    '「日本人は忍耐強い」と「日本人は均質である」という2つのステレオタイプの組み合わせが、'
    '二重に誤った仮定を生む：全ての日本人患者が等しく痛みに耐えるという仮定である。'
    '本研究が示した1.97倍の国内変動はこの仮定を直接的に反駁する。'
    '端的に言えば、「日本人患者」なる存在はいない——'
    'いるのは47の多様な都道府県からの個々の患者だけである。'
)

doc.add_paragraph(
    'いかなる国民集団の疼痛行動についても一枚岩的な特徴づけは'
    '臨床的害のリスクを伴う——この主張は日本をはるかに超えて妥当する。'
    '臨床医が文化的ステレオタイプによって治療選択肢の範囲を狭めるとき、'
    '患者は治療上の不利益を被る。'
    '本研究のデータ——単一の文化的ラベルの下にほぼ2倍の国内変動——は、'
    '普遍的原則の具体的事例研究として機能する：'
    '文化的一般化は個別の臨床判断の基盤としては不十分である。'
)

add_heading_text('痛みは個人的体験であり、客観的モニタリングの開発が待たれる', level=2)
doc.add_paragraph(
    f'IASPは痛みを本質的に主観的な体験と定義している {cite("Raja")}。'
    f'単一国家内でさえ疼痛関連処方がほぼ2倍異なるという本研究の知見は、'
    f'いかなる文化的ラベルも患者の侵害受容状態の直接測定の代替にはなりえないことを示す。'
    f'Onishiらは、正規化脈波容積（NPV）が文化に依存しない'
    f'客観的NANB指標となりうることを示唆した {cite("Onishi")}。'
    f'このような客観的侵害受容モニタリングの臨床実装は、'
    f'文化的バイアスに影響されやすい主観的評価への依存を軽減しうる。'
)

add_heading_text('交絡疾患がCPSPプロキシの地域差を説明する', level=2)
doc.add_paragraph(
    f'方法論的に重要な知見として、神経障害性疼痛薬処方の劇的な地域差'
    f'（未調整d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）が'
    f'交絡疾患プロキシで大部分説明された。'
    f'糖尿病薬処方だけでr = 0.87であり、糖尿病性神経障害との関連を反映する。'
    f'調整後、東北効果は62%減弱し有意でなくなった。'
    f'神経障害性疼痛薬処方をCPSPプロキシとして使用する研究は、'
    f'交絡疾患を考慮しなければならない。'
)

add_heading_text('急性期\u2013慢性期疼痛の連続体', level=2)
doc.add_paragraph(
    f'Phase 1とPhase 2の正の相関（r = 0.29, 調整後）は、'
    f'地域の急性期疼痛管理と慢性疼痛関連処方の緩やかな関連を示唆する。'
    f'生態学的相関は因果関係を確立できないが、'
    f'急性期術後疼痛がCPSPの危険因子であるという文献 {cite("Kehlet")} と整合する。'
)

add_heading_text('強みと限界', level=2)
doc.add_paragraph(
    '本研究の強みとして、集団完全データの使用、急性期と慢性期の新規統合、'
    '透明な交絡調整方法論が挙げられる。'
    '周術期に焦点を当てたことで、医療アクセスが交絡因子とならない利点がある。'
)

doc.add_paragraph('主な限界は生態学的デザインに固有のものである：')
doc.add_paragraph('生態学的誤謬——都道府県レベルの相関は個人レベルの関連を反映しない可能性がある。',
                  style='List Bullet')
doc.add_paragraph('NDBオープンデータには傷病名コードがなく、CPSPを直接同定できない。',
                  style='List Bullet')
doc.add_paragraph('横断的デザインでは時間的順序を区別できない。', style='List Bullet')
doc.add_paragraph('測定されない交絡因子（年齢分布、手術構成）が残余の地域差に寄与しうる。',
                  style='List Bullet')

add_heading_text('今後の展望', level=2)
doc.add_paragraph(
    'NDB特別抽出データを用いた今後の研究では、'
    '個人レベルで手術から新規神経障害性疼痛薬処方への縦断的追跡が可能となり、'
    '直接的なCPSP指標を提供しうる。'
    '文化的ステレオタイプが多文化的臨床環境における鎮痛薬処方に'
    'いかに影響するかを検討する前向き研究は、'
    '「忍耐強い日本人」というラベルが治療上の不利益に直結するかどうかを直接検証しうる。'
)

# ============================================================
# CONCLUSION (integrated into Discussion per PAIN Reports style)
# ============================================================
add_heading_text('結論', level=2)
doc.add_paragraph(
    '日本に根ざした疼痛忍耐の文化的規範（我慢）にもかかわらず、'
    '周術期および慢性疼痛関連処方は都道府県間で最大1.97倍異なる。'
    '交絡疾患は神経障害性疼痛薬処方の地域パターンを大きく修飾する。'
    '「日本人」を疼痛行動に関する均質なカテゴリーとして扱うことは、'
    '不十分な鎮痛を招くリスクがある——'
    'そしてこの原則はあらゆる患者集団に適用されるあらゆる文化的ラベルに等しく当てはまる。'
    '個別化された疼痛評価が、公平な周術期ケアの礎石であり続ける。'
)

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph(
    'NDBオープンデータを公開している厚生労働省に感謝する。'
    '［必要に応じて追加の謝辞を挿入］')

# CONFLICT OF INTEREST
doc.add_paragraph()
add_heading_text('利益相反', level=1)
doc.add_paragraph('著者らに開示すべき利益相反はない。')

# DATA AVAILABILITY
doc.add_paragraph()
add_heading_text('データ利用可能性', level=1)
doc.add_paragraph(
    '本研究で使用したNDBオープンデータは厚生労働省ウェブサイトより公開されている。'
    '解析コードはhttps://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japanで利用可能である。')

doc.add_page_break()

# ============================================================
# REFERENCES (alphabetical, bracketed [n])
# ============================================================
add_heading_text('文献', level=1)
for i, (key, ref_text) in enumerate(refs_sorted, 1):
    doc.add_paragraph(f'[{i}] {ref_text}')

doc.add_page_break()

# ============================================================
# FIGURE LEGENDS (separate page, per PAIN Reports requirements)
# ============================================================
add_heading_text('図の凡例', level=1)

legends = [
    ('Fig. 1. ', '都道府県別 外来神経障害性疼痛薬処方量/手術（未調整）。'
     'プレガバリン＋ミロガバリン＋デュロキセチン＋トラマドール＋ノイロトロピンの合計を'
     '入院手術件数で除した。東北（赤棒）が高値側に集中。破線＝全国平均。'),
    ('Fig. 2. ', '外来神経障害性疼痛薬処方と交絡疾患プロキシの相関。'
     '各点は1都道府県。東北は赤で表示。糖尿病薬が最も強い相関（r = 0.87）を示す。'),
    ('Fig. 3. ', '交絡疾患調整後の都道府県別CPSP指標。'
     '神経障害性疼痛薬処方を4つの交絡プロキシに回帰した残差。'
     '正の値は交絡疾患から予測されるより多い処方を示す。'),
    ('Fig. 4. ', '地域ブロック別 神経障害性疼痛薬処方の比較：(a) 未調整 (b) 調整後。'
     '東北は未調整では最高だが、調整後は中位に移動。エラーバー＝標準偏差。'),
    ('Fig. 5. ', 'Phase 1（急性期鎮痛薬/手術）vs Phase 2（CPSPプロキシ）の統合。'
     '(a) 未調整（r = 0.38, P = 0.008）。(b) 調整後（r = 0.29, P = 0.052）。'),
    ('Supplementary Fig. 1. ', '都道府県別 各指標のZ-scoreヒートマップ。'
     '各行は変数、各列は都道府県。赤＝平均以上、青＝平均以下。東北は赤の垂直線で表示。'),
]

for label, text in legends:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)

# ============================================================
# Save
# ============================================================
outpath = OUTPUT_DIR + 'PainReports_manuscript_JA.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
