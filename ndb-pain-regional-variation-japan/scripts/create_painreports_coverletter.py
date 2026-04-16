#!/usr/bin/env python3
"""Create cover letter for PAIN Reports submission."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Date
p = doc.add_paragraph('[Date]')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# Addressee
doc.add_paragraph('David Yarnitsky, MD')
doc.add_paragraph('Editor-in-Chief, PAIN Reports')
doc.add_paragraph('International Association for the Study of Pain')

doc.add_paragraph()

# Subject
p = doc.add_paragraph()
r = p.add_run('Re: ')
r.bold = True
p.add_run('Submission of Research Article \u2014 ')
r2 = p.add_run('\u201cJapanese Patient\u201d? \u2014 A Patient. '
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures '
    'Challenges the Stereotype of a Stoic Monolith')
r2.italic = True

doc.add_paragraph()

doc.add_paragraph('Dear Professor Yarnitsky,')

doc.add_paragraph()

# Body
doc.add_paragraph(
    'We are pleased to submit the above manuscript for consideration as a Research Article '
    'in PAIN Reports. Following editorial review at PAIN, Professor Karen D. Davis '
    'specifically recommended that our work would be well suited for PAIN Reports, '
    'and we are grateful for this guidance.'
)

doc.add_paragraph(
    'This ecological study leverages Japan\u2019s National Database (NDB) Open Data\u2014'
    'population-complete insurance claims covering approximately 125 million insured individuals\u2014'
    'to map pain-related prescribing across all 47 prefectures and nine regional blocks. '
    'Japan is consistently characterized in the cross-cultural pain literature as stoic, '
    'yet whether this endurance varies regionally within the country has never been examined. '
    'Our analysis reveals a 1.97-fold variation in acute perioperative analgesic prescribing '
    'across prefectures and shows that Tohoku\u2014traditionally perceived as Japan\u2019s most stoic '
    'region\u2014actually prescribes more, not fewer, analgesics than the national average.'
)

doc.add_paragraph(
    'We believe this work is particularly well suited for the \u201cPain around the world\u201d '
    'section of PAIN Reports, as it addresses a locally grounded theme\u2014regional variation '
    'within Japan\u2019s pain culture\u2014with implications of wide significance for international '
    'pain management. The central message\u2014that monolithic cultural stereotypes about any '
    'national population\u2019s pain behavior carry a risk of clinical harm\u2014is relevant to '
    'clinicians worldwide who treat patients from diverse cultural backgrounds.'
)

doc.add_paragraph(
    'The manuscript reports original findings that have not been published previously '
    'and is not under consideration elsewhere. All authors meet the ICMJE criteria for authorship '
    'and have approved the final manuscript. There are no conflicts of interest to declare. '
    'Ethical approval was not required as only publicly available aggregate data were used.'
)

doc.add_paragraph(
    'We confirm that all figures are original and have not been manipulated. '
    'The study is reported following the STROBE and RECORD guidelines; '
    'a completed STROBE checklist is provided as a supplementary file.'
)

doc.add_paragraph(
    'Thank you for considering our manuscript. We look forward to your editorial decision.'
)

doc.add_paragraph()

doc.add_paragraph('Sincerely,')
doc.add_paragraph()
doc.add_paragraph('Tatsuki Onishi')
doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('[Address]')
doc.add_paragraph('E-mail: [email]')

outpath = OUTPUT_DIR + 'PainReports_cover_letter.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
