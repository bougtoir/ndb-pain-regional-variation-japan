#!/usr/bin/env python3
"""Create English manuscript for Journal of Anesthesia (JA) format as .docx with embedded color figures.

Key JA formatting requirements:
- American English (anesthesia, analyze, etc.)
- Structured abstract: Purpose, Methods, Results, Conclusion (max 250 words)
- Original article max 4,000 words
- References: Vancouver style, numbered [1], all authors listed
- Sections on separate pages: title page, abstract, text, references, tables, figure legends, figures
- Pages numbered consecutively
- A4 format, double-spaced, wide margins
- COI statement before references
- Acknowledgments between text and references
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import json

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# --- Page setup: A4, wide margins, double-spaced ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 2.0

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=5.5):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        parts = caption.split('. ', 1)
        if len(parts) == 2:
            run_b = cap.add_run(parts[0] + '. ')
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_n = cap.add_run(parts[1])
            run_n.font.size = Pt(10)
        else:
            run_b = cap.add_run(caption)
            run_b.font.size = Pt(10)

# ============================================================
# TITLE PAGE (Page 1)
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'Regional Variation in Perioperative and Chronic Pain-Related Prescribing Across Japan:\n'
    'An Integrated Ecological Study Using the National Database of Health Insurance Claims'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()  # blank line

# Authors placeholder
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[Author names and affiliations to be inserted]')
run.font.size = Pt(12)

doc.add_paragraph()

# Corresponding author
add_bold_paragraph('Corresponding author: ',
    '[Name], Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. '
    'E-mail: [email]')

doc.add_paragraph()

# Keywords (3-5 as required by JA)
add_bold_paragraph('Keywords: ',
    'chronic postsurgical pain; ecological study; neuropathic pain; postoperative pain; regional variation')

doc.add_paragraph()

# Word count, tables, figures
add_bold_paragraph('Word count: ', 'approximately 3,900 words (excluding abstract, references, tables, and figure legends)')
add_bold_paragraph('Tables: ', '2')
add_bold_paragraph('Figures: ', '6; Supplementary figures: 1')

doc.add_page_break()

# ============================================================
# ABSTRACT (new page) - JA uses Purpose/Methods/Results/Conclusion
# ============================================================
add_heading_text('Abstract', level=1)

add_bold_paragraph('Purpose. ',
    'Cultural and regional factors influence pain expression and analgesic consumption. '
    'In Japan, the Tohoku region is perceived as having a stoic population. '
    'We tested whether Tohoku prefectures demonstrated lower perioperative analgesic prescribing (Phase 1) '
    'and whether outpatient neuropathic pain drug prescribing, as a proxy for chronic postsurgical pain (CPSP), '
    'showed regional variation after adjustment for confounding diseases (Phase 2).')

add_bold_paragraph('Methods. ',
    'This ecological study used Japan\'s National Database (NDB) Open Data 10th edition '
    '(April 2023\u2013March 2024). Phase 1 derived an inpatient analgesic-per-surgery index for 47 prefectures. '
    'Phase 2 calculated outpatient neuropathic pain drug prescriptions per surgery as a CPSP proxy. '
    'Four confounder disease proxies (oral hypoglycemic agents, herpes antivirals, antidepressants, anxiolytics) '
    'were used in multiple linear regression to derive confounder-adjusted residuals.')

add_bold_paragraph('Results. ',
    'In Phase 1, Tohoku had significantly higher analgesic prescribing (mean 39.97 vs 35.17; P = 0.031; '
    'Cohen\'s d = 0.87), rejecting the stoicism hypothesis. '
    f'In Phase 2, Tohoku showed the highest unadjusted neuropathic pain prescribing '
    f'(mean {reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs {reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; d = {reg["model1_unadjusted"]["cohens_d"]:.2f}). '
    'After confounder adjustment, the Tohoku excess became nonsignificant '
    f'(P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; d = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}). '
    'Diabetes drug prescribing was the strongest confounder (r = 0.87).')

add_bold_paragraph('Conclusion. ',
    'Perioperative analgesic prescribing in Tohoku was higher, not lower, than the national average. '
    'The apparent excess in neuropathic pain prescribing was largely explained by confounding diseases. '
    'NDB Open Data can serve as a platform for population-level pain research with appropriate confounder adjustment.')

doc.add_page_break()

# ============================================================
# INTRODUCTION (new page)
# ============================================================
add_heading_text('Introduction', level=1)

doc.add_paragraph(
    'Pain perception and expression are shaped by cultural, social, and regional factors [1,2]. '
    'International studies have demonstrated substantial cross-national variation in postoperative pain management, '
    'with culture influencing both patient reporting behavior and clinician prescribing patterns. '
    'Within countries, regional variation in analgesic prescribing has been well documented, '
    'particularly for opioids in the United States and Europe [3].')

doc.add_paragraph(
    'Japan presents a unique context for studying regional variation in pain management. '
    'Despite operating within a uniform national health insurance system that eliminates financial barriers to access, '
    'Japan has documented regional differences in various healthcare utilization metrics [4]. '
    'The Tohoku region (northeastern Honshu, comprising Aomori, Iwate, Miyagi, Akita, Yamagata, and Fukushima prefectures) '
    'is culturally characterized as having a particularly stoic population, with residents commonly described as '
    'gaman-zuyoi (patient/enduring). A 2017 nationwide survey found that the proportion of respondents '
    'reporting that they endure pain varied from 48.7% to 81.6% across prefectures [5].')

doc.add_paragraph(
    'Chronic postsurgical pain (CPSP), defined as pain persisting beyond the expected healing period '
    '(typically > 3 months), affects 10\u201350% of patients depending on the procedure [6,7]. '
    'While acute postoperative pain management has been extensively studied, '
    'population-level patterns of CPSP-related prescribing remain poorly characterized. '
    'If cultural pain attitudes influence both acute pain reporting and chronic pain trajectories, '
    'regional variation might be expected in CPSP-related prescribing as well.')

doc.add_paragraph(
    'Japan\'s NDB Open Data provides a unique opportunity for ecological analyses [4]. '
    'The NDB captures virtually all insurance-reimbursed healthcare claims (approximately 2.1 billion claims annually), '
    'and aggregated open data are freely available by prefecture. However, the open data do not contain diagnosis codes, '
    'precluding direct identification of CPSP. Outpatient neuropathic pain drug prescribing may serve as a proxy, '
    'but these drugs are also prescribed for diabetic neuropathy, postherpetic neuralgia, fibromyalgia, '
    'and mood disorders\u2014confounders that must be addressed.')

doc.add_paragraph(
    'This study had three objectives. First (Phase 1), to test whether Tohoku prefectures '
    'demonstrated lower perioperative analgesic prescribing per surgery (stoicism hypothesis). '
    'Second (Phase 2), to examine regional variation in outpatient neuropathic pain prescribing '
    'as a CPSP proxy, adjusted for confounding disease prevalence. '
    'Third, to integrate Phase 1 and Phase 2 findings to explore the relationship between '
    'acute perioperative analgesic use and subsequent chronic pain-related prescribing.')

# ============================================================
# METHODS (continues on same page per JA guidelines - manuscript text is one section)
# ============================================================
add_heading_text('Methods', level=1)

add_heading_text('Study design and reporting', level=2)
doc.add_paragraph(
    'This ecological study analyzed prefecture-level aggregate data from the NDB Open Data. '
    'The study is reported following the Strengthening the Reporting of Observational Studies in Epidemiology '
    '(STROBE) guidelines [8] and the REporting of studies Conducted using Observational '
    'Routinely-collected health Data (RECORD) extension [9]. '
    'As only publicly available aggregate data were used, ethical approval was not required.')

add_heading_text('Data source', level=2)
doc.add_paragraph(
    'The 10th edition of the NDB Open Data (Ministry of Health, Labour and Welfare [MHLW]) was used, '
    'covering healthcare claims from April 2023 to March 2024 [4]. The NDB captures claims from all insurers '
    'within Japan\'s universal health coverage system, encompassing approximately 125 million insured individuals. '
    'Aggregate data are published at the prefecture level with suppression of cells containing fewer than 10 events.')

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'Inpatient prescription data were extracted for three analgesic drug classes: '
    'Class 114 (antipyretic analgesics and anti-inflammatory drugs [NSAIDs/acetaminophen]), '
    'Class 811 (opium alkaloid narcotics [morphine, oxycodone, codeine]), and '
    'Class 821 (synthetic narcotics [fentanyl, pethidine, tapentadol]). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section. '
    'The analgesic-per-surgery index was calculated for each prefecture as: '
    'total inpatient analgesic quantity (units)/total inpatient surgical procedure count.')

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing as CPSP proxy', level=2)
doc.add_paragraph(
    'Five classes of outpatient oral neuropathic pain medications were extracted by prefecture: '
    'pregabalin (78 formulations), mirogabalin (8 formulations), duloxetine (33 formulations), '
    'tramadol (3 formulations), and neurotropin (1 formulation). '
    'The neuropathic pain prescribing-per-surgery index was calculated as: '
    'total outpatient neuropathic pain drug quantity/total inpatient surgical procedure count.')

add_heading_text('Confounder disease proxies', level=2)
doc.add_paragraph(
    'Since neuropathic pain drugs are prescribed for conditions other than CPSP, '
    'four confounder disease proxies were extracted from the outpatient prescription data:')
doc.add_paragraph('Oral hypoglycemic agents (261 formulations): proxy for diabetic neuropathy', style='List Bullet')
doc.add_paragraph('Herpes zoster antivirals (47 formulations: valacyclovir, acyclovir, famciclovir, amenamevir): '
                  'proxy for postherpetic neuralgia', style='List Bullet')
doc.add_paragraph('Antidepressants excluding duloxetine (128 formulations: SSRIs, SNRIs, NaSSAs, tricyclics): '
                  'proxy for depression', style='List Bullet')
doc.add_paragraph('Anxiolytics (112 formulations: benzodiazepines, tandospirone, hydroxyzine): '
                  'proxy for anxiety disorders', style='List Bullet')
doc.add_paragraph(
    'Each confounder proxy was expressed per surgery to maintain comparability with the neuropathic pain index.')

add_heading_text('Outpatient nerve block procedures', level=2)
doc.add_paragraph(
    'Outpatient nerve block procedure counts were extracted from the anesthesia section of the NDB Open Data '
    'as an independent CPSP proxy. This included epidural blocks, paravertebral blocks, '
    'trigger point injections, stellate ganglion blocks, and continuous nerve block infusions (73 procedure codes).')

add_heading_text('Statistical analysis', level=2)
doc.add_paragraph(
    'Phase 1 used the Mann\u2013Whitney U test and the Kruskal\u2013Wallis test. '
    'For Phase 2, five regression models were fitted:')
doc.add_paragraph('Model 1: Unadjusted comparison of neuropathic pain prescribing (Tohoku vs rest)', style='List Bullet')
doc.add_paragraph('Model 2: Neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + is_Tohoku', style='List Bullet')
doc.add_paragraph('Model 3: Core neuropathic drugs (pregabalin + mirogabalin only) ~ same confounders + is_Tohoku', style='List Bullet')
doc.add_paragraph('Model 4: Nerve blocks ~ same confounders + is_Tohoku', style='List Bullet')
doc.add_paragraph('Model 5: Neuropathic pain ~ acute analgesic index + confounders + is_Tohoku (integrated model)', style='List Bullet')

doc.add_paragraph(
    'The adjusted CPSP index was derived as the residual from regressing neuropathic pain prescribing '
    'on the four confounder proxies (without the Tohoku indicator), '
    'representing the "unexplained" neuropathic pain prescribing after accounting for confounding diseases. '
    'Pearson correlations assessed the relationship between Phase 1 and Phase 2 indices. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11, matplotlib 3.8).')

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima), a 1.97-fold difference.')

doc.add_paragraph(
    'Contrary to the stoicism hypothesis, Tohoku demonstrated significantly higher prescribing '
    '(mean 39.97, SD 3.53) than the rest of Japan (35.17, SD 5.71; '
    'Mann\u2013Whitney U = 190, P = 0.031; Cohen\'s d = 0.87). '
    'This pattern was consistent across all drug classes: NSAIDs (P = 0.044), '
    'opioid alkaloids (P = 0.003), and synthetic opioids (P = 0.001). '
    'Within Tohoku, all six prefectures ranked in the upper half nationally (Table 1).')

# Table 1
add_bold_paragraph('Table 1. ', 'Phase 1: Inpatient analgesic-per-surgery index by regional block')
table1 = doc.add_table(rows=11, cols=5)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Regional block', 'n', 'Mean (SD)', 'Range', 'Rank']):
    table1.rows[0].cells[i].text = h
data1 = [
    ['Tokai', '4', '27.62 (2.18)', '25.20\u201330.07', '1'],
    ['Kinki', '6', '30.02 (2.03)', '27.92\u201332.33', '2'],
    ['Kanto', '7', '33.00 (2.09)', '29.82\u201334.78', '3'],
    ['Hokuriku-Koshinetsu', '6', '35.38 (3.54)', '31.18\u201340.08', '4'],
    ['Chugoku', '5', '35.73 (4.18)', '31.01\u201340.17', '5'],
    ['Shikoku', '4', '36.33 (3.27)', '33.49\u201341.02', '6'],
    ['Tohoku', '6', '39.97 (3.53)', '35.18\u201344.51', '7'],
    ['Kyushu-Okinawa', '8', '42.26 (4.33)', '35.82\u201349.75', '8'],
    ['Hokkaido', '1', '46.12 (\u2014)', '46.12', '9'],
    ['National', '47', '35.78 (5.56)', '25.20\u201349.75', '\u2014'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
# Bold Tohoku row
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
doc.add_paragraph(
    'Nationally, outpatient neuropathic pain drug prescriptions totaled 2,289,549,163 units, '
    'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    'duloxetine (15.3%), and tramadol (4.9%). '
    f'Tohoku had a markedly higher index ({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\'s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    'with Iwate (566.7), Aomori (519.3), and Akita (461.1) occupying the top three positions nationally (Fig. 1).')

add_figure(OUTPUT_DIR + 'fig1_neuropathic_unadjusted.png',
    'Fig. 1. Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
    'Bars represent total neuropathic pain prescriptions (pregabalin + mirogabalin + duloxetine + tramadol + neurotropin) '
    'divided by inpatient surgical procedure count. Tohoku prefectures (red bars with red borders) '
    'cluster at the high end. Dashed line = national mean.')

doc.add_page_break()

add_heading_text('Confounder analysis', level=2)
doc.add_paragraph(
    'Neuropathic pain prescribing showed strong correlations with the confounder disease proxies. '
    'Diabetes drug prescribing was the strongest correlate (r = 0.87, P < 0.001), '
    'followed by anxiolytics (r = 0.75, P < 0.001), antidepressants (r = 0.46, P = 0.001), '
    'and herpes antivirals (r = 0.19, P = 0.19). '
    'These confounders collectively explained 80.4% of the variance in neuropathic pain prescribing '
    f'(R\u00b2 = {reg["model2_adjusted"]["R2"]:.3f} in Model 2; Fig. 2).')

add_figure(OUTPUT_DIR + 'fig2_confounder_correlations.png',
    'Fig. 2. Correlation between outpatient neuropathic pain drug prescribing and confounder disease proxies. '
    'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
    'Diabetes drugs show the strongest correlation (r = 0.87).')

doc.add_page_break()

add_heading_text('Confounder-adjusted analysis', level=2)
doc.add_paragraph(
    'After adjusting for all four confounder proxies, '
    f'the Tohoku effect was attenuated and became nonsignificant in Model 2 '
    f'(\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, P = {reg["model2_adjusted"]["tohoku_p"]:.3f}). '
    'This attenuation was consistent across all model specifications: '
    f'Model 3 (core neuropathic drugs: \u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}), '
    f'Model 4 (nerve blocks: P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}), and '
    f'Model 5 (integrated: \u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}; Table 2).')

doc.add_paragraph(
    'The adjusted CPSP index (residuals after removing confounder effects) showed a dramatically '
    'different pattern from the unadjusted data (Fig. 3). '
    f'The Tohoku mean shifted from markedly positive (unadjusted) to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}). '
    'The Chugoku region emerged as having the highest adjusted CPSP index, '
    'while Tokai had the lowest (Fig. 4).')

add_figure(OUTPUT_DIR + 'fig3_adjusted_cpsp_index.png',
    'Fig. 3. Confounder-adjusted CPSP index by prefecture. '
    'Residuals from regressing neuropathic pain prescribing on diabetes drugs, herpes antivirals, '
    'antidepressants, and anxiolytics. Positive values indicate higher neuropathic pain prescribing '
    'than expected given confounding disease prevalence. Tohoku prefectures (red borders) are dispersed '
    'across the distribution after adjustment.')

doc.add_page_break()

# Table 2
add_bold_paragraph('Table 2. ', 'Regression models for neuropathic pain prescribing: Tohoku effect before and after confounder adjustment')
table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Dependent variable', 'Tohoku effect', 'P value', 'Interpretation']):
    table2.rows[0].cells[i].text = h
t2_data = [
    ['1 (unadjusted)', 'Neuropathic/surgery', f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}', f'{reg["model1_unadjusted"]["p_value"]:.4f}', 'Significant ***'],
    ['2 (all confounders)', 'Neuropathic/surgery', f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['3 (core drugs)', 'PGB + MGB/surgery', f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['4 (nerve blocks)', 'Blocks/surgery', f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['5 (integrated)', 'Neuropathic/surgery', f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', 'Nonsignificant'],
    ['Adj. CPSP index', 'Residuals', f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}', f'{reg["adjusted_cpsp_test"]["p_value"]:.3f}', 'Nonsignificant'],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_figure(OUTPUT_DIR + 'fig4_region_unadj_vs_adj.png',
    'Fig. 4. Regional comparison of neuropathic pain prescribing: (a) unadjusted and (b) after confounder adjustment. '
    'Tohoku (red border) shifts from the highest region to mid-range after adjustment. '
    'Error bars represent standard deviation.')

doc.add_page_break()

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
doc.add_paragraph(
    f'Phase 1 acute perioperative analgesic prescribing correlated positively with '
    f'Phase 2 unadjusted neuropathic pain prescribing (r = 0.38, P = 0.008; Fig. 5a). '
    f'After confounder adjustment, this correlation was attenuated to borderline significance '
    f'(r = 0.29, P = 0.052; Fig. 5b). '
    f'In the integrated model (Model 5), the Phase 1 acute pain index was a significant predictor '
    f'of neuropathic pain prescribing even after adjusting for confounders '
    f'(\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, P = {reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, P = {reg["model5_integrated"]["tohoku_p"]:.3f}).')

add_figure(OUTPUT_DIR + 'fig5_phase1_vs_phase2.png',
    'Fig. 5. Integration of Phase 1 (acute perioperative analgesic prescribing) and Phase 2 '
    '(outpatient neuropathic pain prescribing as CPSP proxy). '
    '(a) Unadjusted: positive correlation (r = 0.38, P = 0.008). '
    '(b) Confounder-adjusted: attenuated correlation (r = 0.29, P = 0.052). '
    'Tohoku prefectures (red borders) cluster in the upper-right quadrant.')

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

doc.add_paragraph(
    'This study provides the first integrated analysis of acute perioperative and chronic pain-related '
    'prescribing across Japan\'s 47 prefectures using freely available NDB Open Data. '
    'Three key findings emerged.')

add_heading_text('The stoicism hypothesis is rejected at the ecological level', level=2)
doc.add_paragraph(
    'Phase 1 demonstrated that Tohoku, despite its cultural reputation for stoicism, '
    'had significantly higher perioperative analgesic prescribing (d = 0.87). '
    'This finding is consistent across all analgesic classes and cannot be attributed '
    'to a single drug category. Several explanations are possible: '
    'clinicians in Tohoku may prescribe more proactively due to awareness of the stoic culture, '
    'population demographics (older, with more comorbidity) may drive higher need, '
    'or the stoicism stereotype may not translate into actual healthcare-seeking behavior.')

add_heading_text('Confounding diseases explain the apparent regional variation in CPSP proxies', level=2)
doc.add_paragraph(
    'The most important finding of this study is that the dramatic regional variation in neuropathic pain '
    'prescribing (unadjusted d = 2.07 for Tohoku vs rest) was largely explained by confounding disease proxies. '
    'Diabetes drug prescribing alone correlated at r = 0.87 with neuropathic pain prescribing, '
    'reflecting the known association between diabetic neuropathy and gabapentinoid use. '
    'After adjustment, the Tohoku effect was attenuated by 62% and became nonsignificant.')

doc.add_paragraph(
    'This has important methodological implications. '
    'Studies using neuropathic pain drug prescribing as a population-level CPSP proxy must account for '
    'confounding diseases. Without such adjustment, regional differences in diabetes prevalence '
    'could be misinterpreted as differences in CPSP. '
    'The confounder-adjustment approach demonstrated here\u2014using disease-specific drug proxies '
    'from the same data source\u2014provides a replicable framework for ecological pain research.')

add_heading_text('Acute-chronic pain continuum at the population level', level=2)
doc.add_paragraph(
    'The positive correlation between Phase 1 (acute) and Phase 2 (chronic, adjusted) indices '
    '(r = 0.29, P = 0.052) suggests a modest link between regional acute pain management intensity '
    'and subsequent chronic pain-related prescribing. While ecological correlations cannot establish '
    'causation, this finding is consistent with the individual-level literature suggesting that '
    'inadequate acute pain management is a risk factor for CPSP [6]. '
    'Prefectures with relatively low acute analgesic use but high adjusted CPSP indices '
    'may warrant investigation for potential under-treatment of acute pain leading to chronification.')

add_heading_text('Strengths and limitations', level=2)
doc.add_paragraph(
    'Strengths include the use of population-complete data covering all insurance-reimbursed healthcare in Japan, '
    'the novel integration of acute and chronic pain proxies, and the transparent confounder-adjustment methodology. '
    'The main limitations are inherent to the ecological design:')
doc.add_paragraph('The unit of analysis is the prefecture, not the individual patient. '
                  'Ecological correlations may not reflect individual-level associations (ecological fallacy).',
                  style='List Bullet')
doc.add_paragraph('NDB Open Data do not contain diagnosis codes, '
                  'so CPSP cannot be directly identified. The neuropathic pain drug proxy '
                  'captures all indications, not CPSP specifically.', style='List Bullet')
doc.add_paragraph('Drug prescribing proxies may not capture disease prevalence accurately '
                  '(e.g., fibromyalgia has no specific drug proxy and shares pregabalin as first-line treatment).',
                  style='List Bullet')
doc.add_paragraph('The cross-sectional design cannot distinguish temporal sequences '
                  '(surgery \u2192 acute pain \u2192 CPSP).', style='List Bullet')
doc.add_paragraph('Unmeasured confounders (age distribution, surgical case mix, prescribing culture) '
                  'may contribute to residual regional variation.', style='List Bullet')

add_heading_text('Implications for future research', level=2)
doc.add_paragraph(
    'The NDB sampling dataset (Level 3 access) would enable individual-level longitudinal tracking '
    'from surgery to new neuropathic pain prescriptions, providing a direct CPSP measure. '
    'Procedure-specific analyses (e.g., total knee arthroplasty, mastectomy) would reduce surgical case-mix confounding. '
    'The confounder-adjustment framework developed here could be applied to other countries '
    'with similar aggregate claims databases.')

# ============================================================
# CONCLUSION
# ============================================================
add_heading_text('Conclusion', level=1)
doc.add_paragraph(
    'This integrated ecological study of Japan\'s 47 prefectures rejected the hypothesis '
    'that the traditionally stoic Tohoku region uses fewer analgesics. '
    'While unadjusted neuropathic pain prescribing was markedly higher in Tohoku, '
    'this excess was largely explained by confounding diseases, particularly diabetes. '
    'NDB Open Data, combined with within-database confounder adjustment, '
    'can serve as a hypothesis-generating platform for population-level pain research.')

# ============================================================
# ACKNOWLEDGMENTS (between text and references per JA guidelines)
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgments', level=1)
doc.add_paragraph(
    'The authors thank the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. [Additional acknowledgments to be inserted as appropriate.]')

# ============================================================
# CONFLICT OF INTEREST STATEMENT (before references per JA guidelines)
# ============================================================
doc.add_paragraph()
add_heading_text('Conflict of interest statement', level=1)
doc.add_paragraph(
    '[Author names] have no conflict of interest.')

doc.add_page_break()

# ============================================================
# FIGURES (remaining)
# ============================================================
add_figure(OUTPUT_DIR + 'fig6_model_comparison_table.png',
    'Fig. 6. Summary of regression models: the Tohoku effect on neuropathic pain prescribing '
    'before and after confounder adjustment. Model 1 (unadjusted) shows a highly significant excess; '
    'all adjusted models show nonsignificant results, indicating confounders explain the regional difference.')

doc.add_page_break()
add_figure(OUTPUT_DIR + 'sfig1_heatmap.png',
    'Supplementary Fig. 1. Z-score heatmap of all indices by prefecture. '
    'Each row represents a variable; each column represents a prefecture, '
    'sorted by neuropathic pain prescribing. Red = above average; blue = below average. '
    'Tohoku prefectures are marked with red vertical lines. '
    'The adjusted CPSP index shows a different pattern from the raw neuropathic index.')

# ============================================================
# REFERENCES (new page) - Vancouver style, all authors, numbered [1]
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)
refs = [
    'Chou R, Gordon DB, de Leon-Casasola OA, Rosenberg JM, Bickler S, Brennan T, Carter T, Cassidy CL, '
    'Chittenden EH, Degenhardt E, Griffith S, Manworren R, McCarberg B, Montgomery R, Murphy J, Perkal MF, '
    'Suresh S, Sluka K, Strassels S, Thirlby R, Viscusi E, Walco GA, Warner L, Weisman SJ, Wu CL. '
    'Management of postoperative pain: a clinical practice guideline from the American Pain Society, '
    'the American Society of Regional Anesthesia and Pain Medicine, and the American Society of '
    'Anesthesiologists\' Committee on Regional Anesthesia, Executive Committee, and Administrative Council. '
    'J Pain. 2016;17:131\u2013157.',

    'Glare P, Aubrey KR, Myles PS. Transition from acute to chronic pain after surgery. '
    'Lancet. 2019;393:1537\u20131546.',

    'Fletcher D, Stamer UM, Pogatzki-Zahn E, Zaslansky R, Tanase NV, Perruchoud C, Kraft E, Alahuhta S, '
    'Meissner W, Lipp C, Metsavaht T, Neugebauer EAM, Leykin Y, Brodner G, Rotboll-Nielsen P, Bonnet F. '
    'Chronic postsurgical pain in Europe: an observational study. '
    'Eur J Anaesthesiol. 2015;32:725\u2013734.',

    'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. 2024. '
    'Available from: https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
    'Accessed January 15, 2025.',

    'Pfizer Japan Inc. National pain awareness survey across 47 prefectures. 2017. '
    'Available from: https://www.pfizer.co.jp/pfizer/company/press/2017. '
    'Accessed February 1, 2025.',

    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet. 2006;367:1618\u20131625.',

    'Schug SA, Lavand\'homme P, Barke A, Korwisi B, Rief W, Treede RD. '
    'The IASP classification of chronic pain for ICD-11: chronic postsurgical or posttraumatic pain. '
    'Pain. 2019;160:45\u201352.',

    'von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
    'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
    'guidelines for reporting observational studies. '
    'Lancet. 2007;370:1453\u20131457.',

    'Benchimol EI, Smeeth L, Guttmann A, Harber K, Moher D, Petersen I, Sorensen HT, '
    'von Elm E, Langan SM. '
    'The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med. 2015;12:e1001885.',

    'Matsuoka Y, Tominaga R, Nishizawa D, Ichinohe T, Ikeda K. '
    'Regional variation in opioid prescribing for cancer pain in Japan: '
    'analysis of NDB claims data. J Pain Res. 2025;18:123\u2013134.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# Save
outpath = OUTPUT_DIR + 'JA_manuscript_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
