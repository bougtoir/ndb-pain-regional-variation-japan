#!/usr/bin/env python3
"""Create Pain journal 'Summary' file (Table of Contents synopsis).

Per Pain Instructions for Authors:
- Separate file uploaded with manuscript
- 1-2 sentences, 25 words max
- States conclusions of the study
- Used in the Table of Contents
- No first person; do not start with 'This study...'
- Do not merely rephrase the title
- Inform readers of objective, methods, results, and/or conclusions
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# --- Page setup: A4, standard margins ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style: Times New Roman 12pt, double-spaced ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 2.0

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Summary')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()  # blank line

# --- Summary text (25 words max, for Table of Contents) ---
# Word count target: <=25 words
# Content: objective + results + conclusion
summary_text = (
    'Nearly twofold prefecture-level variation in perioperative analgesic prescribing '
    'across Japan challenges the monolithic characterization of Japanese pain behavior '
    'and its potential to bias clinical decisions.'
)

p = doc.add_paragraph(summary_text)

# Verify word count
word_count = len(summary_text.split())
print(f"Summary word count: {word_count} (max 25)")

# --- Save ---
out_path = OUTPUT_DIR + 'Pain_summary.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
