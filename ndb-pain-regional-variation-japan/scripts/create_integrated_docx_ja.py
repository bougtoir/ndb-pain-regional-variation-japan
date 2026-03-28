#!/usr/bin/env python3
"""Create integrated Japanese manuscript (Phase 1 + Phase 2) as .docx with embedded color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import csv
import json

OUTPUT_DIR = '/home/ubuntu/analysis/output/'
doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 2.0

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=6.0):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        parts = caption.split('. ', 1)
        if len(parts) == 2:
            run_b = cap.add_run(parts[0] + '. ')
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_n = cap.add_run(parts[1])
            run_n.font.size = Pt(10)
        else:
            run_b = cap.add_run(caption)
            run_b.font.size = Pt(10)

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    '日本における周術期および慢性疼痛関連処方の地域差：\n'
    'NDBオープンデータを用いた統合的生態学的研究'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
add_bold_paragraph('ランニングタイトル：', '日本における急性期・慢性期疼痛処方の地域差')
add_bold_paragraph('論文種別：', '原著論文')
add_bold_paragraph('投稿先：', 'British Journal of Anaesthesia')
add_bold_paragraph('語数：', '抄録 約300語、本文 約4,500語')
add_bold_paragraph('表：', '4; 図：6; 補足図：1')
add_bold_paragraph('キーワード：', '鎮痛薬; 術後遷延性疼痛; 生態学的研究; 医療サービス研究; 日本; 神経障害性疼痛; オピオイド; 術後疼痛; 地域差')

doc.add_page_break()

# ============================================================
# 抄録
# ============================================================
add_heading_text('抄録', level=1)

add_bold_paragraph('背景：',
    '文化的・地域的要因は疼痛の表現と鎮痛薬消費に影響を及ぼす。'
    '日本では東北地方の住民は伝統的に「我慢強い」と認識されている。'
    '本研究では、東北地方の都道府県が周術期鎮痛薬処方量が低いか（Phase 1）、'
    'および術後遷延性疼痛（CPSP）のプロキシとしての外来神経障害性疼痛薬処方が'
    '交絡疾患調整後にも地域差を示すか（Phase 2）を検証した。')

add_bold_paragraph('方法：',
    '本生態学的研究では、NDBオープンデータ第10回（2023年4月〜2024年3月）を使用した。'
    'Phase 1では47都道府県ごとの入院鎮痛薬/手術指標を算出した。'
    'Phase 2では外来神経障害性疼痛薬処方量（プレガバリン、ミロガバリン、デュロキセチン、'
    'トラマドール、ノイロトロピン）/手術をCPSPプロキシとして算出した。'
    '薬剤の非特異性に対処するため、4つの交絡疾患プロキシ（経口血糖降下薬＝糖尿病性神経障害、'
    '帯状疱疹抗ウイルス薬＝帯状疱疹後神経痛、デュロキセチン以外の抗うつ薬＝うつ病、'
    '抗不安薬＝不安障害）で重回帰分析により調整し、残差を調整済CPSP指標とした。')

add_bold_paragraph('結果：',
    f'Phase 1では、東北は入院鎮痛薬/手術が全国平均を有意に上回り'
    '（平均39.97 vs 35.17; P=0.031; Cohen\'s d=0.87）、我慢強さ仮説は棄却された。'
    f'Phase 2（未調整）では、東北は神経障害性疼痛薬/手術が最高であった'
    f'（平均{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs {reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P<0.001; d={reg["model1_unadjusted"]["cohens_d"]:.2f}）。'
    f'しかし交絡疾患調整後、東北の過剰は有意でなくなった'
    f'（調整済CPSP指標：東北{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs '
    f'非東北{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f'P={reg["adjusted_cpsp_test"]["p_value"]:.3f}; d={reg["adjusted_cpsp_test"]["cohens_d"]:.2f}）。'
    '糖尿病薬処方が最強の交絡因子であった（r=0.87）。'
    'Phase 1の急性期疼痛とPhase 2の慢性疼痛には緩やかな正の相関がみられた（r=0.38, P=0.008）。')

add_bold_paragraph('結論：',
    '東北地方の周術期鎮痛薬処方は全国平均より高く、我慢強さ仮説は棄却された。'
    '外来神経障害性疼痛薬処方の見かけ上の過剰は、交絡疾患（特に糖尿病）によって大部分が説明された。'
    'NDBオープンデータは、適切な交絡調整を行えば、集団レベルの疼痛研究プラットフォームとして機能しうる。')

doc.add_page_break()

# ============================================================
# 緒言
# ============================================================
add_heading_text('緒言', level=1)

doc.add_paragraph(
    '疼痛の知覚と表現は、文化的・社会的・地域的要因により形成される。'
    '国際的な研究では、術後疼痛管理における国家間の大きなばらつきが示されており、'
    '文化は患者の報告行動と臨床医の処方パターンの両方に影響を及ぼす。'
    '各国内でも、特に米国と欧州ではオピオイドを中心とした鎮痛薬処方の地域差が広く報告されている。')

doc.add_paragraph(
    '日本は疼痛管理の地域差を研究するうえで独特の文脈を提供する。'
    'アクセスの経済的障壁を排除する均一な国民皆保険制度を有するにもかかわらず、'
    '日本では様々な医療利用指標において地域差が記録されている。'
    '東北地方（本州北東部：青森、岩手、宮城、秋田、山形、福島の6県）は、'
    '特に我慢強い住民を有する文化として特徴づけられ、住民は一般に「我慢強い」と評される。'
    'ファイザー社の2017年全国調査では、痛みを我慢すると回答した割合は'
    '都道府県間で48.7%〜81.6%と大きく異なっていた。')

doc.add_paragraph(
    '術後遷延性疼痛（CPSP）は、予想される治癒期間（通常3ヶ月以上）を超えて持続する疼痛と定義され、'
    '手術の種類により10〜50%の患者に影響を及ぼす。'
    '急性期術後疼痛管理は広く研究されているが、'
    'CPSP関連処方の集団レベルでのパターンはほとんど解明されていない。'
    '文化的な疼痛態度が急性期の疼痛報告と慢性疼痛の軌跡の両方に影響するならば、'
    'CPSP関連処方にも地域差が予想される。')

doc.add_paragraph(
    'NDBオープンデータは生態学的分析に独自の機会を提供する。'
    'NDBは保険請求されたほぼ全ての医療を収録（年間約21億件）し、'
    '都道府県別の集計データは自由にアクセス可能である。'
    'ただし、オープンデータには傷病名コードが含まれておらず、CPSPの直接的な同定は不可能である。'
    '外来の神経障害性疼痛薬処方がプロキシとなりうるが、'
    'これらの薬剤は糖尿病性神経障害、帯状疱疹後神経痛、線維筋痛症、気分障害にも処方される'
    '——これらの交絡に対処する必要がある。')

doc.add_paragraph(
    '本研究は3つの目的を有する。'
    '第一（Phase 1）、東北地方が周術期鎮痛薬/手術が低いか（我慢強さ仮説）の検証。'
    '第二（Phase 2）、外来神経障害性疼痛薬処方のCPSPプロキシとしての地域差を、'
    '交絡疾患有病率を調整したうえで検討。'
    '第三、Phase 1とPhase 2の統合により、急性期周術期鎮痛薬使用と'
    'その後の慢性疼痛関連処方の関係を探索すること。')

# ============================================================
# 方法
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
doc.add_paragraph(
    '本研究はNDBオープンデータの都道府県別集計データを分析した生態学的研究である。'
    'STROBE声明（横断研究）およびRECORD拡張に準拠して報告する。'
    '公開された集計データのみを使用し個人レベルの情報を含まないため、倫理審査は不要である。')

add_heading_text('データソース', level=2)
doc.add_paragraph(
    'NDBオープンデータ第10回（厚生労働省公表）を使用し、'
    '2023年4月〜2024年3月の診療報酬請求を対象とした。'
    'NDBは日本の国民皆保険制度内の全保険者からの請求を収録し、'
    '約1億2,500万人の被保険者を網羅する。'
    '集計データは都道府県レベルで公表され、10件未満（手技）'
    'または1,000単位未満（処方）のセルは秘匿処理されている。')

add_heading_text('Phase 1：入院周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから以下の3つの鎮痛薬分類を抽出した：'
    '薬効分類114（解熱鎮痛消炎薬＝NSAIDs・アセトアミノフェン）、'
    '薬効分類811（あへんアルカロイド系麻薬＝モルヒネ、オキシコドン、コデイン）、'
    '薬効分類821（合成麻薬＝フェンタニル、ペチジン、タペンタドール）。'
    '入院手術件数はKコード（手術）から抽出した。'
    '各都道府県の鎮痛薬/手術指標＝入院鎮痛薬処方量（単位）÷入院手術件数として算出した。')

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（CPSPプロキシ）', level=2)
doc.add_paragraph(
    '5種類の外来経口神経障害性疼痛薬を都道府県別に抽出した：'
    'プレガバリン（78製剤）、ミロガバリン（8製剤）、デュロキセチン（33製剤）、'
    'トラマドール（3製剤）、ノイロトロピン（1製剤）。'
    '神経障害性疼痛薬/手術指標＝外来神経障害性疼痛薬処方量÷入院手術件数として算出した。')

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '神経障害性疼痛薬はCPSP以外の疾患にも処方されるため、'
    '以下の4つの交絡疾患プロキシを外来処方データから抽出した：')
doc.add_paragraph('経口血糖降下薬（261製剤）：糖尿病性神経障害のプロキシ', style='List Bullet')
doc.add_paragraph('帯状疱疹抗ウイルス薬（47製剤：バラシクロビル、アシクロビル、ファムシクロビル、アメナメビル）：'
                  '帯状疱疹後神経痛のプロキシ', style='List Bullet')
doc.add_paragraph('デュロキセチンを除く抗うつ薬（128製剤：SSRI、SNRI、NaSSA、三環系）：'
                  'うつ病のプロキシ', style='List Bullet')
doc.add_paragraph('抗不安薬（112製剤：ベンゾジアゼピン系、タンドスピロン、ヒドロキシジン）：'
                  '不安障害のプロキシ', style='List Bullet')
doc.add_paragraph('各交絡プロキシは手術あたりで表し、神経障害性疼痛薬指標との比較可能性を維持した。')

add_heading_text('外来神経ブロック手技', level=2)
doc.add_paragraph(
    'NDBオープンデータの麻酔区分から外来神経ブロック算定回数を抽出し、'
    '独立したCPSPプロキシとした。'
    '硬膜外ブロック、傍脊椎神経ブロック、トリガーポイント注射、'
    '星状神経節ブロック、持続神経ブロック注入等73手技コードを含む。')

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    'Phase 1ではMann-Whitney U検定およびKruskal-Wallis検定を使用した。'
    'Phase 2では以下の5つの回帰モデルを適合した：')
doc.add_paragraph('モデル1：未調整の神経障害性疼痛薬処方比較（東北 vs その他）', style='List Bullet')
doc.add_paragraph('モデル2：神経障害性疼痛薬 ~ 糖尿病薬 + 帯状疱疹薬 + 抗うつ薬 + 抗不安薬 + 東北ダミー', style='List Bullet')
doc.add_paragraph('モデル3：コア神経障害性薬（プレガバリン+ミロガバリンのみ） ~ 同交絡因子 + 東北ダミー', style='List Bullet')
doc.add_paragraph('モデル4：神経ブロック ~ 同交絡因子 + 東北ダミー', style='List Bullet')
doc.add_paragraph('モデル5：神経障害性疼痛薬 ~ 急性期鎮痛薬指標 + 交絡因子 + 東北ダミー（統合モデル）', style='List Bullet')

doc.add_paragraph(
    '調整済CPSP指標は、神経障害性疼痛薬処方を4つの交絡プロキシ（東北ダミーなし）に回帰した残差として導出し、'
    '交絡疾患を除いた「説明されない」神経障害性疼痛薬処方を表す。'
    'Pearson相関でPhase 1・Phase 2指標間の関連を評価した。'
    '全解析はPython 3.11（NumPy, SciPy, matplotlib）で実施した。')

doc.add_page_break()

# ============================================================
# 結果
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1：入院周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月に、NDBには7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方が記録された。'
    '全国平均の鎮痛薬/手術指標は35.78（SD 5.56）で、'
    '岐阜（25.20）から鹿児島（49.75）まで1.97倍の差があった。')

doc.add_paragraph(
    '我慢強さ仮説に反して、東北は入院鎮痛薬/手術が有意に高かった'
    '（平均39.97, SD 3.53 vs その他35.17, SD 5.71; '
    'Mann-Whitney U=190, P=0.031; Cohen\'s d=0.87）。'
    'この傾向はNSAIDs（P=0.044）、あへんアルカロイド（P=0.003）、'
    '合成麻薬（P=0.001）の全薬効分類で一貫していた（表1）。')

# 表1
add_bold_paragraph('表1. ', 'Phase 1：地域ブロック別 入院鎮痛薬/手術指標')
table1 = doc.add_table(rows=11, cols=5)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['地域ブロック', 'n', '平均（SD）', '範囲', '順位']):
    table1.rows[0].cells[i].text = h
data1 = [
    ['東海', '4', '27.62 (2.18)', '25.20\u201330.07', '1'],
    ['近畿', '6', '30.02 (2.03)', '27.92\u201332.33', '2'],
    ['関東', '7', '33.00 (2.09)', '29.82\u201334.78', '3'],
    ['北陸・甲信越', '6', '35.38 (3.54)', '31.18\u201340.08', '4'],
    ['中国', '5', '35.73 (4.18)', '31.01\u201340.17', '5'],
    ['四国', '4', '36.33 (3.27)', '33.49\u201341.02', '6'],
    ['東北', '6', '39.97 (3.53)', '35.18\u201344.51', '7'],
    ['九州・沖縄', '8', '42.26 (4.33)', '35.82\u201349.75', '8'],
    ['北海道', '1', '46.12 (\u2014)', '46.12', '9'],
    ['全国', '47', '35.78 (5.56)', '25.20\u201349.75', '\u2014'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
doc.add_paragraph()

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（未調整）', level=2)
doc.add_paragraph(
    '全国の外来神経障害性疼痛薬処方量は合計2,289,549,163単位で、'
    'プレガバリン（40.2%）、ノイロトロピン（20.1%）、ミロガバリン（19.6%）、'
    'デュロキセチン（15.3%）、トラマドール（4.9%）であった。'
    f'未調整では、東北は神経障害性疼痛薬/手術指標が突出して高かった'
    f'（平均{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P<0.001; d={reg["model1_unadjusted"]["cohens_d"]:.2f}）。'
    '岩手（566.7）、青森（519.3）、秋田（461.1）が全国上位3位を占めた（図1）。')

add_figure(OUTPUT_DIR + 'fig1_neuropathic_unadjusted.png',
    '図1. 都道府県別 外来神経障害性疼痛薬処方量/手術（未調整）。'
    'プレガバリン+ミロガバリン+デュロキセチン+トラマドール+ノイロトロピンの合計を入院手術件数で除した。'
    '東北（赤枠の赤棒）が高値側に集中。破線＝全国平均。')

doc.add_page_break()

add_heading_text('交絡因子分析', level=2)
doc.add_paragraph(
    '神経障害性疼痛薬処方は交絡疾患プロキシと強い相関を示した。'
    '糖尿病薬処方が最も強い相関（r=0.87, P<0.001）を示し、'
    '次いで抗不安薬（r=0.75, P<0.001）、抗うつ薬（r=0.46, P=0.001）、'
    '帯状疱疹薬（r=0.19, P=0.19）であった。'
    'これらの交絡因子は神経障害性疼痛薬処方の分散の80.4%を説明した'
    f'（モデル2のR\u00b2={reg["model2_adjusted"]["R2"]:.3f}；図2）。')

add_figure(OUTPUT_DIR + 'fig2_confounder_correlations.png',
    '図2. 外来神経障害性疼痛薬処方と交絡疾患プロキシの相関。'
    '各点は1都道府県。東北は赤枠で表示。糖尿病薬が最も強い相関（r=0.87）を示す。')

doc.add_page_break()

add_heading_text('交絡因子調整後の分析', level=2)
doc.add_paragraph(
    f'4つの交絡プロキシで調整した後、東北の効果は減弱し有意でなくなった'
    f'（モデル2：\u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P={reg["model2_adjusted"]["tohoku_p"]:.3f}）。'
    'この減弱は全モデルで一貫していた：'
    f'モデル3（コア神経障害性薬：\u03b2={reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P={reg["model3_core_neuropathic"]["tohoku_p"]:.3f}）、'
    f'モデル4（神経ブロック：P={reg["model4_nerve_blocks"]["tohoku_p"]:.3f}）、'
    f'モデル5（統合：\u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.3f}）（表2）。')

doc.add_paragraph(
    '調整済CPSP指標（交絡因子除去後の残差）は未調整データとは劇的に異なるパターンを示した（図3）。'
    f'東北の平均は、未調整では顕著な正値から、'
    f'調整後は緩やかな非有意な過剰に変化した'
    f'（{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't={reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P={reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd={reg["adjusted_cpsp_test"]["cohens_d"]:.2f}）。'
    '調整後は中国地方が最高の調整済CPSP指標を示し、東海が最低であった（図4）。')

add_figure(OUTPUT_DIR + 'fig3_adjusted_cpsp_index.png',
    '図3. 交絡疾患調整後の都道府県別CPSP指標。'
    '神経障害性疼痛薬処方を糖尿病薬・帯状疱疹薬・抗うつ薬・抗不安薬に回帰した残差。'
    '正の値は交絡疾患から予測されるより多い神経障害性疼痛薬処方を示す。'
    '東北（赤枠）は調整後、分布全体に分散している。')

doc.add_page_break()

# 表2
add_bold_paragraph('表2. ', '神経障害性疼痛薬処方の回帰モデル：交絡調整前後の東北効果')
table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['モデル', '従属変数', '東北効果', 'P値', '判定']):
    table2.rows[0].cells[i].text = h
t2_data = [
    ['1（未調整）', '神経障害性疼痛薬/手術', f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}', f'{reg["model1_unadjusted"]["p_value"]:.4f}', '有意 ***'],
    ['2（全交絡調整）', '神経障害性疼痛薬/手術', f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', '有意でない'],
    ['3（コア薬）', 'PGB+MGB/手術', f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', '有意でない'],
    ['4（神経ブロック）', 'ブロック/手術', f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', '有意でない'],
    ['5（統合）', '神経障害性疼痛薬/手術', f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', '有意でない'],
    ['調整済CPSP指標', '残差', f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}', f'{reg["adjusted_cpsp_test"]["p_value"]:.3f}', '有意でない'],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
doc.add_paragraph()

add_figure(OUTPUT_DIR + 'fig4_region_unadj_vs_adj.png',
    '図4. 地域ブロック別 神経障害性疼痛薬処方の比較：(A) 未調整 (B) 交絡疾患調整後。'
    '東北（赤枠）は未調整では最高だが、調整後は中位に移動。エラーバー＝標準偏差。')

doc.add_page_break()

add_heading_text('Phase 1\u2013Phase 2 統合', level=2)
doc.add_paragraph(
    f'Phase 1の急性期鎮痛薬/手術はPhase 2の未調整神経障害性疼痛薬/手術と'
    f'正の相関を示した（r=0.38, P=0.008；図5A）。'
    f'交絡調整後、相関は減弱し境界的有意性となった'
    f'（r=0.29, P=0.052；図5B）。'
    f'統合モデル（モデル5）では、Phase 1の急性期疼痛指標は'
    f'交絡調整後も神経障害性疼痛薬処方の有意な予測因子であった'
    f'（\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}）が、'
    f'東北効果は有意でなかった'
    f'（\u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.3f}）。')

add_figure(OUTPUT_DIR + 'fig5_phase1_vs_phase2.png',
    '図5. Phase 1（急性期鎮痛薬/手術）vs Phase 2（外来神経障害性疼痛薬/手術＝CPSPプロキシ）の統合。'
    '(A) 未調整：正の相関（r=0.38, P=0.008）。'
    '(B) 交絡調整後：減弱した相関（r=0.29, P=0.052）。'
    '東北（赤枠）は右上象限に集中。')

doc.add_page_break()

# ============================================================
# 考察
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    '本研究は、NDBオープンデータを用いて日本の47都道府県における'
    '急性期周術期および慢性疼痛関連処方を統合的に分析した初めての研究である。'
    '3つの重要な知見が得られた。')

add_heading_text('我慢強さ仮説は生態学的レベルで棄却される', level=2)
doc.add_paragraph(
    'Phase 1は、文化的に我慢強いと評される東北が、'
    '周術期鎮痛薬/手術が有意に高い（d=0.87）ことを示した。'
    'この知見は全鎮痛薬分類で一貫しており、特定の薬効分類には帰着できない。'
    '複数の説明が可能である：東北の臨床医は我慢強い文化への認識から'
    'より積極的に処方している可能性、人口構成（高齢、合併症多）が'
    'より高い需要を駆動している可能性、'
    'あるいは我慢強さのステレオタイプが実際の受診行動に反映されていない可能性がある。')

add_heading_text('交絡疾患がCPSPプロキシの見かけの地域差を説明する', level=2)
doc.add_paragraph(
    '本研究の最も重要な知見は、神経障害性疼痛薬処方の劇的な地域差'
    '（未調整d=2.07）が、交絡疾患プロキシによって大部分説明されたことである。'
    '糖尿病薬処方だけでr=0.87の相関を示し、'
    'これは糖尿病性神経障害とガバペンチノイド使用の既知の関連を反映する。'
    '調整後、東北効果は62%減弱し有意でなくなった。')

doc.add_paragraph(
    'これは方法論的に重要な含意を持つ。'
    '神経障害性疼痛薬処方をCPSPの集団レベルプロキシとして使用する研究は、'
    '交絡疾患を考慮しなければならない。'
    'そのような調整なしでは、糖尿病有病率の地域差がCPSPの差と誤認される恐れがある。'
    '本研究で示した交絡調整アプローチ——同一データソースの疾患特異的薬剤プロキシを用いる——は、'
    '生態学的疼痛研究の再現可能な枠組みを提供する。')

add_heading_text('集団レベルでの急性期-慢性期疼痛の連続体', level=2)
doc.add_paragraph(
    'Phase 1（急性期）とPhase 2（慢性期、調整済）指標の正の相関'
    '（r=0.29, P=0.052）は、地域の急性期疼痛管理の強度と'
    'その後の慢性疼痛関連処方の間に緩やかな関連があることを示唆する。'
    '生態学的相関は因果関係を確立できないが、'
    'この知見は不十分な急性期疼痛管理がCPSPの危険因子であるという'
    '個人レベルの文献と整合する。'
    '急性期鎮痛薬使用が比較的低いにもかかわらず調整済CPSP指標が高い都道府県は、'
    '急性期疼痛の過小治療が慢性化を招いている可能性について調査に値する。')

add_heading_text('強みと限界', level=2)
doc.add_paragraph(
    '強みとして、日本の全保険請求医療を網羅する集団完全データの使用、'
    '急性期と慢性期の疼痛プロキシの新規統合、'
    '透明な交絡調整方法論が挙げられる。'
    '主な限界は生態学的デザインに固有のものである：')
doc.add_paragraph('分析単位は個人ではなく都道府県であり、生態学的相関は個人レベルの関連を反映しない可能性がある（生態学的誤謬）。',
                  style='List Bullet')
doc.add_paragraph('NDBオープンデータには傷病名コードがないため、CPSPを直接同定できない。'
                  '神経障害性疼痛薬プロキシはCPSP特異的でなく全適応を含む。', style='List Bullet')
doc.add_paragraph('薬剤処方プロキシは疾患有病率を正確に反映しない可能性がある'
                  '（例：線維筋痛症には特異的な薬剤プロキシがなく、プレガバリンを第一選択として共有する）。',
                  style='List Bullet')
doc.add_paragraph('横断的デザインでは時間的順序（手術→急性期疼痛→CPSP）を区別できない。', style='List Bullet')
doc.add_paragraph('測定されない交絡因子（年齢分布、手術構成、処方文化）が残余の地域差に寄与しうる。', style='List Bullet')

add_heading_text('今後の研究への示唆', level=2)
doc.add_paragraph(
    'NDB特別抽出データ（Level 3アクセス）により、個人レベルで手術から新規神経障害性疼痛薬処方への'
    '縦断的追跡が可能となり、直接的なCPSP指標を提供しうる。'
    '手術種別（例：人工膝関節全置換術、乳房切除術）に限定した分析により手術構成交絡を軽減できる。'
    '本研究で開発した交絡調整枠組みは、類似の集計請求データベースを有する他国にも適用可能である。')

# ============================================================
# 結論
# ============================================================
add_heading_text('結論', level=1)
doc.add_paragraph(
    '日本の47都道府県を対象とした本統合的生態学的研究は、'
    '伝統的に我慢強いと評される東北地方の鎮痛薬使用が少ないという仮説を棄却した。'
    '外来神経障害性疼痛薬処方の未調整での顕著な過剰は、'
    '交絡疾患（特に糖尿病）によって大部分が説明された。'
    'NDBオープンデータは、データベース内の交絡調整と組み合わせることで、'
    '集団レベルの疼痛研究における仮説生成プラットフォームとして機能しうる。')

doc.add_page_break()

# Figure 6
add_figure(OUTPUT_DIR + 'fig6_model_comparison_table.png',
    '図6. 回帰モデルの要約：交絡調整前後の東北効果。'
    'モデル1（未調整）は高度に有意な過剰を示すが、全調整モデルでは有意でなくなり、'
    '交絡因子が地域差を説明することを示す。')

# Supplementary Figure
doc.add_page_break()
add_figure(OUTPUT_DIR + 'sfig1_heatmap.png',
    '補足図1. 都道府県別 各指標のZ-scoreヒートマップ。'
    '各行は変数、各列は都道府県（神経障害性疼痛薬/手術で昇順ソート）。'
    '赤＝平均以上、青＝平均以下。東北は赤の垂直線で表示。'
    '調整済CPSP指標は未調整の神経障害性指標とは異なるパターンを示す。')

# ============================================================
# 文献
# ============================================================
doc.add_page_break()
add_heading_text('文献', level=1)
refs = [
    'Chou R, Gordon DB, de Leon-Casasola OA, et al. Management of postoperative pain: a clinical practice guideline. J Pain 2016; 17: 131\u2013157.',
    'Glare P, Aubrey KR, Myles PS. Transition from acute to chronic pain after surgery. Lancet 2019; 393: 1537\u20131546.',
    '厚生労働省. NDBオープンデータ 第10回. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html',
    'Matsuoka Y, Tominaga R, Nishizawa D, et al. Regional variation in opioid prescribing for cancer pain in Japan: analysis of NDB claims data. J Pain Res 2025; 18: 123\u2013134.',
    'ファイザー株式会社. 全国47都道府県 痛みに関する意識調査 2017. https://www.pfizer.co.jp/pfizer/company/press/2017.',
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. Lancet 2006; 367: 1618\u20131625.',
    'Schug SA, Lavand\'homme P, Barke A, et al. The IASP classification of chronic pain for ICD-11: chronic postsurgical or posttraumatic pain. Pain 2019; 160: 45\u201352.',
    'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Lancet 2007; 370: 1453\u20131457.',
    'Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med 2015; 12: e1001885.',
    'Fletcher D, Stamer UM, Pogatzki-Zahn E, et al. Chronic postsurgical pain in Europe: an observational study. Eur J Anaesthesiol 2015; 32: 725\u2013734.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# Save
outpath = OUTPUT_DIR + 'BJA_integrated_manuscript_JA.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
