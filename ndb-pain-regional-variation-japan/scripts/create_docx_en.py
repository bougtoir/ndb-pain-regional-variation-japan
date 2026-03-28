#!/usr/bin/env python3
"""Create English BJA manuscript as .docx with embedded color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = '/home/ubuntu/analysis/output/'
doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 2.0

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=6.0):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_b = cap.add_run(caption.split('.')[0] + '.')
        run_b.bold = True
        run_b.font.size = Pt(10)
        rest = '.'.join(caption.split('.')[1:])
        if rest:
            run_n = cap.add_run(rest)
            run_n.font.size = Pt(10)

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Regional Variation in Perioperative Analgesic Prescribing Across Japan:\nA Nationwide Ecological Study Using the National Database of Health Insurance Claims')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
add_bold_paragraph('Running title: ', 'Regional variation in perioperative analgesic use in Japan')
add_bold_paragraph('Article type: ', 'Original Investigation')
add_bold_paragraph('Target journal: ', 'British Journal of Anaesthesia')
add_bold_paragraph('Word count: ', 'Abstract 247 words; Main text ~3,200 words')
add_bold_paragraph('Tables: ', '3; Figures: 4; Supplementary figures: 1')
add_bold_paragraph('Keywords: ', 'analgesics; health services research; Japan; opioids; postoperative pain; regional variation')

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_text('Abstract', level=1)

add_bold_paragraph('Background: ', 'Cultural and regional factors influence pain expression and analgesic consumption. In Japan, the Tohoku region is traditionally perceived as having a stoic, pain-enduring population. We tested the hypothesis that Tohoku prefectures would demonstrate lower perioperative analgesic prescribing per surgical procedure compared with other regions.')

add_bold_paragraph('Methods: ', 'This ecological study used the 10th edition of Japan\'s National Database of Health Insurance Claims (NDB) Open Data, covering all insurance-reimbursed healthcare from April 2023 to March 2024. Prefecture-level inpatient analgesic prescription quantities (non-steroidal anti-inflammatory drugs [NSAIDs], opioid alkaloids, and synthetic opioids) were divided by inpatient surgical procedure counts to derive an analgesic-per-surgery index for each of the 47 prefectures. The nine standard regional blocks were compared using Kruskal-Wallis and Mann-Whitney U tests.')

add_bold_paragraph('Results: ', 'Across 7,903,515 inpatient surgical procedures and 274,579,851 analgesic prescription units nationally, substantial regional variation was observed (Kruskal-Wallis H=34.10, P<0.001). Contrary to the stoicism hypothesis, Tohoku ranked 7th of 9 regions with a mean analgesic-per-surgery index of 39.97 (SD 3.53) versus the national mean of 35.78, representing 11.7% higher use (Mann-Whitney U=190, P=0.031; Cohen\'s d=0.87). This pattern was consistent across all drug classes: NSAIDs (P=0.044), opioid alkaloids (P=0.003), and synthetic opioids (P=0.001).')

add_bold_paragraph('Conclusions: ', 'Perioperative analgesic prescribing in Japan\'s Tohoku region was significantly higher, not lower, than the national average. Perceived regional stoicism does not translate into reduced analgesic consumption at the ecological level.')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('Introduction', level=1)

doc.add_paragraph('Pain perception and expression are shaped by cultural, social, and regional factors. International studies have demonstrated substantial cross-national variation in postoperative pain management, with culture influencing both patient reporting behaviour and clinician prescribing patterns. Within countries, regional variation in analgesic prescribing has been well documented, particularly for opioids in the United States and Europe.')

doc.add_paragraph('Japan presents a unique context for studying regional variation in pain management. Despite operating within a uniform national health insurance system that eliminates financial barriers to access, Japan has documented regional differences in various healthcare utilisation metrics. The Tohoku region (northeastern Honshu, comprising Aomori, Iwate, Miyagi, Akita, Yamagata, and Fukushima prefectures) is culturally characterised as having a particularly stoic population, with residents commonly described as gaman-zuyoi (patient/enduring). A 2017 nationwide survey by Pfizer Japan found that the proportion of respondents reporting that they endure pain varied from 48.7% to 81.6% across prefectures, with northern prefectures generally scoring higher.')

doc.add_paragraph('Whether this perceived stoicism translates into measurable differences in perioperative analgesic consumption has not been formally tested. Recent studies have begun to characterise regional variation in analgesic use in Japan. Matsuoka and colleagues demonstrated significant prefectural variation in opioid prescribing for cancer pain using National Database (NDB) claims data, with oxycodone injection use ranging from 4.0% in Shikoku to 16.4% in Tokai.')

doc.add_paragraph('Japan\'s NDB Open Data provides a unique opportunity to conduct comprehensive ecological analyses of analgesic prescribing patterns. The NDB captures virtually all insurance-reimbursed healthcare claims in Japan (approximately 2.1 billion claims annually), and the aggregated open data are freely available without institutional access requirements.')

doc.add_paragraph('The primary objective of this study was to test whether prefectures in the Tohoku region demonstrated lower perioperative analgesic prescribing per surgical procedure compared with other regions, as would be predicted by the cultural stoicism hypothesis. The secondary objectives were to characterise the overall pattern of regional variation and to examine whether patterns differed by analgesic class (NSAIDs, opioid alkaloids, and synthetic opioids).')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

add_heading_text('Study design and reporting', level=2)
doc.add_paragraph('This ecological study analysed prefecture-level aggregate data from the NDB Open Data. The study is reported following the Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) guidelines for cross-sectional studies and the REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) extension. As the study used only publicly available aggregate data with no individual-level information, ethical approval was not required.')

add_heading_text('Data source', level=2)
doc.add_paragraph('The 10th edition of the NDB Open Data (published by the Ministry of Health, Labour and Welfare [MHLW]) was used, covering healthcare claims from April 2023 to March 2024. The NDB captures claims from all insurers within Japan\'s universal health coverage system, encompassing approximately 125 million insured individuals. Aggregate data are published at the prefecture level with suppression of cells containing fewer than 10 events (for procedure data) or 1,000 units (for prescription data) to prevent re-identification.')

add_heading_text('Outcome: analgesic prescription quantity', level=2)
doc.add_paragraph('Inpatient prescription data were extracted from the NDB Open Data file for inpatient oral medications by prefecture and drug efficacy classification. Three analgesic drug classes were identified using the Japanese drug efficacy classification system:')
doc.add_paragraph('Class 114: Antipyretic analgesics and anti-inflammatory drugs, encompassing NSAIDs and acetaminophen', style='List Bullet')
doc.add_paragraph('Class 811: Opium alkaloid narcotics, including morphine, oxycodone, and codeine', style='List Bullet')
doc.add_paragraph('Class 821: Synthetic narcotics, including fentanyl, pethidine, and tapentadol', style='List Bullet')

add_heading_text('Exposure: surgical procedure volume', level=2)
doc.add_paragraph('Inpatient surgical procedure data were extracted from the K Surgery section (procedure counts by section and prefecture). The total surgical procedure count for each prefecture was calculated as the sum of all procedure items across all surgical sections. Suppressed cells were treated as zero.')

add_heading_text('Analgesic-per-surgery index', level=2)
p = doc.add_paragraph('The primary metric was calculated for each prefecture as: ')
run = p.add_run('Analgesic-per-surgery index = Total analgesic prescription quantity (units) / Total inpatient surgical procedure count')
run.italic = True

add_heading_text('Regional classification', level=2)
doc.add_paragraph('Prefectures were grouped into nine standard regional blocks: Hokkaido (n=1), Tohoku (n=6), Kanto (n=7), Hokuriku-Koshinetsu (n=6), Tokai (n=4), Kinki (n=6), Chugoku (n=5), Shikoku (n=4), and Kyushu-Okinawa (n=8).')

add_heading_text('Statistical analysis', level=2)
doc.add_paragraph('Descriptive statistics (mean, standard deviation [SD], median, range) were calculated for each prefecture and regional block. The primary comparison was between Tohoku (n=6) and the rest of Japan (n=41). Given the small sample size, the Mann-Whitney U test was used for the primary comparison. Cohen\'s d was calculated as the effect size. The Kruskal-Wallis test was used for the overall comparison across all nine regional blocks. Subgroup analyses were conducted for each analgesic class separately. Statistical significance was set at P<0.05 (two-sided). All analyses were performed using Python 3.11 with pandas, NumPy, SciPy, and matplotlib.')

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

add_heading_text('National overview', level=2)
doc.add_paragraph('During the study period (April 2023 to March 2024), the NDB recorded 7,903,515 inpatient surgical procedures and 274,579,851 analgesic prescription units (oral medications) across Japan\'s 47 prefectures. Of the total analgesic quantity, NSAIDs and acetaminophen (Class 114) accounted for 95.6% (262,451,388 units), opioid alkaloids (Class 811) for 4.1% (11,368,144 units), and synthetic opioids (Class 821) for 0.3% (760,319 units). The national mean analgesic-per-surgery index was 35.78 (SD 5.56), with substantial variation ranging from 25.20 (Gifu) to 49.75 (Kagoshima), a 1.97-fold difference (Table 1).')

add_heading_text('Regional block comparison', level=2)
doc.add_paragraph('Significant variation was observed across the nine regional blocks (Kruskal-Wallis H=34.10, P<0.001; Table 1, Figure 1). A broad geographical gradient was apparent: Tokai (27.62, SD 2.18) and Kinki (30.02, SD 2.03) had the lowest indices, while Hokkaido (46.12) and Kyushu-Okinawa (42.26, SD 4.33) had the highest.')

# --- Table 1 ---
add_bold_paragraph('Table 1. ', 'Prefecture-level analgesic-per-surgery index: descriptive statistics by regional block')
table1 = doc.add_table(rows=10, cols=7)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Regional block', 'n', 'Mean', 'SD', 'Median', 'Range', 'Rank']
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h
data1 = [
    ['Tokai', '4', '27.62', '2.18', '27.60', '25.20-30.07', '1'],
    ['Kinki', '6', '30.02', '2.03', '29.22', '27.92-32.33', '2'],
    ['Kanto', '7', '33.00', '2.09', '33.78', '29.82-34.78', '3'],
    ['Hokuriku-Koshinetsu', '6', '35.38', '3.54', '35.06', '31.18-40.08', '4'],
    ['Chugoku', '5', '35.73', '4.18', '34.63', '31.01-40.17', '5'],
    ['Shikoku', '4', '36.33', '3.27', '35.40', '33.49-41.02', '6'],
    ['Tohoku', '6', '39.97', '3.53', '40.02', '35.18-44.51', '7'],
    ['Kyushu-Okinawa', '8', '42.26', '4.33', '42.80', '35.82-49.75', '8'],
    ['Hokkaido', '1', '46.12', '\u2014', '46.12', '46.12', '9'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
# Bold Tohoku row
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

# --- Figure 1 ---
add_figure(OUTPUT_DIR + 'fig2_regional_comparison.png',
           'Figure 1. Analgesic-per-surgery index by regional block. Bars represent mean values; error bars represent standard deviation. The dashed line indicates the national mean (35.78). Tohoku (red) ranks 7th of 9 regions.')

doc.add_page_break()

add_heading_text('Primary analysis: Tohoku versus rest of Japan', level=2)
doc.add_paragraph('Contrary to the stoicism hypothesis, the Tohoku region demonstrated significantly higher analgesic prescribing per surgery compared with the rest of Japan (Table 3). The mean analgesic-per-surgery index was 39.97 (SD 3.53) in Tohoku versus 35.17 (SD 5.71) in the rest of Japan, a difference of 4.80 units (95% CI: 1.48 to 8.12). This difference was statistically significant by the Mann-Whitney U test (U=190, P=0.031) and borderline by the t-test (t=1.99, P=0.053). The effect size was large (Cohen\'s d=0.87). The one-sided test for lower Tohoku use was non-significant (P=0.986), effectively excluding the stoicism hypothesis at this ecological level.')

doc.add_paragraph('Tohoku ranked 7th of 9 regions, with mean analgesic use 11.7% above the national average. Within Tohoku, all six prefectures ranked in the upper half nationally: Akita (44.51, rank 44/47), Aomori (42.99, rank 41/47), Iwate (40.89, rank 38/47), Yamagata (39.16, rank 33/47), Miyagi (37.10, rank 30/47), and Fukushima (35.18, rank 25/47).')

# --- Table 2 ---
add_bold_paragraph('Table 2. ', 'Tohoku prefectures: individual analgesic-per-surgery indices and national rankings')
table2 = doc.add_table(rows=8, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Prefecture', 'Surgeries', 'Analgesic qty', 'Index', 'Rank (/47)']):
    table2.rows[0].cells[i].text = h
data2 = [
    ['Akita', '58,679', '2,611,773', '44.51', '44'],
    ['Aomori', '72,664', '3,123,606', '42.99', '41'],
    ['Iwate', '59,572', '2,435,978', '40.89', '38'],
    ['Yamagata', '62,433', '2,444,620', '39.16', '33'],
    ['Miyagi', '143,501', '5,323,830', '37.10', '30'],
    ['Fukushima', '109,901', '3,866,645', '35.18', '25'],
    ['Tohoku total', '506,750', '19,806,452', '39.97 (mean)', '\u2014'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for paragraph in table2.rows[7].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True

doc.add_paragraph()

# --- Figure 2 ---
add_figure(OUTPUT_DIR + 'fig1_prefecture_bar.png',
           'Figure 2. Prefecture-level analgesic-per-surgery index for all 47 prefectures. Red bars indicate Tohoku prefectures; blue bars indicate all other prefectures. The dashed line indicates the national mean. All six Tohoku prefectures rank in the upper half nationally.')

doc.add_page_break()

# --- Table 3 ---
add_heading_text('Drug class subgroup analysis', level=2)
doc.add_paragraph('The pattern of higher Tohoku analgesic use was consistent across all three drug classes:')
doc.add_paragraph('NSAIDs/acetaminophen (Class 114): Tohoku mean 37.90 (SD 2.99) versus rest 33.66 (SD 5.49); Mann-Whitney U=186, P=0.044.', style='List Bullet')
doc.add_paragraph('Opioid alkaloids (Class 811): Tohoku mean 1.91 (SD 0.35) versus rest 1.43 (SD 0.27); Mann-Whitney U=211, P=0.003.', style='List Bullet')
doc.add_paragraph('Synthetic opioids (Class 821): Tohoku mean 0.16 (SD 0.06) versus rest 0.09 (SD 0.04); Mann-Whitney U=217, P=0.001.', style='List Bullet')
doc.add_paragraph('Tohoku prefectures prescribed 34% more opioid alkaloids and 86% more synthetic opioids per surgery compared with the rest of Japan.')

add_bold_paragraph('Table 3. ', 'Primary comparison: Tohoku versus rest of Japan')
table3 = doc.add_table(rows=6, cols=5)
table3.style = 'Light Shading Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Parameter', 'Tohoku (n=6)', 'Rest (n=41)', 'Diff (95% CI)', 'P value']):
    table3.rows[0].cells[i].text = h
data3 = [
    ['All analgesics', '39.97 (3.53)', '35.17 (5.71)', '4.80 (1.48\u20138.12)', '0.031'],
    ['NSAIDs (114)', '37.90 (2.99)', '33.66 (5.49)', '4.23 (1.09\u20137.38)', '0.044'],
    ['Opioid alkaloids (811)', '1.91 (0.35)', '1.43 (0.27)', '0.48 (0.21\u20130.76)', '0.003'],
    ['Synthetic opioids (821)', '0.16 (0.06)', '0.09 (0.04)', '0.07 (0.03\u20130.12)', '0.001'],
    ['', '', '', '', 'Mann-Whitney U'],
]
for r, row_data in enumerate(data3):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

doc.add_paragraph()

# --- Figure 3 ---
add_figure(OUTPUT_DIR + 'fig3_scatter.png',
           'Figure 3. Relationship between surgical volume and analgesic-per-surgery index. Each point represents one prefecture. Diamond markers (red) indicate Tohoku prefectures. No clear relationship between surgical volume and analgesic intensity is observed.')

doc.add_page_break()

# --- Figure 4 ---
add_figure(OUTPUT_DIR + 'fig4_boxplot.png',
           'Figure 4. Distribution of analgesic-per-surgery index by regional block (box plots). Boxes represent interquartile range with median line. Tohoku (red) shows a relatively compact distribution above the national median.')

doc.add_paragraph()

# Supplementary Figure
add_figure(OUTPUT_DIR + 'fig5_drug_class_breakdown.png',
           'Supplementary Figure 1. Drug class breakdown: per-surgery prescribing quantity by prefecture for NSAIDs/acetaminophen (Class 114), opioid alkaloids (Class 811), and synthetic opioids (Class 821). Red bars indicate Tohoku prefectures.',
           width=6.5)

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

add_heading_text('Principal findings', level=2)
doc.add_paragraph('This nationwide ecological study using Japan\'s NDB Open Data found that perioperative analgesic prescribing in the Tohoku region was significantly higher, not lower, than in the rest of Japan. The Tohoku region ranked 7th of 9 regional blocks, with a mean analgesic-per-surgery index 11.7% above the national average and a large effect size (Cohen\'s d=0.87). This pattern was consistent across NSAIDs, opioid alkaloids, and synthetic opioids. These findings directly contradict the hypothesis that the perceived cultural stoicism of Tohoku residents would manifest as lower perioperative analgesic consumption.')

add_heading_text('Comparison with existing evidence', level=2)
doc.add_paragraph('Our findings align with, and extend, previous work on regional variation in analgesic prescribing in Japan. Matsuoka and colleagues reported substantial prefectural variation in cancer pain opioid prescribing using NDB data. The broad geographical gradient we observed\u2014with higher use in the northern and southern peripheries and lower use in central Japan\u2014parallels patterns reported in regional healthcare utilisation studies.')
doc.add_paragraph('The 2017 Pfizer Japan survey found that northern prefectures reported higher rates of self-described pain endurance. Our data suggest that this self-reported stoicism does not suppress perioperative analgesic prescribing\u2014and may paradoxically coexist with higher prescribing.')

add_heading_text('Possible explanations', level=2)
doc.add_paragraph('First, the ecological fallacy must be acknowledged: individual-level pain endurance may not aggregate into prefecture-level prescribing differences because prescribing is primarily determined by clinician behaviour, institutional protocols, and drug availability rather than patient stoicism alone.')
doc.add_paragraph('Second, the geographical gradient may reflect differences in surgical case mix. Northern prefectures have older populations with higher rates of musculoskeletal and oncological surgery, which typically require more intensive analgesia.')
doc.add_paragraph('Third, regional variation in prescribing culture and practice patterns among clinicians may be the dominant driver. The higher opioid use in Tohoku may reflect specific institutional champions or training programmes that promote multimodal opioid-inclusive analgesia.')
doc.add_paragraph('Fourth, the higher use in peripheral regions may partly reflect longer hospital stays, which mechanically increase cumulative analgesic prescribing per surgical admission.')

add_heading_text('Implications', level=2)
doc.add_paragraph('Cultural stereotypes about regional pain tolerance should not influence clinical analgesic prescribing. The nearly two-fold variation in analgesic-per-surgery index across prefectures suggests substantial scope for standardisation. The particularly large regional variation in opioid prescribing warrants attention in the context of Japan\'s evolving opioid policies.')

add_heading_text('Strengths and limitations', level=2)
doc.add_paragraph('The NDB captures virtually all insurance-reimbursed healthcare in Japan, providing near-complete population coverage. However, as an ecological study, individual-level inferences are not possible. The analgesic-per-surgery index is not linked at the individual level. The NDB Open Data provide only oral prescription data; parenteral analgesics could not be included. We could not adjust for case mix, length of stay, or institutional factors. The analysis is limited to a single fiscal year.')

add_heading_text('Conclusions', level=1)
doc.add_paragraph('This nationwide ecological study found that perioperative analgesic prescribing in Japan\'s Tohoku region was significantly higher, not lower, than the national average, with large and consistent effects across NSAIDs and opioids. The perceived cultural stoicism of Tohoku residents does not translate into reduced analgesic consumption at the prefecture level. Nearly two-fold variation across prefectures highlights the need for standardisation of perioperative pain management. Cultural stereotypes should not substitute for evidence-based, individualised analgesic prescribing.')

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_text('References', level=1)
refs = [
    'Peacock S, Patel S. Cultural influences on pain. Rev Pain 2008; 1: 6-9',
    'Callister LC. Cultural influences on pain perceptions and behaviors. Home Health Care Manag Pract 2003; 15: 207-11',
    'Zaslansky R, Rothaug J, Chapman CR, et al. PAIN OUT: the making of an international acute pain registry. Eur J Pain 2015; 19: 490-502',
    'Schwenkglenks M, Gerbershagen HJ, Taylor RS, et al. Correlates of satisfaction with pain treatment in the acute postoperative period. Pain 2014; 155: 1401-11',
    'Deyo RA, Hallvik SE, Hildebran C, et al. Association between initial opioid prescribing patterns and subsequent long-term use. J Gen Intern Med 2017; 32: 21-7',
    'Nowakowska M, van Staa T, M\u00f6lter A, et al. Opioid analgesic prescribing in England: a regional analysis of variation. Br J Pain 2022; 16: 49-59',
    'OECD. OECD Reviews of Health Care Quality: Japan 2015. Paris: OECD Publishing, 2015',
    'Hendry J. Understanding Japanese Society. 4th ed. London: Routledge, 2013',
    'Pfizer Japan Inc. 47 Prefectures Pain Awareness Survey. Tokyo: Pfizer Japan, 2017',
    'Matsuoka H, Maeda I, Yoshioka A, et al. Regional variation in opioid prescribing for cancer pain in Japan. J Pain Symptom Manage 2025; 69: 125-34',
    'Shoji T, Murata S, Moriyama T, et al. Variation in multimodal analgesia protocols for TKA across Japanese institutions. J Arthroplasty 2025; 40: 456-63',
    'Ministry of Health, Labour and Welfare. 10th NDB Open Data. 2024. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html',
    'von Elm E, Altman DG, Egger M, et al. The STROBE statement. Lancet 2007; 370: 1453-7',
    'Benchimol EI, Smeeth L, Guttmann A, et al. The RECORD statement. PLoS Med 2015; 12: e1001885',
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain. J Pain 2009; 10: 1187-204',
    'Ministry of Health, Labour and Welfare. 2023 Patient Survey. Tokyo: MHLW, 2024',
    'Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan. Lancet 2011; 378: 1174-82',
    'OECD. Health at a Glance 2023: OECD Indicators. Paris: OECD Publishing, 2023',
    'International Narcotics Control Board. Narcotic Drugs: Estimated World Requirements for 2024. Vienna: United Nations, 2024',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# ============================================================
# SAVE
# ============================================================
outpath = OUTPUT_DIR + 'BJA_manuscript_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
