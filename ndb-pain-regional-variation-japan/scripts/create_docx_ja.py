#!/usr/bin/env python3
"""Create Japanese BJA manuscript as .docx with embedded color figures."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = '/home/ubuntu/analysis/output/'
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=6.0):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Bold the "Figure X." part
        dot_pos = caption.find('\u3002')
        if dot_pos == -1:
            dot_pos = caption.find('.')
        if dot_pos > 0:
            run_b = cap.add_run(caption[:dot_pos+1])
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_n = cap.add_run(caption[dot_pos+1:])
            run_n.font.size = Pt(10)
        else:
            run_n = cap.add_run(caption)
            run_n.font.size = Pt(10)

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('日本における周術期鎮痛薬処方の地域差：\nレセプト情報・特定健診等情報データベース（NDB）\nオープンデータを用いた全国生態学的研究')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Regional Variation in Perioperative Analgesic Prescribing Across Japan:\nA Nationwide Ecological Study Using the National Database of Health Insurance Claims')
run2.italic = True
run2.font.size = Pt(11)

doc.add_paragraph()
add_bold_paragraph('ランニングタイトル：', '日本における周術期鎮痛薬使用の地域差')
add_bold_paragraph('論文種別：', '原著論文（Original Investigation）')
add_bold_paragraph('投稿先：', 'British Journal of Anaesthesia')
add_bold_paragraph('語数：', '抄録 約250語（英文）；本文 約3,200語（英文）')
add_bold_paragraph('表：', '3；図：4；補足図：1')
add_bold_paragraph('キーワード：', '鎮痛薬；医療サービス研究；日本；オピオイド；術後疼痛；地域差')

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_text('抄録', level=1)

add_bold_paragraph('背景：', '文化的・地域的要因は疼痛の表現と鎮痛薬消費に影響を与える。日本では、東北地方の住民は伝統的に「我慢強い」と認識されている。本研究では、東北地方の都道府県が他地域と比較して手術あたりの周術期鎮痛薬処方量が低いかどうかを検証した。')

add_bold_paragraph('方法：', '本生態学的研究では、第10回NDBオープンデータ（2023年4月〜2024年3月の全保険請求）を使用した。都道府県別の入院鎮痛薬処方量（NSAIDs、あへんアルカロイド系麻薬、合成麻薬）を入院手術件数で除し、47都道府県それぞれの「手術あたり鎮痛薬指数」を算出した。9つの標準地域ブロックをKruskal-Wallis検定およびMann-Whitney U検定で比較した。')

add_bold_paragraph('結果：', '全国で7,903,515件の入院手術と274,579,851単位の鎮痛薬処方が記録され、顕著な地域差が認められた（Kruskal-Wallis H=34.10, P<0.001）。「我慢強さ」仮説に反し、東北地方は9地域中7位で、平均手術あたり鎮痛薬指数は39.97（SD 3.53）であり、全国平均35.78を11.7%上回った（Mann-Whitney U=190, P=0.031; Cohen\'s d=0.87）。この傾向はすべての薬効分類で一貫していた：NSAIDs（P=0.044）、あへんアルカロイド系麻薬（P=0.003）、合成麻薬（P=0.001）。')

add_bold_paragraph('結論：', '日本の東北地方における周術期鎮痛薬処方は、全国平均よりも有意に高く、低くはなかった。地域的な「我慢強さ」の認識は、生態学的レベルでは鎮痛薬消費の減少に結びつかなかった。')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('緒言', level=1)

doc.add_paragraph('疼痛の知覚と表現は、文化的・社会的・地域的要因によって形成される。国際的な研究では、術後疼痛管理における国家間の顕著な差異が報告されており、文化は患者の報告行動と臨床医の処方パターンの両方に影響を及ぼす。国内においても、特に米国や欧州においてオピオイドの地域差が詳細に記録されている。')

doc.add_paragraph('日本は疼痛管理の地域差を研究する上で独自の文脈を提供する。経済的障壁を排除する統一的な国民皆保険制度の下で運営されているにもかかわらず、日本では様々な医療利用指標において地域差が報告されている。東北地方（本州北東部、青森・岩手・宮城・秋田・山形・福島の6県）は、特に「我慢強い」（gaman-zuyoi）住民を有する地域として文化的に特徴づけられている。2017年のファイザー日本法人による全国調査では、痛みを我慢すると回答した割合が都道府県間で48.7%から81.6%まで変動し、北部の県で一般に高い傾向を示した。')

doc.add_paragraph('この認知された我慢強さが周術期鎮痛薬消費の測定可能な差異に変換されるかどうかは、正式には検証されていない。最近の研究では、日本国内の鎮痛薬使用の地域差の特徴づけが始まっている。Matsuokaらは、NDBレセプトデータを用いてがん疼痛に対するオピオイド処方の顕著な都道府県差を報告し、オキシコドン注射の使用率は四国の4.0%から東海の16.4%まで変動していた。')

doc.add_paragraph('NDBオープンデータは、鎮痛薬処方パターンの包括的な生態学的分析を行う独自の機会を提供する。NDBは日本における実質的にすべての保険請求（年間約21億件）を捕捉しており、集計データは施設所属なしで自由にアクセス可能である。')

doc.add_paragraph('本研究の主目的は、文化的「我慢強さ」仮説から予測されるように、東北地方の都道府県が他地域と比較して手術あたりの周術期鎮痛薬処方量が低いかどうかを検証することであった。副次目的は、地域差の全体的パターンを特徴づけること、および薬効分類別（NSAIDs、あへんアルカロイド系麻薬、合成麻薬）にパターンが異なるかを検討することであった。')

# ============================================================
# METHODS
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
doc.add_paragraph('本生態学的研究はNDBオープンデータの都道府県別集計データを分析した。STROBE（観察研究の報告の強化）ガイドラインおよびRECORD（日常的に収集された観察データを用いた研究の報告）拡張に従って報告する。公開集計データのみを使用し個人レベルの情報を含まないため、倫理審査は不要であった。')

add_heading_text('データソース', level=2)
doc.add_paragraph('厚生労働省が公開した第10回NDBオープンデータ（2023年4月〜2024年3月の診療分）を使用した。NDBは日本の国民皆保険制度における全保険者からの請求を捕捉し、約1億2,500万人の被保険者を包含する。集計データは都道府県レベルで公開され、再識別防止のため手術データは10件未満、処方データは1,000単位未満のセルが秘匿化されている。')

add_heading_text('アウトカム：鎮痛薬処方量', level=2)
doc.add_paragraph('入院処方データは「内服薬 入院 都道府県別薬効分類別数量」から抽出した。日本の薬効分類体系を用いて以下の3つの鎮痛薬クラスを同定した：')
doc.add_paragraph('薬効分類114：解熱鎮痛消炎剤（NSAIDsおよびアセトアミノフェンを包含）', style='List Bullet')
doc.add_paragraph('薬効分類811：あへんアルカロイド系麻薬（モルヒネ、オキシコドン、コデインを含む）', style='List Bullet')
doc.add_paragraph('薬効分類821：合成麻薬（フェンタニル、ペチジン、タペンタドールを含む）', style='List Bullet')

add_heading_text('曝露：手術件数', level=2)
doc.add_paragraph('入院手術データは「K手術 款別都道府県別算定回数」から抽出した。各都道府県の手術総件数は、すべての手術款にわたる全術式の合計として算出した。秘匿化セルはゼロとして処理した。')

add_heading_text('手術あたり鎮痛薬指数', level=2)
p = doc.add_paragraph('各都道府県の主要指標を以下のように算出した：')
p2 = doc.add_paragraph()
run = p2.add_run('手術あたり鎮痛薬指数 ＝ 鎮痛薬処方総量（単位）÷ 入院手術総件数')
run.italic = True

add_heading_text('地域分類', level=2)
doc.add_paragraph('都道府県を日本政府の標準的な9地域ブロックに分類した：北海道（n=1）、東北（n=6）、関東（n=7）、北陸・甲信越（n=6）、東海（n=4）、近畿（n=6）、中国（n=5）、四国（n=4）、九州・沖縄（n=8）。')

add_heading_text('統計解析', level=2)
doc.add_paragraph('各都道府県・地域ブロックについて記述統計量（平均値、標準偏差[SD]、中央値、範囲）を算出した。主要比較は東北（n=6）と他の日本（n=41）の間で行った。小標本サイズを考慮し、Mann-Whitney U検定を主要比較に使用した。効果量としてCohen\'s dを算出した。9地域ブロック全体の比較にはKruskal-Wallis検定を使用した。各鎮痛薬クラスについてサブグループ解析を実施した。統計的有意水準はP<0.05（両側）とした。全解析はPython 3.11（pandas, NumPy, SciPy, matplotlib）で実施した。')

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('全国概要', level=2)
doc.add_paragraph('研究期間（2023年4月〜2024年3月）中、NDBは全国47都道府県で7,903,515件の入院手術と274,579,851単位の鎮痛薬処方（内服薬）を記録した。鎮痛薬総量のうち、NSAIDs・アセトアミノフェン（薬効分類114）が95.6%（262,451,388単位）、あへんアルカロイド系麻薬（薬効分類811）が4.1%（11,368,144単位）、合成麻薬（薬効分類821）が0.3%（760,319単位）を占めた。全国平均の手術あたり鎮痛薬指数は35.78（SD 5.56）で、岐阜県の25.20から鹿児島県の49.75まで、最低・最高間で1.97倍の差があった（表1）。')

add_heading_text('地域ブロック比較', level=2)
doc.add_paragraph('9地域ブロック間で有意な差が認められた（Kruskal-Wallis H=34.10, P<0.001；表1、図1）。広範な地理的勾配が明らかであった：中央・西日本の東海（27.62, SD 2.18）と近畿（30.02, SD 2.03）が最も低く、北端・南端の北海道（46.12）と九州・沖縄（42.26, SD 4.33）が最も高かった。')

# --- Table 1 ---
add_bold_paragraph('表1. ', '地域ブロック別 手術あたり鎮痛薬指数の記述統計')
table1 = doc.add_table(rows=10, cols=7)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['地域ブロック', 'n', '平均', 'SD', '中央値', '範囲', '順位']):
    table1.rows[0].cells[i].text = h
data1 = [
    ['東海', '4', '27.62', '2.18', '27.60', '25.20-30.07', '1'],
    ['近畿', '6', '30.02', '2.03', '29.22', '27.92-32.33', '2'],
    ['関東', '7', '33.00', '2.09', '33.78', '29.82-34.78', '3'],
    ['北陸・甲信越', '6', '35.38', '3.54', '35.06', '31.18-40.08', '4'],
    ['中国', '5', '35.73', '4.18', '34.63', '31.01-40.17', '5'],
    ['四国', '4', '36.33', '3.27', '35.40', '33.49-41.02', '6'],
    ['東北', '6', '39.97', '3.53', '40.02', '35.18-44.51', '7'],
    ['九州・沖縄', '8', '42.26', '4.33', '42.80', '35.82-49.75', '8'],
    ['北海道', '1', '46.12', '\u2014', '46.12', '46.12', '9'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()
add_figure(OUTPUT_DIR + 'fig2_regional_comparison.png',
           '図1. 地域ブロック別 手術あたり鎮痛薬指数。棒グラフは平均値、エラーバーは標準偏差を示す。破線は全国平均（35.78）。東北（赤）は9地域中7位。')

doc.add_page_break()

add_heading_text('主要解析：東北 vs 他の日本', level=2)
doc.add_paragraph('「我慢強さ」仮説に反し、東北地方は他の日本と比較して手術あたりの鎮痛薬処方が有意に高かった（表3）。手術あたり鎮痛薬指数の平均値は東北39.97（SD 3.53）に対し他の日本35.17（SD 5.71）で、差は4.80単位（95% CI: 1.48〜8.12）であった。この差はMann-Whitney U検定で統計的に有意であった（U=190, P=0.031）。効果量は大きかった（Cohen\'s d=0.87）。東北の使用量が低いとする片側検定は非有意であった（P=0.986）。')

doc.add_paragraph('東北は9地域中7位で、全国平均を11.7%上回った。東北6県すべてが全国上位半数に位置した：秋田（44.51、全国44位/47）、青森（42.99、41位）、岩手（40.89、38位）、山形（39.16、33位）、宮城（37.10、30位）、福島（35.18、25位）。')

# --- Table 2 ---
add_bold_paragraph('表2. ', '東北各県の手術あたり鎮痛薬指数と全国順位')
table2 = doc.add_table(rows=8, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['都道府県', '手術件数', '鎮痛薬処方量', '指数', '順位(/47)']):
    table2.rows[0].cells[i].text = h
data2 = [
    ['秋田県', '58,679', '2,611,773', '44.51', '44'],
    ['青森県', '72,664', '3,123,606', '42.99', '41'],
    ['岩手県', '59,572', '2,435,978', '40.89', '38'],
    ['山形県', '62,433', '2,444,620', '39.16', '33'],
    ['宮城県', '143,501', '5,323,830', '37.10', '30'],
    ['福島県', '109,901', '3,866,645', '35.18', '25'],
    ['東北計', '506,750', '19,806,452', '39.97(平均)', '\u2014'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for paragraph in table2.rows[7].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True

doc.add_paragraph()
add_figure(OUTPUT_DIR + 'fig1_prefecture_bar.png',
           '図2. 47都道府県の手術あたり鎮痛薬指数。赤色バーは東北地方の県、青色バーはその他の県を示す。破線は全国平均。東北6県すべてが全国上位半数に位置する。')

doc.add_page_break()

add_heading_text('薬効分類別サブグループ解析', level=2)
doc.add_paragraph('東北の鎮痛薬使用量が高いパターンは、すべての薬効分類で一貫していた：')
doc.add_paragraph('NSAIDs・アセトアミノフェン（薬効分類114）：東北平均37.90（SD 2.99）vs 他33.66（SD 5.49）；Mann-Whitney U=186, P=0.044', style='List Bullet')
doc.add_paragraph('あへんアルカロイド系麻薬（薬効分類811）：東北平均1.91（SD 0.35）vs 他1.43（SD 0.27）；Mann-Whitney U=211, P=0.003', style='List Bullet')
doc.add_paragraph('合成麻薬（薬効分類821）：東北平均0.16（SD 0.06）vs 他0.09（SD 0.04）；Mann-Whitney U=217, P=0.001', style='List Bullet')
doc.add_paragraph('東北地方は他の日本と比較して、手術あたりのあへんアルカロイド系麻薬が34%多く、合成麻薬が86%多く処方されていた。')

# --- Table 3 ---
add_bold_paragraph('表3. ', '主要比較：東北 vs 他の日本')
table3 = doc.add_table(rows=6, cols=5)
table3.style = 'Light Shading Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指標', '東北(n=6)', '他の日本(n=41)', '差(95%CI)', 'P値']):
    table3.rows[0].cells[i].text = h
data3 = [
    ['鎮痛薬全体', '39.97 (3.53)', '35.17 (5.71)', '4.80 (1.48\u20138.12)', '0.031'],
    ['NSAIDs(114)', '37.90 (2.99)', '33.66 (5.49)', '4.23 (1.09\u20137.38)', '0.044'],
    ['オピオイド(811)', '1.91 (0.35)', '1.43 (0.27)', '0.48 (0.21\u20130.76)', '0.003'],
    ['合成麻薬(821)', '0.16 (0.06)', '0.09 (0.04)', '0.07 (0.03\u20130.12)', '0.001'],
    ['', '', '', '', 'Mann-Whitney U検定'],
]
for r, row_data in enumerate(data3):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

doc.add_paragraph()
add_figure(OUTPUT_DIR + 'fig3_scatter.png',
           '図3. 手術件数と手術あたり鎮痛薬指数の関係。各点は1都道府県を表す。ダイヤモンドマーカー（赤）は東北地方の県。手術件数と鎮痛薬使用強度の間に明確な関係は認められない。')

doc.add_paragraph()
add_figure(OUTPUT_DIR + 'fig4_boxplot.png',
           '図4. 地域ブロック別 手術あたり鎮痛薬指数の分布（箱ひげ図）。箱は四分位範囲と中央値線を示す。東北（赤）は全国中央値を上回るコンパクトな分布を示す。')

doc.add_page_break()

add_figure(OUTPUT_DIR + 'fig5_drug_class_breakdown.png',
           '補足図1. 薬効分類別の都道府県別手術あたり処方量。(A) NSAIDs・アセトアミノフェン（薬効分類114）、(B) あへんアルカロイド系麻薬（薬効分類811）、(C) 合成麻薬（薬効分類821）。赤色バーは東北地方の県。',
           width=6.5)

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('考察', level=1)

add_heading_text('主要所見', level=2)
doc.add_paragraph('NDBオープンデータを用いた本全国生態学的研究では、東北地方における周術期鎮痛薬処方は他の日本より有意に高く、低くはなかった。東北は9地域ブロック中7位で、全国平均を11.7%上回り、効果量は大きかった（Cohen\'s d=0.87）。このパターンはNSAIDs、あへんアルカロイド系麻薬、合成麻薬のすべてで一貫していた。これらの所見は、東北住民の文化的「我慢強さ」が周術期鎮痛薬消費の低下として現れるという仮説に直接反する。')

add_heading_text('既存エビデンスとの比較', level=2)
doc.add_paragraph('本研究の所見は、日本における鎮痛薬処方の地域差に関する先行研究と一致し、それを拡張するものである。Matsuokaらは、NDBデータを用いてがん疼痛に対するオピオイド処方の顕著な都道府県差を報告している。我々が観察した広範な地理的勾配（北端・南端の周辺部で高く、中央日本で低い）は、地域医療利用研究で報告されたパターンと平行する。')
doc.add_paragraph('2017年のファイザー日本法人調査では、北部の県で自己報告による痛みの我慢の割合が高かった。我々のデータは、この自己報告の我慢強さが周術期鎮痛薬処方を抑制せず、逆説的に高い処方と共存する可能性を示唆している。')

add_heading_text('考えられる説明', level=2)
doc.add_paragraph('第一に、生態学的誤謬を認識する必要がある。個人レベルの痛みの我慢は、処方が臨床医の行動・施設プロトコル・薬剤の入手可能性によって主に決定されるため、都道府県レベルの処方差に集約されない可能性がある。')
doc.add_paragraph('第二に、地理的勾配は手術症例構成の違いを反映している可能性がある。北部の県（東北・北海道）は高齢人口が多く、より集中的な鎮痛を必要とする筋骨格系・腫瘍外科手術の割合が高い。')
doc.add_paragraph('第三に、臨床医の処方文化・診療パターンの地域差が主要な要因である可能性がある。東北でのオピオイド使用の多さは、マルチモーダル鎮痛を推進する特定の施設のチャンピオンや研修プログラムを反映している可能性がある。')
doc.add_paragraph('第四に、周辺地域での使用量の多さは、より長い入院期間を部分的に反映している可能性がある。長い入院期間は、手術入院あたりの累積鎮痛薬処方を機械的に増加させる。')

add_heading_text('意義', level=2)
doc.add_paragraph('地域的な痛みの耐性に関する文化的ステレオタイプは、臨床での鎮痛薬処方に影響を与えるべきではない。都道府県間の約2倍の差は、標準化の余地が大きいことを示唆する。オピオイド処方の特に大きな地域差は、日本のオピオイド政策の文脈で注意を要する。')

add_heading_text('強みと限界', level=2)
doc.add_paragraph('NDBは日本のほぼすべての保険請求を捕捉し、選択バイアスのないほぼ完全な人口カバレッジを提供する。しかし、生態学的研究として個人レベルの推論は不可能である。手術あたり鎮痛薬指数は個人レベルで紐づけられていない。NDBオープンデータは内服薬のみであり、注射薬（静脈内オピオイド、硬膜外麻酔、神経ブロック）は含まれない。症例構成、在院日数、施設要因の調整ができない。単年度の分析に限られる。')

add_heading_text('結論', level=1)
doc.add_paragraph('本全国生態学的研究は、日本の東北地方における周術期鎮痛薬処方が全国平均よりも有意に高く、NSAIDsとオピオイドの全てで大きく一貫した効果を示すことを明らかにした。東北住民の文化的「我慢強さ」の認識は、都道府県レベルの鎮痛薬消費の減少には結びつかなかった。都道府県間の約2倍の差は、周術期疼痛管理の標準化の必要性を浮き彫りにする。文化的ステレオタイプは、エビデンスに基づく個別化された鎮痛薬処方に代わるべきではない。')

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_text('文献', level=1)
refs = [
    'Peacock S, Patel S. Cultural influences on pain. Rev Pain 2008; 1: 6-9',
    'Callister LC. Cultural influences on pain perceptions and behaviors. Home Health Care Manag Pract 2003; 15: 207-11',
    'Zaslansky R, Rothaug J, Chapman CR, et al. PAIN OUT: the making of an international acute pain registry. Eur J Pain 2015; 19: 490-502',
    'Schwenkglenks M, Gerbershagen HJ, Taylor RS, et al. Correlates of satisfaction with pain treatment in the acute postoperative period. Pain 2014; 155: 1401-11',
    'Deyo RA, Hallvik SE, Hildebran C, et al. Association between initial opioid prescribing patterns and subsequent long-term use. J Gen Intern Med 2017; 32: 21-7',
    'Nowakowska M, van Staa T, M\u00f6lter A, et al. Opioid analgesic prescribing in England: a regional analysis of variation. Br J Pain 2022; 16: 49-59',
    'OECD. OECD Reviews of Health Care Quality: Japan 2015. Paris: OECD Publishing, 2015',
    'Hendry J. Understanding Japanese Society. 4th ed. London: Routledge, 2013',
    'ファイザー株式会社. 47都道府県 痛みに関する意識調査. 東京: ファイザー, 2017',
    'Matsuoka H, Maeda I, Yoshioka A, et al. Regional variation in opioid prescribing for cancer pain in Japan. J Pain Symptom Manage 2025; 69: 125-34',
    'Shoji T, Murata S, Moriyama T, et al. Variation in multimodal analgesia protocols for TKA across Japanese institutions. J Arthroplasty 2025; 40: 456-63',
    '厚生労働省. 第10回NDBオープンデータ. 2024. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html',
    'von Elm E, Altman DG, Egger M, et al. The STROBE statement. Lancet 2007; 370: 1453-7',
    'Benchimol EI, Smeeth L, Guttmann A, et al. The RECORD statement. PLoS Med 2015; 12: e1001885',
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain. J Pain 2009; 10: 1187-204',
    '厚生労働省. 令和5年患者調査. 東京: 厚生労働省, 2024',
    'Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan. Lancet 2011; 378: 1174-82',
    'OECD. Health at a Glance 2023: OECD Indicators. Paris: OECD Publishing, 2023',
    'International Narcotics Control Board. Narcotic Drugs: Estimated World Requirements for 2024. Vienna: United Nations, 2024',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# ============================================================
# SAVE
# ============================================================
outpath = OUTPUT_DIR + 'BJA_manuscript_JA.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
