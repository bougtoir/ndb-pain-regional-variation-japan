#!/usr/bin/env python3
"""Create STROBE checklist for PAIN Reports submission."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# Landscape A4
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(9)

h = doc.add_heading('STROBE Statement\u2014Checklist of items for cross-sectional studies (ecological design)', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)

# STROBE items
items = [
    ('', 'Item No', 'Recommendation', 'Reported on page / section'),
    ('Title and abstract', '1', '(a) Indicate the study\u2019s design with a commonly used term in the title or abstract\n'
     '(b) Provide in the abstract an informative and balanced summary of what was done and what was found',
     '(a) Title: "Regional Heterogeneity..." implies ecological design; Abstract line 1: "ecological study"\n'
     '(b) Abstract: structured (Introduction, Methods, Results, Conclusions)'),

    ('Introduction', '', '', ''),
    ('Background/rationale', '2', 'Explain the scientific background and rationale for the investigation being reported',
     'Introduction, paragraphs 1\u20133'),
    ('Objectives', '3', 'State specific objectives, including any prespecified hypotheses',
     'Introduction, paragraph 4: three objectives stated'),

    ('Methods', '', '', ''),
    ('Study design', '4', 'Present key elements of study design early in the paper',
     'Methods: "This ecological study analyzed prefecture-level aggregate data..."'),
    ('Setting', '5', 'Describe the setting, locations, and relevant dates',
     'Methods, Data source: NDB Open Data, April 2023\u2013March 2024, 47 prefectures'),
    ('Participants', '6', 'Give the eligibility criteria, and the sources and methods of selection of participants',
     'Methods, Data source: population-complete insurance claims, ~125 million insured; '
     'Methods, Regional classification: nine blocks'),
    ('Variables', '7', 'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
     'Methods, Phase 1 (analgesic classes), Phase 2 (neuropathic pain drugs), Confounder proxies (4 classes)'),
    ('Data sources/measurement', '8', 'For each variable, give sources of data and details of assessment methods',
     'Methods: NDB drug classification codes, K Surgery section procedure codes'),
    ('Bias', '9', 'Describe any efforts to address potential sources of bias',
     'Methods, Confounder disease proxies: four disease proxies extracted; '
     'Discussion, Strengths and limitations'),
    ('Study size', '10', 'Explain how the study size was arrived at',
     'Results: 7,903,515 surgical procedures, 274,579,851 analgesic units across 47 prefectures (population-complete)'),
    ('Quantitative variables', '11', 'Explain how quantitative variables were handled in the analyses',
     'Methods: analgesic-per-surgery index, neuropathic pain prescribing-per-surgery index, '
     'confounder proxies per surgery, adjusted CPSP index (residuals)'),
    ('Statistical methods', '12', '(a) Describe all statistical methods\n'
     '(b) Describe any methods used to examine subgroups and interactions\n'
     '(c) Explain how missing data were addressed\n'
     '(d) If applicable, describe analytical methods taking account of sampling strategy\n'
     '(e) Describe any sensitivity analyses',
     '(a) Methods, Statistical analysis: Kruskal\u2013Wallis, Mann\u2013Whitney U, five regression models\n'
     '(b) Tohoku vs non-Tohoku comparisons; regional block analysis\n'
     '(c) NDB suppresses cells <10; population-complete data minimizes missingness\n'
     '(d) N/A (population-complete, not sampled)\n'
     '(e) Models 2\u20135 as sensitivity analyses with different outcome definitions'),

    ('Results', '', '', ''),
    ('Participants', '13', 'Report numbers of individuals at each stage of study',
     'Results, Phase 1: 7,903,515 procedures, 274,579,851 analgesic units'),
    ('Descriptive data', '14', 'Give characteristics of study participants and information on exposures and potential confounders',
     'Results: Table 1 (regional summary), confounder correlations (Fig. 2)'),
    ('Outcome data', '15', 'Report numbers of outcome events or summary measures',
     'Results: analgesic-per-surgery index range 25.20\u201349.75; neuropathic pain index by prefecture (Fig. 1)'),
    ('Main results', '16', 'Give unadjusted estimates and, if applicable, confounder-adjusted estimates',
     'Results: unadjusted (d = 2.81, P < 0.001) and adjusted (d = 0.44, P = 0.323); Table 2, Table 3'),
    ('Other analyses', '17', 'Report other analyses done\u2014e.g., analyses of subgroups and interactions',
     'Results: Phase 1\u2013Phase 2 integration (Fig. 5), Z-score heatmap (SFig. 1)'),

    ('Discussion', '', '', ''),
    ('Key results', '18', 'Summarise key results with reference to study objectives',
     'Discussion, paragraph 1'),
    ('Limitations', '19', 'Discuss limitations of the study',
     'Discussion, Strengths and limitations'),
    ('Interpretation', '20', 'Give a cautious overall interpretation of results',
     'Discussion: regional variation, confounders, clinical implications, conclusion'),
    ('Generalisability', '21', 'Discuss the generalisability (external validity) of the study results',
     'Discussion: clinical implications extend beyond Japan; '
     'Implications and future directions'),

    ('Other information', '', '', ''),
    ('Funding', '22', 'Give the source of funding and the role of the funders',
     'Acknowledgments: no specific funding disclosed'),
]

table = doc.add_table(rows=len(items), cols=4)
set_table_borders(table)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(1.5)
    row.cells[2].width = Cm(12.0)
    row.cells[3].width = Cm(9.0)

for i, (section_name, item_no, recommendation, page_ref) in enumerate(items):
    row = table.rows[i]
    row.cells[0].text = section_name
    row.cells[1].text = item_no
    row.cells[2].text = recommendation
    row.cells[3].text = page_ref

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'

    # Header row and section headers
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
    elif item_no == '':
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

doc.add_paragraph()
note = doc.add_paragraph(
    'Reference: von Elm E, et al. The STROBE statement: guidelines for reporting observational studies. '
    'Lancet 2007;370:1453\u20131457. '
    'Extended with RECORD recommendations: Benchimol EI, et al. PLoS Med 2015;12:e1001885.')
note.runs[0].font.size = Pt(8)
note.runs[0].font.italic = True

outpath = OUTPUT_DIR + 'PainReports_STROBE_checklist.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
