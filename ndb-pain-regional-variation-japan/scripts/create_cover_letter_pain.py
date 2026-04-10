#!/usr/bin/env python3
"""Create cover letter for Pain (IASP) journal submission."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# --- Page setup: A4, standard margins ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.15

# ============================================================
# DATE
# ============================================================
date_p = doc.add_paragraph()
date_p.add_run('[Date]')

doc.add_paragraph()

# ============================================================
# ADDRESSEE
# ============================================================
doc.add_paragraph('Editor-in-Chief')
doc.add_paragraph('Pain')
doc.add_paragraph('International Association for the Study of Pain (IASP)')

doc.add_paragraph()

# ============================================================
# SUBJECT LINE
# ============================================================
subj = doc.add_paragraph()
run_subj = subj.add_run(
    'Re: Submission of Original Article \u2014 '
    '\u201c\u201cJapanese Patient\u201d? \u2014 A Patient. '
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures '
    'Challenges the Stereotype of a Stoic Monolith\u201d')
run_subj.bold = True

doc.add_paragraph()

# ============================================================
# SALUTATION
# ============================================================
doc.add_paragraph('Dear Editor,')

doc.add_paragraph()

# ============================================================
# BODY
# ============================================================

# Para 1: What we are submitting and why Pain
doc.add_paragraph(
    'We are pleased to submit the above-titled manuscript for consideration as an Original Article in Pain. '
    'We believe this work speaks directly to a principle at the heart of the revised IASP definition of pain '
    '(Raja et al., Pain, 2020): that pain is always a personal experience, '
    'and that verbal report is only one of several behaviors by which it may be expressed. '
    'Our study provides population-level evidence that challenges the monolithic characterization '
    'of an entire nation\u2019s pain behavior\u2014and, by extension, '
    'any cultural stereotype applied to clinical pain assessment.')

# Para 2: What we found (key results)
doc.add_paragraph(
    'Using Japan\u2019s National Database (NDB) Open Data\u2014population-complete insurance claims '
    'covering approximately 125 million insured individuals\u2014we mapped pain-related prescribing '
    'across all 47 prefectures and nine regional blocks. '
    'We found 1.97-fold variation in the inpatient analgesic-per-surgery index, '
    'despite Japan\u2019s uniform national health insurance system and standardized drug pricing. '
    'The Tohoku region, traditionally perceived as the most stoic part of Japan, '
    'actually prescribed more, not fewer, analgesics than the national average '
    '(P = 0.031; Cohen\u2019s d = 0.87)\u2014a finding that directly contradicts '
    'the popular stereotype. After adjustment for confounding diseases '
    '(diabetic neuropathy, postherpetic neuralgia, depression, anxiety), '
    'Tohoku\u2019s apparent excess in neuropathic pain prescribing became nonsignificant, '
    'with diabetes drug prescribing identified as the strongest confounder (r = 0.87).')

# Para 3: Why this matters (clinical harm from stereotypes)
doc.add_paragraph(
    'We wish to draw the Editor\u2019s attention to the broader clinical message of this work. '
    'If nearly twofold variation in pain-related prescribing exists within a country '
    'internationally stereotyped as uniformly stoic, '
    'then the assumption that national-level cultural labels predict individual pain behavior is untenable. '
    'This is not merely an academic observation: '
    'when cultural stereotypes narrow the range of treatment options that clinicians consider for a patient, '
    'clinical harm may result. '
    'This principle applies not only to Japanese patients treated abroad, '
    'but to any patient whose pain assessment is filtered through a cultural label rather than '
    'individualized evaluation. '
    'We believe this universal message aligns with Pain\u2019s mission to advance pain research '
    'and improve the care of people in pain worldwide.')

# Para 4: Methodological strengths
doc.add_paragraph(
    'The perioperative setting of our Phase 1 analysis offers a methodological advantage '
    'over community-based pain surveys: because all surgical patients are hospitalized, '
    'healthcare access\u2014a major confounding factor in outpatient studies\u2014is effectively neutralized. '
    'Furthermore, the two-phase design (acute perioperative analgesics followed by '
    'chronic postsurgical pain proxy with confounder adjustment) '
    'provides complementary perspectives on the same underlying question. '
    'We acknowledge the ecological design as a limitation, '
    'but submit that population-level evidence is a necessary first step '
    'toward individual-level investigations, and that our findings establish '
    'a clear rationale for such studies.')

# Para 5: Declarations
doc.add_paragraph(
    'This manuscript has not been published previously, is not under consideration elsewhere, '
    'and its publication is approved by all authors. '
    'The authors declare no conflicts of interest. '
    'As the study used only publicly available aggregate data, '
    'ethical approval was not required. '
    'The NDB Open Data and analysis code are publicly available (URLs provided in the manuscript).')

# Para 6: Suggested reviewers (optional but helpful)
doc.add_paragraph(
    'Should the Editor find it helpful, we would be pleased to suggest potential reviewers '
    'with expertise in cross-cultural pain research and population-level pain epidemiology. '
    'We would also be grateful if the manuscript could be considered for '
    'accompanying Editorial Commentary, given the breadth of its clinical implications.')

# Para 7: Closing
doc.add_paragraph(
    'We thank you for your consideration and look forward to your response.')

doc.add_paragraph()

# ============================================================
# SIGN-OFF
# ============================================================
doc.add_paragraph('Sincerely,')

doc.add_paragraph()

signoff = doc.add_paragraph()
signoff.add_run('[Corresponding author name]\n').bold = False
signoff.add_run('[Title, Department]\n')
signoff.add_run('[Institution]\n')
signoff.add_run('[Address]\n')
signoff.add_run('E-mail: [email]\n')
signoff.add_run('Tel: [phone]')

# ============================================================
# SAVE
# ============================================================
out_path = OUTPUT_DIR + 'Pain_cover_letter.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
