#!/usr/bin/env python3
"""Create Japanese manuscript for Pain (IASP) format.

NEW NARRATIVE (v2):
  "Japanese are stoic (cross-cultural evidence) -> Is there regional variation WITHIN Japan?
   -> Exploratory analysis -> Cultural stereotypes narrow treatment options for ALL patients"

Key changes from v1:
- Introduction starts from cross-cultural pain literature, not Tohoku stereotype
- Study is framed as exploratory, not hypothesis-testing
- Discussion includes homogeneity myth, pain as individual experience, objective NANB monitoring
- Broader message: cultural stereotypes harm all patients, not just Japanese
- References updated to 23 (includes Befu, Burgess, IASP, Onishi NPV)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import json

OUTPUT_DIR = '/home/ubuntu/analysis/output/'

doc = Document()

# --- Page setup: A4, wide margins, double-spaced ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing = 2.0

# Load regression summary
with open(OUTPUT_DIR + 'cpsp_regression_summary.json', 'r') as f:
    reg = json.load(f)

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_figure(fig_path, caption, width=5.5):
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        parts = caption.split('. ', 1)
        if len(parts) == 2:
            run_b = cap.add_run(parts[0] + '. ')
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_n = cap.add_run(parts[1])
            run_n.font.size = Pt(10)
        else:
            run_b = cap.add_run(caption)
            run_b.font.size = Pt(10)

# ============================================================
# TITLE PAGE (Page 1)
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    '\u201c日本人患者\u201dか？\u2014\u2014ただの患者だ。\n'
    '47都道府県における疼痛関連処方の地域間異質性は\n'
    '「我慢強い単一民族」というステレオタイプに疑義を呈する'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# English title (JA requires both)
title_en = doc.add_paragraph()
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_en = title_en.add_run(
    '\u201cJapanese Patient\u201d? \u2014 A Patient.\n'
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures\n'
    'Challenges the Stereotype of a Stoic Monolith'
)
run_en.font.size = Pt(12)
run_en.italic = True

doc.add_paragraph()

# Authors placeholder
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[著者名・所属を挿入]')
run.font.size = Pt(12)

doc.add_paragraph()

# Corresponding author
add_bold_paragraph('責任著者：',
    '[氏名]、[所属]麻酔科学教室、'
    '[住所]、[郵便番号]、日本　'
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

# Keywords (3-5 as required by JA)
add_bold_paragraph('キーワード：',
    '文化的疼痛行動; 生態学的研究; 我慢; 神経障害性疼痛; 地域差; 処方パターン')

doc.add_paragraph()

# Manuscript info
add_bold_paragraph('論文種別：', '原著論文 (Research Paper)')
add_bold_paragraph('投稿先：', 'Pain')
add_bold_paragraph('ページ数：', '[フォーマット後に決定]')
add_bold_paragraph('図数：', '7（6本文 + 1補足）')
add_bold_paragraph('表数：', '2')

doc.add_page_break()

# ============================================================
# ABSTRACT (new page)
# ============================================================
add_heading_text('抄録', level=1)

# Pain uses unstructured abstract (single paragraph)
doc.add_paragraph(
    '異文化研究は日本人が西洋人と比べて痛みに対し忍耐強いことを一貫して示してきたが、'
    'この文化的疼痛忍耐が日本国内で地域的に異なるかどうかは未解明であった。'
    '本生態学的研究では、日本のNDBオープンデータ（集団完全保険請求、2023年4月〜2024年3月）を用い、'
    '47都道府県・9地域ブロックにおける疼痛関連処方を検討した。'
    'Phase 1では、入院鎮痛薬/手術指標が都道府県間で1.97倍の差を示した'
    '（岐阜25.20〜鹿児島49.75; Kruskal–Wallis P < 0.001）。'
    '周術期は全患者が入院中であるため、医療アクセスが交絡因子とならない方法論上の利点がある。'
    '伝統的に我慢強いとされる東北地方は9地域中7位であり'
    '（平均39.97 vs 全国35.78; P = 0.031; Cohen\u2019s d = 0.87）、鎮痛薬使用はむしろ多かった。'
    f'Phase 2では、外来神経障害性疼痛薬処方（術後遷延性疼痛プロキシ）において'
    f'東北が未調整で最高値を示したが'
    f'（平均{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}）、'
    f'交絡疾患（糖尿病、帯状癒疹後神経痛、うつ病、不安障害）の調整後は有意でなくなった'
    f'（P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}）。'
    '糖尿病薬処方が最強の交絡因子であった（r = 0.87）。'
    'これらの知見は、日本の疼痛文化が一枚岩でないことを実証する：'
    '鎮痛薬処方の約2倍の国内変動は、「日本人」を疼痛行動の均質なカテゴリーとする仮定に疑義を呈する。'
    'あらゆる国民集団の疼痛行動を一枚岩的に特徴づけることは臨床的害悪のリスクを伴う'
    '——文化的ステレオタイプが治療選択肢の範囲を狭めるとき、患者は治療上の不利益を被る。'
    '文化的ステレオタイプに基づく仮定ではなく、客観的侵害受容モニタリングに基づく個別化された疼痛評価が、'
    '周術期ケアの標準であるべきである。')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('緒言', level=1)

# Para 1: Cross-cultural pain literature (foundation)
doc.add_paragraph(
    '疼痛は普遍的な体験であるが、その表現・忍耐・管理は文化によって深く形成される [1,2]。'
    'Zborowskiが米国内の民族集団間で疼痛行動が著しく異なることを観察した先駆的研究以来 [3]、'
    '文化的規範が個人の疼痛報告・受療行動・忍耐に影響を及ぼすことを示す'
    '膨大な文献が蓄積されてきた [4]。'
    'これらの異文化間差異には実際的含意がある：'
    '術後の疼痛評価、鎮痛薬処方、ひいては患者転帰に影響を及ぼすからである。')

# Para 2: Japanese stoicism specifically (gaman)
doc.add_paragraph(
    '世界の文化の中で、日本は痛みに対して特に忍耐強い傾向を持つと一貫して特徴づけられている。'
    'Hobaraは、日本人回答者が疼痛行動（うめく、泣く等）を'
    'ヨーロッパ系アメリカ人よりも有意に不適切と評価したことを報告した [5]。'
    'Fengらは、EQ-5Dを用いて日本人参加者が英国参加者よりも'
    '痛み・不快感を回避するために時間を犠牲にする意向が著しく低いことを示した：'
    '中等度から重度の痛みへの効用値減少は、英国で0.65に対し日本では0.15に過ぎなかった [6]。'
    'この文化的に媒介された疼痛忍耐は、日本語の「我慢」（\u6211\u6162、忍耐の徳）'
    'という概念に集約され、核心的な文化的価値と広く認識されている [7]。'
    'このような忍耐は、集団レベルでの疼痛の過小報告と鎮痛薬の過小利用をもたらしうる。'
    '国際的には、このステレオタイプは臨床的リスクを伴う：'
    '日本人患者に不慣れな臨床医は、忍耐的な態度を鎮痛薬の必要性が低いと誤解する可能性がある。')

# Para 3: Within-country regional variation (gap)
doc.add_paragraph(
    '文化的疼痛行動に関するほとんどの研究は国家間比較である。'
    'しかし、国内にも相当な文化的異質性が存在する。'
    'CohenとNisbettは、米国南部の「名誉の文化」が'
    '北部とは測定可能な行動的・生理学的反応の差を生み出すことを実証した [8]'
    '——現実の帰結を伴う国内文化差である。'
    '日本でも地域文化のアイデンティティは強い：'
    '東北地方は伝統的に忍耐強く逞しいと認識され、'
    '他の地域にもそれぞれ独自の文化的特性がある [9]；'
    '痛みの擬態語「うずく」ですら、地域によって使用パターンが異なることが報告されている [10]。'
    'ファイザー日本の全国調査では、痛みを「我慢する」と回答した慢性疼痛患者の割合は'
    '都道府県間で48.7%〜81.6%と大きく異なっていた [11]。'
    'しかし、このような態度の差が疼痛関連医療利用の測定可能な差異として'
    '反映されるかどうかは未解明のままである。')

# Para 4: Japan as natural experiment + NDB
doc.add_paragraph(
    '日本は疼痛行動の地域差を研究する上で理想的な環境を提供する。'
    '国民皆保険制度がアクセスの経済的障壁を排除し、'
    '標準化された薬価が市場主導の処方変動を除去し、'
    'NDB（レセプト情報・特定健診等情報データベース）が'
    '保険請求されたほぼ全ての医療を収録する（年間約21億件） [12]。'
    'さらに、近年の研究は疼痛関連指標における有意な地域差を報告している：'
    'Wakaizumiらは都道府県間で高インパクト慢性疼痛有病率に最大1.6倍の差を認め [13]、'
    'Matsuokaらはがん疼痛に対するオピオイド処方に4倍の地域差を報告した [14]。'
    'しかし、都道府県レベルでの周術期および慢性疼痛関連処方を検討した研究は存在しない。')

doc.add_paragraph(
    '周術期という研究対象は、地域在住者を対象とした疼痛調査に対して方法論上の明確な優位性を有する。'
    '手術患者は全員がすでに入院しているため、外来研究における主要な交絡因子である'
    '医療へのアクセスが事実上中立化される。'
    '地理的遠隔性、社会経済的地位、受療行動の違いにかかわらず、'
    'すべての患者が直接の臨床観察下にあり、鎮痛薬の投与を受ける機会は均等である。'
    'したがって、周術期処方の地域差は、医療アクセスの差異よりも、'
    '臨床慣行と疼痛表現の差異をより直接的に反映するものと考えられる。')

# Para 5: Objectives (exploratory)
doc.add_paragraph(
    '本探索的研究の目的は3つである。'
    '第一（Phase 1）、47都道府県・9地域ブロックにおける'
    '急性期周術期鎮痛薬処方の地域差をマッピングすること。'
    '第二（Phase 2）、交絡疾患（糖尿病性神経障害、帯状疱疹後神経痛、うつ病、不安障害）を'
    '調整した上で、外来神経障害性疼痛薬処方のCPSPプロキシとしての地域差を検討すること。'
    '第三、Phase 1とPhase 2の統合により、集団レベルでの急性期周術期鎮痛薬使用と'
    'その後の慢性疼痛関連処方の関連を探索すること。')

# ============================================================
# METHODS
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
doc.add_paragraph(
    '本研究はNDBオープンデータの都道府県別集計データを分析した生態学的研究である。'
    'STROBE声明（横断研究）[15] およびRECORD拡張 [16] に準拠して報告する。'
    '公開された集計データのみを使用し個人レベルの情報を含まないため、倫理審査は不要である。')

add_heading_text('データソース', level=2)
doc.add_paragraph(
    'NDBオープンデータ第10回（厚生労働省公表）を使用し、'
    '2023年4月〜2024年3月の診療報酬請求を対象とした [12]。'
    'NDBは日本の国民皆保険制度内の全保険者からの請求を収録し、'
    '約1億2,500万人の被保険者を網羅する。'
    '集計データは都道府県レベルで公表され、10件未満のセルは秘匿処理されている。')

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県を以下の9つの標準地域ブロックに分類した：'
    '北海道（1）、東北（6：青森、岩手、宮城、秋田、山形、福島）、'
    '関東（7）、北陸・甲信越（6）、東海（4）、近畿（6）、'
    '中国（5）、四国（4）、九州・沖縄（8）。'
    'この分類は厚生労働省が使用する標準的な行政区分に従う。')

add_heading_text('Phase 1：急性期周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから以下の3つの鎮痛薬分類を抽出した：'
    '薬効分類114（解熱鎮痛消炎薬＝NSAIDs・アセトアミノフェン）、'
    '薬効分類811（あへんアルカロイド系麻薬＝モルヒネ、オキシコドン、コデイン）、'
    '薬効分類821（合成麻薬＝フェンタニル、ペチジン、タペンタドール）。'
    '入院手術件数はKコード（手術）から抽出した。'
    '各都道府県の鎮痛薬/手術指標＝入院鎮痛薬処方量（単位）÷入院手術件数として算出した。')

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（CPSPプロキシ）', level=2)
doc.add_paragraph(
    '5種類の外来経口神経障害性疼痛薬を都道府県別に抽出した：'
    'プレガバリン（78製剤）、ミロガバリン（8製剤）、デュロキセチン（33製剤）、'
    'トラマドール（3製剤）、ノイロトロピン（1製剤）。'
    '神経障害性疼痛薬/手術指標＝外来神経障害性疼痛薬処方量÷入院手術件数として算出した。'
    'この指標はCPSPのプロキシとして使用するが、これらの薬剤が'
    '他の神経障害性疼痛疾患にも処方されることを認識している。')

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '神経障害性疼痛薬はCPSP以外の疾患にも処方されるため、'
    '以下の4つの交絡疾患プロキシを外来処方データから抽出した：')
doc.add_paragraph('経口血糖降下薬（261製剤）：糖尿病性神経障害のプロキシ', style='List Bullet')
doc.add_paragraph('帯状疱疹抗ウイルス薬（47製剤：バラシクロビル、アシクロビル、ファムシクロビル、アメナメビル）：'
                  '帯状疱疹後神経痛のプロキシ', style='List Bullet')
doc.add_paragraph('デュロキセチンを除く抗うつ薬（128製剤：SSRI、SNRI、NaSSA、三環系）：'
                  'うつ病のプロキシ', style='List Bullet')
doc.add_paragraph('抗不安薬（112製剤：ベンゾジアゼピン系、タンドスピロン、ヒドロキシジン）：'
                  '不安障害のプロキシ', style='List Bullet')
doc.add_paragraph('各交絡プロキシは手術あたりで表し、神経障害性疼痛薬指標との比較可能性を維持した。')

add_heading_text('外来神経ブロック手技', level=2)
doc.add_paragraph(
    'NDBオープンデータの麻酔区分から外来神経ブロック算定回数を抽出し、'
    '独立したCPSPプロキシとした。'
    '硬膜外ブロック、傍脊椎神経ブロック、トリガーポイント注射、'
    '星状神経節ブロック、持続神経ブロック注入等73手技コードを含む。')

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    'Phase 1の地域差はKruskal-Wallis検定で9地域ブロック間を比較し、'
    '事後検定としてBonferroni補正付きMann-Whitney U検定を用いた。'
    'Phase 2では以下の5つの回帰モデルを適合した：')
doc.add_paragraph('モデル1：未調整の地域間神経障害性疼痛薬処方比較', style='List Bullet')
doc.add_paragraph('モデル2：神経障害性疼痛薬 ~ 糖尿病薬 + 帯状疱疹薬 + 抗うつ薬 + 抗不安薬 + 地域ブロック', style='List Bullet')
doc.add_paragraph('モデル3：コア神経障害性薬（プレガバリン＋ミロガバリンのみ）~ 同交絡因子', style='List Bullet')
doc.add_paragraph('モデル4：神経ブロック ~ 同交絡因子', style='List Bullet')
doc.add_paragraph('モデル5：神経障害性疼痛薬 ~ 急性期鎮痛薬指標 + 交絡因子（統合モデル）', style='List Bullet')

doc.add_paragraph(
    '調整済CPSP指標は、神経障害性疼痛薬処方を4つの交絡プロキシに回帰した残差として導出し、'
    '交絡疾患を除いた「説明されない」神経障害性疼痛薬処方を表す。'
    'PearsonおよびSpearman相関で指標間の関連を評価した。'
    '全解析はPython 3.11（NumPy 1.24, SciPy 1.11, matplotlib 3.8）で実施した。')

doc.add_page_break()

# ============================================================
# RESULTS
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1：急性期周術期鎮痛薬処方の地域差', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月に、NDBには7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方が記録された。'
    '全国平均の鎮痛薬/手術指標は35.78（SD 5.56）で、'
    '岐阜（25.20）から鹿児島（49.75）まで1.97倍の差があった'
    '（9地域間Kruskal-Wallis P < 0.001）。')

doc.add_paragraph(
    '顕著な地域集積が観察された。東海・近畿（西日本）が最低で、'
    '九州・沖縄・北海道が最高であった（表1）。'
    '注目すべき知見として、日本で最も我慢強い地域と文化的に認識される東北は、'
    '9地域中7位であった（平均39.97, SD 3.53 vs 東北以外35.17; '
    'Mann-Whitney U = 190, P = 0.031; Cohen\'s d = 0.87）。'
    '東北6県全てが全国上位半分に位置した。'
    'このパターンは全薬効分類で一貫していた：NSAIDs（P = 0.044）、'
    'あへんアルカロイド（P = 0.003）、合成麻薬（P = 0.001）。')

# Table 1
add_bold_paragraph('表1. ', 'Phase 1：地域ブロック別 入院鎮痛薬/手術指標')
table1 = doc.add_table(rows=11, cols=5)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['地域ブロック', 'n', '平均（SD）', '範囲', '順位']):
    table1.rows[0].cells[i].text = h
data1 = [
    ['東海', '4', '27.62 (2.18)', '25.20\u201330.07', '1'],
    ['近畿', '6', '30.02 (2.03)', '27.92\u201332.33', '2'],
    ['関東', '7', '33.00 (2.09)', '29.82\u201334.78', '3'],
    ['北陸・甲信越', '6', '35.38 (3.54)', '31.18\u201340.08', '4'],
    ['中国', '5', '35.73 (4.18)', '31.01\u201340.17', '5'],
    ['四国', '4', '36.33 (3.27)', '33.49\u201341.02', '6'],
    ['東北', '6', '39.97 (3.53)', '35.18\u201344.51', '7'],
    ['九州・沖縄', '8', '42.26 (4.33)', '35.82\u201349.75', '8'],
    ['北海道', '1', '46.12 (\u2014)', '46.12', '9'],
    ['全国', '47', '35.78 (5.56)', '25.20\u201349.75', '\u2014'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        table1.rows[r+1].cells[c].text = val
# Bold Tohoku row
for cell in table1.rows[7].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（未調整）', level=2)
doc.add_paragraph(
    '全国の外来神経障害性疼痛薬処方量は合計2,289,549,163単位で、'
    'プレガバリン（40.2%）、ノイロトロピン（20.1%）、ミロガバリン（19.6%）、'
    'デュロキセチン（15.3%）、トラマドール（4.9%）であった。'
    f'東北は未調整で指標が突出して高かった'
    f'（平均{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\u2019s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）。'
    '岩手（566.7）、青森（519.3）、秋田（461.1）が全国上位3位を占めた（Fig. 1）。')

add_figure(OUTPUT_DIR + 'fig1_neuropathic_unadjusted.png',
    'Fig. 1. 都道府県別 外来神経障害性疼痛薬処方量/手術（未調整）。'
    'プレガバリン＋ミロガバリン＋デュロキセチン＋トラマドール＋ノイロトロピンの合計を'
    '入院手術件数で除した。東北（赤枠の赤棒）が高値側に集中。破線＝全国平均。')

doc.add_page_break()

add_heading_text('交絡因子分析', level=2)
doc.add_paragraph(
    '神経障害性疼痛薬処方は交絡疾患プロキシと強い相関を示した。'
    '糖尿病薬処方が最も強い相関（r = 0.87, P < 0.001）を示し、'
    '次いで抗不安薬（r = 0.75, P < 0.001）、抗うつ薬（r = 0.46, P = 0.001）、'
    '帯状疱疹薬（r = 0.19, P = 0.19）であった。'
    'これらの交絡因子は神経障害性疼痛薬処方の分散の80.4%を説明した'
    f'（モデル2のR\u00b2 = {reg["model2_adjusted"]["R2"]:.3f}；Fig. 2）。')

add_figure(OUTPUT_DIR + 'fig2_confounder_correlations.png',
    'Fig. 2. 外来神経障害性疼痛薬処方と交絡疾患プロキシの相関。'
    '各点は1都道府県。東北は赤枠で表示。糖尿病薬が最も強い相関（r = 0.87）を示す。')

doc.add_page_break()

add_heading_text('交絡因子調整後の分析', level=2)
doc.add_paragraph(
    f'4つの交絡プロキシで調整した後、東北の効果は減弱し有意でなくなった'
    f'（モデル2：\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model2_adjusted"]["tohoku_p"]:.3f}）。'
    'この減弱は全モデルで一貫していた：'
    f'モデル3（コア神経障害性薬：\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}）、'
    f'モデル4（神経ブロック：P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}）、'
    f'モデル5（統合：\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}）（表2）。')

doc.add_paragraph(
    '調整済CPSP指標（交絡因子除去後の残差）は未調整データとは劇的に異なるパターンを示した（Fig. 3）。'
    f'東北の平均は、未調整では顕著な正値から、'
    f'調整後は緩やかな非有意な過剰に変化した'
    f'（{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}）。'
    '調整後は中国地方が最高の調整済CPSP指標を示し、東海が最低であった（Fig. 4）。')

add_figure(OUTPUT_DIR + 'fig3_adjusted_cpsp_index.png',
    'Fig. 3. 交絡疾患調整後の都道府県別CPSP指標。'
    '神経障害性疼痛薬処方を糖尿病薬・帯状疱疹薬・抗うつ薬・抗不安薬に回帰した残差。'
    '正の値は交絡疾患から予測されるより多い神経障害性疼痛薬処方を示す。'
    '東北（赤枠）は調整後、分布全体に分散している。')

doc.add_page_break()

# Table 2
add_bold_paragraph('表2. ', '神経障害性疼痛薬処方の回帰モデル：交絡調整の効果')
table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['モデル', '従属変数', '東北効果', 'P値', '判定']):
    table2.rows[0].cells[i].text = h
t2_data = [
    ['1（未調整）', '神経障害性疼痛薬/手術', f'd = {reg["model1_unadjusted"]["cohens_d"]:.2f}', f'{reg["model1_unadjusted"]["p_value"]:.4f}', '有意 ***'],
    ['2（全交絡調整）', '神経障害性疼痛薬/手術', f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}', f'{reg["model2_adjusted"]["tohoku_p"]:.3f}', '有意でない'],
    ['3（コア薬）', 'PGB+MGB/手術', f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}', f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}', '有意でない'],
    ['4（神経ブロック）', 'ブロック/手術', f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}', f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}', '有意でない'],
    ['5（統合）', '神経障害性疼痛薬/手術', f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}', f'{reg["model5_integrated"]["tohoku_p"]:.3f}', '有意でない'],
    ['調整済CPSP指標', '残差', f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}', f'{reg["adjusted_cpsp_test"]["p_value"]:.3f}', '有意でない'],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

add_figure(OUTPUT_DIR + 'fig4_region_unadj_vs_adj.png',
    'Fig. 4. 地域ブロック別 神経障害性疼痛薬処方の比較：(a) 未調整 (b) 交絡疾患調整後。'
    '東北（赤枠）は未調整では最高だが、調整後は中位に移動。エラーバー＝標準偏差。')

doc.add_page_break()

add_heading_text('Phase 1\u2013Phase 2 統合', level=2)
doc.add_paragraph(
    f'Phase 1の急性期鎮痛薬/手術はPhase 2の未調整神経障害性疼痛薬/手術と'
    f'正の相関を示した（r = 0.38, P = 0.008；Fig. 5a）。'
    f'交絡調整後、相関は減弱し境界的有意性となった'
    f'（r = 0.29, P = 0.052；Fig. 5b）。'
    f'統合モデル（モデル5）では、Phase 1の急性期疼痛指標は'
    f'交絡調整後も神経障害性疼痛薬処方の有意な予測因子であった'
    f'（\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P = {reg["model5_integrated"]["acute_pain_p"]:.3f}）が、'
    f'東北効果は有意でなかった'
    f'（\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}）。')

add_figure(OUTPUT_DIR + 'fig5_phase1_vs_phase2.png',
    'Fig. 5. Phase 1（急性期鎮痛薬/手術）vs Phase 2（外来神経障害性疼痛薬/手術＝CPSPプロキシ）の統合。'
    '(a) 未調整：正の相関（r = 0.38, P = 0.008）。'
    '(b) 交絡調整後：減弱した相関（r = 0.29, P = 0.052）。'
    '東北（赤枠）は右上象限に集中。')

doc.add_page_break()

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    '本研究は、自由に利用可能なNDBオープンデータを用いて、'
    '日本の47都道府県全てにおける周術期および慢性疼痛関連処方をマッピングした初の研究である。'
    '本探索的分析から3つの主要な知見が得られた。')

add_heading_text('忍耐強い文化の中にも相当な地域差が存在する', level=2)
doc.add_paragraph(
    '日本の痛みに対する文化的忍耐は十分に文献化されているにもかかわらず [5,6]、'
    '急性期周術期鎮痛薬処方には都道府県間で1.97倍の差があり、'
    '地域ブロック間にも有意差が認められた。'
    'これはCohenとNisbettによる米国の「名誉の文化」研究と類似する：'
    '単一国家内でも地域文化規範が測定可能に異なる行動学的帰結を生む [8]。'
    '本研究の知見は、日本の疼痛文化が一枚岩ではないことを示唆する——'
    '地域の人口構成、医療インフラ、地域の臨床慣行が、'
    '国民共有の文化規範の表面下に異質性を生み出している。')

add_heading_text('臨床的含意：一枚岩的な文化的ステレオタイプの危険性', level=2)
doc.add_paragraph(
    '本研究で示された1.97倍の日本国内変動は、国内の文脈にとどまらない臨床的意義を有する。'
    '豊富なエビデンスが、民族的・文化的ステレオタイプが'
    '臨床医の疼痛評価と鎮痛薬処方に影響を及ぼすことを示している。'
    'Andersonらは、米国において人種的・民族的マイノリティが'
    '急性・慢性・がん・緩和ケアの全領域で一貫して不十分な疼痛管理を受けることを示した [17]。'
    'CampbellとEdwardsは、患者の文化的疼痛行動に対する臨床医の期待が'
    '系統的な過少治療または過剰治療につながりうることを指摘した [18]。'
    'Roggerらは、文化的フレーミングが患者の報告だけでなく、'
    '臨床医の疼痛の解釈と対応にも影響することを強調した [2]。')

doc.add_paragraph(
    'このリスクは、日本が文化的・民族的に均質な社会であるという広く流布した認識によってさらに増幅される。'
    '日本人論（「日本人の独自性に関する理論」）の言説は、'
    '日本人が均一な行動規範を共有する単一の集団を構成するという概念を長く広めてきた [19]。'
    'しかし、Befuが名付けたこの「均質性の覇権」は、'
    '経験的事実ではなくイデオロギー的構築物である [19]。'
    'Burgessは、この「均質な日本」という「幻想」が'
    '社会政策と公衆認識に実質的影響を及ぼすことを論証した [20]。'
    '臨床的文脈では、「日本人は忍耐強い」と「日本人は均質である」という'
    '2つのステレオタイプの組み合わせが、二重に誤った仮定を生む：'
    '全ての日本人患者が等しく痛みに耐え、より少ない鎮痛で済むという仮定である。'
    '本研究が示した鎮痛薬処方の1.97倍の国内変動は、'
    'この仮定を直接的に反駁する。')

doc.add_paragraph(
    '端的に言えば、「日本人患者」なる存在はいない——'
    '国籍だけから疼痛行動を予測できるような存在は。'
    'いるのは47の多様な都道府県からの個々の患者だけであり、'
    'それぞれが異なる人口構成、臨床環境、疼痛関連処方文化を持つ。')

doc.add_paragraph(
    '海外で治療を受ける日本人患者に対して、臨床医は二重の誤解のもとに行動している可能性がある：'
    '日本文化は一様に忍耐強く、その忍耐は鎮痛薬の必要性の低さを意味する、と。'
    '本研究のデータは両方の前提に疑義を呈する。'
    '疼痛関連処方は都道府県間でほぼ2倍異なり、'
    'ステレオタイプ上最も忍耐強いとされる地域が実際にはより少ない鎮痛薬を使用しているわけではない。'
    '国内でこの規模の変動が存在するならば、国レベルのステレオタイプを'
    '個々の臨床判断に適用することは、エビデンスに基づかず安全でもない。')

doc.add_paragraph(
    'したがって我々は、いかなる国民集団の疼痛行動についても一枚岩的な特徴づけを行うことが'
    '臨床的害のリスクを伴うと主張する——そしてこの主張は日本をはるかに超えて妥当する。'
    '臨床医が文化的ステレオタイプによって患者に対する治療選択肢の範囲を狭めるとき、'
    '患者は治療上の不利益を被る。'
    '「忍耐強い日本人」というラベルは鎮痛薬の過少処方につながりうる。'
    '逆に、他の集団に対するステレオタイプは過剰処方、'
    '特定症状の見落とし、鑑別診断における早期閉鎖につながりうる。'
    '本研究のデータ——単一の文化的ラベルの下にほぼ2倍の国内変動を示す——は、'
    '普遍的原則の具体的事例研究として機能する：'
    '文化的一般化は、いかに便利であっても、個別の臨床判断の基盤としては不十分である。')

doc.add_paragraph(
    '本研究の背後にある願いは、国籍や民族にかかわらず、'
    '文化的ステレオタイプによって治療選択肢を狭められるという不利益を被る患者が'
    '一人でも少なくなることである。'
    '海外で治療を受ける日本人患者に対して、臨床医は疼痛忍耐の規範が不均質であり、'
    '忍耐的な態度が鎮痛薬の必要性の低さを確実に示すものではないことを認識すべきである。'
    'しかし同じ論理はあらゆる文化のあらゆる患者に当てはまる：'
    '文化的均質性の仮定は——いかに便利であっても——'
    '疼痛医学においても他のいかなる領域においても科学的に擁護しえない。'
    '文化的ステレオタイプに基づく仮定ではなく、個別化された疼痛評価こそが、'
    '公平な周術期ケアの礎石であり続ける。')

add_heading_text('痛みは個人的体験であり、客観的侵害受容モニタリングの開発が待たれる', level=2)
doc.add_paragraph(
    '根本的に、痛みは個人的な体験である。'
    '国際疼痛学会（IASP）は痛みを「実際のまたは潜在的な組織損傷に関連する、'
    'またはそれに似た不快な感覚的・情動的体験」と定義しており——'
    'この定義は本質的に主観的である [21]。'
    '本研究の生態学的データはこの原理を集団レベルで例証している：'
    '同一の言語、保険制度、広い文化的遺産を共有する単一国家内でさえ、'
    '疼痛関連処方はほぼ2倍異なる。'
    'この異質性は、疾患負荷や臨床慣行の差異だけでなく、'
    '疼痛知覚と表現の還元不可能な個別性をも反映していると考えられる。')

doc.add_paragraph(
    'この個別性がもたらす臨床的帰結は、'
    '「忍耐強い日本人」であれ他の何であれ、いかなる文化的ラベルも'
    '患者の侵害受容状態の直接測定の代替にはなりえないということである。'
    'しかし、そのような測定を行うツールは未だ発展途上にある。'
    '侵害受容と抗侵害受容のバランス（NANB）は生理学的な量であり、'
    '原理的には術中・術後に客観的にモニタリングすることが可能であり、'
    '患者の自己報告や臨床医の推論を完全にバイパスしうる。'
    'Onishiらは、標準的なパルスオキシメトリから導出される'
    '正規化脈波容積（NPV）が、オピオイド効果部位濃度（変動係数62.4%）よりも'
    '有意に低い個体間変動（変動係数36.3%）で自発呼吸の回復を予測することを示し、'
    'NPVが文化に依存しない客観的NANB指標となりうることを示唆した [22]。')

doc.add_paragraph(
    'このような客観的侵害受容モニタリングシステムの開発と臨床実装は、'
    '周術期疼痛管理の風景を根本的に変えうる。'
    'BIS（bispectral index）が催眠深度を定量化するように、'
    'NANBをベッドサイドで定量化できるようになれば、'
    '鎮痛薬の調整はもはや文化的バイアスに影響されやすい主観的疼痛スコアにも、'
    'どの患者が痛みに「耐える」かについての集団レベルのステレオタイプにも依存しなくなる。'
    'そのようなシステムが臨床的に利用可能になるまで、'
    '本研究の知見は、文化的一般化が個人の侵害受容の現実にとって'
    '不十分な代替指標であることを改めて示すものである。')

add_heading_text('交絡疾患がCPSPプロキシの見かけの地域差を説明する', level=2)
doc.add_paragraph(
    '方法論的に最も重要な知見は、神経障害性疼痛薬処方の劇的な地域差'
    f'（東北 vs その他の未調整d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）が、交絡疾患プロキシによって大部分説明されたことである。'
    '糖尿病薬処方だけでr = 0.87の相関を示し、'
    'これは糖尿病性神経障害とガバペンチノイド使用の既知の関連を反映する。'
    '調整後、東北効果は62%減弱し有意でなくなった。')

doc.add_paragraph(
    'これは生態学的疼痛研究に重要な含意を持つ。'
    '神経障害性疼痛薬処方を集団レベルのCPSPプロキシとして使用する研究は、'
    '交絡疾患を考慮しなければならない。'
    'そのような調整なしでは、糖尿病有病率の地域差がCPSPの差と誤認される恐れがある。'
    '本研究で示したデータベース内交絡調整——同一データソースの疾患特異的薬剤プロキシを用いる——は、'
    '再現可能な枠組みを提供する。')

add_heading_text('集団レベルでの急性期\u2013慢性期疼痛の連続体', level=2)
doc.add_paragraph(
    'Phase 1（急性期）とPhase 2（慢性期、調整済）指標の正の相関'
    '（r = 0.29, P = 0.052）は、地域の急性期疼痛管理の強度と'
    'その後の慢性疼痛関連処方の間に緩やかな関連があることを示唆する。'
    '生態学的相関は因果関係を確立できないが、'
    'この知見は急性期術後疼痛の強度がCPSPの危険因子であるという'
    '個人レベルの文献 [23] と整合する。'
    '急性期鎮痛薬使用が比較的低いにもかかわらず調整済CPSP指標が高い都道府県は、'
    '急性期疼痛の過小治療が慢性化を招いている可能性について調査に値する。')

add_heading_text('強みと限界', level=2)
doc.add_paragraph(
    '本研究の強みとして、日本の全保険請求医療を網羅する集団完全データの使用、'
    '急性期と慢性期の疼痛プロキシの新規統合、透明な交絡調整方法論、'
    '仮説生成を可能にする探索的デザインが挙げられる。'
    'さらに、周術期に焦点を当てた本研究に特有の強みとして、'
    '医療アクセスが鎮痛薬処方データの交絡因子とならない点がある：'
    'Phase 1の全患者は定義上入院患者であり、'
    '地域在住者研究を制約する医療アクセスの異質性が排除される。')

doc.add_paragraph('主な限界は生態学的デザインに固有のものである：')
doc.add_paragraph('分析単位は個人ではなく都道府県であり、'
                  '生態学的相関は個人レベルの関連を反映しない可能性がある（生態学的誤謬）。',
                  style='List Bullet')
doc.add_paragraph('NDBオープンデータには傷病名コードがないため、CPSPを直接同定できない。'
                  '神経障害性疼痛薬プロキシはCPSP特異的でなく全適応を含む。', style='List Bullet')
doc.add_paragraph('薬剤処方プロキシは疾患有病率を正確に反映しない可能性がある'
                  '（例：線維筋痛症には特異的な薬剤プロキシがなく、プレガバリンを第一選択として共有する）。',
                  style='List Bullet')
doc.add_paragraph('横断的デザインでは時間的順序（手術→急性期疼痛→CPSP）を区別できない。', style='List Bullet')
doc.add_paragraph('測定されない交絡因子（年齢分布、手術構成、処方文化）が残余の地域差に寄与しうる。', style='List Bullet')

add_heading_text('今後の展望', level=2)
doc.add_paragraph(
    '本探索的研究は、日本のNDBオープンデータが集団レベルの疼痛研究における'
    '仮説生成プラットフォームとして機能しうることを実証する。'
    'NDB特別抽出データ（Level 3アクセス）を用いた今後の研究では、'
    '個人レベルで手術から新規神経障害性疼痛薬処方への縦断的追跡が可能となり、'
    '直接的なCPSP指標を提供しうる。'
    '手術種別（例：人工膝関節全置換術、乳房切除術）に限定した分析により'
    '手術構成交絡を軽減できる。'
    '処方データと患者報告疼痛アウトカムの連結により、'
    '地域の処方差が疼痛体験・疼痛表現・臨床慣行のいずれの差異を反映するかが明らかになろう。')

doc.add_paragraph(
    '国際的な観点からは、本研究の知見は文化的にきめ細やかだが個別化された'
    '疼痛管理の必要性を強調する。'
    '文化的ステレオタイプが多文化的臨床環境における日本人患者への鎮痛薬処方に'
    'いかに影響するかを検討する前向き研究は、'
    '「忍耐強い日本人」というラベルが治療上の不利益に直結するかどうかを直接検証しうる。'
    'そのような研究は、グローバルな人口移動と医療ツーリズムが'
    '文化的背景に不慣れな医療システムでの治療を受ける患者を増加させている中で、'
    'ますます喫緊のものとなっている。')

# ============================================================
# CONCLUSION
# ============================================================
add_heading_text('結論', level=1)
doc.add_paragraph(
    '日本に根ざした疼痛忍耐の文化的規範（我慢）にもかかわらず、'
    '周術期および慢性疼痛関連処方は都道府県間で最大1.97倍異なる。'
    '交絡疾患（特に糖尿病）は、神経障害性疼痛薬処方の見かけの地域パターンを大きく修飾する。'
    'これらの知見は、日本の疼痛文化が一枚岩でないことを実証する。'
    '「日本人」を疼痛行動に関する均質なカテゴリーとして扱うことは、'
    '海外で治療を受ける日本人患者への不十分な鎮痛を招くリスクがある'
    '——そしてこの原則は、あらゆる患者集団に適用されるあらゆる文化的ラベルに等しく当てはまる。'
    '文化的ステレオタイプに基づく仮定ではなく、個別化された疼痛評価が、'
    'あらゆる臨床環境における公平な周術期ケアを担保すべきである。')

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph(
    'NDBオープンデータを公開している厚生労働省に感謝する。'
    '［必要に応じて追加の謝辞を挿入］')

# ============================================================
# CONFLICT OF INTEREST STATEMENT
# ============================================================
doc.add_paragraph()
add_heading_text('利益相反', level=1)
doc.add_paragraph(
    '著者らに開示すべき利益相反はない。')

# DATA AVAILABILITY STATEMENT
doc.add_paragraph()
add_heading_text('データ利用可能性', level=1)
doc.add_paragraph(
    '本研究で使用したNDBオープンデータは厚生労働省ウェブサイト'
    '（https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html）より公開されている。'
    '解析コードはhttps://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japanで利用可能である。')

doc.add_page_break()

# ============================================================
# FIGURES (remaining)
# ============================================================
add_figure(OUTPUT_DIR + 'fig6_model_comparison_table.png',
    'Fig. 6. 回帰モデルの要約：交絡調整前後の東北効果。'
    'モデル1（未調整）は高度に有意な過剰を示すが、全調整モデルでは有意でなくなり、'
    '交絡因子が地域差を説明することを示す。')

doc.add_page_break()
add_figure(OUTPUT_DIR + 'sfig1_heatmap.png',
    'Supplementary Fig. 1. 都道府県別 各指標のZ-scoreヒートマップ。'
    '各行は変数、各列は都道府県（神経障害性疼痛薬/手術で昇順ソート）。'
    '赤＝平均以上、青＝平均以下。東北は赤の垂直線で表示。'
    '調整済CPSP指標は未調整の神経障害性指標とは異なるパターンを示す。')

# ============================================================
# REFERENCES - 23 references matching English v2
# ============================================================
doc.add_page_break()
add_heading_text('文献', level=1)
refs = [
    # [1] Callister 2003 - Cultural influences on pain
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract. 2003;15:207\u2013211.',

    # [2] Rogger 2023 - Cultural framing and acute pain
    'Rogger R, Bello C, Romero CS, Urman RD, Luedi MM, Filipovic MG. '
    'Cultural framing and the impact on acute pain and pain services. '
    'Curr Pain Headache Rep. 2023;27:429\u2013436.',

    # [3] Zborowski 1969 - People in Pain
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass; 1969.',

    # [4] Okolo 2024 - Cultural variability review
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
    'a review of cross-cultural studies. Int J Sci Res Arch. 2024;11:2550\u20132556.',

    # [5] Hobara 2005 - Japanese vs Euro-American pain behavior beliefs
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain. 2005;9:389\u2013393.',

    # [6] Feng 2017 - Japan vs Europe EQ-5D
    'Feng Y, Herdman M, van Nooten F, Cleeland C, Parkin D, Ikeda S, Igarashi A, Devlin NJ. '
    'An exploration of differences between Japan and two European countries in the self-reporting '
    'and valuation of pain and discomfort on the EQ-5D. Qual Life Res. 2017;26:2067\u20132078.',

    # [7] Hayashi 2022 - PCS variation by country
    'Hayashi K, Ikemoto T, Shiro Y, Arai YC, Marcuzzi A, Costa D, Wrigley PJ. '
    'A systematic review of the variation in Pain Catastrophizing Scale reference scores '
    'based on language version and country in patients with chronic primary (non-specific) pain. '
    'Pain Ther. 2022;11:751\u2013780.',

    # [8] Cohen & Nisbett 1996 - Culture of honor
    'Cohen D, Nisbett RE, Bowdle BF, Schwarz N. '
    'Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d '
    'J Pers Soc Psychol. 1996;70:945\u2013960.',

    # [9] Kumagai 2020 - Media reproducing Tohoku stereotypes
    '\u7066\u8c37 \u58ee. \u30e1\u30c7\u30a3\u30a2\u306b\u3088\u308b\u6771\u5317\u30a4\u30e1\u30fc\u30b8\u306e\u518d\u751f\u7523'
    '\u2014\u2014\u300c\u79d8\u5bc6\u306e\u30b1\u30f3\u30df\u30f3SHOW\u300d\u306e\u6771\u5317\u5fa9\u8208\u30b3\u30fc\u30ca\u30fc\u3092\u4e2d\u5fc3\u306b. '
    '\u3053\u3068\u3070. 2020;41:21\u201338.',

    # [10] Takeda & Yarimizu 2016 - Regional differences in pain expression "uzuku"
    '\u7af9\u7530 \u664b\u4e5f, \u9461\u6c34 \u517c\u8c9e. \u75db\u307f\u8868\u73fe\u300c\u30a6\u30ba\u30af\u300d\u306e\u5730\u57df\u5dee. '
    '\u56fd\u7acb\u56fd\u8a9e\u7814\u7a76\u6240\u8ad6\u96c6. 2016;10:85\u2013107.',

    # [11] Pfizer 2017 - Prefecture pain tolerance survey
    '\u30d5\u30a1\u30a4\u30b6\u30fc\u682a\u5f0f\u4f1a\u793e. \u516847\u90fd\u9053\u5e9c\u770c \u75db\u307f\u306b\u95a2\u3059\u308b\u610f\u8b58\u8abf\u67fb 2017. '
    'https://www.pfizer.co.jp/pfizer/company/press/2017\uff082025\u5e742\u67081\u65e5\u30a2\u30af\u30bb\u30b9\uff09.',

    # [12] NDB Open Data
    '\u539a\u751f\u52b4\u50cd\u7701. NDB\u30aa\u30fc\u30d7\u30f3\u30c7\u30fc\u30bf \u7b2c10\u56de. 2024. '
    'https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html\uff082025\u5e741\u670815\u65e5\u30a2\u30af\u30bb\u30b9\uff09.',

    # [13] Wakaizumi 2024 - Geographic variation in HICP
    'Wakaizumi K, Tanaka C, Shinohara Y, Wu Y, Takaoka S, Kawate M, Oka H, Matsudaira K. '
    'Geographical variation in high-impact chronic pain and psychological associations at the regional level: '
    'a multilevel analysis of a large-scale internet-based cross-sectional survey. '
    'Front Public Health. 2024;12:1482177.',

    # [14] Matsuoka 2025 - Regional opioid prescribing for cancer pain
    'Matsuoka Y, et al. Population-based claims study of regional and hospital function differences '
    'in opioid prescribing for cancer patients who died in hospital in Japan. '
    'Jpn J Clin Oncol. 2025;hyaf149.',

    # [15] STROBE
    'von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP. '
    'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
    'guidelines for reporting observational studies. '
    'Lancet. 2007;370:1453\u20131457.',

    # [16] RECORD
    'Benchimol EI, Smeeth L, Guttmann A, Harber K, Moher D, Petersen I, Sorensen HT, '
    'von Elm E, Langan SM. '
    'The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med. 2015;12:e1001885.',

    # [17] Anderson 2009 - Racial/ethnic disparities in pain
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
    'causes and consequences of unequal care. J Pain. 2009;10:1187\u20131204.',

    # [18] Campbell & Edwards 2012 - Ethnic differences in pain management
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag. 2012;2:219\u2013230.',

    # [19] Befu 2001 - Hegemony of Homogeneity
    'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
    'Melbourne: Trans Pacific Press; 2001.',

    # [20] Burgess 2010 - Illusion of homogeneous Japan
    'Burgess C. The \u201cillusion\u201d of homogeneous Japan and national character: '
    'discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. '
    'Asia Pac J. 2010;8(9):1\u201322.',

    # [21] IASP definition of pain
    'Raja SN, Carr DB, Cohen M, Finnerup NB, Flor H, Gibson S, Keefe FJ, Mogil JS, '
    'Ringkamp M, Sluka KA, Song XJ, Stevens B, Sullivan MD, Tutelman PR, Ushida T, Vader K. '
    'The revised International Association for the Study of Pain definition of pain: '
    'concepts, challenges, and compromises. Pain. 2020;161:1976\u20131982.',

    # [22] Onishi et al. 2024 - NPV as objective NANB monitor
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
    'and quantification of nociception anti-nociception balance compared to opioid effect site concentration: '
    'a prospective, observational study. F1000Research. 2024;13:233.',

    # [23] Kehlet 2006 - Persistent postsurgical pain
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet. 2006;367:1618\u20131625.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

# Save
outpath = OUTPUT_DIR + 'Pain_manuscript_JA.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
