#!/usr/bin/env python3
"""Create cover letter for Journal of Epidemiology submission.

JE requires:
  a) Journal name
  b) Paper title
  c) Principal findings and significance
  d) Statement that all authors approved and work not submitted elsewhere
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JE_DIR = os.path.join(OUTPUT_DIR, 'je')
os.makedirs(JE_DIR, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.15

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run('[Date]')

doc.add_paragraph()

# Addressee
doc.add_paragraph('Editor-in-Chief')
doc.add_paragraph('Journal of Epidemiology')
doc.add_paragraph('Japan Epidemiological Association')

doc.add_paragraph()

# Subject
p = doc.add_paragraph()
r = p.add_run('Re: ')
r.bold = True
p.add_run(
    'Submission of original article \u2014 '
    '"Regional Variation in Pain-Related Prescribing Across '
    'Japan\u2019s 47 Prefectures: An Ecological Study Using '
    'the National Database of Health Insurance Claims"'
)

doc.add_paragraph()

# Salutation
doc.add_paragraph('Dear Editor-in-Chief,')

doc.add_paragraph()

# Body
body_paras = [
    'I am pleased to submit the above-titled manuscript for consideration as an '
    'Original Article in the Journal of Epidemiology.',

    'This ecological study uses the population-complete National Database of Health '
    'Insurance Claims (NDB Open Data, 10th edition; April 2023\u2013March 2024) to map '
    'pain-related prescribing variation across all 47 prefectures of Japan. To our '
    'knowledge, this is the first study to comprehensively quantify prefecture-level '
    'variation in perioperative analgesic and neuropathic pain drug prescribing using '
    'these freely available national data.',

    'Our principal findings are as follows: '
    '(1) Acute perioperative analgesic prescribing varied 1.97-fold across prefectures, '
    'persisting after age-sex standardisation\u2014indicating supply-sensitive practice variation '
    'rather than demographic differences alone. '
    '(2) The apparent regional clustering of neuropathic pain drug prescribing was largely '
    'explained by confounding disease proxies, particularly diabetes drug prescribing '
    '(r=0.87); after adjustment, inter-regional differences were substantially '
    'attenuated and became nonsignificant. '
    '(3) Most notably, prescribing variation was dissociated from symptom burden: '
    'symptom prevalence from the Comprehensive Survey of Living Conditions (CSLC 2022) '
    'varied only 1.29-fold and showed no correlation with prescribing (r=0.03, P=0.85), '
    'formally meeting the criteria for Wennberg unwarranted variation. '
    '(4) Acute perioperative prescribing was a significant predictor of chronic pain-related '
    'prescribing even after confounder adjustment.',

    'These findings are of direct relevance to the readership of the Journal of '
    'Epidemiology. The study demonstrates that Japan\u2019s NDB Open Data\u2014a freely '
    'accessible resource\u2014can reveal clinically important practice variation that would '
    'be invisible in individual-level or single-institution studies. The triangulation of '
    'NDB prescribing data with the CSLC household survey provides a novel demand\u2013supply '
    'framework for evaluating unwarranted variation. The methodological '
    'lesson\u2014that ecological studies of neuropathic pain drugs must account for diabetic '
    'neuropathy and other confounding diseases\u2014is broadly relevant to researchers using '
    'claims data for pain epidemiology.',

    'The manuscript follows the STROBE and RECORD reporting guidelines. '
    'As only publicly available aggregate data were used, ethical approval was not required.',

    'The author approved the final manuscript. This work has not been published '
    'previously and is not under consideration for publication elsewhere.',

    'Parts of data processing and manuscript preparation were assisted by generative AI '
    '(Claude, Anthropic). The author takes full responsibility for the accuracy and '
    'content of the manuscript.',

    'Thank you for considering this manuscript for publication in the Journal of Epidemiology.',
]

for text in body_paras:
    doc.add_paragraph(text)

doc.add_paragraph()

# Closing
doc.add_paragraph('Sincerely,')

doc.add_paragraph()

doc.add_paragraph('Tatsuki Onishi, MD')
doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('[Address], [City], [Postal code], Japan')
doc.add_paragraph('E-mail: [email]')

outpath = os.path.join(JE_DIR, 'JE_cover_letter.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
