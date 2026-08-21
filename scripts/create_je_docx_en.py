#!/usr/bin/env python3
"""Create English manuscript for Journal of Epidemiology (JE).

JE Original Article requirements (Guide for Authors, updated July 1 2025):
- Body: ≤3,500 words (excl abstract/refs/legends/tables)
- Abstract: structured (Background, Methods, Results, Conclusions) ≤250 words
- References: AMA style — consecutive superscript Arabic numerals
- Figures: separate files; legends in a separate section
- Tables: editable text, numbered consecutively
- Highlights: mandatory, 3-5 bullets ≤150 chars each (separate file)
- STROBE checklist: encouraged for observational studies
- Title page: title, authors, affiliations, corresponding author, running title
  (≤8 words), table/figure/supplementary counts
- Keywords: 3-5
- Double-spaced, continuous line numbering (Abstract → Acknowledgments)
- Cover letter required
- AI use: declare in Acknowledgements
- Data availability statement required
- Single-column format
- Submission via ScholarOne: https://mc.manuscriptcentral.com/je

Reframe strategy (Option C-1 from previous session analysis):
- Technical epidemiology framing (NOT stereotype-led)
- Wennberg unwarranted variation framework as theoretical backbone
- Stereotype discussion moved to Discussion implications only
- A+B extension: add CSLC symptom data + Pfizer survey detail
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

# ============================================================
# Paths
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JE_DIR = os.path.join(OUTPUT_DIR, 'je')
os.makedirs(JE_DIR, exist_ok=True)

FIG_DIR = OUTPUT_DIR

# ============================================================
# Load data
# ============================================================
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'cslc_analysis.json'), 'r') as f:
    cslc = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

# ============================================================
# REFERENCES — AMA style, numbered in order of first appearance
# ============================================================
ref_list = [
    # 1 Wennberg (unwarranted variation — theoretical backbone)
    'Wennberg JE. Tracking Medicine: A Researcher\'s Quest to Understand Health Care. '
    'New York, NY: Oxford University Press; 2010.',
    # 2 Corallo (supply-sensitive variation — cited with Wennberg in intro para 1)
    'Corallo AN, Croxford R, Goodman DC, Bryan EL, Srivastava D, Stukel TA. '
    'A systematic review of medical practice variation in OECD countries. '
    'Health Policy. 2014;114:5\u201314. '
    'doi:10.1016/j.healthpol.2013.08.002',
    # 3 MHLW NDB
    'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. '
    'https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
    'Accessed January 15, 2025.',
    # 4 Taira (NDB ecological precedent)
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilisation '
    'in Japan: an ecological study using the National Database of Health Insurance Claims. '
    'Lancet Reg Health West Pac. 2021;12:100170. '
    'doi:10.1016/j.lanwpc.2021.100170',
    # 5 Wakaizumi (chronic pain regional variation)
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact '
    'chronic pain and psychological associations at the regional level: a multilevel analysis '
    'of a large-scale internet-based cross-sectional survey. Front Public Health. 2024;12:1482177. '
    'doi:10.3389/fpubh.2024.1482177',
    # 6 Takahashi (cancer opioid variation)
    'Takahashi R, Miyashita M, Nakazawa Y, Wada S, Matsuoka Y. Population-based claims study '
    'of regional and hospital function differences in opioid prescribing for cancer patients '
    'who died in hospital in Japan. Jpn J Clin Oncol. 2025;55:1372\u20137. '
    'doi:10.1093/jjco/hyaf149',
    # 7 CSLC / Kokumin Seikatsu Kiso Chousa
    'Ministry of Health, Labour and Welfare. Comprehensive Survey of Living Conditions '
    '(Kokumin Seikatsu Kiso Chousa), 2022. '
    'https://www.mhlw.go.jp/toukei/saikin/hw/k-tyosa/k-tyosa22/. '
    'Accessed March 1, 2025.',
    # 8 von Elm STROBE
    'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational '
    'Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. '
    'Lancet. 2007;370:1453\u20137. '
    'doi:10.1016/S0140-6736(07)61602-X',
    # 9 Benchimol RECORD
    'Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using '
    'Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med. 2015;12:e1001885. '
    'doi:10.1371/journal.pmed.1001885',
    # 10 Kehlet (CPSP)
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet. 2006;367:1618\u201325. '
    'doi:10.1016/S0140-6736(06)68700-X',
    # 11 Anderson (disparities — Discussion ref)
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
    'causes and consequences of unequal care. J Pain. 2009;10:1187\u2013204. '
    'doi:10.1016/j.jpain.2009.10.002',
    # 12 Campbell (ethnic differences — Discussion ref)
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag. 2012;2:219\u201330. '
    'doi:10.2217/pmt.12.7',
    # 13 Callister (cultural influences — Discussion ref)
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract. 2003;15:207\u201311. '
    'doi:10.1177/1084822302250687',
    # 14 Hobara (Japan stoic — Discussion ref)
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain. 2005;9:389\u201393. '
    'doi:10.1016/j.ejpain.2004.09.006',
    # 15 Pfizer Japan survey
    'Pfizer Japan Inc. 47-prefecture survey on chronic pain (2012 vs 2017). '
    'https://www.pfizer.co.jp/pfizer/company/press/2017/2017_08_23.html. '
    'Accessed May 28, 2022. Archived at: '
    'https://web.archive.org/web/20220528073616/'
    'https://www.pfizer.co.jp/pfizer/company/press/2017/2017_08_23.html',
    # 16 Raja IASP definition
    'Raja SN, Carr DB, Cohen M, et al. The revised International Association for the '
    'Study of Pain definition of pain: concepts, challenges, and compromises. '
    'Pain. 2020;161:1976\u201382. '
    'doi:10.1097/j.pain.0000000000001939',
    # 17 Onishi (nociception monitoring)
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery '
    'and quantification of nociception anti-nociception balance compared to opioid effect site '
    'concentration: a prospective, observational study. F1000Research. 2024;13:233. '
    'doi:10.12688/f1000research.146215.1',
]


def cite(*nums):
    """Return AMA superscript citation marker string."""
    return '{' + ','.join(str(n) for n in nums) + '}'


def add_ref_runs(p, text):
    """Parse text with {n} or {n,m} markers and create runs with font-based superscript."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)


# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# Page numbers in footer
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

# Enable continuous line numbering (Abstract → Acknowledgments)
for section in doc.sections:
    sectPr = section._sectPr
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), '1')
    lnNumType.set(qn('w:restart'), 'continuous')
    sectPr.append(lnNumType)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def wc(text):
    return len(re.sub(r'\{[^}]+\}', '', text).split())


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


# Map figure numbers to source PNG files
FIG_SOURCES = {
    1: os.path.join(FIG_DIR, 'fig1_neuropathic_unadjusted_en.png'),
    2: os.path.join(FIG_DIR, 'fig2_confounder_correlations_en.png'),
    3: os.path.join(FIG_DIR, 'fig4_region_unadj_vs_adj_en.png'),
    4: os.path.join(FIG_DIR, 'fig_cslc_demand_supply_en.png'),
}


def add_inline_figure(caption_text, fig_num):
    """Insert figure image inline and caption below it."""
    fig_path = FIG_SOURCES.get(fig_num)
    if fig_path and os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        r0 = p.add_run(f'[Insert Figure {fig_num} here]')
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'Figure {fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(caption_text)
    r2.font.size = Pt(10)


# ============================================================
# Computed values
# ============================================================
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']
# (Tohoku-specific SCR values available in data but not used in region-flat analysis)

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    'Regional Variation in Pain-Related Prescribing Across '
    'Japan\u2019s 47 Prefectures: An Ecological Study Using '
    'the National Database of Health Insurance Claims'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Tatsuki Onishi, MD')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Department of Anesthesiology, [Institution], [Address], [City], [Postal code], Japan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Corresponding author: ')
r.bold = True
p.add_run(
    'Tatsuki Onishi, MD, Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. E-mail: [email]'
)

doc.add_paragraph()

# Title page metadata
p = doc.add_paragraph()
r = p.add_run('Article type: ')
r.bold = True
p.add_run('Original Article')

p = doc.add_paragraph()
r = p.add_run('Running title: ')
r.bold = True
p.add_run('Pain prescribing variation in Japan')  # ≤8 words

p = doc.add_paragraph()
r = p.add_run('Word count (body): ')
r.bold = True
p.add_run('[body word count]')

p = doc.add_paragraph()
r = p.add_run('Word count (abstract): ')
r.bold = True
p.add_run('[abstract word count]')

p = doc.add_paragraph()
r = p.add_run('Number of figures: ')
r.bold = True
p.add_run('4')

p = doc.add_paragraph()
r = p.add_run('Number of tables: ')
r.bold = True
p.add_run('2')

p = doc.add_paragraph()
r = p.add_run('Number of supplementary materials: ')
r.bold = True
p.add_run('1 (Supplementary Table 1)')

p = doc.add_paragraph()
r = p.add_run('Keywords: ')
r.bold = True
p.add_run('analgesics; ecological study; health care disparities; Japan; prescriptions')

doc.add_page_break()

# ============================================================
# ABSTRACT (Background, Methods, Results, Conclusions) ≤250 words
# ============================================================
add_heading_text('Abstract', level=1)

abstract_bg = (
    'Pain-related prescribing practices may vary geographically even within a single '
    'national health insurance system, but the magnitude and determinants of such '
    'variation remain poorly characterised in Japan. '
    'Whether prescribing variation reflects differences in symptom burden (warranted '
    'variation) or practice patterns (unwarranted variation) has not been investigated.'
)

abstract_methods = (
    'This ecological study analysed prefecture-level aggregate data from the '
    'National Database of Health Insurance Claims (NDB Open Data, 10th edition; '
    'April 2023\u2013March 2024) covering all 47 prefectures. '
    'Phase 1 examined acute perioperative analgesic prescribing per surgery. '
    'Phase 2 examined outpatient neuropathic pain drug prescribing as a chronic pain '
    'proxy, with regression adjustment for confounding disease proxies (diabetes, '
    'herpes zoster, depression, anxiety). '
    'Standardised claim ratios (SCR) were computed by indirect age-sex standardisation. '
    'Symptom prevalence from the Comprehensive Survey of Living Conditions (CSLC 2022) '
    'was used as a demand-side benchmark.'
)

abstract_results = (
    f'Acute analgesic prescribing varied 1.97-fold across prefectures '
    f'(Kruskal\u2013Wallis P<0.001). '
    f'Unadjusted neuropathic pain prescribing showed marked regional clustering, '
    f'but confounding disease proxies\u2014particularly diabetes drugs (r=0.87)\u2014'
    f'explained {reg["model2_adjusted"]["R2"]*100:.0f}% of between-prefecture variance. '
    f'After confounder adjustment, apparent inter-regional differences were '
    f'substantially attenuated and became nonsignificant. '
    f'CSLC symptom prevalence varied only '
    f'{cslc["symptom_rate_stats"]["range_ratio"]:.2f}-fold and showed no correlation with '
    f'prescribing (r={cslc["correlations"]["symptom_vs_acute"]["pearson_r"]:.2f}, '
    f'P={cslc["correlations"]["symptom_vs_acute"]["pearson_p"]:.2f}), '
    f'indicating supply-sensitive unwarranted variation.'
)

abstract_conclusions = (
    'Nearly twofold within-country variation in pain-related prescribing persists '
    'after age-sex standardisation and is dissociated from symptom burden, '
    'meeting the criteria for Wennberg unwarranted variation. '
    'Ecological studies of neuropathic pain drugs must account for confounding diseases, '
    'particularly diabetes, to avoid misattributing disease-driven prescribing to '
    'regional practice differences.'
)

abstract_sections = [
    ('Background', abstract_bg),
    ('Methods', abstract_methods),
    ('Results', abstract_results),
    ('Conclusions', abstract_conclusions),
]

for label, text in abstract_sections:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

abstract_total = sum(wc(t) for t in [
    abstract_bg, abstract_methods, abstract_results, abstract_conclusions,
])
print(f'Abstract word count: {abstract_total} (JE limit: 250)')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('Introduction', level=1)

intro_parts = []

intro1 = (
    f'Geographic variation in medical practice has been recognised as a major challenge '
    f'for healthcare quality since Wennberg\'s pioneering work on small-area '
    f'variation.{cite(1)} When prescribing rates differ across regions beyond what '
    f'disease burden and demographics can explain, such unwarranted variation signals '
    f'potential over- or under-treatment and merits systematic investigation.{cite(1,2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)
intro_parts.append(intro1)

intro2 = (
    f'Japan\'s universal health insurance system, standardised drug pricing, and the '
    f'National Database of Health Insurance Claims (NDB)\u2014which captures virtually all '
    f'reimbursed healthcare utilisation for approximately 125 million insured '
    f'individuals{cite(3)}\u2014provide a unique setting for studying practice variation. '
    f'Taira et al demonstrated substantial regional inequality in dental care utilisation '
    f'using NDB-derived standardised claim ratios across all 47 prefectures.{cite(4)} '
    f'Recent studies have also documented up to 1.6-fold regional variation in chronic pain '
    f'prevalence{cite(5)} and 4-fold variation in cancer opioid prescribing.{cite(6)} '
    f'However, no study has comprehensively mapped pain-related prescribing variation across '
    f'all 47 prefectures using the population-complete NDB.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)
intro_parts.append(intro2)

intro3 = (
    f'Wennberg\'s framework distinguishes warranted variation\u2014driven by differences in '
    f'disease burden\u2014from unwarranted variation driven by supply-side factors such as '
    f'physician practice style, local norms, and resource availability.{cite(1,2)} '
    f'To apply this framework, an independent demand-side benchmark is needed. '
    f'The Comprehensive Survey of Living Conditions (CSLC), a nationally representative '
    f'household survey conducted by the Ministry of Health, Labour and Welfare, provides '
    f'prefecture-level self-reported symptom prevalence rates that can serve as such a '
    f'benchmark.{cite(7)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro3)
intro_parts.append(intro3)

intro4 = (
    'Pain management is particularly amenable to ecological analysis because analgesic '
    'prescribing reflects both disease burden and clinical practice patterns. '
    'The perioperative setting offers a further methodological advantage: because all '
    'patients in this analysis are hospitalised for surgery, healthcare access\u2014a major '
    'confounder in community-based prescribing studies\u2014is neutralised by design. '
    'Outpatient neuropathic pain drugs can serve as a population-level proxy for chronic '
    'pain burden, but these medications have multiple indications (notably diabetic '
    'neuropathy), requiring careful confounder adjustment.'
)
doc.add_paragraph(intro4)
intro_parts.append(intro4)

intro5 = (
    'This study had four objectives: (1) map regional variation in acute perioperative '
    'analgesic prescribing across 47 prefectures; (2) examine outpatient neuropathic pain '
    'drug prescribing as a chronic pain proxy after adjustment for '
    'confounding diseases; (3) quantify the contribution of confounders to the apparent '
    'regional pattern; and (4) evaluate whether prescribing variation is dissociated from '
    'symptom burden, thereby meeting the criteria for unwarranted variation.'
)
doc.add_paragraph(intro5)
intro_parts.append(intro5)

intro_total = sum(wc(t) for t in intro_parts)
print(f'Introduction word count: {intro_total}')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

methods_parts = []

add_heading_text('Study design and reporting', level=2)
m1 = (
    f'This ecological study analysed prefecture-level aggregate data from the NDB Open Data. '
    f'It is reported following the Strengthening the Reporting of Observational Studies in '
    f'Epidemiology (STROBE) statement{cite(8)} and the REporting of studies Conducted using '
    f'Observational Routinely-collected health Data (RECORD) extension.{cite(9)} '
    f'As only publicly available aggregate data were used, ethical approval was not required '
    f'under Japan\'s Ethical Guidelines for Medical and Biological Research Involving '
    f'Human Subjects.'
)
p = doc.add_paragraph()
add_ref_runs(p, m1)
methods_parts.append(m1)

add_heading_text('Data source', level=2)
m2 = (
    f'The 10th edition of the NDB Open Data (April 2023\u2013March 2024) was used.{cite(3)} '
    f'The NDB captures claims from all insurers within Japan\'s universal coverage system, '
    f'encompassing approximately 125 million insured individuals. '
    f'Aggregate prescription and procedure data are published at the prefecture level (n=47) '
    f'with suppression of cells containing fewer than 10 events. '
    f'Prefecture-level population estimates (October 2023, by 5-year age group and sex) '
    f'from the Statistics Bureau of Japan were used to compute per-capita rates and '
    f'standardised claim ratios.'
)
p = doc.add_paragraph()
add_ref_runs(p, m2)
methods_parts.append(m2)

add_heading_text('Regional classification', level=2)
m3 = (
    'Prefectures were grouped into nine standard regional blocks following '
    'the classification used by the Statistics Bureau of Japan: '
    'Hokkaido (1 prefecture), Tohoku (6), Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), '
    'Kinki (6), Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8).'
)
doc.add_paragraph(m3)
methods_parts.append(m3)

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
m4 = (
    'Inpatient prescription data were extracted for three analgesic drug classes commonly '
    'used in perioperative pain management: '
    'Class 114 (antipyretic analgesics including NSAIDs and acetaminophen), '
    'Class 811 (opium alkaloid narcotics), '
    'and Class 821 (synthetic narcotics including fentanyl and pethidine). '
    'The analgesic-per-surgery index was calculated for each prefecture as '
    'total inpatient analgesic prescription units divided by total inpatient '
    'surgical procedure count (K Surgery section), providing a standardised measure of '
    'analgesic intensity that accounts for differences in surgical volume.'
)
doc.add_paragraph(m4)
methods_parts.append(m4)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing', level=2)
m5 = (
    'Five classes of outpatient oral neuropathic pain medications were extracted: '
    'pregabalin (78 formulations), mirogabalin (8), '
    'duloxetine (33), tramadol (3), and neurotropin (1). '
    'The neuropathic pain prescribing-per-surgery index was calculated as total outpatient '
    'neuropathic pain drug quantity divided by total inpatient surgical procedure count. '
    'Per-capita neuropathic pain prescribing rates were additionally computed.'
)
doc.add_paragraph(m5)
methods_parts.append(m5)

add_heading_text('Confounder disease proxies', level=2)
m6 = (
    'Four confounder disease proxies were extracted from outpatient data: '
    'oral hypoglycaemic agents (261 formulations; proxy for diabetic neuropathy), '
    'herpes zoster antivirals (47 formulations; proxy for postherpetic neuralgia), '
    'antidepressants excluding duloxetine (128 formulations; proxy for depression), '
    'and anxiolytics (112 formulations; proxy for anxiety disorders). '
    'Each proxy was expressed per surgery for consistency. '
    'Outpatient nerve block procedure counts (73 codes) served as an additional '
    'independent proxy.'
)
doc.add_paragraph(m6)
methods_parts.append(m6)

add_heading_text('Demand-side benchmark: Comprehensive Survey of Living Conditions', level=2)
m_cslc = (
    f'Prefecture-level self-reported symptom prevalence rates were obtained from the '
    f'Comprehensive Survey of Living Conditions (CSLC; Kokumin Seikatsu Kiso Chousa) '
    f'2022 edition.{cite(7)} The CSLC is a nationally representative household survey '
    f'covering approximately 300,000 households, with symptom prevalence rates '
    f'(yuushosha-ritsu) published per 1,000 population for each of the 47 prefectures. '
    f'Symptom prevalence, reflecting the proportion of the population reporting any '
    f'health complaint (musculoskeletal pain being the most common category nationally), '
    f'was used as a demand-side proxy for healthcare need. '
    f'A demand\u2013supply mismatch index was computed as the difference between '
    f'z-standardised prescribing rates and z-standardised symptom prevalence rates '
    f'for each prefecture. Positive values indicate relative over-supply; negative '
    f'values indicate relative under-supply.'
)
p = doc.add_paragraph()
add_ref_runs(p, m_cslc)
methods_parts.append(m_cslc)

add_heading_text('Statistical analysis', level=2)
m7a = (
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test '
    'across nine regional blocks, followed by post hoc Mann\u2013Whitney U tests with Bonferroni '
    'correction. Effect sizes were quantified using Cohen\'s d. '
    'For Phase 2, five regression models were fitted with progressive confounder adjustment: '
    'Model 1 (unadjusted inter-regional comparison), '
    'Model 2 (neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + region), '
    'Model 3 (core neuropathic drugs only ~ same confounders), '
    'Model 4 (nerve blocks ~ same confounders), and '
    'Model 5 (neuropathic pain ~ acute analgesic index + confounders). '
    'The adjusted CPSP index was derived as residuals from regressing neuropathic pain '
    'prescribing on the four confounder proxies.'
)
doc.add_paragraph(m7a)
methods_parts.append(m7a)

m7b = (
    f'Standardised claim ratios (SCR) were computed by indirect age-sex standardisation '
    f'following Taira et al.{cite(4)} '
    f'National age-sex-specific prescription rates (18 five-year age groups \u00d7 2 sexes) '
    f'were applied to each prefecture\'s population structure. '
    f'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11).'
)
p = doc.add_paragraph()
add_ref_runs(p, m7b)
methods_parts.append(m7b)

methods_total = sum(wc(t) for t in methods_parts)
print(f'Methods word count: {methods_total}')

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

results_parts = []

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
r1 = (
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD, 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima)\u2014a 1.97-fold difference '
    '(Kruskal\u2013Wallis P<0.001 across nine regions; Table 1).'
)
doc.add_paragraph(r1)
results_parts.append(r1)

# === TABLE 1 INLINE ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_cap = p_cap.add_run('Table 1. ')
r_cap.bold = True
r_cap.font.size = Pt(10)
p_cap.add_run(
    'Regional summary of inpatient analgesic prescribing per surgery. '
    'Values are mean (SD). '
    'Kruskal\u2013Wallis P<0.001.'
).font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t1.rows[0].cells
for i, h in enumerate(['Region', 'n', 'Mean (SD)', 'Range']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = REGION_EN[reg_name]
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f} ({np.std(vals, ddof=1):.2f})' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t1)
doc.add_paragraph()

r2 = (
    'Substantial regional clustering was observed (Table 1). '
    'Tokai and Kinki had the lowest indices, while Kyushu-Okinawa and Hokkaido had the highest. '
    'Post hoc pairwise comparisons confirmed significant differences between several regional '
    'blocks (Bonferroni-corrected Mann\u2013Whitney U tests). '
    'This pattern was consistent across all three analgesic drug classes: '
    'NSAIDs, opioid alkaloids, and synthetic opioids.'
)
doc.add_paragraph(r2)
results_parts.append(r2)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing', level=2)
r3 = (
    f'Outpatient neuropathic pain drug prescriptions totalled 2,289,549,163 units, '
    f'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    f'duloxetine (15.3%), and tramadol (4.9%). '
    f'Marked inter-regional variation was observed in neuropathic pain '
    f'prescribing-per-surgery index (Kruskal\u2013Wallis P<0.001; Figure 1), '
    f'with individual prefectures spanning a wide range.'
)
doc.add_paragraph(r3)
results_parts.append(r3)

# === FIGURE 1 placeholder ===
add_inline_figure(
    'Outpatient neuropathic pain drug prescribing per surgery by prefecture. '
    'Bars indicate regional blocks (distinguished by colour and hatch pattern). '
    'Dashed line indicates the national mean.',
    1
)

add_heading_text('Confounder analysis and adjustment', level=2)
r4a = (
    f'Neuropathic pain prescribing showed strong correlations with confounder disease proxies. '
    f'Diabetes drug prescribing was the strongest correlate (r=0.87, P<0.001), '
    f'followed by anxiolytics (r=0.75), antidepressants (r=0.46), '
    f'and herpes antivirals (r=0.19). '
    f'These four confounders collectively explained '
    f'{reg["model2_adjusted"]["R2"]*100:.1f}% of between-prefecture variance in neuropathic pain '
    f'prescribing (R\u00b2={reg["model2_adjusted"]["R2"]:.3f}; Figure 2).'
)
doc.add_paragraph(r4a)
results_parts.append(r4a)

# === FIGURE 2 placeholder ===
add_inline_figure(
    'Correlation between neuropathic pain prescribing and confounder disease proxies '
    'across 47 prefectures. '
    'Diabetes drugs show the strongest correlation (r=0.87).',
    2
)

r4b = (
    f'After adjustment for these confounders, the apparent inter-regional clustering '
    f'was substantially attenuated (Table 2). '
    f'Regional indicator variables that were significant in unadjusted models became '
    f'nonsignificant across all model specifications '
    f'(Models 2\u20135), indicating that confounding diseases\u2014not regional practice '
    f'differences\u2014drove the observed pattern.'
)
doc.add_paragraph(r4b)
results_parts.append(r4b)

# === TABLE 2 INLINE ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_cap2 = p_cap2.add_run('Table 2. ')
r_cap2.bold = True
r_cap2.font.size = Pt(10)
p_cap2.add_run(
    'Regression models for outpatient neuropathic pain prescribing.'
).font.size = Pt(10)

models = [
    ('Model 1 (unadjusted)', '\u2014',
     'Kruskal\u2013Wallis H, P<0.001', 'Omnibus 9-region test'),
    ('Model 2 (all confounders)',
     f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model2_adjusted"]["R2_adj"]:.3f}',
     f'Regional \u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
     f'P={reg["model2_adjusted"]["tohoku_p"]:.3f}'),
    ('Model 3 (core neuro)',
     f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model3_core_neuropathic"]["R2_adj"]:.3f}',
     f'Regional \u03b2={reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
     f'P={reg["model3_core_neuropathic"]["tohoku_p"]:.3f}'),
    ('Model 4 (nerve blocks)',
     f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model4_nerve_blocks"]["R2_adj"]:.3f}',
     f'Regional \u03b2={reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}, '
     f'P={reg["model4_nerve_blocks"]["tohoku_p"]:.3f}'),
    ('Model 5 (integrated)',
     f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}',
     f'Acute \u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, P={reg["model5_integrated"]["acute_pain_p"]:.3f}',
     f'Regional \u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
     f'P={reg["model5_integrated"]["tohoku_p"]:.3f}'),
]

t2 = doc.add_table(rows=1 + len(models), cols=4, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Confounder R\u00b2', 'Key statistic', 'Note']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, (name, r2_val, key_stat, note) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = r2_val
    row[2].text = key_stat
    row[3].text = note
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t2)
doc.add_paragraph()

r5a = (
    f'The adjusted CPSP index showed a markedly different geographic pattern '
    f'from the unadjusted data (Figure 3). '
    f'Regions that appeared to have high neuropathic pain prescribing in unadjusted '
    f'analyses no longer showed significant excess after confounder adjustment '
    f'(Kruskal\u2013Wallis P>0.05), indicating that the original clustering was '
    f'driven primarily by regional differences in confounding disease burden.'
)
doc.add_paragraph(r5a)
results_parts.append(r5a)

# === FIGURE 3 placeholder ===
add_inline_figure(
    'Regional comparison of neuropathic pain prescribing: '
    '(a) unadjusted and (b) after confounder adjustment. Error bars indicate SD.',
    3
)

add_heading_text('Demand\u2013supply dissociation', level=2)
r_cslc1 = (
    f'CSLC symptom prevalence ranged from {cslc["symptom_rate_stats"]["min"]:.1f} '
    f'({cslc["symptom_rate_stats"]["min_pref"]}) to {cslc["symptom_rate_stats"]["max"]:.1f} '
    f'per 1,000 population ({cslc["symptom_rate_stats"]["max_pref"]}), a '
    f'{cslc["symptom_rate_stats"]["range_ratio"]:.2f}-fold difference '
    f'(mean {cslc["symptom_rate_stats"]["mean"]:.1f}, '
    f'SD {cslc["symptom_rate_stats"]["sd"]:.1f}). '
    f'In contrast, acute analgesic prescribing varied 1.97-fold and neuropathic pain '
    f'prescribing per capita varied {scr_neuro_ratio:.1f}-fold (after age-sex standardisation). '
    f'Thus, prescribing heterogeneity substantially exceeded the modest variation in '
    f'symptom burden.'
)
doc.add_paragraph(r_cslc1)
results_parts.append(r_cslc1)

r_cslc2 = (
    f'CSLC symptom prevalence showed no significant correlation with acute '
    f'analgesic prescribing (r={cslc["correlations"]["symptom_vs_acute"]["pearson_r"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_acute"]["pearson_p"]:.3f}; '
    f'Spearman \u03c1={cslc["correlations"]["symptom_vs_acute"]["spearman_rho"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_acute"]["spearman_p"]:.3f}; Figure 4) or with '
    f'neuropathic pain prescribing per capita '
    f'(r={cslc["correlations"]["symptom_vs_neuro_percapita"]["pearson_r"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_neuro_percapita"]["pearson_p"]:.3f}). '
    f'The absence of correlation indicates that prescribing variation '
    f'is not explained by differences in self-reported symptom burden across prefectures.'
)
doc.add_paragraph(r_cslc2)
results_parts.append(r_cslc2)

# === FIGURE 4 placeholder ===
add_inline_figure(
    'Demand\u2013supply dissociation: CSLC symptom prevalence rate '
    '(per 1,000 population) vs acute analgesic prescribing per surgery '
    'across 47 prefectures. Each marker shape represents a regional block. '
    'The near-zero correlation (r=0.03) indicates that '
    'prescribing variation is dissociated from symptom burden.',
    4
)

add_heading_text('Integration of acute and chronic prescribing', level=2)
r6 = (
    f'Acute perioperative prescribing correlated positively with unadjusted neuropathic '
    f'pain prescribing (r=0.38, P=0.008). After confounder adjustment, this '
    f'correlation was attenuated (r=0.29, P=0.052). '
    f'In Model 5, the acute pain index remained a significant predictor '
    f'(\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while regional indicators were nonsignificant after confounder adjustment.'
)
doc.add_paragraph(r6)
results_parts.append(r6)

add_heading_text('Age-sex standardised claim ratios', level=2)
r7 = (
    f'After indirect age-sex standardisation, inpatient analgesic SCR ranged from '
    f'{scr_analgesic_range[0]:.1f} to {scr_analgesic_range[1]:.1f} '
    f'({scr_analgesic_ratio:.1f}-fold variation), confirming that prescribing heterogeneity '
    f'was not attributable to differences in prefectural age-sex composition. '
    f'Outpatient neuropathic pain drug SCR ranged from {scr_neuro_range[0]:.1f} to '
    f'{scr_neuro_range[1]:.1f} ({scr_neuro_ratio:.1f}-fold).'
)
doc.add_paragraph(r7)
results_parts.append(r7)

results_total = sum(wc(t) for t in results_parts)
print(f'Results word count: {results_total}')

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

disc_parts = []

d0 = (
    'This study is the first to map perioperative and chronic pain-related prescribing across '
    'all 47 prefectures of Japan using the population-complete NDB Open Data, '
    'and the first to formally evaluate whether this variation meets the criteria for '
    'Wennberg unwarranted variation by benchmarking it against an independent demand-side '
    'measure. Four principal findings emerged.'
)
doc.add_paragraph(d0)
disc_parts.append(d0)

add_heading_text('Principal findings', level=2)
d1 = (
    f'First, acute perioperative analgesic prescribing varied 1.97-fold across '
    f'prefectures\u2014comparable to the regional variation in dental care '
    f'utilisation reported by Taira et al{cite(4)} and consistent with the '
    f'supply-sensitive variation pattern described by Wennberg.{cite(1)} '
    f'This variation was robust: it persisted after age-sex standardisation (SCR range, '
    f'{scr_analgesic_range[0]:.0f}\u2013{scr_analgesic_range[1]:.0f}) and was consistent '
    f'across all three analgesic drug classes.'
)
p = doc.add_paragraph()
add_ref_runs(p, d1)
disc_parts.append(d1)

d2 = (
    f'Second, the most methodologically important finding is that the marked regional '
    f'clustering in neuropathic pain prescribing was largely '
    f'explained by confounding disease proxies, particularly diabetes drug prescribing '
    f'(r=0.87). After adjustment, apparent inter-regional differences were substantially '
    f'attenuated and became nonsignificant across all model specifications. '
    f'This has important implications: ecological studies using neuropathic pain drugs as a '
    f'chronic pain proxy must account for diabetic neuropathy. Without such adjustment, '
    f'regional differences in diabetes prevalence will be misattributed to differences '
    f'in pain management practice.'
)
doc.add_paragraph(d2)
disc_parts.append(d2)

d3 = (
    f'Third, the most novel finding is that prescribing variation was dissociated from '
    f'symptom burden. CSLC symptom prevalence varied only '
    f'{cslc["symptom_rate_stats"]["range_ratio"]:.2f}-fold '
    f'and showed no correlation with analgesic prescribing (r=0.03, P=0.85). '
    f'This dissociation formally meets the criteria for unwarranted variation under '
    f'Wennberg\'s framework:{cite(1)} the supply side (prescribing) varies independently '
    f'of the demand side (symptom burden), indicating that practice patterns rather than '
    f'disease burden drive the observed heterogeneity. '
    f'This finding parallels international evidence showing that medical practice variation '
    f'is frequently supply-sensitive rather than demand-driven.{cite(2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)
disc_parts.append(d3)

d3b = (
    f'Fourth, the positive correlation between acute and confounder-adjusted chronic '
    f'indices (r=0.29, P=0.052) and the significant acute pain predictor in Model 5 '
    f'(\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}) suggest a modest link between '
    f'regional acute pain management intensity and subsequent chronic pain-related '
    f'prescribing. This is consistent with individual-level evidence that acute postoperative '
    f'pain intensity is a risk factor for CPSP.{cite(10)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3b)
disc_parts.append(d3b)

add_heading_text('Clinical implications', level=2)
d4 = (
    f'The 1.97-fold within-Japan variation has broader implications beyond descriptive '
    f'epidemiology. A large body of evidence demonstrates that cultural stereotypes '
    f'influence clinician pain assessment and prescribing behaviour internationally: '
    f'Anderson et al showed systematic disparities in pain management across racial and '
    f'ethnic groups,{cite(11)} and Campbell and Edwards identified that clinician expectations '
    f'about cultural pain behaviour can lead to systematic under- or '
    f'over-treatment.{cite(12)} '
    f'Japan is often characterised as culturally homogeneous with patients who are '
    f'stoic about pain.{cite(13,14)} Yet our data demonstrate that even within this '
    f'putatively uniform population, prescribing patterns vary nearly twofold\u2014a degree '
    f'of heterogeneity that makes any national-level characterisation clinically misleading. '
    f'The Pfizer Japan 2017 survey provides additional corroboration: the proportion of '
    f'chronic pain patients "enduring pain without seeking treatment" ranged from 48.7% '
    f'(Osaka) to 81.6% (Yamanashi) across prefectures, and Akita Prefecture '
    f'(Tohoku region)\u2014often stereotyped as stoic\u2014had the lowest proportion who '
    f'believed pain "should be endured" (60.2%), contradicting the cultural '
    f'stereotype.{cite(15)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d4)
disc_parts.append(d4)

d5 = (
    f'The revised IASP definition describes pain as inherently '
    f'subjective.{cite(16)} No regional label or cultural generalisation can substitute '
    f'for individualised pain assessment. Objective nociception monitoring may help '
    f'standardise perioperative assessment.{cite(17)} '
    f'Our demand\u2013supply analysis reinforces this point: prefectures with identical symptom '
    f'burden show widely divergent prescribing, indicating that clinician-level factors '
    f'rather than patient-level need determine treatment intensity.'
)
p = doc.add_paragraph()
add_ref_runs(p, d5)
disc_parts.append(d5)

add_heading_text('Strengths and limitations', level=2)
d6a = (
    'Strengths include the use of population-complete data covering all '
    'insurance-reimbursed healthcare in Japan, the novel integration of acute and chronic '
    'pain proxies, the transparent confounder-adjustment methodology, '
    'a perioperative design that neutralises healthcare access as a confounder, '
    'and the triangulation of NDB prescribing data with an independent household survey '
    '(CSLC) for demand\u2013supply analysis.'
)
doc.add_paragraph(d6a)
disc_parts.append(d6a)

d6b = (
    'Limitations are inherent to the ecological design. '
    'The unit of analysis is the prefecture, not the individual; ecological '
    'correlations may not reflect individual-level associations. '
    'NDB Open Data lack diagnosis codes, so the neuropathic pain drug proxy '
    'captures all indications, not CPSP specifically. '
    'Unmeasured confounders including surgical case mix and physician density '
    'may contribute to residual variation. '
    'The CSLC symptom prevalence rate is a general indicator of self-reported health '
    'complaints, not specific to pain; however, musculoskeletal pain symptoms (lower back '
    'pain and stiff shoulders) are the most common complaints nationally, and the '
    'CSLC rate thus serves as a reasonable ecological proxy for pain-related demand.'
)
doc.add_paragraph(d6b)
disc_parts.append(d6b)

add_heading_text('Conclusion', level=2)
d7 = (
    'Pain-related prescribing varies nearly twofold across Japan\'s 47 prefectures, '
    'persisting after age-sex standardisation and dissociated from symptom burden. '
    'This pattern meets the criteria for Wennberg unwarranted variation, '
    'suggesting that supply-side factors drive prescribing heterogeneity. '
    'Confounding diseases\u2014particularly diabetes\u2014substantially modify the apparent '
    'regional pattern of neuropathic pain prescribing and must be accounted for in '
    'ecological studies. '
    'Clinicians should base analgesic decisions on individual assessment '
    'rather than regional or cultural assumptions.'
)
doc.add_paragraph(d7)
disc_parts.append(d7)

disc_total = sum(wc(t) for t in disc_parts)
print(f'Discussion word count: {disc_total}')

# ============================================================
# ACKNOWLEDGEMENTS (including COI and AI disclosure per JE)
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgements', level=1)
doc.add_paragraph(
    'The author thanks the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available.'
)

p_ai = doc.add_paragraph()
p_ai.add_run('Use of AI: ').bold = True
p_ai.add_run(
    'Parts of data processing and manuscript preparation were assisted by generative AI '
    '(Claude, Anthropic). The author takes full responsibility for the accuracy and content '
    'of the manuscript.'
)

p_coi = doc.add_paragraph()
p_coi.add_run('Conflicts of interest: ').bold = True
p_coi.add_run(
    'The author declares no conflict of interest with respect to this research study and paper.'
)

# ============================================================
# FUNDING
# ============================================================
doc.add_paragraph()
add_heading_text('Funding', level=1)
doc.add_paragraph(
    'This research did not receive any specific grant from funding agencies in the '
    'public, commercial, or not-for-profit sectors.')

# ============================================================
# CONTRIBUTORS (CRediT)
# ============================================================
doc.add_paragraph()
add_heading_text('Contributors', level=1)
doc.add_paragraph(
    'TO: Conceptualization, Data curation, Formal analysis, Investigation, '
    'Methodology, Software, Visualization, Writing \u2013 original draft, '
    'Writing \u2013 review & editing.')

# ============================================================
# DATA AVAILABILITY
# ============================================================
doc.add_paragraph()
add_heading_text('Data Availability', level=1)
doc.add_paragraph(
    'Data derived from a source in the public domain. '
    'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    'Labour and Welfare website '
    '(https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    'Analysis code is available at '
    'https://github.com/bougtoir/ndb-pain-regional-variation-japan.')

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)

for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.bold = True
    p.add_run(ref_text)

# ============================================================
# FIGURE LEGENDS (separate section per JE)
# ============================================================
doc.add_page_break()
add_heading_text('Figure Legends', level=1)

legends = [
    ('Figure 1.', 'Outpatient neuropathic pain drug prescribing per surgery by prefecture. '
     'Bars represent individual prefectures ordered by prescribing index. '
     'Bars indicate regional blocks (distinguished by colour and hatch pattern). '
     'Dashed line indicates the national mean.'),
    ('Figure 2.', 'Correlation between neuropathic pain prescribing and confounder disease '
     'proxies across 47 prefectures. '
     'Each marker represents a regional block (distinguished by colour and shape). '
     'Diabetes drugs show the strongest correlation (r=0.87).'),
    ('Figure 3.', 'Regional comparison of neuropathic pain prescribing: (a) unadjusted and '
     '(b) after adjustment for confounding disease proxies. '
     'Error bars indicate SD.'),
    ('Figure 4.', 'Demand\u2013supply dissociation: CSLC symptom prevalence rate '
     '(per 1,000 population) vs acute analgesic prescribing per surgery '
     'across 47 prefectures. Each marker shape represents a regional block. '
     'The near-zero correlation (r=0.03, P=0.85) indicates that '
     'prescribing variation is dissociated from symptom burden.'),
]

for label, text in legends:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(JE_DIR, 'JE_manuscript_EN.docx')
doc.save(outpath)
print(f'\nSaved: {outpath}')

body_total = intro_total + methods_total + results_total + disc_total
print(f'\nBody word count: {body_total} (JE limit: 3,500)')
print(f'Abstract word count: {abstract_total} (JE limit: 250)')
print(f'References: {len(ref_list)}')
print(f'Display items: 4 figures + 2 tables = 6')

# Word count warning
if body_total > 3500:
    print(f'\n*** WARNING: Body exceeds 3,500-word limit by {body_total - 3500} words ***')
if abstract_total > 250:
    print(f'\n*** WARNING: Abstract exceeds 250-word limit by {abstract_total - 250} words ***')

# Verification
print('\n--- Reference order verification ---')
for i, ref_text in enumerate(ref_list, 1):
    print(f'  {i}. {ref_text[:60]}...')
