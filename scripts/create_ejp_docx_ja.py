#!/usr/bin/env python3
"""Create Japanese manuscript for European Journal of Pain (EJP) Original Article format.

Same content as EN version but with Japanese text.
EJP uses author-date references (Harvard style).
Tables and figures uploaded as separate files (not embedded in main text).
Figure/table legends appended at end of manuscript.
Abstract includes Significance statement.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import csv
import numpy as np
from collections import defaultdict

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
EJP_DIR = os.path.join(OUTPUT_DIR, 'ejp')
os.makedirs(EJP_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_JA = {
    '北海道': '北海道', '東北': '東北', '関東': '関東',
    '北陸・甲信越': '北陸・甲信越', '東海': '東海', '近畿': '近畿',
    '中国': '中国', '四国': '四国', '九州・沖縄': '九州・沖縄',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

# References (same Harvard style as EN)
refs = {
    'Anderson2009': '(Anderson et al., 2009)',
    'Befu2001': '(Befu, 2001)',
    'Benchimol2015': '(Benchimol et al., 2015)',
    'Burgess2010': '(Burgess, 2010)',
    'Callister2003': '(Callister, 2003)',
    'Campbell2012': '(Campbell and Edwards, 2012)',
    'Cohen1996': '(Cohen et al., 1996)',
    'Feng2017': '(Feng et al., 2017)',
    'Hobara2005': '(Hobara, 2005)',
    'Kehlet2006': '(Kehlet et al., 2006)',
    'Kumagai2020': '(Kumagai, 2020)',
    'Matsuoka2025': '(Matsuoka et al., 2025)',
    'MHLW2024': '(厚生労働省, 2024)',
    'Okolo2024': '(Okolo et al., 2024)',
    'Onishi2024': '(Onishi and Onishi, 2024)',
    'Pfizer2017': '(ファイザー, 2017)',
    'Raja2020': '(Raja et al., 2020)',
    'Rogger2023': '(Rogger et al., 2023)',
    'Takeda2016': '(武田・鑓水, 2016)',
    'vonElm2007': '(von Elm et al., 2007)',
    'Wakaizumi2024': '(Wakaizumi et al., 2024)',
    'Zborowski1969': '(Zborowski, 1969)',
}

def cite(*keys):
    parts = [refs[k].strip('()') for k in keys]
    return '(' + '; '.join(parts) + ')'

# References bibliography (same as EN — references are always in English)
refs_bib = [
    'Anderson, K.O., Green, C.R., Payne, R. (2009). Racial and ethnic disparities in pain: causes and consequences of unequal care. J Pain 10, 1187\u20131204.',
    'Befu, H. (2001). Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. Melbourne: Trans Pacific Press.',
    'Benchimol, E.I., Smeeth, L., Guttmann, A., Harron, K., Moher, D., Petersen, I., Sorensen, H.T., von Elm, E., Langan, S.M. (2015). The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med 12, e1001885.',
    'Burgess, C. (2010). The \u201cillusion\u201d of homogeneous Japan and national character: discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. Asia Pac J 8(9), 1\u201322.',
    'Callister, L.C. (2003). Cultural influences on pain perceptions and behaviors. Home Health Care Manag Pract 15, 207\u2013211.',
    'Campbell, C.M., Edwards, R.R. (2012). Ethnic differences in pain and pain management. Pain Manag 2, 219\u2013230.',
    'Cohen, D., Nisbett, R.E., Bowdle, B.F., Schwarz, N. (1996). Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d J Pers Soc Psychol 70, 945\u2013960.',
    'Feng, Y., Herdman, M., van Nooten, F., Cleeland, C., Parkin, D., Ikeda, S., Igarashi, A., Devlin, N.J. (2017). An exploration of differences between Japan and two European countries in the self-reporting and valuation of pain and discomfort on the EQ-5D. Qual Life Res 26, 2067\u20132078.',
    'Hobara, M. (2005). Beliefs about appropriate pain behavior: cross-cultural and sex differences between Japanese and Euro-Americans. Eur J Pain 9, 389\u2013393.',
    'Kehlet, H., Jensen, T.S., Woolf, C.J. (2006). Persistent postsurgical pain: risk factors and prevention. Lancet 367, 1618\u20131625.',
    '熊谷慎介. (2020). 東北のイメージを再生産するメディア表象：「秘密のケンミンSHOW」東北復興コーナーを事例に. ことば 41, 21\u201338.',
    'Matsuoka, Y., Morishima, T., Sato, A., Ogawa, T., Miyashiro, I. (2025). Population-based claims study of regional and hospital function differences in opioid prescribing for cancer patients who died in hospital in Japan. Jpn J Clin Oncol 55, hyaf149.',
    '厚生労働省. (2024). NDBオープンデータ 第10回. https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. 2025年1月15日閲覧.',
    'Okolo, C.A., Olorunsogo, T., Babawarun, O. (2024). Cultural variability in pain perception: a review of cross-cultural studies. Int J Sci Res Arch 11, 2550\u20132556.',
    'Onishi, T., Onishi, Y. (2024). Normalized pulse volume as a superior predictor of respiration recovery and quantification of nociception anti-nociception balance compared to opioid effect site concentration: a prospective, observational study. F1000Research 13, 233.',
    'ファイザー株式会社. (2017). 47都道府県 慢性疼痛に関する調査. https://www.pfizer.co.jp/pfizer/company/press/2017. 2025年2月1日閲覧.',
    'Raja, S.N., Carr, D.B., Cohen, M., Finnerup, N.B., Flor, H., Gibson, S., Keefe, F.J., Mogil, J.S., Ringkamp, M., Sluka, K.A., Song, X.J., Stevens, B., Sullivan, M.D., Tutelman, P.R., Ushida, T., Vader, K. (2020). The revised International Association for the Study of Pain definition of pain: concepts, challenges, and compromises. Pain 161, 1976\u20131982.',
    'Rogger, R., Bello, C., Romero, C.S., Urman, R.D., Luedi, M.M., Filipovic, M.G. (2023). Cultural framing and the impact on acute pain and pain services. Curr Pain Headache Rep 27, 429\u2013436.',
    '武田加奈子, 鑓水兼貴. (2016). 痛みの表現「うずく」における地域差. 国語研プロジェクトレビュー 10, 85\u2013107.',
    'von Elm, E., Altman, D.G., Egger, M., Pocock, S.J., Gotzsche, P.C., Vandenbroucke, J.P. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 370, 1453\u20131457.',
    'Wakaizumi, K., Tanaka, C., Shinohara, Y., Wu, Y., Takaoka, S., Kawate, M., Oka, H., Matsudaira, K. (2024). Geographical variation in high-impact chronic pain and psychological associations at the regional level: a multilevel analysis of a large-scale internet-based cross-sectional survey. Front Public Health 12, 1482177.',
    'Zborowski, M. (1969). People in Pain. San Francisco: Jossey-Bass.',
]

# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)





# ============================================================
# TITLE PAGE
# ============================================================
title_ja = doc.add_paragraph()
title_ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_ja.add_run(
    '47都道府県における疼痛関連処方の地域間異質性は\n'
    '「我慢強い一枚岩の患者像」に疑義を呈する：生態学的研究'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

title_en = doc.add_paragraph()
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_en = title_en.add_run(
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures '
    'Challenges a Stoic Monolithic Patient Stereotype: an ecological study'
)
run_en.font.size = Pt(12)
run_en.italic = True

doc.add_paragraph()

add_bold_paragraph('ランニングタイトル：', '日本における疼痛処方の地域差')
doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('大西 辰紀')
run.font.size = Pt(12)

doc.add_paragraph()

add_bold_paragraph('責任著者：',
    '大西 辰紀、[所属]麻酔科学教室、[住所]、[郵便番号]、日本　'
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

add_bold_paragraph('キーワード：',
    '文化的疼痛行動; 生態学的研究; 我慢; 神経障害性疼痛; 地域差; 処方パターン')

doc.add_paragraph()

add_bold_paragraph('論文種別：', 'Original Article')
add_bold_paragraph('投稿先：', 'European Journal of Pain')
add_bold_paragraph('図数：', '5（＋補足1）')
add_bold_paragraph('表数：', '3（本文中にインライン配置）')

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_text('抄録', level=1)

abstract_bg = (
    '異文化研究は日本人が痛みに対し忍耐強いことを一貫して示してきたが、'
    'この文化的忍耐が日本国内で地域的に異なるかどうかは未解明であった。'
)
abstract_methods = (
    '本生態学的研究では、NDBオープンデータ'
    '（集団完全保険請求、2023年4月〜2024年3月）を用い、'
    '47都道府県・9地域ブロックの疼痛関連処方を検討した。'
    'Phase 1は急性期周術期鎮痛薬処方、'
    'Phase 2は交絡疾患調整済みの外来神経障害性疼痛薬処方'
    '（慢性術後疼痛[CPSP]プロキシ）を分析した。'
)
abstract_results = (
    f'Phase 1では、鎮痛薬/手術指標が都道府県間で1.97倍の差を示した'
    f'（岐阜25.20〜鹿児島49.75; Kruskal\u2013Wallis P < 0.001）。'
    f'伝統的に我慢強いとされる東北は鎮痛薬使用がむしろ多かった'
    f'（平均39.97 vs 全国35.78; Cohen\u2019s d = 0.87）。'
    f'Phase 2では、東北の神経障害性疼痛薬処方は未調整で最高値'
    f'（d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）だったが、'
    f'交絡疾患調整後は有意でなくなった'
    f'（P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}）。'
    f'糖尿病薬処方が最強の交絡因子（r = 0.87）であった。'
)
abstract_conclusions = (
    '約2倍の日本国内変動は、「日本人」を疼痛行動の均質なカテゴリーとする仮定に疑義を呈する。'
    '文化的ステレオタイプに基づく仮定ではなく、個別化された疼痛評価が周術期ケアの標準であるべきである。'
)
abstract_significance = (
    '本研究は、全人口をカバーするレセプトデータを用いて47都道府県の疼痛関連処方の国内変動を'
    '定量化した初の全国規模の研究である。'
    '一様に忍耐強いとされてきた単一国家内に約2倍の変動を実証することで、'
    '疼痛耐性に関する文化的一般化が集団レベルでは信頼できないことを示す経験的根拠を提供する。'
    '本知見は、日本人および他の東アジア患者の疼痛評価・治療において、'
    '民族的ステレオタイプを排すべきことを世界の臨床医に促すものである。'
)

p = doc.add_paragraph()
r = p.add_run('背景：')
r.bold = True
p.add_run(abstract_bg)

p = doc.add_paragraph()
r = p.add_run('方法：')
r.bold = True
p.add_run(abstract_methods)

p = doc.add_paragraph()
r = p.add_run('結果：')
r.bold = True
p.add_run(abstract_results)

p = doc.add_paragraph()
r = p.add_run('結論：')
r.bold = True
p.add_run(abstract_conclusions)

p = doc.add_paragraph()
r = p.add_run('Significance：')
r.bold = True
p.add_run(abstract_significance)

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('緒言', level=1)

doc.add_paragraph(
    f'疼痛は普遍的な経験であるが、その表現と管理は文化によって深く影響される'
    f'{cite("Callister2003","Rogger2023")}。Zborowskiが民族集団間で疼痛行動が異なることを'
    f'観察して以来{cite("Zborowski1969")}、文化規範が疼痛報告・治療希求行動・鎮痛薬使用に'
    f'影響することは膨大な文献で確立されている{cite("Okolo2024")}。'
    f'日本は一貫して忍耐強いとされ、Hobaraは日本人回答者が痛み行動を「不適切」と評価する'
    f'程度が米系欧州人より高いことを示し{cite("Hobara2005")}、Fengらは日本人がEQ-5Dにおいて'
    f'痛みを回避するための時間トレードを著しく嫌がることを示した{cite("Feng2017")}。'
    f'「我慢」に象徴されるこの文化的忍耐は、忍耐的な表出が鎮痛薬必要量の低さを意味するという'
    f'臨床的誤解を招くリスクを伴う。'
)

doc.add_paragraph(
    f'多くの研究は国家間の疼痛行動を比較してきた。しかし、国内にも実質的な異質性が存在する：'
    f'CohenとNisbettは米国南部の「名誉の文化」が北部とは測定可能な行動差を生むことを示した{cite("Cohen1996")}。'
    f'日本では地域アイデンティティが強く、東北は伝統的に忍耐強いとされ{cite("Kumagai2020")}、'
    f'痛み表現「うずく」も地域的使用パターンを示す{cite("Takeda2016")}。'
    f'ファイザーの調査では慢性疼痛患者の「痛みに耐えている」割合が都道府県間で48.7%〜81.6%に'
    f'分布していた{cite("Pfizer2017")}。しかし、こうした差が医療利用の測定可能な差に繋がるかは不明であった。'
)

doc.add_paragraph(
    f'日本の国民皆保険制度、標準化された薬価、そしてNDB'
    f'——保険請求の実質全数を捕捉する{cite("MHLW2024")}——は理想的な研究基盤を提供する。'
    f'近年、慢性疼痛有病率に最大1.6倍{cite("Wakaizumi2024")}、'
    f'癌性オピオイド処方に4倍の地域差{cite("Matsuoka2025")}が報告されている。'
    f'周術期設定には入院という条件が医療アクセスの交絡を中和するという追加的利点がある。'
)

doc.add_paragraph(
    f'本探索的研究の目的は3つであった：'
    f'(1) 47都道府県における急性期周術期鎮痛薬処方の地域差のマッピング、'
    f'(2) 交絡疾患調整後の外来神経障害性疼痛薬処方のCPSPプロキシとしての検討、'
    f'(3) 急性期と慢性期の知見の集団レベルでの統合。'
)

# ============================================================
# METHODS
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインおよび報告', level=2)
doc.add_paragraph(
    f'本研究はNDBオープンデータの都道府県レベル集計データを分析した生態学的研究である。'
    f'STROBE声明{cite("vonElm2007")}およびRECORD拡張{cite("Benchimol2015")}に準拠して報告する。'
    f'公開された集計データのみを使用したため、倫理審査は不要であった。'
)

add_heading_text('データソース', level=2)
doc.add_paragraph(
    f'NDBオープンデータ第10回（2023年4月〜2024年3月）を使用した{cite("MHLW2024")}。'
    f'NDBは日本の国民皆保険制度下の全保険者からの請求を捕捉し、約1億2500万人の被保険者を包含する。'
    f'集計データは都道府県レベルで公表され、10件未満のセルは非表示となる。'
)

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県を9つの標準地域ブロックに分類した：'
    '北海道（1）、東北（6：青森、岩手、宮城、秋田、山形、福島）、'
    '関東（7）、北陸・甲信越（6）、東海（4）、近畿（6）、'
    '中国（5）、四国（4）、九州・沖縄（8）。'
)

add_heading_text('Phase 1：急性期周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '入院処方データから3つの鎮痛薬分類を抽出した：'
    '114類（解熱鎮痛消炎剤/NSAIDs/アセトアミノフェン）、'
    '811類（アヘンアルカロイド系麻薬）、821類（合成麻薬）。'
    '入院手術件数はK手術セクションから抽出した。'
    '鎮痛薬/手術指標は各都道府県について以下のように算出した：'
    '入院鎮痛薬処方総量（単位）/入院手術総件数。'
)

add_heading_text('Phase 2：CPSPプロキシとしての外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    '5種類の経口神経障害性疼痛薬を抽出した：'
    'プレガバリン（78規格）、ミロガバリン（8）、デュロキセチン（33）、'
    'トラマドール（3）、ノイロトロピン（1）。'
    '神経障害性疼痛薬処方/手術指標として外来神経障害性疼痛薬総量/入院手術総件数を算出した。'
)

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph('外来データから4つの交絡疾患プロキシを抽出した：')
doc.add_paragraph('経口血糖降下薬（261規格）：糖尿病性ニューロパチーのプロキシ', style='List Bullet')
doc.add_paragraph('帯状疱疹抗ウイルス薬（47規格）：帯状疱疹後神経痛のプロキシ', style='List Bullet')
doc.add_paragraph('抗うつ薬（デュロキセチン除く、128規格）：うつ病のプロキシ', style='List Bullet')
doc.add_paragraph('抗不安薬（112規格）：不安障害のプロキシ', style='List Bullet')
doc.add_paragraph(
    '各プロキシは手術あたりで表した。外来神経ブロック処置件数（73コード）を独立したCPSPプロキシとした。'
)

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    'Phase 1の地域差はKruskal\u2013Wallis検定で評価し、'
    '事後検定としてBonferroni補正付きMann\u2013Whitney U検定を行った。'
    'Phase 2では5つの回帰モデルを適合した：'
)
doc.add_paragraph('モデル1：未調整地域比較', style='List Bullet')
doc.add_paragraph('モデル2：神経障害性疼痛薬 ~ 糖尿病 + 帯状疱疹 + 抗うつ薬 + 抗不安薬 + 地域', style='List Bullet')
doc.add_paragraph('モデル3：コア神経障害性疼痛薬（プレガバリン + ミロガバリン）~ 同交絡因子', style='List Bullet')
doc.add_paragraph('モデル4：神経ブロック ~ 同交絡因子', style='List Bullet')
doc.add_paragraph('モデル5：神経障害性疼痛薬 ~ 急性期鎮痛薬指標 + 交絡因子（統合）', style='List Bullet')
doc.add_paragraph(
    '調整済みCPSP指標は、神経障害性疼痛薬処方を4つの交絡プロキシに回帰した残差として導出した。'
    'PearsonおよびSpearman相関により指標間関係を評価した。'
    '全解析にPython 3.11（NumPy 1.24, SciPy 1.11, matplotlib 3.8）を使用した。'
)

# ============================================================
# RESULTS
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('Phase 1：急性期周術期鎮痛薬処方の地域差', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月のNDBには、47都道府県で計7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方が記録された。'
    '全国平均鎮痛薬/手術指標は35.78（SD 5.56）であり、'
    '岐阜の25.20から鹿児島の49.75まで1.97倍の差があった'
    '（9地域間Kruskal\u2013Wallis P < 0.001；表1）。'
)

# === TABLE 1 inline ===
region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# [表1は別ファイルとしてアップロード]

doc.add_paragraph(
    '明確な地域クラスタリングが観察された。東海・近畿（西日本）が最低値、'
    '九州・沖縄と北海道が最高値を示した。'
    '伝統的に最も忍耐強いとされる東北は、9地域中7位で平均指標39.97（SD 3.53）であり、'
    '非東北平均35.17を有意に上回った'
    '（Mann\u2013Whitney U = 190, P = 0.031; Cohen\u2019s d = 0.87）。'
    '6東北県全てが全国上位半分にランクした。'
    'このパターンは薬効群間で一貫していた：NSAIDs（P = 0.044）、'
    'オピオイドアルカロイド（P = 0.003）、合成麻薬（P = 0.001）。'
)

add_heading_text('Phase 2：外来神経障害性疼痛薬処方（未調整）', level=2)
doc.add_paragraph(
    f'全国の外来神経障害性疼痛薬処方は計2,289,549,163単位で、'
    f'プレガバリン（40.2%）、ノイロトロピン（20.1%）、ミロガバリン（19.6%）、'
    f'デュロキセチン（15.3%）、トラマドール（4.9%）から構成された。'
    f'東北は著しく高い指標を示した（{reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\u2019s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）。'
    f'岩手（566.7）、青森（519.3）、秋田（461.1）が全国上位3位を占めた（図1）。'
)

# [図1は別ファイルとしてアップロード]

add_heading_text('交絡因子分析', level=2)
doc.add_paragraph(
    f'神経障害性疼痛薬処方は交絡プロキシと強い相関を示した。'
    f'糖尿病薬処方が最強（r = 0.87, P < 0.001）であり、'
    f'抗不安薬（r = 0.75）、抗うつ薬（r = 0.46）、帯状疱疹抗ウイルス薬（r = 0.19）が続いた。'
    f'これらは分散の{reg["model2_adjusted"]["R2"]*100:.1f}%を説明した'
    f'（モデル2のR\u00b2 = {reg["model2_adjusted"]["R2"]:.3f}；図2）。'
)

# [図2は別ファイルとしてアップロード]

add_heading_text('交絡因子調整後の分析', level=2)
doc.add_paragraph(
    f'4交絡因子調整後、東北効果は減弱し有意でなくなった（モデル2：'
    f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, P = {reg["model2_adjusted"]["tohoku_p"]:.3f}）。'
    f'他のモデルでも一貫していた：'
    f'モデル3（\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}）、'
    f'モデル4（P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}）、'
    f'モデル5（\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}；表2）。'
)

# [表2は別ファイルとしてアップロード]

unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attn = (1 - adj_d / unadj_d) * 100

doc.add_paragraph(
    f'調整済みCPSP指標は未調整データとは劇的に異なるパターンを示した（図3）。'
    f'東北平均は強い正値から控えめな有意でない過剰へと変化した'
    f'（{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}）。'
    f'中国地方が最高の調整済み指標、東海が最低であった（図4）。'
)

# [図3は別ファイルとしてアップロード]
# [図4は別ファイルとしてアップロード]

add_heading_text('Phase 1\u2013Phase 2 統合', level=2)
doc.add_paragraph(
    f'急性期処方は未調整の神経障害性疼痛薬処方と正の相関を示した'
    f'（r = 0.38, P = 0.008；図5a）。交絡調整後は減弱した'
    f'（r = 0.29, P = 0.052；図5b）。'
    f'モデル5では急性期指標は有意な予測因子であった'
    f'（\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, P = {reg["model5_integrated"]["acute_pain_p"]:.3f}）が、'
    f'東北効果は有意でなかった'
    f'（\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, P = {reg["model5_integrated"]["tohoku_p"]:.3f}）。'
    f'Zスコアヒートマップにより多変数の異質パターンを確認した（補足図1）。'
    f'調整後、東北効果は{attn:.0f}%減弱し有意でなくなった（表3）。'
)

# [図5は別ファイルとしてアップロード]
# [表3は別ファイルとしてアップロード]

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    f'本研究は、自由に利用可能なNDBオープンデータを活用し、日本の全47都道府県における'
    f'周術期および慢性疼痛関連処方を初めてマッピングしたものである。'
    f'この探索的分析から3つの主要知見が浮かび上がった。'
)

add_heading_text('忍耐的文化の中の地域差', level=2)
doc.add_paragraph(
    f'日本の痛みに対する忍耐的文化が広く記録されているにもかかわらず'
    f'{cite("Hobara2005","Feng2017")}、急性期鎮痛薬処方に都道府県間で1.97倍の差、'
    f'地域ブロック間で有意差が認められた。'
    f'これは他領域の知見と類似する：CohenとNisbettの「名誉の文化」研究は、'
    f'単一国家内でも地域的文化規範が測定可能な行動差を生むことを示した{cite("Cohen1996")}。'
    f'本知見は、日本の疼痛文化が一枚岩ではなく、地域人口動態・医療インフラ・'
    f'地域の臨床慣行が全国的に共有される文化規範の表面下に異質性を生み出していることを示唆する。'
)

add_heading_text('臨床的示唆：一枚岩的文化ステレオタイプの危険性', level=2)
doc.add_paragraph(
    f'本研究で示された日本国内1.97倍の変動は、国内にとどまらない直接的な臨床的関連性を持つ。'
    f'民族的・文化的ステレオタイプが臨床医の疼痛評価と鎮痛薬処方に影響することは多くのエビデンスが示す。'
    f'Andersonらは米国で人種的・民族的マイノリティが急性・慢性・癌・緩和ケア全てにおいて不十分な'
    f'疼痛管理を受けることを示した{cite("Anderson2009")}。'
    f'CampbellとEdwardsは臨床医の文化的期待が体系的な治療不足・過剰をもたらしうることを示した{cite("Campbell2012")}。'
    f'Roggerらは文化的枠組みが患者の報告だけでなく臨床医の解釈・対応にも影響することを強調した{cite("Rogger2023")}。'
)

doc.add_paragraph(
    f'このリスクは日本が文化的・民族的に均質な社会だという広範な認識によって増幅される。'
    f'日本人論（日本人の独自性理論）は日本人が行動規範を共有する均質な集団であるという観念を長く'
    f'推進してきた{cite("Befu2001")}。しかし、この「均質性の覇権」は経験的事実ではなくイデオロギー的構築物である。'
    f'Burgessはこの均質な日本という「幻想」が社会政策と世論に具体的な影響を持つことを示した{cite("Burgess2010")}。'
    f'臨床的文脈では、「日本人は忍耐強い」と「日本人は均質である」という2つのステレオタイプの組み合わせが'
    f'二重に誤解を招く前提を作り出す。本研究の1.97倍の国内変動はこの前提を直接的に否定する。'
)

doc.add_paragraph(
    f'端的に言えば、国籍だけから疼痛行動を予測できる「日本人患者」なるものは存在しない——'
    f'存在するのは、人口動態・臨床環境・処方文化がそれぞれ異なる47都道府県からの個々の患者だけである。'
    f'海外で治療される日本人患者に対し、臨床医は日本文化が一律に忍耐的であり、'
    f'その忍耐が低い鎮痛薬必要量に繋がるという二重の誤解の下で行動している可能性がある。'
    f'本データは両前提に疑義を呈する。'
)

doc.add_paragraph(
    f'したがって、いかなる国民集団の疼痛行動を一枚岩的に特徴づけることも臨床的害のリスクを伴う'
    f'——この議論は日本に限らない。文化的ステレオタイプが患者の治療選択肢を狭めるとき、'
    f'患者は治療的不利益を被る{cite("Raja2020")}。本研究の願望は、国籍や民族にかかわらず、'
    f'文化的ステレオタイプによって治療選択肢が制約される不利益を被る患者がより少なくなることである。'
    f'個別化された疼痛評価こそが公正な周術期ケアの礎である。'
)

add_heading_text('疼痛は個人的経験である', level=2)
doc.add_paragraph(
    f'根本的に、疼痛は個人的経験である。改訂されたIASPの定義は疼痛を'
    f'「実際のまたは潜在的な組織損傷に伴う、またはそれに類似した不快な感覚的・情動的体験」'
    f'と記述する——本質的に主観的な定義である{cite("Raja2020")}。'
    f'本研究の生態学的データは集団レベルでこの原理を例証する：'
    f'同一言語・保険制度・文化的遺産を共有する単一国家内でも、疼痛関連処方は約2倍変動する。'
    f'文化的ラベルは患者の侵害受容状態の直接測定の代替にはなりえない。'
    f'正規化脈波容積のような客観的侵害受容モニタリング{cite("Onishi2024")}は、'
    f'文化的背景を問わず周術期評価を標準化する助けとなりうる。'
)

add_heading_text('交絡因子は東北の見かけの過剰を説明する', level=2)
doc.add_paragraph(
    f'方法論的に最も重要な知見は、神経障害性疼痛薬処方の劇的な地域差'
    f'（東北 vs 非東北の未調整 d = {reg["model1_unadjusted"]["cohens_d"]:.2f}）が、'
    f'交絡疾患プロキシによって大部分説明されたことである。'
    f'糖尿病薬処方だけで r = 0.87の相関を示し、ガバペンチノイドを要する糖尿病性ニューロパチーの'
    f'高い有病率を反映している。調整後、東北効果は{attn:.0f}%減弱し有意でなくなった（表3）。'
    f'これは生態学的疼痛研究に重要な示唆を持つ：神経障害性疼痛薬処方を集団レベルのCPSPプロキシとして'
    f'使用する研究は交絡疾患を考慮しなければならない。そのような調整なしには、糖尿病有病率の地域差が'
    f'CPSP差と誤認されうる。'
)

add_heading_text('集団レベルの急性-慢性疼痛連続体', level=2)
doc.add_paragraph(
    f'Phase 1（急性期）とPhase 2（慢性期、調整済み）の正の相関（r = 0.29, P = 0.052）は、'
    f'地域の急性期疼痛管理強度とその後の慢性疼痛関連処方との間の控えめなリンクを示唆する。'
    f'生態学的相関は因果関係を確立できないが、急性期術後痛の強度がCPSPの危険因子であるという'
    f'個人レベルのエビデンスと整合する{cite("Kehlet2006")}。'
)

add_heading_text('強みと限界', level=2)
doc.add_paragraph(
    f'強みには、日本の全保険償還医療を網羅する集団完全データの使用、急性期・慢性期疼痛プロキシの'
    f'新規統合、透明な交絡調整方法論、仮説生成を可能にする探索的デザインが含まれる。'
    f'周術期に焦点を当てた固有の強みとして、Phase 1の全患者が定義上入院患者であるため、'
    f'医療アクセスが交絡しない点がある。'
)

doc.add_paragraph(
    f'主な限界は生態学的デザインに固有のものである。分析単位は個人ではなく都道府県であり、'
    f'生態学的相関は個人レベルの関連を反映しない可能性がある（生態学的誤謬）。'
    f'NDBオープンデータは診断コードを含まないため、CPSPを直接同定できない。'
    f'薬剤処方プロキシは疾患有病率を正確に捕捉しない可能性がある。'
    f'横断的デザインは時間的順序（手術→急性期疼痛→CPSP）を区別できず、'
    f'年齢分布・手術症例構成・処方文化などの未測定交絡因子が残余の地域差に寄与しうる。'
)

add_heading_text('結論', level=2)
doc.add_paragraph(
    f'日本の文化的に根付いた疼痛忍耐規範（我慢）にもかかわらず、'
    f'周術期および慢性疼痛関連処方は都道府県間で最大1.97倍変動する。'
    f'交絡疾患、特に糖尿病は神経障害性疼痛薬処方の見かけの地域パターンを大きく修飾する。'
    f'これらの知見は日本の疼痛文化が一枚岩ではないことを示す。'
    f'「日本人」を疼痛行動の均一カテゴリーとして扱うことは、海外で治療される日本人患者の'
    f'不十分な鎮痛につながるリスクがあり——同じ原理は全ての患者集団に適用されるあらゆる'
    f'文化的ラベルに当てはまる。'
    f'個別化された疼痛評価こそが全臨床環境における公正な周術期ケアの保証となるべきである。'
)

# ============================================================
# ACKNOWLEDGMENTS etc.
# ============================================================
doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph(
    'NDBオープンデータを公開している厚生労働省に感謝する。'
    'データ処理および原稿作成の一部に生成AI（Claude, Anthropic）の支援を受けた。'
    '原稿の正確性と内容について著者が全責任を負う。')

doc.add_paragraph()
add_heading_text('著者貢献', level=1)
doc.add_paragraph(
    '大西辰紀：概念化、データキュレーション、正式分析、調査、方法論、ソフトウェア、'
    '可視化、原稿執筆（初稿）、原稿執筆（査読・編集）。')

doc.add_paragraph()
add_heading_text('利益相反', level=1)
doc.add_paragraph('著者に申告すべき利益相反はない。')

doc.add_paragraph()
add_heading_text('資金源', level=1)
doc.add_paragraph('本研究は公的・商業的・非営利セクターからの特定の助成金を受けていない。')

doc.add_paragraph()
add_heading_text('データ可用性', level=1)
doc.add_paragraph(
    f'本研究で使用したNDBオープンデータは厚生労働省ウェブサイト'
    f'（https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html）から公開されている。'
    f'分析コードは https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan で利用可能。')

# References
doc.add_page_break()
add_heading_text('文献', level=1)
for bib in refs_bib:
    doc.add_paragraph(bib)

# ============================================================
# TABLE LEGENDS
# ============================================================
doc.add_page_break()
add_heading_text('表の説明', level=1)

p = doc.add_paragraph()
r = p.add_run('表1. ')
r.bold = True
p.add_run('Phase 1：9地域ブロック別の入院鎮痛薬/手術指標の概要。'
    '値は鎮痛薬/手術指標（各都道府県の入院鎮痛薬処方単位合計÷入院手術件数合計）を表す。'
    '9地域間のKruskal\u2013Wallis検定：P < 0.001。')

p = doc.add_paragraph()
r = p.add_run('表2. ')
r.bold = True
p.add_run('Phase 2：東北指標・交絡因子調整を含む回帰モデル。'
    '*** P < 0.001; ns = 有意でない。'
    'モデル2〜5：東北指標（二値）と交絡疾患プロキシによる重回帰分析。'
    '調整済CPSP：神経障害性疼痛薬処方を4交絡プロキシに回帰した残差。')

p = doc.add_paragraph()
r = p.add_run('表3. ')
r.bold = True
p.add_run('交絡因子調整による東北地域指標の変化。'
    '交絡因子：経口血糖降下薬（糖尿病プロキシ）、帯状疱疹抗ウイルス薬、'
    '抗うつ薬（デュロキセチン除く）、抗不安薬。')

# ============================================================
# FIGURE LEGENDS
# ============================================================
doc.add_page_break()
add_heading_text('図の説明', level=1)

p = doc.add_paragraph()
r = p.add_run('図1. ')
r.bold = True
p.add_run('都道府県別の外来神経障害性疼痛薬処方/手術指標（未調整）。'
    '棒グラフは総神経障害性疼痛薬処方量（プレガバリン＋ミロガバリン＋デュロキセチン＋'
    'トラマドール＋ノイロトロピン）を入院手術件数で除した値。'
    '東北県（赤棒・赤枠）が高値側に集中。破線＝全国平均。')

p = doc.add_paragraph()
r = p.add_run('図2. ')
r.bold = True
p.add_run('外来神経障害性疼痛薬処方と交絡疾患プロキシの相関。'
    '各点は1都道府県を表す。東北県は赤枠で示す。糖尿病薬が最強の相関（r = 0.87）。')

p = doc.add_paragraph()
r = p.add_run('図3. ')
r.bold = True
p.add_run('交絡因子調整済みCPSP指標（都道府県別）。糖尿病薬・帯状疱疹抗ウイルス薬・'
    '抗うつ薬・抗不安薬による回帰の残差。東北県（赤枠）は調整後に分布全体に分散。')

p = doc.add_paragraph()
r = p.add_run('図4. ')
r.bold = True
p.add_run('神経障害性疼痛薬処方の地域比較：(a) 未調整、(b) 交絡因子調整後。'
    '東北（赤枠）は調整後に最高位から中間に移動。誤差棒＝標準偏差。')

p = doc.add_paragraph()
r = p.add_run('図5. ')
r.bold = True
p.add_run('Phase 1（急性期鎮痛薬処方）とPhase 2（神経障害性疼痛薬処方＝CPSPプロキシ）の統合。'
    '(a) 未調整：正の相関（r = 0.38, P = 0.008）。(b) 調整後：減弱した相関（r = 0.29, P = 0.052）。')

# ============================================================
# SUPPLEMENTARY MATERIAL
# ============================================================
doc.add_page_break()
add_heading_text('補足資料', level=1)
p = doc.add_paragraph()
r = p.add_run('補足図1. ')
r.bold = True
p.add_run('全指標の都道府県別Zスコアヒートマップ。各行は変数、各列は都道府県（神経障害性疼痛薬処方順にソート）。'
    '赤＝平均以上、青＝平均以下。東北県は赤い垂直線で示す。')

# SAVE
outpath = os.path.join(EJP_DIR, 'EJP_manuscript_JA.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
