#!/usr/bin/env python3
"""Create English manuscript for JoPPP (Journal of Pharmaceutical Policy and Practice).

JoPPP requirements (Research Article):
- Abstract: structured (Background, Methods, Results, Conclusion) <= 300 words
- Body: <= 6,000 words including tables, references, figure/table captions and footnotes
- Section order: Background, Methods, Results, Discussion, Conclusion
- References: Vancouver style -- numbered in order of first appearance
- Figures: uploaded as separate files (not embedded in the manuscript text)
- Tables: supplied as editable files and placed where first cited
- Keywords: between 4 and 10
- STROBE checklist for observational studies
- Disclosure statement / competing interests
- AI use acknowledgement
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import os

try:
    from latex2mathml.converter import convert as latex_to_mathml
    from math_ml2omml import convert as mathml_to_omml
    _HAS_OMML = True
except Exception:
    _HAS_OMML = False
    import warnings
    warnings.warn(
        'latex2mathml or math_ml2omml not available; OMML equations will not be embedded. '
        'Install both packages to generate Word-native equations.'
    )

import json
import re
import csv
import numpy as np
from scipy import stats
from collections import defaultdict

# ============================================================
# Paths
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JOPP_DIR = os.path.join(OUTPUT_DIR, 'joppp')
os.makedirs(JOPP_DIR, exist_ok=True)

FIG_DIR = OUTPUT_DIR

# ============================================================
# Load data
# ============================================================
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

class_rows = []
with open(os.path.join(OUTPUT_DIR, 'analgesic_class_summary.csv'), 'r', encoding='utf-8-sig') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('class_code', 'class_name_en'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        class_rows.append(r)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

PREF_EN = {
    '北海道': 'Hokkaido', '青森県': 'Aomori', '岩手県': 'Iwate', '宮城県': 'Miyagi',
    '秋田県': 'Akita', '山形県': 'Yamagata', '福島県': 'Fukushima', '茨城県': 'Ibaraki',
    '栃木県': 'Tochigi', '群馬県': 'Gunma', '埼玉県': 'Saitama', '千葉県': 'Chiba',
    '東京都': 'Tokyo', '神奈川県': 'Kanagawa', '新潟県': 'Niigata', '富山県': 'Toyama',
    '石川県': 'Ishikawa', '福井県': 'Fukui', '山梨県': 'Yamanashi', '長野県': 'Nagano',
    '岐阜県': 'Gifu', '静岡県': 'Shizuoka', '愛知県': 'Aichi', '三重県': 'Mie',
    '滋賀県': 'Shiga', '京都府': 'Kyoto', '大阪府': 'Osaka', '兵庫県': 'Hyogo',
    '奈良県': 'Nara', '和歌山県': 'Wakayama', '鳥取県': 'Tottori', '島根県': 'Shimane',
    '岡山県': 'Okayama', '広島県': 'Hiroshima', '山口県': 'Yamaguchi', '徳島県': 'Tokushima',
    '香川県': 'Kagawa', '愛媛県': 'Ehime', '高知県': 'Kochi', '福岡県': 'Fukuoka',
    '佐賀県': 'Saga', '長崎県': 'Nagasaki', '熊本県': 'Kumamoto', '大分県': 'Oita',
    '宮崎県': 'Miyazaki', '鹿児島県': 'Kagoshima', '沖縄県': 'Okinawa',
}

# ============================================================
# REFERENCES — Vancouver style, numbered in order of first appearance
# ============================================================
ref_list = [
    # 1 Callister
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract 2003;15:207\u201311. '
    'doi:10.1177/1084822302250687',
    # 2 Rogger
    'Rogger R, Bello C, Romero CS, et al. '
    'Cultural framing and the impact on acute pain and pain services. '
    'Curr Pain Headache Rep 2023;27:429\u201336. '
    'doi:10.1007/s11916-023-01125-2',
    # 3 Zborowski
    'Zborowski M. People in Pain. San Francisco: Jossey-Bass, 1969.',
    # 4 Okolo
    'Okolo CA, Olorunsogo T, Babawarun O. Cultural variability in pain perception: '
    'a review of cross-cultural studies. Int J Sci Res Arch 2024;11:2550\u20136. '
    'doi:10.30574/ijsra.2024.11.1.0339',
    # 5 Hobara
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain 2005;9:389\u201393. '
    'doi:10.1016/j.ejpain.2004.09.006',
    # 6 Feng
    'Feng Y, Herdman M, van Nooten F, et al. An exploration of differences between Japan '
    'and two European countries in the self-reporting and valuation of pain and discomfort '
    'on the EQ-5D. Qual Life Res 2017;26:2067\u201378. '
    'doi:10.1007/s11136-017-1541-5',
    # 7 Cohen
    'Cohen D, Nisbett RE, Bowdle BF, et al. Insult, aggression, and the southern '
    'culture of honor: an "experimental ethnography." J Pers Soc Psychol 1996;70:945\u201360. '
    'doi:10.1037/0022-3514.70.5.945',
    # 8 Kumagai
    'Kumagai S. Media representations reproducing images of Tohoku: the Tohoku '
    'reconstruction corner in "Secret Kenmin SHOW." Kotoba 2020;41:21\u201338. doi:10.20741/kotoba.41.0_21 [in Japanese]',
    # 9 Takeda
    'Takeda K, Yarimizu K. Regional differences in the pain expression uzuku. '
    'NINJAL Research Papers 2016;10:221\u201343. doi:10.15084/00000816 [in Japanese]',
    # 10 Pfizer
    'Pfizer Japan Inc. 47-prefecture survey on chronic pain (2012 vs 2017 comparison) [in Japanese]. Research Research, 2017. '
    'https://www.lisalisa50.com/research20171014_2.html (accessed 21 Aug 2026).',
    # 11 MHLW
    'Ministry of Health, Labour and Welfare. NDB Open Data, 10th edition. '
    'https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html '
    '(accessed 15 Jan 2025).',
    # 12 Wakaizumi
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact '
    'chronic pain and psychological associations at the regional level: a multilevel analysis '
    'of a large-scale internet-based cross-sectional survey. Front Public Health 2024;12:1482177. '
    'doi:10.3389/fpubh.2024.1482177',
    # 13 Matsuoka
    'Matsuoka Y, Morishima T, Sato A, et al. Population-based claims study '
    'of regional and hospital function differences in opioid prescribing for cancer patients '
    'who died in hospital in Japan. Jpn J Clin Oncol 2025;55:hyaf149. doi:10.1093/jjco/hyaf149',
    # 14 Taira
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilisation '
    'in Japan: an ecological study using the National Database of Health Insurance Claims. '
    'Lancet Reg Health West Pac 2021;12:100170. '
    'doi:10.1016/j.lanwpc.2021.100170',
    # 15 von Elm (STROBE)
    'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational '
    'Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. '
    'Lancet 2007;370:1453\u20137. '
    'doi:10.1016/S0140-6736(07)61602-X',
    # 16 Benchimol (RECORD)
    'Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using '
    'Observational Routinely-collected health Data (RECORD) statement. '
    'PLoS Med 2015;12:e1001885. '
    'doi:10.1371/journal.pmed.1001885',
    # 17 Anderson
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain: '
    'causes and consequences of unequal care. J Pain 2009;10:1187\u2013204. '
    'doi:10.1016/j.jpain.2009.10.002',
    # 18 Campbell
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag 2012;2:219\u201330. '
    'doi:10.2217/pmt.12.7',
    # 19 Befu
    'Befu H. Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
    'Melbourne: Trans Pacific Press, 2001.',
    # 20 Burgess
    'Burgess C. The "illusion" of homogeneous Japan and national character: '
    'discourse as a tool to transcend the "myth" vs. "reality" binary. '
    'Asia Pac J 2010;8(9):1\u201322. doi:10.1017/s1557466010009381',
    # 21 Raja (IASP)
    'Raja SN, Carr DB, Cohen M, et al. The revised International Association for the '
    'Study of Pain definition of pain: concepts, challenges, and compromises. '
    'Pain 2020;161:1976\u201382. '
    'doi:10.1097/j.pain.0000000000001939',
    # 22 Kehlet
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet 2006;367:1618\u201325. '
    'doi:10.1016/S0140-6736(06)68700-X',
]


def cite(*nums):
    """Return Vancouver superscript citation marker string."""
    return '{' + ','.join(str(n) for n in nums) + '}'


def add_ref_runs(p, text):
    """Parse text with {n} or {n,m} markers and create runs with font-based superscript."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)


# ============================================================
# Convert inline math symbols to Word OMML equations
# ============================================================
BETA_CHAR = '\u03b2'
SUP2_CHAR = '\u00b2'
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def make_omml(latex):
    """Render a LaTeX math fragment as a docx OMML element."""
    mathml = latex_to_mathml(latex, display='inline')
    omml = mathml_to_omml(mathml)
    omml = omml.replace('<m:oMath>', f'<m:oMath xmlns:m="{MATH_NS}">', 1)
    return parse_xml(omml)


# Each tuple is (regex pattern, replacement generating a LaTeX string)
EQ_PATTERNS = [
    (r'R' + SUP2_CHAR + r'\s*=\s*([\-0-9.]+)', lambda m: 'R^2=' + m.group(1)),
    (BETA_CHAR + r'\s*=\s*([\-0-9.]+)', lambda m: r'\beta=' + m.group(1)),
    (BETA_CHAR + r'(?!\s*=)', lambda m: r'\beta'),
    (r'\br\s*=\s*([\-0-9.]+)', lambda m: 'r=' + m.group(1)),
    (r'\bd\s*=\s*([\-0-9.]+)', lambda m: 'd=' + m.group(1)),
    (r'\bP\s*(=|<)\s*(0?\.[0-9]+|<0\.001)', lambda m: 'P' + m.group(1) + m.group(2)),
]


def split_text_by_pattern(items, pattern, latex_fn):
    """Split every string item in items, replacing matches with equation markers."""
    regex = re.compile(pattern)
    out = []
    for item in items:
        if not isinstance(item, str):
            out.append(item)
            continue
        last = 0
        for m in regex.finditer(item):
            if m.start() > last:
                out.append(item[last:m.start()])
            out.append(('eq', latex_fn(m)))
            last = m.end()
        if last < len(item):
            out.append(item[last:])
    return out


def process_text_for_equations(text):
    """Return a list of (str|('eq', latex)) pieces for one text run."""
    pieces = [text]
    for pattern, fn in EQ_PATTERNS:
        pieces = split_text_by_pattern(pieces, pattern, fn)
    return pieces


def clone_run_font(source_run, target_run):
    """Copy basic font properties from source to target run."""
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.superscript is not None:
        target_run.font.superscript = source_run.font.superscript
    if source_run.font.subscript is not None:
        target_run.font.subscript = source_run.font.subscript
    if source_run.bold:
        target_run.font.bold = True
    if source_run.italic:
        target_run.font.italic = True
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def process_paragraph(para):
    """Replace math expressions inside one paragraph with OMML equations while preserving run order."""
    runs = list(para.runs)
    # Preserve formatting for all runs first; otherwise iterating over the
    # underlying element list while modifying it can reorder runs.
    sequence = []
    for run in runs:
        text = run.text
        if not text:
            sequence.append((run, None))
            continue
        pieces = process_text_for_equations(text)
        if len(pieces) == 1 and isinstance(pieces[0], str):
            sequence.append((run, pieces[0]))
            continue
        for piece in pieces:
            sequence.append((run, piece))

    # Remove all existing run elements, then rebuild in the correct order.
    for r in list(para._p.findall(qn('w:r'))):
        para._p.remove(r)

    for src, piece in sequence:
        if piece is None:
            continue
        if isinstance(piece, str):
            new_run = para.add_run(piece)
            clone_run_font(src, new_run)
        else:
            try:
                omml = make_omml(piece[1])
                # Office Math ML (m:oMath) is a paragraph-level sibling of w:r,
                # not a child of a run; append directly to the paragraph element.
                para._p.append(omml)
            except Exception:
                # Fallback: keep original LaTeX-like text if conversion fails
                new_run = para.add_run(piece[1])
                clone_run_font(src, new_run)


def embed_equations_in_doc(document):
    """Process paragraphs recursively, including those inside tables."""
    if not _HAS_OMML:
        return
    def _process_container(container):
        for para in container.paragraphs:
            process_paragraph(para)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_container(cell)
    _process_container(document)


# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# Page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def wc(text):
    return len(re.sub(r'\{[^}]+\}', '', text).split())


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def add_inline_figure(caption_text, fig_num):
    """Insert figure caption only; figures are supplied as individual TIFF files."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'Figure {fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(caption_text)
    r2.font.size = Pt(10)


# ============================================================
# Computed values
# ============================================================
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_analgesic_tohoku = scr['analgesic_inpatient']['scr_tohoku_mean']
scr_analgesic_non_tohoku = scr['analgesic_inpatient']['scr_non_tohoku_mean']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']
scr_neuro_tohoku = scr['neuropathic_outpatient']['scr_tohoku_mean']
scr_neuro_non_tohoku = scr['neuropathic_outpatient']['scr_non_tohoku_mean']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])


def fmt_p(p):
    return '<0.001' if p < 0.001 else f'={p:.3f}'


# --- Phase 1: acute perioperative analgesic prescribing ---
acute_vals = np.array([r['acute_analgesic_per_surgery'] for r in rows])
surgery_counts = np.array([r['surgery_count'] for r in rows])

total_surgeries = int(surgery_counts.sum())
total_analgesic_units = int((acute_vals * surgery_counts).sum())
acute_mean = float(np.mean(acute_vals))
acute_sd = float(np.std(acute_vals, ddof=1))

acute_min_row = min(rows, key=lambda r: r['acute_analgesic_per_surgery'])
acute_max_row = max(rows, key=lambda r: r['acute_analgesic_per_surgery'])
acute_min_pref = PREF_EN[acute_min_row['pref_name']]
acute_max_pref = PREF_EN[acute_max_row['pref_name']]
acute_min_val = float(acute_min_row['acute_analgesic_per_surgery'])
acute_max_val = float(acute_max_row['acute_analgesic_per_surgery'])
acute_fold_ratio = acute_max_val / acute_min_val

acute_kruskal = stats.kruskal(*region_data.values())
acute_kruskal_p = float(acute_kruskal.pvalue)

tohoku_acute = np.array([r['acute_analgesic_per_surgery'] for r in rows if r['is_tohoku']])
non_tohoku_acute = np.array([r['acute_analgesic_per_surgery'] for r in rows if not r['is_tohoku']])
acute_mw = stats.mannwhitneyu(tohoku_acute, non_tohoku_acute, alternative='two-sided')
acute_mw_u = float(acute_mw.statistic)
acute_mw_p = float(acute_mw.pvalue)
pooled_sd = np.sqrt(((len(tohoku_acute) - 1) * np.var(tohoku_acute, ddof=1) +
                     (len(non_tohoku_acute) - 1) * np.var(non_tohoku_acute, ddof=1)) /
                    (len(tohoku_acute) + len(non_tohoku_acute) - 2))
acute_d = float((np.mean(tohoku_acute) - np.mean(non_tohoku_acute)) / pooled_sd)
tohoku_acute_mean = float(np.mean(tohoku_acute))
tohoku_acute_sd = float(np.std(tohoku_acute, ddof=1))
non_tohoku_acute_mean = float(np.mean(non_tohoku_acute))

region_means = {reg: np.mean(vals) for reg, vals in region_data.items()}
tohoku_rank = sorted(region_means.items(), key=lambda x: x[1]).index(
    ('東北', region_means['東北'])) + 1
_suffix = {1: 'st', 2: 'nd', 3: 'rd'}
tohoku_rank_text = f"{tohoku_rank}{_suffix.get(tohoku_rank % 10, 'th') if not (11 <= tohoku_rank <= 13) else 'th'}"

neuro_min = min(r['neuropathic_per_surgery'] for r in rows)
neuro_max = max(r['neuropathic_per_surgery'] for r in rows)
neuro_fold_ratio = neuro_max / neuro_min

# --- Phase 2: outpatient neuropathic pain prescribing ---
neuro_total_out = sum(r['neuropathic_total_out'] for r in rows)
drug_totals = {
    'pregabalin': sum(r['pregabalin_out'] for r in rows),
    'neurotropin': sum(r['neurotropin_out'] for r in rows),
    'mirogabalin': sum(r['mirogabalin_out'] for r in rows),
    'duloxetine': sum(r['duloxetine_out'] for r in rows),
    'tramadol': sum(r['tramadol_out'] for r in rows),
}
drug_pcts = {k: v / neuro_total_out * 100 for k, v in drug_totals.items()}

top3_neuro = sorted(rows, key=lambda r: r['neuropathic_per_surgery'], reverse=True)[:3]
top3_neuro_info = [
    (PREF_EN[r['pref_name']], float(r['neuropathic_per_surgery'])) for r in top3_neuro
]

# --- Confounder correlations ---
diabetes_r, diabetes_p = stats.pearsonr(
    [r['neuropathic_per_surgery'] for r in rows],
    [r['diabetes_per_surgery'] for r in rows])
anxiolytic_r, anxiolytic_p = stats.pearsonr(
    [r['neuropathic_per_surgery'] for r in rows],
    [r['anxiolytic_per_surgery'] for r in rows])
antidep_r, antidep_p = stats.pearsonr(
    [r['neuropathic_per_surgery'] for r in rows],
    [r['antidep_per_surgery'] for r in rows])
herpes_r, herpes_p = stats.pearsonr(
    [r['neuropathic_per_surgery'] for r in rows],
    [r['herpes_per_surgery'] for r in rows])

# --- Phase 1-2 integration correlations ---
acute_neuro_r, acute_neuro_p = stats.pearsonr(
    acute_vals, [r['neuropathic_per_surgery'] for r in rows])
acute_adj_r, acute_adj_p = stats.pearsonr(
    acute_vals, [r['adjusted_cpsp_index'] for r in rows])

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    "Regional variation in acute and neuropathic pain medication prescribing "
    "across Japan's 47 prefectures: a population-level ecological study"
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Tatsuki Onishi, MD')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Department of Anesthesiology, [Institution], [Address], [City], [Postal code], Japan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Corresponding author: ')
r.bold = True
p.add_run(
    'Tatsuki Onishi, MD, Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. E-mail: [email]'
)

doc.add_paragraph()

# Word count line (placeholder; filled after document is built)
p = doc.add_paragraph()
r = p.add_run('Word count: ')
r.bold = True
word_count_run = p.add_run('__WORDCOUNT__')

doc.add_page_break()

# ============================================================
# ABSTRACT (Background, Methods, Results, Conclusion) ≤300 words
# ============================================================
add_heading_text('Abstract', level=1)

abstract_bg = (
    'Clinicians frequently rely on cultural labels when estimating analgesic requirements, '
    'yet whether pain-related prescribing varies within Japan has not been examined '
    'at the population level.'
)

abstract_methods = (
    'This ecological study used Japan\u2019s National Database of Health Insurance Claims '
    '(NDB Open Data, 10th edition; population-complete, April 2023\u2013March 2024) to map '
    'pain-related prescribing across all 47 prefectures. '
    'Phase 1 examined acute perioperative analgesic prescribing per surgery. '
    'Phase 2 examined outpatient neuropathic pain drug prescribing as a population-level proxy '
    'for chronic neuropathic pain-related prescribing, adjusted for comorbidity proxies using '
    'multiple regression. '
    'Standardised claim ratios (SCR) confirmed findings after age-sex standardisation.'
)

abstract_results = (
    f'Crude acute analgesic prescribing varied {acute_fold_ratio:.2f}-fold across prefectures '
    f'(Kruskal\u2013Wallis P{fmt_p(acute_kruskal_p)}). '
    f'Tohoku had a higher unadjusted analgesic-per-surgery index than non-Tohoku prefectures '
    f'(Cohen\u2019s d={acute_d:.2f}); however, after age-sex standardisation the Tohoku inpatient '
    f'analgesic standardised claim ratio (SCR) was lower than the non-Tohoku mean '
    f'({scr_analgesic_tohoku:.1f} vs {scr_analgesic_non_tohoku:.1f}). '
    f'Unadjusted outpatient neuropathic pain prescribing also showed a large Tohoku excess '
    f'(d={unadj_d:.2f}), but diabetes and other comorbidity proxies attenuated it by '
    f'{attenuation:.0f}% (P={reg["adjusted_cpsp_test"]["p_value"]:.2f}). '
    f'SCR confirmed {scr_neuro_ratio:.2f}-fold variation in outpatient neuropathic pain '
    f'prescribing after age-sex standardisation.'
)

abstract_conclusions = (
    'Crude acute analgesic prescribing varied nearly twofold across prefectures, but the '
    'unadjusted Tohoku excess was not evident after age-sex standardisation. Neuropathic pain '
    'prescribing varied even more widely and was partly accounted for by comorbidity proxies. '
    'These findings suggest that cultural labels are not a reliable basis for estimating '
    'population analgesic prescribing, and that individual clinical assessment remains essential.'
)

for label, text in [
    ('Background', abstract_bg),
    ('Methods', abstract_methods),
    ('Results', abstract_results),
    ('Conclusion', abstract_conclusions),
]:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

abstract_total = sum(wc(t) for t in [abstract_bg, abstract_methods, abstract_results,
                                      abstract_conclusions])
print(f'Abstract word count: {abstract_total} (JoPPP limit: 300)')

doc.add_paragraph('Keywords: pain; analgesics; cultural stereotype; regional variation; Japan; NDB Open Data')

doc.add_page_break()

# ============================================================
# BACKGROUND
# ============================================================
add_heading_text('Background', level=1)

intro_parts = []

intro1 = (
    f'When prescribing analgesics, clinicians routinely make assumptions about how much pain '
    f'a patient can tolerate. These assumptions are often anchored to cultural '
    f'labels.{cite(1,2)} Since Zborowski\u2019s landmark observation that ethnic groups '
    f'differ in pain behaviour,{cite(3)} a large body of literature has established that '
    f'cultural norms influence pain reporting, treatment-seeking, and analgesic use.{cite(4)} '
    f'Japan is a case in point: Hobara found that Japanese respondents rated pain behaviours '
    f'as less appropriate than Euro-Americans,{cite(5)} and Feng et al showed that Japanese '
    f'participants were far less willing to trade time to avoid pain on the EQ-5D than '
    f'respondents in the UK and the Netherlands.{cite(6)} '
    f'The resulting cultural label\u2014\u201cJapanese patients are stoic\u201d\u2014carries '
    f'clinical risk whenever a clinician\u2014whether surgeon, internist, or '
    f'anaesthesiologist\u2014assumes that a stoic presentation indicates lower analgesic need.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)
intro_parts.append(intro1)

intro2 = (
    f'Most cross-cultural pain studies compare behaviour between nations, treating each country '
    f'as a homogeneous unit. However, substantial cultural heterogeneity exists within countries. '
    f'Cohen and Nisbett demonstrated that the \u201cculture of honour\u201d in the southern '
    f'United States produces measurably different behavioural responses from the North, including '
    f'physiological stress responses.{cite(7)} '
    f'In Japan, regional cultural identities remain strong despite the discourse of national '
    f'homogeneity. The Tohoku region (northeastern Honshu) is traditionally perceived as '
    f'embodying stoic endurance,{cite(8)} and even the Japanese word for throbbing pain '
    f'(uzuku) shows distinct regional usage patterns that map onto historical dialect '
    f'boundaries.{cite(9)} A 2017 Pfizer Japan survey found that 74.1% of chronic pain '
    f'respondents reported enduring pain, with prefecture-level rates ranging from 68.3% '
    f'(Kanagawa) to 81.6% (Tochigi).{cite(10)} '
    f'Yet whether such differences translate into measurable differences in healthcare '
    f'utilisation at the population level remains unknown.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)
intro_parts.append(intro2)

intro3 = (
    f'Japan\u2019s universal health insurance system, standardised drug pricing, and the '
    f'National Database of Health Insurance Claims (NDB)\u2014capturing virtually all '
    f'reimbursed healthcare utilisation for approximately 125 million insured '
    f'individuals{cite(11)}\u2014provide a unique setting for ecological analysis of '
    f'regional prescribing variation. Recent work has documented up to 1.6-fold regional '
    f'variation in high-impact chronic pain prevalence using a large internet survey{cite(12)} and '
    f'up to 4-fold variation in cancer opioid prescribing using hospital claims data.{cite(13)} '
    f'Taira et al demonstrated substantial regional inequality in dental care utilisation '
    f'using NDB-derived standardised claim ratios across all 47 prefectures.{cite(14)} '
    f'However, no study has applied this population-complete ecological framework '
    f'to pain-related prescribing.'
)
p = doc.add_paragraph()
add_ref_runs(p, intro3)
intro_parts.append(intro3)

intro4 = (
    'The perioperative setting offers a further methodological advantage: because all '
    'patients in Phase 1 are hospitalised for surgery, geographic differences in outpatient '
    'healthcare access are unlikely to be the main driver of the Phase 1 findings, although '
    'surgical case-mix and hospital practice patterns remain unmeasured. '
    'This study had three objectives: (1) map regional variation in acute perioperative '
    'analgesic prescribing across 47 prefectures; (2) examine outpatient neuropathic pain '
    'prescribing as a population-level proxy for chronic neuropathic pain-related '
    'prescribing after adjustment for comorbidity proxies; and (3) integrate acute and '
    'chronic pain findings at the population level.'
)
doc.add_paragraph(intro4)
intro_parts.append(intro4)

intro_total = sum(wc(t) for t in intro_parts)
print(f'Background word count: {intro_total}')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

methods_parts = []

add_heading_text('Study design and reporting', level=2)
m1 = (
    f'This ecological study analysed prefecture-level aggregate data from the NDB Open Data. '
    f'It is reported following the Strengthening the Reporting of Observational Studies in '
    f'Epidemiology (STROBE) statement{cite(15)} and the REporting of studies Conducted using '
    f'Observational Routinely-collected health Data (RECORD) extension.{cite(16)} '
    f'As only publicly available aggregate data were used, ethical approval was not required '
    f'under Japan\u2019s Ethical Guidelines for Medical and Biological Research Involving '
    f'Human Subjects.'
)
p = doc.add_paragraph()
add_ref_runs(p, m1)
methods_parts.append(m1)

add_heading_text('Data source', level=2)
m2 = (
    f'The 10th edition of the NDB Open Data (April 2023\u2013March 2024) was used.{cite(11)} '
    f'The NDB captures claims from all insurers within Japan\u2019s universal coverage system, '
    f'encompassing approximately 125 million insured individuals. '
    f'Aggregate prescription and procedure data are published at the prefecture level (n=47) '
    f'with suppression of cells containing fewer than ten events. '
    f'Prefecture-level population estimates (October 2023, by five-year age group and sex) '
    f'from the Statistics Bureau of Japan were used to compute per-capita rates and '
    f'standardised claim ratios (SCR).'
)
p = doc.add_paragraph()
add_ref_runs(p, m2)
methods_parts.append(m2)

add_heading_text('Regional classification', level=2)
m3 = (
    'Prefectures were grouped into nine standard regional blocks following '
    'the classification used by the Statistics Bureau of Japan: '
    'Hokkaido (1 prefecture), Tohoku (6: Aomori, Iwate, Miyagi, Akita, Yamagata, '
    'Fukushima), Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), '
    'Kinki (6), Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8). '
    'Tohoku was designated as the primary region of interest a priori, based on its '
    'traditional cultural characterisation as embodying patient endurance (gaman).'
)
doc.add_paragraph(m3)
methods_parts.append(m3)

add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
m4 = (
    'Inpatient prescription data were extracted for three analgesic drug classes commonly '
    'used in perioperative pain management: '
    'Class 114 (antipyretic analgesics including NSAIDs and acetaminophen), '
    'Class 811 (opium alkaloid narcotics including morphine and codeine), '
    'and Class 821 (synthetic narcotics including fentanyl and pethidine). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section '
    '(one of the principal procedure categories in the NDB classification of hospital claims) '
    'of the claims data. '
    'The analgesic-per-surgery index was calculated for each prefecture as '
    'total inpatient analgesic prescription units divided by total inpatient '
    'surgical procedure count. This ratio provides a standardised measure of analgesic '
    'intensity that accounts for differences in surgical volume between prefectures. '
    'Separate sub-analyses were conducted for each drug class to assess consistency.'
)
doc.add_paragraph(m4)
methods_parts.append(m4)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing', level=2)
m5 = (
    'Five classes of outpatient oral neuropathic pain medications were extracted: '
    'pregabalin, mirogabalin, duloxetine, tramadol, and neurotropin. '
    'These agents are first-line or commonly used medications for neuropathic pain in Japan '
    'and serve as a population-level proxy for chronic neuropathic pain-related prescribing. '
    'Because individual diagnosis codes are unavailable in the aggregate open data, the proxy '
    'cannot distinguish postsurgical neuropathy from diabetic, postherpetic, or other '
    'neuropathic conditions. '
    'The neuropathic pain prescribing-per-surgery index was calculated as total outpatient '
    'neuropathic pain drug quantity divided by total inpatient surgical procedure count, '
    'using surgery count as a denominator that normalises for healthcare system capacity. '
    'Per-capita neuropathic pain prescribing rates (units per thousand population) were '
    'additionally computed using prefecture population data.'
)
doc.add_paragraph(m5)
methods_parts.append(m5)

add_heading_text('Confounder disease proxies', level=2)
m6 = (
    'Four confounder disease proxies were extracted from outpatient data: '
    'oral hypoglycaemic agents (proxy for diabetic neuropathy, the most '
    'common cause of neuropathic pain in Japan), '
    'herpes zoster antivirals (proxy for postherpetic neuralgia), '
    'antidepressants excluding duloxetine (proxy for depression, '
    'because some antidepressants are co-prescribed with neuropathic pain medications), '
    'and anxiolytics (proxy for anxiety disorders). '
    'Each proxy was expressed per surgery to maintain consistency with the primary outcome. '
    'Outpatient nerve block procedure counts served as an additional independent '
    'neuropathic-pain-related proxy.'
)
doc.add_paragraph(m6)
methods_parts.append(m6)

add_heading_text('Statistical analysis', level=2)
m7a = (
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test '
    'across nine regional blocks, followed by post-hoc Mann\u2013Whitney U tests with Bonferroni '
    'correction for pairwise comparisons. Effect sizes were quantified using Cohen\u2019s d. '
    'For Phase 2, five regression models were fitted to examine the Tohoku regional effect '
    'with progressive confounder adjustment: '
    'Model 1 (unadjusted Tohoku vs non-Tohoku comparison), '
    'Model 2 (neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + Tohoku indicator), '
    'Model 3 (core neuropathic drugs only ~ same confounders), '
    'Model 4 (nerve blocks ~ same confounders), and '
    'Model 5 (neuropathic pain ~ acute analgesic index + confounders).'
)
doc.add_paragraph(m7a)
methods_parts.append(m7a)

m7b = (
    'The adjusted neuropathic pain prescribing index was derived as residuals from regressing '
    'neuropathic pain prescribing on the four confounder proxies. '
    'Standardised claim ratios (SCR) were computed by indirect age-sex standardisation '
    f'following Taira et al.{cite(14)} '
    'National age-sex-specific prescription rates (18 five-year age groups \u00d7 2 sexes) '
    'from the NDB were applied to each prefecture\u2019s population structure. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11).'
)
p = doc.add_paragraph()
add_ref_runs(p, m7b)
methods_parts.append(m7b)

add_heading_text('Patient and public involvement', level=2)
m_ppi = (
    'Patients or members of the public were not involved in the design, conduct, '
    'reporting, or dissemination plans of this research.'
)
doc.add_paragraph(m_ppi)
methods_parts.append(m_ppi)

methods_total = sum(wc(t) for t in methods_parts)
print(f'Methods word count: {methods_total}')

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

results_parts = []

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
r1 = (
    f'During April 2023\u2013March 2024, the NDB recorded {total_surgeries:,} inpatient surgical procedures '
    f'and {total_analgesic_units:,} analgesic prescription units across 47 prefectures. '
    f'The national mean analgesic-per-surgery index was {acute_mean:.2f} (SD {acute_sd:.2f}), '
    f'ranging from {acute_min_val:.2f} ({acute_min_pref}) to {acute_max_val:.2f} ({acute_max_pref})\u2014a '
    f'{acute_fold_ratio:.2f}-fold difference (Kruskal\u2013Wallis P{fmt_p(acute_kruskal_p)} across nine regions; Table 1).'
)
doc.add_paragraph(r1)
results_parts.append(r1)

# === TABLE 1 INLINE ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_cap = p_cap.add_run('Table 1. ')
r_cap.bold = True
r_cap.font.size = Pt(10)
p_cap.add_run(
    'Regional summary of inpatient analgesic prescribing per surgery across nine regional blocks. '
    'Values are unadjusted analgesic-per-surgery index (mean \u00b1 SD) and do not account for '
    'age-sex or case-mix differences. '
    f'Kruskal\u2013Wallis P{fmt_p(acute_kruskal_p)}.'
).font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t1.rows[0].cells
for i, h in enumerate(['Region', 'n (prefectures)', 'Mean \u00b1 SD', 'Range']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = REGION_EN[reg_name]
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f} \u00b1 {np.std(vals, ddof=1):.2f}' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t1)
doc.add_paragraph()

r2 = (
    f'Substantial regional clustering was observed. Tokai and Kinki (western Japan) had the lowest '
    f'indices, while Kyushu-Okinawa and Hokkaido had the highest. '
    f'Tohoku, culturally perceived as Japan\u2019s most stoic region, '
    f'ranked {tohoku_rank_text} of the nine regions with a mean index of {tohoku_acute_mean:.2f} (SD {tohoku_acute_sd:.2f}), '
    f'significantly above the non-Tohoku mean of {non_tohoku_acute_mean:.2f} '
    f'(Mann\u2013Whitney U={acute_mw_u:.0f}, P{fmt_p(acute_mw_p)}; Cohen\u2019s d={acute_d:.2f}). '
    f'All six Tohoku prefectures ranked in the upper half nationally.'
)
doc.add_paragraph(r2)
results_parts.append(r2)

class_text_parts = []
for cr in class_rows:
    class_text_parts.append(
        f"{cr['class_name_en']} "
        f"({cr['tohoku_mean']:.2f} vs {cr['non_tohoku_mean']:.2f}, "
        f"P{fmt_p(cr['p_value'])})"
    )
r2b = (
    f'Sub-analyses by drug class showed that the Tohoku excess was consistent across all three '
    f'acute analgesic categories: {class_text_parts[0]}, {class_text_parts[1]}, and {class_text_parts[2]}. '
    f'The largest absolute class difference was observed among opium alkaloid narcotics.'
)
doc.add_paragraph(r2b)
results_parts.append(r2b)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
r3 = (
    f'Nationally, outpatient neuropathic pain drug prescriptions totalled {int(neuro_total_out):,} units, '
    f'comprising pregabalin ({drug_pcts["pregabalin"]:.1f}%), neurotropin ({drug_pcts["neurotropin"]:.1f}%), '
    f'mirogabalin ({drug_pcts["mirogabalin"]:.1f}%), duloxetine ({drug_pcts["duloxetine"]:.1f}%), '
    f'and tramadol ({drug_pcts["tramadol"]:.1f}%). '
    f'Tohoku had a markedly higher neuropathic pain prescribing-per-surgery index '
    f'({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; P{fmt_p(reg["model1_unadjusted"]["p_value"])}; '
    f'd={reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    f'with {top3_neuro_info[0][0]} ({top3_neuro_info[0][1]:.1f}), {top3_neuro_info[1][0]} ({top3_neuro_info[1][1]:.1f}), and {top3_neuro_info[2][0]} ({top3_neuro_info[2][1]:.1f}) '
    f'occupying the top three nationally (Figure 1).'
)
doc.add_paragraph(r3)
results_parts.append(r3)

# === FIGURE 1 placeholder ===
add_inline_figure(
    'Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
    'Tohoku prefectures (red) cluster at the high end. Dashed line = national mean.',
    1
)

add_heading_text('Confounder analysis and adjustment', level=2)
r4a = (
    f'Neuropathic pain prescribing showed strong correlations with confounder disease proxies. '
    f'Diabetes drug prescribing was the strongest correlate (r={diabetes_r:.2f}, P{fmt_p(diabetes_p)}), '
    f'followed by anxiolytics (r={anxiolytic_r:.2f}, P{fmt_p(anxiolytic_p)}), antidepressants (r={antidep_r:.2f}, P{fmt_p(antidep_p)}), '
    f'and herpes antivirals (r={herpes_r:.2f}, P{fmt_p(herpes_p)}). '
    f'These four confounders collectively accounted for '
    f'{reg["model2_adjusted"]["R2"]*100:.1f}% of between-prefecture variance in neuropathic pain '
    f'prescribing (R\u00b2={reg["model2_adjusted"]["R2"]:.3f} in Model 2; Figure 2).'
)
doc.add_paragraph(r4a)
results_parts.append(r4a)

# === FIGURE 2 placeholder ===
add_inline_figure(
    'Correlation between neuropathic pain prescribing and confounder disease proxies. '
    'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
    f'Diabetes drugs show the strongest correlation (r={diabetes_r:.2f}).',
    2
)

r4b = (
    f'After adjustment for all four confounders, the Tohoku effect was attenuated '
    f'and became nonsignificant in Model 2 '
    f'(\u03b2={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
    f'P={reg["model2_adjusted"]["tohoku_p"]:.2f}). '
    f'This was consistent across specifications: '
    f'Model 3 (core neuropathic drugs only; '
    f'\u03b2={reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P={reg["model3_core_neuropathic"]["tohoku_p"]:.2f}), '
    f'Model 4 (nerve blocks; P={reg["model4_nerve_blocks"]["tohoku_p"]:.2f}), and '
    f'Model 5 (integrated; \u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.2f}; Table 2).'
)
doc.add_paragraph(r4b)
results_parts.append(r4b)

# === TABLE 2 INLINE ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_cap2 = p_cap2.add_run('Table 2. ')
r_cap2.bold = True
r_cap2.font.size = Pt(10)
p_cap2.add_run(
    'Regression models for outpatient neuropathic pain prescribing with Tohoku '
    'indicator and confounder adjustment.'
).font.size = Pt(10)

models = [
    ('Model 1 (unadjusted)', '\u2014', '\u2014',
     f'd={reg["model1_unadjusted"]["cohens_d"]:.2f}',
     f'P{fmt_p(reg["model1_unadjusted"]["p_value"])}'),
    ('Model 2 (all confounders)',
     f'{reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}', ''),
    ('Model 3 (core neuro)',
     f'{reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}', ''),
    ('Model 4 (nerve blocks)',
     f'{reg["model4_nerve_blocks"]["tohoku_coef"]:.1f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}', ''),
    ('Model 5 (integrated)',
     f'{reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.3f}',
     f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}', ''),
]

t2 = doc.add_table(rows=1 + len(models), cols=5, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Tohoku \u03b2', 'Tohoku P', 'Fit', 'Note']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, (name, beta, pval, fit, note) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = beta
    row[2].text = pval
    row[3].text = fit
    row[4].text = note
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t2)
doc.add_paragraph()

r5a = (
    f'The adjusted neuropathic pain prescribing index showed a markedly different geographic pattern '
    f'from the unadjusted data. '
    f'The Tohoku mean shifted from markedly positive to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs '
    f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f'P={reg["adjusted_cpsp_test"]["p_value"]:.2f}; '
    f'd={reg["adjusted_cpsp_test"]["cohens_d"]:.2f}; Figure 3). '
    f'Chugoku emerged as having the highest adjusted index, while Tokai had the lowest.'
)
doc.add_paragraph(r5a)
results_parts.append(r5a)

# === FIGURE 3 placeholder ===
add_inline_figure(
    'Regional comparison of neuropathic pain prescribing: (a) unadjusted and '
    '(b) after confounder adjustment. Tohoku (red) shifts from the highest region to '
    'mid-range after adjustment. Error bars = SD.',
    3
)

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
r6 = (
    f'Acute perioperative prescribing correlated positively with unadjusted neuropathic '
    f'pain prescribing (r={acute_neuro_r:.2f}, P{fmt_p(acute_neuro_p)}). After confounder adjustment, this '
    f'correlation was attenuated (r={acute_adj_r:.2f}, P{fmt_p(acute_adj_p)}). '
    f'In Model 5, the acute pain index remained independently associated with chronic neuropathic '
    f'pain prescribing (\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P={reg["model5_integrated"]["tohoku_p"]:.2f}). '
    f'After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}%.'
)
doc.add_paragraph(r6)
results_parts.append(r6)

add_heading_text('Age-sex standardised claim ratios', level=2)
r7 = (
    f'After indirect age-sex standardisation, inpatient analgesic SCR ranged from '
    f'{scr_analgesic_range[0]:.1f} to {scr_analgesic_range[1]:.1f} '
    f'({scr_analgesic_ratio:.2f}-fold variation), confirming that the prescribing heterogeneity '
    f'was not attributable to differences in prefectural age-sex composition. '
    f'The geographic pattern differed from the crude index: Tohoku\u2019s inpatient analgesic SCR '
    f'({scr_analgesic_tohoku:.1f}) was below the non-Tohoku mean ({scr_analgesic_non_tohoku:.1f}), '
    f'whereas Hokkaido and Kyushu-Okinawa were the highest. '
    f'Outpatient neuropathic pain drug SCR ranged from {scr_neuro_range[0]:.1f} to '
    f'{scr_neuro_range[1]:.1f} ({scr_neuro_ratio:.2f}-fold), and Tohoku remained elevated '
    f'(mean {scr_neuro_tohoku:.1f} vs {scr_neuro_non_tohoku:.1f} for non-Tohoku), consistent '
    f'with the crude analysis.'
)
doc.add_paragraph(r7)
results_parts.append(r7)

results_total = sum(wc(t) for t in results_parts)
print(f'Results word count: {results_total}')

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

disc_parts = []

d1 = (
    'To our knowledge, this is the first study to map perioperative and chronic pain-related '
    'prescribing across all 47 prefectures of Japan, using freely available NDB Open Data '
    'that capture virtually all reimbursed healthcare use. Taken together, the findings do not '
    'support the idea that a single cultural label reliably indicates analgesic prescribing. '
    'Three principal findings carry implications for clinicians and pharmaceutical policy.'
)
doc.add_paragraph(d1)
disc_parts.append(d1)

add_heading_text('Within-country heterogeneity challenges cultural stereotypes', level=2)
d2 = (
    f'Despite Japan\u2019s well-documented cultural stoicism,{cite(5,6)} '
    f'we found {acute_fold_ratio:.2f}-fold variation in acute perioperative analgesic prescribing. '
    f'This parallels Cohen and Nisbett\u2019s finding that the US \u201cculture of '
    f'honour\u201d produces regional behavioural differences within a single '
    f'nation.{cite(7)} Japan\u2019s pain culture is not monolithic; regional '
    f'demographics, healthcare infrastructure, and clinical practices generate '
    f'heterogeneity beyond what a single national label can capture. '
    f'The crude Tohoku excess in perioperative analgesic prescribing was attenuated after '
    f'age-sex standardisation, indicating that demographic structure contributes to regional '
    f'differences. Thus the data do not warrant a simple cultural interpretation.'
)
p = doc.add_paragraph()
add_ref_runs(p, d2)
disc_parts.append(d2)

add_heading_text('Implications for clinical practice', level=2)
d3 = (
    f'The {acute_fold_ratio:.2f}-fold within-Japan variation is relevant for any physician '
    f'who writes an analgesic prescription\u2014not only pain specialists but also surgeons, '
    f'internists, and general practitioners. Evidence suggests that '
    f'cultural stereotypes influence clinician pain assessment and prescribing behaviour. '
    f'Anderson et al showed that racial and ethnic minorities in the United States consistently '
    f'receive less adequate pain management across acute, chronic, cancer, and palliative '
    f'settings.{cite(17)} Campbell and Edwards identified that clinician expectations about '
    f'a patient\u2019s cultural pain behaviour can lead to systematic under- or '
    f'over-treatment.{cite(18)} Rogger et al emphasised that cultural framing affects not '
    f'only patient reporting but also how clinicians interpret and respond to pain cues.{cite(2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)
disc_parts.append(d3)

d3b = (
    f'The broader lesson is that within-country heterogeneity in pain prescribing '
    f'challenges the use of national cultural stereotypes in clinical practice\u2014'
    f'not only in Japan but in any country where cultural generalisations guide prescribing '
    f'decisions. The nihonjinron discourse (theories of Japanese '
    f'uniqueness) has long promoted the notion that Japanese people constitute a uniform '
    f'population sharing a single set of behavioural norms.{cite(19)} Yet this '
    f'\u201chegemony of homogeneity,\u201d as Befu termed it, is an ideological construct '
    f'rather than an empirical fact. Burgess showed that the \u201cillusion\u201d of '
    f'homogeneous Japan has tangible consequences for social policy and public '
    f'perception.{cite(20)} In the clinical context, the combination of two '
    f'stereotypes\u2014\u201cJapanese are stoic\u201d and \u201cJapanese are '
    f'homogeneous\u201d\u2014creates a doubly misleading assumption: that all Japanese '
    f'patients will tolerate pain equally and require less analgesia. '
    f'Our finding of {acute_fold_ratio:.2f}-fold within-Japan variation does not support this assumption. '
    f'If this degree of heterogeneity exists within a society widely regarded as culturally '
    f'uniform, broad cultural labels are likely to be equally unreliable in other national '
    f'pain contexts.'
)
p = doc.add_paragraph()
add_ref_runs(p, d3b)
disc_parts.append(d3b)

d3c = (
    f'Pain is an inherently subjective experience. The revised definition from the '
    f'International Association for the Study of Pain (IASP) describes it as '
    f'\u201can unpleasant sensory and emotional experience associated with, or resembling '
    f'that associated with, actual or potential tissue damage\u201d;{cite(21)} a cultural label '
    f'cannot substitute for individual assessment. Put simply, these ecological data do not '
    f'support using nationality or region as a basis for inferring analgesic requirement; '
    f'prescribing decisions should be individualised.'
)
p = doc.add_paragraph()
add_ref_runs(p, d3c)
disc_parts.append(d3c)

add_heading_text('Policy implications', level=2)
d3d = (
    'These findings also have implications for pharmaceutical policy. Japan\u2019s uniform '
    'national fee schedule and reimbursement rules are intended to reduce unwarranted variation, '
    'yet we observed nearly twofold variation in acute analgesic intensity and larger variation '
    'in neuropathic pain prescribing. Distinguishing warranted variation (driven by clinical need '
    'and comorbidity) from unwarranted variation (driven by prescriber habit or formulary access) '
    'is essential. Future policy could include routine feedback on within-region prescribing, '
    'adherence to perioperative analgesic protocols, and regional benchmarking to ensure equitable '
    'access to pain management. Japan\u2019s NDB Open Data can support routine post-market surveillance '
    'of analgesic utilisation and help identify regions where apparently high or low prescribing '
    'warrants further investigation before policy interventions are designed.'
)
doc.add_paragraph(d3d)
disc_parts.append(d3d)

add_heading_text('Confounders account for Tohoku\u2019s apparent excess', level=2)
d4 = (
    f'The most important methodological finding is that the large regional variation '
    f'in neuropathic pain prescribing (unadjusted d={unadj_d:.2f} for Tohoku vs rest) '
    f'was largely accounted for by comorbidity proxies. '
    f'Diabetes drug prescribing alone correlated at r={diabetes_r:.2f} with neuropathic pain '
    f'prescribing, reflecting the known high burden of diabetic neuropathy treated with '
    f'gabapentinoids. After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}% '
    f'and became nonsignificant. '
    f'This has important implications for ecological pain research: studies using neuropathic '
    f'pain drug prescribing as a population-level proxy for chronic neuropathic pain-related '
    f'prescribing must account for comorbidity proxies. Without such adjustment, regional '
    f'differences in diabetic neuropathy burden could be misinterpreted as differences in '
    f'postsurgical neuropathic pain. The within-database adjustment demonstrated here\u2014using '
    f'disease-specific drug proxies from the same data source\u2014provides a replicable approach '
    f'for other countries with national claims databases.'
)
doc.add_paragraph(d4)
disc_parts.append(d4)

add_heading_text('A population-level acute\u2013chronic pain continuum', level=2)
d5 = (
    f'The positive correlation between Phase 1 (acute) and Phase 2 (chronic, adjusted) '
    f'indices (r={acute_adj_r:.2f}, P{fmt_p(acute_adj_p)}) suggests an association between regional acute pain '
    f'management intensity and chronic pain-related prescribing recorded in the same period. '
    f'While the cross-sectional design cannot determine temporal order, this ecological '
    f'association is consistent with individual-level evidence that greater acute postoperative '
    f'pain is associated with a higher risk of chronic postsurgical pain (CPSP).{cite(22)} '
    f'The contribution of the acute pain index in Model 5 '
    f'(\u03b2={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}), after adjustment for confounders, '
    f'raises the possibility that acute pain management patterns are statistically linked '
    f'to observed chronic prescribing, but it does not demonstrate a causal pathway.'
)
p = doc.add_paragraph()
add_ref_runs(p, d5)
disc_parts.append(d5)

add_heading_text('Strengths and limitations', level=2)
d6a = (
    'Strengths of this study include the use of population-complete data covering all '
    'insurance-reimbursed healthcare in Japan, the novel integration of acute and chronic '
    'pain proxies within a single analytical framework, the transparent '
    'confounder-adjustment methodology, and the use of a perioperative design that '
    'reduces the role of outpatient healthcare access as a confounder (all Phase 1 patients are inpatients '
    'by definition).'
)
doc.add_paragraph(d6a)
disc_parts.append(d6a)

d6b = (
    'The main limitations are inherent to the ecological design. '
    'The unit of analysis is the prefecture, not the individual patient; ecological '
    'correlations may not reflect individual-level associations (ecological fallacy). '
    'NDB Open Data lack diagnosis codes, so the neuropathic pain drug proxy '
    'captures all indications, not CPSP specifically. '
    'Unmeasured confounders such as surgical case mix, physician density, and '
    'regional prescribing culture may contribute to residual variation.'
)
doc.add_paragraph(d6b)
disc_parts.append(d6b)

add_heading_text('Conclusion', level=1)
d7 = (
    'Cultural labels alone are not a reliable guide to analgesic prescribing. Crude acute perioperative '
    f'prescribing varied {acute_fold_ratio:.2f}-fold across prefectures, but the unadjusted Tohoku excess was '
    f'attenuated after age-sex standardisation and adjustment for comorbidity proxies. '
    f'Outpatient neuropathic pain prescribing showed even larger variation '
    f'({neuro_fold_ratio:.2f}-fold unadjusted, {scr_neuro_ratio:.2f}-fold after age-sex standardisation), '
    f'much of which was accounted for by confounding disease burden. Clinicians and policymakers '
    f'should avoid assuming that regional or national identity is a reliable indicator of analgesic '
    f'requirement; individualised pain assessment remains essential.'
)
doc.add_paragraph(d7)
disc_parts.append(d7)

disc_total = sum(wc(t) for t in disc_parts)
print(f'Discussion word count: {disc_total}')

# ============================================================
# CONTRIBUTORS
# ============================================================
doc.add_paragraph()
add_heading_text('Contributors', level=1)
doc.add_paragraph(
    'TO conceived and designed the study, obtained and verified the underlying data, '
    'performed all analyses, created visualisations, and wrote the manuscript. '
    'TO had full access to all data and is the guarantor.')

# ============================================================
# FUNDING
# ============================================================
doc.add_paragraph()
add_heading_text('Funding', level=1)
doc.add_paragraph('None.')

# ============================================================
# COMPETING INTERESTS
# ============================================================
doc.add_paragraph()
add_heading_text('Disclosure statement', level=1)
doc.add_paragraph('The author reports no conflicts of interest.')

# ============================================================
# ETHICS APPROVAL
# ============================================================
doc.add_paragraph()
add_heading_text('Ethics approval', level=1)
doc.add_paragraph(
    'This study used only publicly available aggregate data from the NDB Open Data. '
    'Ethical approval was not required under Japan\u2019s Ethical Guidelines for '
    'Medical and Biological Research Involving Human Subjects.')

# ============================================================
# DATA AVAILABILITY STATEMENT
# ============================================================
doc.add_paragraph()
add_heading_text('Data availability statement', level=1)
doc.add_paragraph(
    'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    'Labour and Welfare website '
    '(https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    'Analysis code is available at '
    'https://github.com/bougtoir/ndb-pain-regional-variation-japan.')

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgements', level=1)
doc.add_paragraph(
    'The author thanks the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. '
    'Parts of data processing and manuscript preparation were assisted by generative AI '
    '(Claude, Anthropic). The author takes full responsibility for the accuracy and content '
    'of the manuscript.')

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)

for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i} ')
    run_num.bold = True
    p.add_run(ref_text)

# SAVE
# ============================================================
outpath = os.path.join(JOPP_DIR, 'JoPPP_manuscript_EN.docx')

body_total = intro_total + methods_total + results_total + disc_total

# Total manuscript word count (including references, tables and figure captions)
def count_doc_words(doc):
    total = 0
    word_re = re.compile(r'\S+')
    for p in doc.paragraphs:
        total += len(word_re.findall(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(word_re.findall(cell.text))
    return total

# Compute word count before OMML conversion so that inline math text is counted.
# The placeholder is replaced with the final count line so the title-page field is not blank.
total_wc = count_doc_words(doc)
body_only = total_wc - 1  # exclude the __WORDCOUNT__ placeholder line
replacement_text = f'{body_only} (including references, tables and figure captions; JoPPP limit <= 6,000)'
final_wc = body_only + len(re.findall(r'\S+', replacement_text))
word_count_run.text = f'{final_wc} (including references, tables and figure captions; JoPPP limit <= 6,000)'

# Convert inline math (β, R², r, d, P) to Word OMML equations last, just before saving.
embed_equations_in_doc(doc)

doc.save(outpath)
print(f'\nSaved: {outpath}')

print(f'\nBody word count: {body_total}')
print(f'Total word count (incl. references, tables and captions): {final_wc} (JoPPP limit <= 6,000)')
print(f'Abstract word count: {abstract_total} (JoPPP limit: 300)')
print(f'References: {len(ref_list)}')
print(f'Display items: 3 figures + 2 tables = 5')

# Verification
print('\n--- Vancouver reference order verification ---')
for i, ref_text in enumerate(ref_list, 1):
    print(f'  [{i}] {ref_text[:70]}...')
