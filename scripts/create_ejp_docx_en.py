#!/usr/bin/env python3
"""Create English manuscript for European Journal of Pain (EJP) Original Article format.

Key EJP requirements:
- Original Article: ~5,000 words body (excl. abstract/refs/legends)
- Abstract: structured (Background, Methods, Results, Conclusions) + "Significance" statement
- References: author-date (Harvard) style — (Author, Year) in text, alphabetical bibliography
- Tables: uploaded as separate files (not embedded in main text)
- Figures: uploaded as separate files (not embedded in main text)
- Figure/table legends appended at end of manuscript
- STROBE checklist required for observational studies
- Author contributions paragraph
- Title page/checklist (separate file — use EJP template)
- Double-spaced, continuous line numbering
- Editor-in-Chief: Didier Bouhassira
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
EJP_DIR = os.path.join(OUTPUT_DIR, 'ejp')
os.makedirs(EJP_DIR, exist_ok=True)

# Load regression summary
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

# Load integrated results for tables
rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '北海道': 'Hokkaido', '東北': 'Tohoku', '関東': 'Kanto',
    '北陸・甲信越': 'Hokuriku-Koshinetsu', '東海': 'Tokai', '近畿': 'Kinki',
    '中国': 'Chugoku', '四国': 'Shikoku', '九州・沖縄': 'Kyushu-Okinawa',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

# ============================================================
# REFERENCES — alphabetical by first-author surname
# EJP uses author-date (Harvard) style
# Format: Author1, Author2. (Year). Title. Journal Volume, Pages.
# In-text: (Author, Year) or (Author and Author, Year) or (Author et al., Year)
# ============================================================
refs = {
    'Anderson2009': {
        'cite': 'Anderson et al., 2009',
        'bib': 'Anderson, K.O., Green, C.R., Payne, R. (2009). Racial and ethnic disparities in pain: '
               'causes and consequences of unequal care. J Pain 10, 1187\u20131204.',
    },
    'Befu2001': {
        'cite': 'Befu, 2001',
        'bib': 'Befu, H. (2001). Hegemony of Homogeneity: An Anthropological Analysis of Nihonjinron. '
               'Melbourne: Trans Pacific Press.',
    },
    'Benchimol2015': {
        'cite': 'Benchimol et al., 2015',
        'bib': 'Benchimol, E.I., Smeeth, L., Guttmann, A., Harron, K., Moher, D., Petersen, I., '
               'Sorensen, H.T., von Elm, E., Langan, S.M. (2015). '
               'The REporting of studies Conducted using Observational Routinely-collected health Data '
               '(RECORD) statement. PLoS Med 12, e1001885.',
    },
    'Burgess2010': {
        'cite': 'Burgess, 2010',
        'bib': 'Burgess, C. (2010). The \u201cillusion\u201d of homogeneous Japan and national character: '
               'discourse as a tool to transcend the \u201cmyth\u201d vs. \u201creality\u201d binary. '
               'Asia Pac J 8(9), 1\u201322.',
    },
    'Callister2003': {
        'cite': 'Callister, 2003',
        'bib': 'Callister, L.C. (2003). Cultural influences on pain perceptions and behaviors. '
               'Home Health Care Manag Pract 15, 207\u2013211.',
    },
    'Campbell2012': {
        'cite': 'Campbell and Edwards, 2012',
        'bib': 'Campbell, C.M., Edwards, R.R. (2012). Ethnic differences in pain and pain management. '
               'Pain Manag 2, 219\u2013230.',
    },
    'Cohen1996': {
        'cite': 'Cohen et al., 1996',
        'bib': 'Cohen, D., Nisbett, R.E., Bowdle, B.F., Schwarz, N. (1996). '
               'Insult, aggression, and the southern culture of honor: an \u201cexperimental ethnography.\u201d '
               'J Pers Soc Psychol 70, 945\u2013960.',
    },
    'Feng2017': {
        'cite': 'Feng et al., 2017',
        'bib': 'Feng, Y., Herdman, M., van Nooten, F., Cleeland, C., Parkin, D., Ikeda, S., Igarashi, A., Devlin, N.J. (2017). '
               'An exploration of differences between Japan and two European countries in the self-reporting '
               'and valuation of pain and discomfort on the EQ-5D. Qual Life Res 26, 2067\u20132078.',
    },
    'Hobara2005': {
        'cite': 'Hobara, 2005',
        'bib': 'Hobara, M. (2005). Beliefs about appropriate pain behavior: cross-cultural and sex differences '
               'between Japanese and Euro-Americans. Eur J Pain 9, 389\u2013393.',
    },
    'Kehlet2006': {
        'cite': 'Kehlet et al., 2006',
        'bib': 'Kehlet, H., Jensen, T.S., Woolf, C.J. (2006). Persistent postsurgical pain: risk factors and prevention. '
               'Lancet 367, 1618\u20131625.',
    },
    'Kumagai2020': {
        'cite': 'Kumagai, 2020',
        'bib': 'Kumagai, S. (2020). Media representations reproducing images of Tohoku: '
               'the Tohoku reconstruction corner in \u201cSecret Kenmin SHOW.\u201d '
               'Kotoba 41, 21\u201338. [in Japanese]',
    },
    'Matsuoka2025': {
        'cite': 'Matsuoka et al., 2025',
        'bib': 'Matsuoka, Y., Morishima, T., Sato, A., Ogawa, T., Miyashiro, I. (2025). '
               'Population-based claims study of regional and hospital function differences '
               'in opioid prescribing for cancer patients who died in hospital in Japan. '
               'Jpn J Clin Oncol 55, hyaf149.',
    },
    'MHLW2024': {
        'cite': 'Ministry of Health, Labour and Welfare, 2024',
        'bib': 'Ministry of Health, Labour and Welfare. (2024). NDB Open Data, 10th edition. '
               'Available: https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
               'Accessed January 15, 2025.',
    },
    'Okolo2024': {
        'cite': 'Okolo et al., 2024',
        'bib': 'Okolo, C.A., Olorunsogo, T., Babawarun, O. (2024). Cultural variability in pain perception: '
               'a review of cross-cultural studies. Int J Sci Res Arch 11, 2550\u20132556.',
    },
    'Onishi2024': {
        'cite': 'Onishi and Onishi, 2024',
        'bib': 'Onishi, T., Onishi, Y. (2024). Normalized pulse volume as a superior predictor of respiration recovery '
               'and quantification of nociception anti-nociception balance compared to opioid effect site concentration: '
               'a prospective, observational study. F1000Research 13, 233.',
    },
    'Pfizer2017': {
        'cite': 'Pfizer Japan, 2017',
        'bib': 'Pfizer Japan Inc. (2017). 47-prefecture survey on chronic pain. '
               'Available: https://www.pfizer.co.jp/pfizer/company/press/2017. '
               'Accessed February 1, 2025.',
    },
    'Raja2020': {
        'cite': 'Raja et al., 2020',
        'bib': 'Raja, S.N., Carr, D.B., Cohen, M., Finnerup, N.B., Flor, H., Gibson, S., Keefe, F.J., Mogil, J.S., '
               'Ringkamp, M., Sluka, K.A., Song, X.J., Stevens, B., Sullivan, M.D., Tutelman, P.R., Ushida, T., Vader, K. (2020). '
               'The revised International Association for the Study of Pain definition of pain: '
               'concepts, challenges, and compromises. Pain 161, 1976\u20131982.',
    },
    'Rogger2023': {
        'cite': 'Rogger et al., 2023',
        'bib': 'Rogger, R., Bello, C., Romero, C.S., Urman, R.D., Luedi, M.M., Filipovic, M.G. (2023). '
               'Cultural framing and the impact on acute pain and pain services. '
               'Curr Pain Headache Rep 27, 429\u2013436.',
    },
    'Takeda2016': {
        'cite': 'Takeda and Yarimizu, 2016',
        'bib': 'Takeda, K., Yarimizu, K. (2016). Regional differences in the pain expression uzuku. '
               'NINJAL Research Papers 10, 85\u2013107. [in Japanese]',
    },
    'vonElm2007': {
        'cite': 'von Elm et al., 2007',
        'bib': 'von Elm, E., Altman, D.G., Egger, M., Pocock, S.J., Gotzsche, P.C., Vandenbroucke, J.P. (2007). '
               'The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: '
               'guidelines for reporting observational studies. '
               'Lancet 370, 1453\u20131457.',
    },
    'Wakaizumi2024': {
        'cite': 'Wakaizumi et al., 2024',
        'bib': 'Wakaizumi, K., Tanaka, C., Shinohara, Y., Wu, Y., Takaoka, S., Kawate, M., Oka, H., Matsudaira, K. (2024). '
               'Geographical variation in high-impact chronic pain and psychological associations at the regional level: '
               'a multilevel analysis of a large-scale internet-based cross-sectional survey. '
               'Front Public Health 12, 1482177.',
    },
    'Zborowski1969': {
        'cite': 'Zborowski, 1969',
        'bib': 'Zborowski, M. (1969). People in Pain. San Francisco: Jossey-Bass.',
    },
}

def cite(*keys):
    """Return author-date citation string, e.g. '(Author et al., 2009; Author, 2001)'."""
    parts = [refs[k]['cite'] for k in keys]
    return '(' + '; '.join(parts) + ')'


# ============================================================
# Document creation
# ============================================================
doc = Document()

# --- Page setup: A4, margins, double-spaced ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# --- Page numbers (footer, centered) ---
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def add_bold_paragraph(bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)





# Word-count helper
def wc(text):
    return len(text.split())


# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    'Regional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures '
    'Challenges a Stoic Monolithic Patient Stereotype: an ecological study'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# Running title
add_bold_paragraph('Running title: ', 'Pain prescribing variation in Japan')

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Tatsuki Onishi')
run.font.size = Pt(12)

doc.add_paragraph()

add_bold_paragraph('Corresponding author: ',
    'Tatsuki Onishi, Department of Anesthesiology, [Institution], '
    '[Address], [City], [Postal code], Japan. '
    'E-mail: [email]; Tel: [phone]; Fax: [fax]')

doc.add_paragraph()

add_bold_paragraph('Keywords: ',
    'cultural pain behavior; ecological study; gaman; neuropathic pain; regional variation; prescribing patterns')

doc.add_paragraph()

add_bold_paragraph('Number of text pages: ', '[to be determined after formatting]')
add_bold_paragraph('Number of figures: ', '5 (+ 1 supplementary)')
add_bold_paragraph('Number of tables: ', '3')

doc.add_paragraph()

# Category
add_bold_paragraph('Category: ', 'Original Article')

doc.add_page_break()

# ============================================================
# ABSTRACT (structured: Background, Methods, Results, Conclusions + Significance)
# ============================================================
add_heading_text('Abstract', level=1)

abstract_bg = (
    'Cross-cultural studies characterize Japanese people as stoic toward pain, '
    'yet whether this endurance varies regionally within Japan remains unexplored.'
)
abstract_methods = (
    'This ecological study used Japan\u2019s National Database (NDB) Open Data '
    '(population-complete insurance claims, April 2023\u2013March 2024) to map pain-related prescribing '
    'across 47 prefectures and nine regional blocks. Phase 1 examined acute perioperative analgesic prescribing; '
    'Phase 2 examined outpatient neuropathic pain prescribing as a chronic postsurgical pain (CPSP) proxy, '
    'adjusted for confounding diseases.'
)
abstract_results = (
    f'In Phase 1, the inpatient analgesic-per-surgery index '
    f'varied 1.97-fold (Gifu 25.20 to Kagoshima 49.75; Kruskal\u2013Wallis P < 0.001). '
    f'Contrary to popular perception, Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014'
    f'prescribed more, not fewer, analgesics (mean 39.97 vs national 35.78; Cohen\u2019s d = 0.87). '
    f'In Phase 2, Tohoku had the highest unadjusted neuropathic pain index '
    f'(d = {reg["model1_unadjusted"]["cohens_d"]:.2f}), but this excess became nonsignificant '
    f'after adjustment for confounding diseases '
    f'(P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}), '
    f'with diabetes drug prescribing as the strongest confounder (r = 0.87).'
)
abstract_conclusions = (
    'Nearly twofold within-country variation challenges the assumption '
    'that \u201cJapanese\u201d constitutes a uniform category for pain behavior. '
    'Individualized pain assessment should replace culturally stereotyped assumptions '
    'in perioperative care.'
)
abstract_significance = (
    'This is the first nationwide study to quantify within-country variation in pain-related '
    'prescribing across all 47 Japanese prefectures, using claims data covering the entire population. '
    'By demonstrating nearly twofold variation within a single nation commonly characterized as '
    'uniformly stoic, this work provides empirical evidence that cultural generalizations about '
    'pain tolerance are unreliable at the population level. '
    'These findings urge clinicians worldwide to abandon ethnic stereotyping '
    'when assessing and treating pain in Japanese and other East Asian patients.'
)

p = doc.add_paragraph()
r = p.add_run('Background: ')
r.bold = True
p.add_run(abstract_bg)

p = doc.add_paragraph()
r = p.add_run('Methods: ')
r.bold = True
p.add_run(abstract_methods)

p = doc.add_paragraph()
r = p.add_run('Results: ')
r.bold = True
p.add_run(abstract_results)

p = doc.add_paragraph()
r = p.add_run('Conclusions: ')
r.bold = True
p.add_run(abstract_conclusions)

p = doc.add_paragraph()
r = p.add_run('Significance: ')
r.bold = True
p.add_run(abstract_significance)

abstract_total = wc(abstract_bg) + wc(abstract_methods) + wc(abstract_results) + wc(abstract_conclusions) + wc(abstract_significance)
print(f'Abstract word count: {abstract_total}')

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('Introduction', level=1)

intro_parts = []

intro1 = (
    f'Pain is a universal experience, yet its expression and management are profoundly shaped '
    f'by culture {cite("Callister2003","Rogger2023")}. Since Zborowski\u2019s observation that ethnic groups '
    f'differ in pain behavior {cite("Zborowski1969")}, a large literature has established that cultural norms '
    f'influence pain reporting, treatment-seeking, and analgesic use {cite("Okolo2024")}. '
    f'Japan is consistently characterized as stoic: Hobara found that Japanese respondents rated pain '
    f'behaviors as less appropriate than Euro-Americans {cite("Hobara2005")}, and Feng et al. showed that '
    f'Japanese participants were far less willing to trade time to avoid pain on the EQ-5D {cite("Feng2017")}. '
    f'This culturally mediated endurance, encapsulated by the concept of gaman (\u6211\u6162), '
    f'carries clinical risk when clinicians assume stoic presentation indicates lower analgesic need.'
)
doc.add_paragraph(intro1)
intro_parts.append(intro1)

intro2 = (
    f'Most studies compare pain behavior between nations. However, substantial heterogeneity exists '
    f'within countries: Cohen and Nisbett showed that the \u201cculture of honor\u201d in the US South '
    f'produces measurably different behavioral responses from the North {cite("Cohen1996")}. '
    f'In Japan, regional identities are strong\u2014the Tohoku region is traditionally perceived as '
    f'stoic {cite("Kumagai2020")}, and even the word for throbbing pain (uzuku) shows regional usage '
    f'patterns {cite("Takeda2016")}. A Pfizer Japan survey found the proportion of chronic pain patients '
    f'\u201cenduring pain\u201d ranged from 48.7% to 81.6% across prefectures {cite("Pfizer2017")}. '
    f'Yet whether such differences translate into measurable healthcare utilization differences remains unknown.'
)
doc.add_paragraph(intro2)
intro_parts.append(intro2)

intro3 = (
    f'Japan\u2019s universal health insurance, standardized drug pricing, and the NDB\u2014capturing '
    f'virtually all reimbursed utilization {cite("MHLW2024")}\u2014provide an ideal setting. '
    f'Recent work has documented up to 1.6-fold regional variation in chronic pain prevalence '
    f'{cite("Wakaizumi2024")} and 4-fold variation in cancer opioid prescribing {cite("Matsuoka2025")}. '
    f'The perioperative setting offers a further advantage: because all patients are hospitalized, '
    f'healthcare access\u2014a major confounder in outpatient studies\u2014is neutralized.'
)
doc.add_paragraph(intro3)
intro_parts.append(intro3)

intro4 = (
    f'This exploratory study had three objectives: (1) map regional variation in acute perioperative '
    f'analgesic prescribing across 47 prefectures; (2) examine outpatient neuropathic pain prescribing '
    f'as a chronic postsurgical pain (CPSP) proxy after confounder adjustment; and (3) integrate '
    f'acute and chronic findings at the population level.'
)
doc.add_paragraph(intro4)
intro_parts.append(intro4)

intro_total = sum(wc(p) for p in intro_parts)
print(f'Introduction word count: {intro_total}')

# ============================================================
# METHODS
# ============================================================
add_heading_text('Methods', level=1)

methods_parts = []

m1 = (
    f'This ecological study analyzed prefecture-level aggregate data from the NDB Open Data. '
    f'It is reported following the Strengthening the Reporting of Observational Studies in Epidemiology '
    f'(STROBE) statement {cite("vonElm2007")} and the REporting of studies Conducted using Observational '
    f'Routinely-collected health Data (RECORD) extension {cite("Benchimol2015")}. '
    f'As only publicly available aggregate data were used, ethical approval was not required.'
)
add_heading_text('Study design and reporting', level=2)
doc.add_paragraph(m1)
methods_parts.append(m1)

m2 = (
    f'The 10th edition of the NDB Open Data (April 2023\u2013March 2024) was used {cite("MHLW2024")}. '
    f'The NDB captures claims from all insurers within Japan\u2019s universal coverage system, '
    f'encompassing approximately 125 million insured individuals. '
    f'Aggregate data are published at the prefecture level with suppression of cells <10 events.'
)
add_heading_text('Data source', level=2)
doc.add_paragraph(m2)
methods_parts.append(m2)

m3 = (
    'Prefectures were grouped into nine standard regional blocks: '
    'Hokkaido (1), Tohoku (6: Aomori, Iwate, Miyagi, Akita, Yamagata, Fukushima), '
    'Kanto (7), Hokuriku-Koshinetsu (6), Tokai (4), Kinki (6), '
    'Chugoku (5), Shikoku (4), and Kyushu-Okinawa (8).'
)
add_heading_text('Regional classification', level=2)
doc.add_paragraph(m3)
methods_parts.append(m3)

m4 = (
    'Inpatient prescription data were extracted for three analgesic drug classes: '
    'Class 114 (antipyretic analgesics/NSAIDs/acetaminophen), '
    'Class 811 (opium alkaloid narcotics), and Class 821 (synthetic narcotics). '
    'Inpatient surgical procedure counts were extracted from the K Surgery section. '
    'The analgesic-per-surgery index was calculated for each prefecture as: '
    'total inpatient analgesic quantity (units)/total inpatient surgical procedure count.'
)
add_heading_text('Phase 1: Acute perioperative analgesic prescribing', level=2)
doc.add_paragraph(m4)
methods_parts.append(m4)

m5 = (
    'Five classes of outpatient oral neuropathic pain medications were extracted: '
    'pregabalin (78 formulations), mirogabalin (8), duloxetine (33), '
    'tramadol (3), and neurotropin (1). '
    'The neuropathic pain prescribing-per-surgery index was calculated as: '
    'total outpatient neuropathic pain drug quantity/total inpatient surgical procedure count.'
)
add_heading_text('Phase 2: Outpatient neuropathic pain prescribing as CPSP proxy', level=2)
doc.add_paragraph(m5)
methods_parts.append(m5)

m6a = 'Four confounder disease proxies were extracted from outpatient data:'
add_heading_text('Confounder disease proxies', level=2)
doc.add_paragraph(m6a)
methods_parts.append(m6a)
doc.add_paragraph('Oral hypoglycemic agents (261 formulations): proxy for diabetic neuropathy',
                  style='List Bullet')
doc.add_paragraph('Herpes zoster antivirals (47 formulations): proxy for postherpetic neuralgia',
                  style='List Bullet')
doc.add_paragraph('Antidepressants excluding duloxetine (128 formulations): proxy for depression',
                  style='List Bullet')
doc.add_paragraph('Anxiolytics (112 formulations): proxy for anxiety disorders',
                  style='List Bullet')
m6b = (
    'Each proxy was expressed per surgery. Outpatient nerve block procedure counts '
    '(73 codes) served as an independent CPSP proxy.'
)
doc.add_paragraph(m6b)
methods_parts.append(m6b)

m7a = (
    'Regional differences in Phase 1 were assessed using the Kruskal\u2013Wallis test, '
    'followed by post hoc Mann\u2013Whitney U tests with Bonferroni correction. '
    'For Phase 2, five regression models were fitted:'
)
add_heading_text('Statistical analysis', level=2)
doc.add_paragraph(m7a)
methods_parts.append(m7a)
doc.add_paragraph('Model 1: Unadjusted regional comparison', style='List Bullet')
doc.add_paragraph('Model 2: Neuropathic pain ~ diabetes + herpes + antidepressants + anxiolytics + region',
                  style='List Bullet')
doc.add_paragraph('Model 3: Core neuropathic drugs (pregabalin + mirogabalin) ~ same confounders',
                  style='List Bullet')
doc.add_paragraph('Model 4: Nerve blocks ~ same confounders', style='List Bullet')
doc.add_paragraph('Model 5: Neuropathic pain ~ acute analgesic index + confounders (integrated)',
                  style='List Bullet')
m7b = (
    'The adjusted CPSP index was derived as residuals from regressing neuropathic pain prescribing '
    'on the four confounder proxies. '
    'Pearson and Spearman correlations assessed inter-index relationships. '
    'All analyses used Python 3.11 (NumPy 1.24, SciPy 1.11, matplotlib 3.8).'
)
doc.add_paragraph(m7b)
methods_parts.append(m7b)

methods_total = sum(wc(p) for p in methods_parts)
print(f'Methods word count: {methods_total}')

# ============================================================
# RESULTS
# ============================================================
add_heading_text('Results', level=1)

results_parts = []

add_heading_text('Phase 1: Regional variation in acute perioperative analgesic prescribing', level=2)
r1 = (
    'During April 2023\u2013March 2024, the NDB recorded 7,903,515 inpatient surgical procedures '
    'and 274,579,851 analgesic prescription units across 47 prefectures. '
    'The national mean analgesic-per-surgery index was 35.78 (SD 5.56), '
    'ranging from 25.20 (Gifu) to 49.75 (Kagoshima)\u2014a 1.97-fold difference '
    '(Kruskal\u2013Wallis P < 0.001 across nine regions; Table 1).'
)
doc.add_paragraph(r1)
results_parts.append(r1)

# === TABLE 1 inline ===
region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# [Table 1 uploaded as separate file]

r2 = (
    'Substantial regional clustering was observed. Tokai and Kinki (western Japan) had the lowest '
    'indices, while Kyushu-Okinawa and Hokkaido had the highest. '
    'Tohoku, culturally perceived as Japan\u2019s most stoic region, '
    'ranked seventh of nine with a mean index of 39.97 (SD 3.53), '
    'significantly above the non-Tohoku mean of 35.17 '
    '(Mann\u2013Whitney U = 190, P = 0.031; Cohen\u2019s d = 0.87). '
    'All six Tohoku prefectures ranked in the upper half nationally. '
    'This pattern was consistent across drug classes: NSAIDs (P = 0.044), '
    'opioid alkaloids (P = 0.003), and synthetic opioids (P = 0.001).'
)
doc.add_paragraph(r2)
results_parts.append(r2)

add_heading_text('Phase 2: Outpatient neuropathic pain prescribing (unadjusted)', level=2)
r3 = (
    f'Nationally, outpatient neuropathic pain drug prescriptions totaled 2,289,549,163 units, '
    f'comprising pregabalin (40.2%), neurotropin (20.1%), mirogabalin (19.6%), '
    f'duloxetine (15.3%), and tramadol (4.9%). '
    f'Tohoku had a markedly higher index ({reg["model1_unadjusted"]["tohoku_mean"]:.1f} vs '
    f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}; '
    f'P < 0.001; Cohen\u2019s d = {reg["model1_unadjusted"]["cohens_d"]:.2f}), '
    f'with Iwate (566.7), Aomori (519.3), and Akita (461.1) occupying the top three nationally (Fig. 1).'
)
doc.add_paragraph(r3)
results_parts.append(r3)

# [Figure 1 uploaded as separate file]

add_heading_text('Confounder analysis', level=2)
r4 = (
    f'Neuropathic pain prescribing showed strong correlations with confounder proxies. '
    f'Diabetes drug prescribing was the strongest (r = 0.87, P < 0.001), '
    f'followed by anxiolytics (r = 0.75), antidepressants (r = 0.46), '
    f'and herpes antivirals (r = 0.19). '
    f'These collectively explained {reg["model2_adjusted"]["R2"]*100:.1f}% of variance '
    f'(R\u00b2 = {reg["model2_adjusted"]["R2"]:.3f} in Model 2; Fig. 2).'
)
doc.add_paragraph(r4)
results_parts.append(r4)

# [Figure 2 uploaded as separate file]

add_heading_text('Confounder-adjusted analysis', level=2)
r5 = (
    f'After adjustment for all four confounders, '
    f'the Tohoku effect was attenuated and became nonsignificant in Model 2 '
    f'(\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}, P = {reg["model2_adjusted"]["tohoku_p"]:.3f}). '
    f'This was consistent across specifications: '
    f'Model 3 (\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model3_core_neuropathic"]["tohoku_p"]:.3f}), '
    f'Model 4 (P = {reg["model4_nerve_blocks"]["tohoku_p"]:.3f}), and '
    f'Model 5 (\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, '
    f'P = {reg["model5_integrated"]["tohoku_p"]:.3f}; Table 2).'
)
doc.add_paragraph(r5)
results_parts.append(r5)

# [Table 2 uploaded as separate file]

r6 = (
    f'The adjusted CPSP index showed a dramatically different pattern from the unadjusted data (Fig. 3). '
    f'The Tohoku mean shifted from markedly positive to a modest, nonsignificant excess '
    f'({reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f} vs {reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}; '
    f't = {reg["adjusted_cpsp_test"]["t_statistic"]:.3f}, P = {reg["adjusted_cpsp_test"]["p_value"]:.3f}; '
    f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.2f}). '
    f'Chugoku emerged as having the highest adjusted index, while Tokai had the lowest (Fig. 4).'
)
doc.add_paragraph(r6)
results_parts.append(r6)

# === FIGURE 3 inline ===
# [Figure 3 uploaded as separate file]
# [Figure 4 uploaded as separate file]

add_heading_text('Phase 1\u2013Phase 2 integration', level=2)
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100
r7 = (
    f'Acute perioperative prescribing correlated positively with unadjusted neuropathic pain prescribing '
    f'(r = 0.38, P = 0.008; Fig. 5a). After confounder adjustment, this was attenuated '
    f'(r = 0.29, P = 0.052; Fig. 5b). '
    f'In Model 5, acute pain index remained a significant predictor '
    f'(\u03b2 = {reg["model5_integrated"]["acute_pain_coef"]:.2f}, P = {reg["model5_integrated"]["acute_pain_p"]:.3f}), '
    f'while the Tohoku effect remained nonsignificant '
    f'(\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}, P = {reg["model5_integrated"]["tohoku_p"]:.3f}). '
    f'A comprehensive Z-score heatmap confirmed the heterogeneous multi-variable pattern (Supplementary Fig. 1). '
    f'After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}% and became nonsignificant (Table 3).'
)
doc.add_paragraph(r7)
results_parts.append(r7)

# [Figure 5 uploaded as separate file]
# [Table 3 uploaded as separate file]

results_total = sum(wc(p) for p in results_parts)
print(f'Results word count: {results_total}')

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('Discussion', level=1)

disc_parts = []

d1 = (
    f'This study is the first to map perioperative and chronic pain-related prescribing across '
    f'all 47 prefectures of Japan, leveraging freely available NDB Open Data. '
    f'Three principal findings emerged from this exploratory analysis.'
)
doc.add_paragraph(d1)
disc_parts.append(d1)

add_heading_text('Regional variation within a stoic culture', level=2)
d2 = (
    f'Despite Japan\u2019s well-documented cultural stoicism toward pain {cite("Hobara2005","Feng2017")}, '
    f'we found a 1.97-fold variation in acute perioperative analgesic prescribing across prefectures '
    f'and significant differences across regional blocks. '
    f'This parallels findings from other domains: Cohen and Nisbett\u2019s work on the \u201cculture of honor\u201d '
    f'in the United States demonstrated that even within a single nation, regional cultural norms '
    f'produce measurably different behavioral outcomes {cite("Cohen1996")}. '
    f'Our finding suggests that Japan\u2019s pain culture is not monolithic; '
    f'regional demographics, healthcare infrastructure, and local clinical practices '
    f'generate heterogeneity beneath the surface of a nationally shared cultural norm.'
)
doc.add_paragraph(d2)
disc_parts.append(d2)

add_heading_text('Clinical implications: the danger of monolithic cultural stereotypes', level=2)
d3 = (
    f'The 1.97-fold within-Japan variation documented here has direct clinical relevance '
    f'beyond the domestic context. A large body of evidence demonstrates that ethnic and cultural '
    f'stereotypes influence clinician pain assessment and analgesic prescribing. '
    f'Anderson et al. showed that racial and ethnic minorities in the United States consistently '
    f'receive less adequate pain management across acute, chronic, cancer, and palliative settings '
    f'{cite("Anderson2009")}. Campbell and Edwards identified that clinician expectations about '
    f'a patient\u2019s cultural pain behavior can lead to systematic under- or over-treatment '
    f'{cite("Campbell2012")}. Rogger et al. emphasized that cultural framing affects not only '
    f'patient reporting but also how clinicians interpret and respond to pain cues {cite("Rogger2023")}.'
)
doc.add_paragraph(d3)
disc_parts.append(d3)

d3b = (
    f'This risk is compounded by the widespread perception of Japan as a culturally and ethnically '
    f'homogeneous society. The nihonjinron discourse (theories of Japanese uniqueness) '
    f'has long promoted the notion that Japanese people constitute a uniform population '
    f'sharing a single set of behavioral norms {cite("Befu2001")}. '
    f'Yet this \u201chegemony of homogeneity,\u201d as Befu termed it, is an ideological construct '
    f'rather than an empirical fact. '
    f'Burgess showed that this \u201cillusion\u201d of a homogeneous Japan '
    f'has tangible consequences for social policy and public perception {cite("Burgess2010")}. '
    f'In the clinical context, the combination of two stereotypes\u2014\u201cJapanese are stoic\u201d '
    f'and \u201cJapanese are homogeneous\u201d\u2014creates a doubly misleading assumption: '
    f'that all Japanese patients will tolerate pain equally and require less analgesia. '
    f'Our finding of 1.97-fold within-Japan variation in analgesic prescribing '
    f'directly refutes this assumption.'
)
doc.add_paragraph(d3b)
disc_parts.append(d3b)

d3c = (
    f'Put simply, there is no such entity as \u201cthe Japanese patient\u201d '
    f'whose pain behavior can be predicted from nationality alone\u2014'
    f'there are only individual patients from 47 diverse prefectures, '
    f'each with distinct demographic profiles, clinical environments, '
    f'and pain-related prescribing cultures. '
    f'For Japanese patients treated outside Japan, clinicians may be operating under '
    f'a dual misapprehension: that Japanese culture is uniformly stoic, '
    f'and that this stoicism translates into lower analgesic need. '
    f'Our data challenge both premises.'
)
doc.add_paragraph(d3c)
disc_parts.append(d3c)

d3d = (
    f'We therefore argue that the monolithic characterization of any national population\u2019s '
    f'pain behavior carries a risk of clinical harm\u2014and this argument extends well beyond Japan. '
    f'Whenever clinicians allow cultural stereotypes to narrow the range of treatment options '
    f'considered for a patient, the patient is placed at a therapeutic disadvantage {cite("Raja2020")}. '
    f'The aspiration behind this study is that fewer patients, regardless of nationality or ethnicity, '
    f'should suffer the disadvantage of having their treatment options constrained by cultural stereotypes. '
    f'Individualized pain assessment, rather than culturally stereotyped assumptions, '
    f'remains the cornerstone of equitable perioperative care.'
)
doc.add_paragraph(d3d)
disc_parts.append(d3d)

add_heading_text('Pain as an individual experience', level=2)
d3e = (
    f'Fundamentally, pain is a personal experience. The revised IASP definition '
    f'describes pain as \u201can unpleasant sensory and emotional experience associated with, '
    f'or resembling that associated with, actual or potential tissue damage\u201d\u2014'
    f'a definition that is inherently subjective {cite("Raja2020")}. '
    f'Our ecological data illustrate this principle at the population level: '
    f'even within a single nation sharing the same language, insurance system, '
    f'and broad cultural heritage, pain-related prescribing varies nearly twofold. '
    f'No cultural label can substitute for direct measurement of a patient\u2019s nociceptive state. '
    f'Objective nociception monitoring, such as normalized pulse volume as a culture-independent '
    f'indicator of the nociception\u2013anti-nociception balance {cite("Onishi2024")}, '
    f'may help standardize perioperative assessment regardless of cultural background.'
)
doc.add_paragraph(d3e)
disc_parts.append(d3e)

add_heading_text('Confounders explain Tohoku\u2019s apparent excess', level=2)
d5 = (
    f'The most methodologically important finding is that the dramatic regional variation in '
    f'neuropathic pain prescribing '
    f'(unadjusted d = {reg["model1_unadjusted"]["cohens_d"]:.2f} for Tohoku vs rest) was largely explained '
    f'by confounding disease proxies. '
    f'Diabetes drug prescribing alone correlated at r = 0.87 with neuropathic pain prescribing, '
    f'reflecting the known high prevalence of diabetic neuropathy requiring gabapentinoids. '
    f'After adjustment, the Tohoku effect was attenuated by {attenuation:.0f}% and became nonsignificant (Table 3). '
    f'This has important implications for ecological pain research: '
    f'studies using neuropathic pain drug prescribing as a population-level CPSP proxy must account for '
    f'confounding diseases. Without such adjustment, regional differences in diabetes prevalence '
    f'could be misinterpreted as differences in CPSP. '
    f'The within-database confounder adjustment demonstrated here\u2014using disease-specific drug proxies '
    f'from the same data source\u2014provides a replicable framework.'
)
doc.add_paragraph(d5)
disc_parts.append(d5)

add_heading_text('A population-level acute\u2013chronic pain continuum', level=2)
d6 = (
    f'The positive correlation between Phase 1 (acute) and Phase 2 (chronic, adjusted) indices '
    f'(r = 0.29, P = 0.052) suggests a modest link between regional acute pain management intensity '
    f'and subsequent chronic pain-related prescribing. While ecological correlations cannot establish '
    f'causation, this finding is consistent with individual-level evidence that '
    f'the intensity of acute postoperative pain is a risk factor for CPSP {cite("Kehlet2006")}. '
    f'Prefectures with relatively low acute analgesic use but high adjusted CPSP indices '
    f'may warrant investigation for potential under-treatment of acute pain leading to chronification.'
)
doc.add_paragraph(d6)
disc_parts.append(d6)

add_heading_text('Strengths and limitations', level=2)
d7 = (
    f'Strengths of this study include the use of population-complete data covering all insurance-reimbursed healthcare '
    f'in Japan, the novel integration of acute and chronic pain proxies, the transparent confounder-adjustment methodology, '
    f'and the exploratory rather than confirmatory design that allows hypothesis generation. '
    f'A further strength specific to the perioperative focus is that healthcare access does not confound '
    f'the analgesic prescribing data: all patients in Phase 1 are inpatients by definition, '
    f'eliminating the access-to-care heterogeneity that limits community-based pain studies.'
)
doc.add_paragraph(d7)
disc_parts.append(d7)

d7b = (
    f'The main limitations are inherent to the ecological design. '
    f'The unit of analysis is the prefecture, not the individual patient; '
    f'ecological correlations may not reflect individual-level associations (ecological fallacy). '
    f'NDB Open Data do not contain diagnosis codes, so CPSP cannot be directly identified; '
    f'the neuropathic pain drug proxy captures all indications, not CPSP specifically. '
    f'Drug prescribing proxies may not capture disease prevalence accurately '
    f'(e.g., fibromyalgia has no specific drug proxy and shares pregabalin as first-line treatment). '
    f'The cross-sectional design cannot distinguish temporal sequences (surgery \u2192 acute pain \u2192 CPSP), '
    f'and unmeasured confounders such as age distribution, surgical case mix, '
    f'and prescribing culture may contribute to residual regional variation.'
)
doc.add_paragraph(d7b)
disc_parts.append(d7b)

add_heading_text('Implications and future directions', level=2)
d7c = (
    f'This exploratory study demonstrates that Japan\u2019s NDB Open Data can serve as '
    f'a hypothesis-generating platform for population-level pain research. '
    f'Future studies using the NDB sampling dataset (Level 3 access) could enable '
    f'individual-level longitudinal tracking from surgery to new neuropathic pain prescriptions, '
    f'providing a direct CPSP measure. Procedure-specific analyses (e.g., total knee arthroplasty, mastectomy) '
    f'would reduce surgical case-mix confounding. '
    f'From an international perspective, prospective studies examining how cultural '
    f'stereotypes influence analgesic prescribing for Japanese patients in multicultural clinical '
    f'settings would directly test whether the \u201cstoic Japanese\u201d label translates into therapeutic disadvantage. '
    f'Such research is increasingly urgent as global migration and medical tourism expose more '
    f'patients to healthcare systems unfamiliar with their cultural background.'
)
doc.add_paragraph(d7c)
disc_parts.append(d7c)

add_heading_text('Conclusion', level=2)
d8 = (
    f'Despite Japan\u2019s culturally ingrained norm of pain endurance (gaman), '
    f'perioperative and chronic pain-related prescribing varies up to 1.97-fold across prefectures. '
    f'Confounding diseases, particularly diabetes, substantially modify the apparent regional pattern '
    f'of neuropathic pain prescribing. '
    f'These findings demonstrate that Japan\u2019s pain culture is not monolithic. '
    f'Treating \u201cJapanese\u201d as a uniform category for pain behavior risks inadequate analgesia '
    f'for Japanese patients treated abroad\u2014and the same principle applies to every cultural label '
    f'applied to any patient population. '
    f'Individualized pain assessment should replace culturally stereotyped assumptions '
    f'to ensure equitable perioperative care across all clinical settings.'
)
doc.add_paragraph(d8)
disc_parts.append(d8)

disc_total = sum(wc(p) for p in disc_parts)
print(f'Discussion word count (incl. conclusion): {disc_total}')

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_paragraph()
add_heading_text('Acknowledgments', level=1)
doc.add_paragraph(
    'The authors thank the Ministry of Health, Labour and Welfare for making '
    'the NDB Open Data publicly available. '
    'Parts of data processing and manuscript preparation were assisted by generative AI (Claude, Anthropic). '
    'The authors take full responsibility for the accuracy and content of the manuscript.')

doc.add_paragraph()
add_heading_text('Author contributions', level=1)
doc.add_paragraph(
    'Tatsuki Onishi: Conceptualization, Data curation, Formal analysis, Investigation, '
    'Methodology, Software, Visualization, Writing \u2013 original draft, Writing \u2013 review & editing.')

doc.add_paragraph()
add_heading_text('Conflict of interest statement', level=1)
doc.add_paragraph('The authors have no conflicts of interest to declare.')

doc.add_paragraph()
add_heading_text('Funding sources', level=1)
doc.add_paragraph('This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.')

doc.add_paragraph()
add_heading_text('Data availability statement', level=1)
doc.add_paragraph(
    f'The NDB Open Data used in this study are publicly available from the Ministry of Health, '
    f'Labour and Welfare website (https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html). '
    f'Analysis code is available at https://github.com/bougtoir/wip/tree/main/ndb-pain-regional-variation-japan.')

# ============================================================
# REFERENCES (alphabetical by first-author surname, Harvard style)
# ============================================================
doc.add_page_break()
add_heading_text('References', level=1)
sorted_keys = sorted(refs.keys(), key=lambda k: refs[k]['bib'].lower())
for key in sorted_keys:
    doc.add_paragraph(refs[key]['bib'])

# ============================================================
# TABLE LEGENDS
# ============================================================
doc.add_page_break()
add_heading_text('Table legends', level=1)

p = doc.add_paragraph()
r = p.add_run('Table 1. ')
r.bold = True
p.add_run('Phase 1: Regional summary of inpatient analgesic prescribing per surgery across nine '
    'regional blocks. Values represent the analgesic-per-surgery index: total inpatient analgesic '
    'prescription units divided by total inpatient surgical procedure count for each prefecture. '
    'Kruskal\u2013Wallis test across nine regions: P < 0.001.')

p = doc.add_paragraph()
r = p.add_run('Table 2. ')
r.bold = True
p.add_run('Phase 2: Regression models for outpatient neuropathic pain prescribing with Tohoku '
    'indicator and confounder adjustment. *** P < 0.001; ns = not significant. '
    'Models 2\u20135: multiple linear regression with Tohoku indicator (binary) and confounder disease proxies. '
    'Adj CPSP: adjusted CPSP index derived as residuals from regressing neuropathic pain prescribing '
    'on four confounder proxies.')

p = doc.add_paragraph()
r = p.add_run('Table 3. ')
r.bold = True
p.add_run('Effect of confounder adjustment on Tohoku regional indicators. '
    'Confounders: oral hypoglycemic agents (diabetes proxy), herpes zoster antivirals, '
    'antidepressants (excluding duloxetine), and anxiolytics.')

# ============================================================
# FIGURE LEGENDS
# ============================================================
doc.add_page_break()
add_heading_text('Figure legends', level=1)

p = doc.add_paragraph()
r = p.add_run('Figure 1. ')
r.bold = True
p.add_run('Outpatient neuropathic pain drug prescribing per surgery by prefecture (unadjusted). '
    'Bars represent total neuropathic pain prescriptions (pregabalin + mirogabalin + duloxetine + '
    'tramadol + neurotropin) divided by inpatient surgical procedure count. '
    'Tohoku prefectures (red bars with red borders) cluster at the high end. Dashed line = national mean.')

p = doc.add_paragraph()
r = p.add_run('Figure 2. ')
r.bold = True
p.add_run('Correlation between outpatient neuropathic pain drug prescribing and confounder disease proxies. '
    'Each dot represents one prefecture. Tohoku prefectures are marked with red borders. '
    'Diabetes drugs show the strongest correlation (r = 0.87).')

p = doc.add_paragraph()
r = p.add_run('Figure 3. ')
r.bold = True
p.add_run('Confounder-adjusted CPSP index by prefecture. '
    'Residuals from regressing neuropathic pain prescribing on diabetes drugs, herpes antivirals, '
    'antidepressants, and anxiolytics. Positive values indicate higher neuropathic pain prescribing '
    'than expected given confounding disease prevalence. Tohoku prefectures (red borders) are dispersed '
    'across the distribution after adjustment.')

p = doc.add_paragraph()
r = p.add_run('Figure 4. ')
r.bold = True
p.add_run('Regional comparison of neuropathic pain prescribing: (a) unadjusted and (b) after '
    'confounder adjustment. Tohoku (red border) shifts from the highest region to mid-range after '
    'adjustment. Error bars represent standard deviation.')

p = doc.add_paragraph()
r = p.add_run('Figure 5. ')
r.bold = True
p.add_run('Integration of Phase 1 (acute perioperative analgesic prescribing) and Phase 2 '
    '(outpatient neuropathic pain prescribing as CPSP proxy). '
    '(a) Unadjusted: positive correlation (r = 0.38, P = 0.008). '
    '(b) Confounder-adjusted: attenuated correlation (r = 0.29, P = 0.052). '
    'Tohoku prefectures (red borders) cluster in the upper-right quadrant.')

# ============================================================
# SUPPLEMENTARY MATERIAL
# ============================================================
doc.add_page_break()
add_heading_text('Supplementary material', level=1)

p = doc.add_paragraph()
r = p.add_run('Supplementary Figure 1. ')
r.bold = True
p.add_run('Z-score heatmap of all indices by prefecture. '
    'Each row represents a variable; each column represents a prefecture, '
    'sorted by neuropathic pain prescribing. Red = above average; blue = below average. '
    'Tohoku prefectures are marked with red vertical lines. '
    'The adjusted CPSP index shows a different pattern from the raw neuropathic index.')

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(EJP_DIR, 'EJP_manuscript_EN.docx')
doc.save(outpath)
print(f'\nSaved: {outpath}')

# Total body word count (excl. abstract, refs, legends)
body_total = intro_total + methods_total + results_total + disc_total
print(f'\nBody word count: {body_total} (EJP limit: ~5,000)')
print(f'Abstract word count: {abstract_total} (EJP limit: 250)')

# Print reference list for verification
print('\nReference list (alphabetical):')
for key in sorted_keys:
    print(f'  {refs[key]["cite"]}: {refs[key]["bib"][:80]}...')
