#!/usr/bin/env python3
"""Create STROBE checklist for JE submission.

STROBE Statement items for cross-sectional / ecological studies,
plus RECORD extension items for routinely collected data.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JE_DIR = os.path.join(OUTPUT_DIR, 'je')
os.makedirs(JE_DIR, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)

h = doc.add_heading('STROBE Statement\u2014Checklist of Items for Cross-Sectional Studies', level=1)
for run in h.runs:
    run.font.name = 'Arial'

p = doc.add_paragraph()
p.add_run('Combined with RECORD extension items for routinely collected health data').italic = True

# STROBE items for cross-sectional study
items = [
    ('Title and abstract', '1', '(a) Indicate the study\u2019s design with a commonly used term in the title or the abstract',
     'Title: "An Ecological Study"; Abstract: "ecological study"'),
    ('', '', '(b) Provide in the abstract an informative and balanced summary of what was done and what was found',
     'Abstract (Background, Methods, Results, Conclusions)'),

    ('Introduction', '', '', ''),
    ('Background/rationale', '2', 'Explain the scientific background and rationale for the investigation being reported',
     'Introduction, paragraphs 1\u20133'),
    ('Objectives', '3', 'State specific objectives, including any prespecified hypotheses',
     'Introduction, paragraph 5 (4 objectives including demand\u2013supply dissociation)'),

    ('Methods', '', '', ''),
    ('Study design', '4', 'Present key elements of study design early in the paper',
     'Methods, "Study design and reporting"'),
    ('Setting', '5', 'Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection',
     'Methods, "Data source" (April 2023\u2013March 2024)'),
    ('Participants', '6', 'Give the eligibility criteria, and the sources and methods of selection of participants',
     'Methods, "Data source" (NDB captures all insured individuals)'),
    ('Variables', '7', 'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
     'Methods, "Phase 1," "Phase 2," "Confounder disease proxies"'),
    ('Data sources/measurement', '8', 'For each variable of interest, give sources of data and details of methods of assessment (measurement)',
     'Methods, "Phase 1," "Phase 2" (specific drug codes and formulation counts); "Demand-side benchmark: CSLC"'),
    ('Bias', '9', 'Describe any efforts to address potential sources of bias',
     'Methods, "Confounder disease proxies"; Discussion, "Strengths and limitations"'),
    ('Study size', '10', 'Explain how the study size was arrived at',
     'Methods, "Data source" (population-complete; n=47 prefectures)'),
    ('Quantitative variables', '11', 'Explain how quantitative variables were handled in the analyses',
     'Methods, "Statistical analysis" (SCR, regression models)'),
    ('Statistical methods', '12',
     '(a) Describe all statistical methods, including those used to control for confounding',
     'Methods, "Statistical analysis" (Kruskal\u2013Wallis, Mann\u2013Whitney, regression Models 1\u20135, SCR)'),
    ('', '', '(b) Describe any methods used to examine subgroups and interactions',
     'Methods, "Phase 1" (sub-analyses by drug class)'),
    ('', '', '(c) Explain how missing data were addressed',
     'NDB suppresses cells <10 events; no individual-level missing data'),
    ('', '', '(d) If applicable, describe analytical methods taking account of sampling strategy',
     'N/A (population-complete data, not sampled)'),
    ('', '', '(e) Describe any sensitivity analyses',
     'Models 2\u20135 represent progressive confounder adjustment (sensitivity)'),

    ('Results', '', '', ''),
    ('Participants', '13', 'Report numbers of individuals at each stage of study',
     'Results, Phase 1 paragraph (7,903,515 surgeries; 274,579,851 prescription units)'),
    ('Descriptive data', '14',
     '(a) Give characteristics of study participants and information on exposures and potential confounders',
     'Results, Table 1 (regional summary); confounder correlations'),
    ('', '', '(b) Indicate number of participants with missing data for each variable of interest',
     'N/A (aggregate data; cells <10 suppressed by source)'),
    ('Outcome data', '15', 'Report numbers of outcome events or summary measures',
     'Results, Phase 1 and Phase 2 paragraphs'),
    ('Main results', '16',
     '(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision',
     'Results, Table 2 (Models 1\u20135 with \u03b2 coefficients and P values); confounder r values'),
    ('', '', '(b) Report category boundaries when continuous variables were categorized',
     'Nine regional blocks defined in Methods'),
    ('', '', '(c) If relevant, consider translating estimates of relative risk into absolute risk for a meaningful time period',
     'N/A (ecological indices, not individual risk)'),
    ('Other analyses', '17', 'Report other analyses done\u2014e.g., analyses of subgroups and interactions, and sensitivity analyses',
     'Results, "Age-sex standardised claim ratios"; "Demand\u2013supply dissociation"; "Integration of acute and chronic prescribing"'),

    ('Discussion', '', '', ''),
    ('Key results', '18', 'Summarise key results with reference to study objectives',
     'Discussion, "Principal findings" paragraphs 1\u20134 (including CSLC dissociation)'),
    ('Limitations', '19', 'Discuss limitations of the study, taking into account sources of potential bias or imprecision',
     'Discussion, "Strengths and limitations"'),
    ('Interpretation', '20', 'Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence',
     'Discussion, "Clinical implications"'),
    ('Generalisability', '21', 'Discuss the generalisability (external validity) of the study results',
     'Discussion, "Clinical implications" (international parallels)'),

    ('Other information', '', '', ''),
    ('Funding', '22', 'Give the source of funding and the role of the funders for the present study and, if applicable, for the original study on which the present article is based',
     'Funding section'),
]

# Create table
t = doc.add_table(rows=1, cols=4, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(['Section/Topic', 'Item No.', 'STROBE Recommendation', 'Reported on page/section']):
    cell = t.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(8)

# Set column widths
widths = [Cm(3.5), Cm(1.5), Cm(14.0), Cm(7.5)]
for row_cells in [t.rows[0].cells]:
    for i, width in enumerate(widths):
        row_cells[i].width = width

prev_section = ''
for section_topic, item_no, recommendation, location in items:
    row = t.add_row()
    cells = row.cells
    # Section header rows
    if item_no == '' and recommendation == '':
        cells[0].text = section_topic
        for par in cells[0].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(8)
        continue

    cells[0].text = section_topic
    cells[1].text = item_no
    cells[2].text = recommendation
    cells[3].text = location
    for cell in cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(8)

outpath = os.path.join(JE_DIR, 'JE_STROBE_checklist.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
