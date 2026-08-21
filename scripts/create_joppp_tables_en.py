#!/usr/bin/env python3
"""Create separate editable tables DOCX for JoPPP submission (English)."""
import os
import json
import csv
import numpy as np
from scipy import stats
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
JOPP_DIR = os.path.join(OUTPUT_DIR, 'joppp')
os.makedirs(JOPP_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except Exception:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_EN = {
    '\u5317\u6d77\u9053': 'Hokkaido', '\u6771\u5317': 'Tohoku', '\u95a2\u6771': 'Kanto',
    '\u5317\u9678\u30fb\u7532\u4fe1\u8d8a': 'Hokuriku-Koshinetsu', '\u6771\u6d77': 'Tokai', '\u8fd1\u757f': 'Kinki',
    '\u4e2d\u56fd': 'Chugoku', '\u56db\u56fd': 'Shikoku', '\u4e5d\u5dde\u30fb\u6c96\u7e04': 'Kyushu-Okinawa',
}
REGION_ORDER = ['\u5317\u6d77\u9053', '\u6771\u5317', '\u95a2\u6771', '\u5317\u9678\u30fb\u7532\u4fe1\u8d8a', '\u6771\u6d77', '\u8fd1\u757f', '\u4e2d\u56fd', '\u56db\u56fd', '\u4e5d\u5dde\u30fb\u6c96\u7e04']

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.5


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def fmt_p(p):
    if p < 0.001:
        return '< 0.001'
    return f'{p:.3f}'


# ============================================================
# TABLE 1
# ============================================================
region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

region_vals_list = [region_data[r] for r in REGION_ORDER]
acute_kruskal = stats.kruskal(*region_vals_list)
kruskal_p_text = '< 0.001' if acute_kruskal.pvalue < 0.001 else f'{acute_kruskal.pvalue:.3f}'

cap_p = doc.add_paragraph()
cap_r = cap_p.add_run('Table 1. ')
cap_r.bold = True
cap_r.font.size = Pt(10)

cap_p.add_run(
    'Regional summary of inpatient analgesic prescribing per surgery across nine regional blocks. '
    'Values are analgesic-per-surgery index (mean \u00b1 SD). '
    f'Kruskal\u2013Wallis P = {kruskal_p_text}.'
).font.size = Pt(10)

table1 = doc.add_table(rows=1 + len(REGION_ORDER) + 1, cols=5)
set_table_borders(table1)
t1_headers = ['Region', 'N prefectures', 'Mean', 'SD', 'Range']
for i, h in enumerate(t1_headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = table1.rows[idx + 1]
    row.cells[0].text = REGION_EN[reg_name]
    row.cells[1].text = str(len(vals))
    row.cells[2].text = f'{np.mean(vals):.2f}'
    if len(vals) > 1:
        row.cells[3].text = f'{np.std(vals, ddof=1):.2f}'
    else:
        row.cells[3].text = '\u2014'
    row.cells[4].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

all_vals = [r['acute_analgesic_per_surgery'] for r in rows]
nat_row = table1.rows[len(REGION_ORDER) + 1]
nat_row.cells[0].text = 'National'
nat_row.cells[1].text = str(len(rows))
nat_row.cells[2].text = f'{np.mean(all_vals):.2f}'
nat_row.cells[3].text = f'{np.std(all_vals, ddof=1):.2f}'
nat_row.cells[4].text = f'{min(all_vals):.2f}\u2013{max(all_vals):.2f}'
for cell in nat_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

# ============================================================
# TABLE 2
# ============================================================
doc.add_page_break()
cap_p = doc.add_paragraph()
cap_r = cap_p.add_run('Table 2. ')
cap_r.bold = True
cap_r.font.size = Pt(10)
cap_p.add_run(
    'Regression models for outpatient neuropathic pain prescribing with Tohoku indicator and confounder adjustment.'
).font.size = Pt(10)

table2 = doc.add_table(rows=7, cols=5)
set_table_borders(table2)
t2_headers = ['Model', 'Dependent variable', 'Tohoku coefficient / effect size', 'P value', 'Significance']
for i, h in enumerate(t2_headers):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t2_data = [
    ['Model 1', 'Unadjusted t-test',
     f"d = {reg['model1_unadjusted']['cohens_d']:.3f}",
     fmt_p(reg['model1_unadjusted']['p_value']), '***'],
    ['Model 2', 'Neuropathic pain drugs / surgery (fully adjusted)',
     f"\u03b2 = {reg['model2_adjusted']['tohoku_coef']:.1f}",
     fmt_p(reg['model2_adjusted']['tohoku_p']), 'ns'],
    ['Model 3', 'Core neuropathic drugs (PGB+MGB) (fully adjusted)',
     f"\u03b2 = {reg['model3_core_neuropathic']['tohoku_coef']:.1f}",
     fmt_p(reg['model3_core_neuropathic']['tohoku_p']), 'ns'],
    ['Model 4', 'Nerve blocks / surgery (fully adjusted)',
     f"\u03b2 = {reg['model4_nerve_blocks']['tohoku_coef']:.2f}",
     fmt_p(reg['model4_nerve_blocks']['tohoku_p']), 'ns'],
    ['Model 5', 'Neuropathic pain drugs (acute + confounder adj.)',
     f"\u03b2 = {reg['model5_integrated']['tohoku_coef']:.1f}",
     fmt_p(reg['model5_integrated']['tohoku_p']), 'ns'],
    ['Adj CPSP', 'Confounder-removed residual',
     f"d = {reg['adjusted_cpsp_test']['cohens_d']:.3f}",
     fmt_p(reg['adjusted_cpsp_test']['p_value']), 'ns'],
]
for r_idx, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r_idx + 1].cells[c].text = val
        for paragraph in table2.rows[r_idx + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

for cell in table2.rows[1].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_r = note_p.add_run(
    '*** P < 0.001; ns = not significant. '
    'Models 2\u20135: multiple linear regression with Tohoku indicator (binary) and confounder disease proxies. '
    'Adj CPSP: adjusted CPSP index derived as residuals from regressing neuropathic pain prescribing '
    'on four confounder proxies.'
)
note_r.font.size = Pt(8)
note_r.font.italic = True

outpath = os.path.join(JOPP_DIR, 'JoPPP_tables_EN.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
