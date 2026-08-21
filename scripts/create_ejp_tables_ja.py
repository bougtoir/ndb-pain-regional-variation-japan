#!/usr/bin/env python3
"""Create separate editable tables DOCX for EJP submission (Japanese).

EJP requires figures and tables to be uploaded as separate files,
not embedded in the main manuscript text.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import csv
import numpy as np
from collections import defaultdict

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
EJP_DIR = os.path.join(OUTPUT_DIR, 'ejp')
os.makedirs(EJP_DIR, exist_ok=True)

with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_JA = {
    '\u5317\u6d77\u9053': '\u5317\u6d77\u9053', '\u6771\u5317': '\u6771\u5317', '\u95a2\u6771': '\u95a2\u6771',
    '\u5317\u9678\u30fb\u7532\u4fe1\u8d8a': '\u5317\u9678\u30fb\u7532\u4fe1\u8d8a', '\u6771\u6d77': '\u6771\u6d77', '\u8fd1\u757f': '\u8fd1\u757f',
    '\u4e2d\u56fd': '\u4e2d\u56fd', '\u56db\u56fd': '\u56db\u56fd', '\u4e5d\u5dde\u30fb\u6c96\u7e04': '\u4e5d\u5dde\u30fb\u6c96\u7e04',
}
REGION_ORDER = ['\u5317\u6d77\u9053', '\u6771\u5317', '\u95a2\u6771', '\u5317\u9678\u30fb\u7532\u4fe1\u8d8a', '\u6771\u6d77', '\u8fd1\u757f', '\u4e2d\u56fd', '\u56db\u56fd', '\u4e5d\u5dde\u30fb\u6c96\u7e04']

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.5


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


# ============================================================
# \u88681
# ============================================================
cap_p = doc.add_paragraph()
cap_r = cap_p.add_run('\u88681. ')
cap_r.bold = True
cap_r.font.size = Pt(10)
cap_p.add_run('Phase 1\uff1a9\u5730\u57df\u30d6\u30ed\u30c3\u30af\u5225\u306e\u5165\u9662\u93ae\u75db\u85ac/\u624b\u8853\u6307\u6a19\u306e\u6982\u8981\u3002').font.size = Pt(10)

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

table1 = doc.add_table(rows=1 + len(REGION_ORDER) + 1, cols=5)
set_table_borders(table1)
t1_headers = ['\u5730\u57df', '\u90fd\u9053\u5e9c\u770c\u6570', '\u5e73\u5747', 'SD', '\u7bc4\u56f2']
for i, h in enumerate(t1_headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = table1.rows[idx + 1]
    row.cells[0].text = REGION_JA[reg_name]
    row.cells[1].text = str(len(vals))
    row.cells[2].text = f'{np.mean(vals):.2f}'
    row.cells[3].text = f'{np.std(vals):.2f}'
    row.cells[4].text = f'{min(vals):.2f}\u2013{max(vals):.2f}'
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

all_vals = [r['acute_analgesic_per_surgery'] for r in rows]
nat_row = table1.rows[len(REGION_ORDER) + 1]
nat_row.cells[0].text = '\u5168\u56fd'
nat_row.cells[1].text = str(len(rows))
nat_row.cells[2].text = f'{np.mean(all_vals):.2f}'
nat_row.cells[3].text = f'{np.std(all_vals):.2f}'
nat_row.cells[4].text = f'{min(all_vals):.2f}\u2013{max(all_vals):.2f}'
for cell in nat_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_r = note_p.add_run(
    '\u5024\u306f\u93ae\u75db\u85ac/\u624b\u8853\u6307\u6a19\uff08\u5404\u90fd\u9053\u5e9c\u770c\u306e\u5165\u9662\u93ae\u75db\u85ac\u51e6\u65b9\u5358\u4f4d\u5408\u8a08\u00f7\u5165\u9662\u624b\u8853\u4ef6\u6570\u5408\u8a08\uff09\u3092\u8868\u3059\u3002'
    '9\u5730\u57df\u9593\u306eKruskal\u2013Wallis\u691c\u5b9a\uff1aP < 0.001\u3002')
note_r.font.size = Pt(8)
note_r.font.italic = True

# ============================================================
# \u88682
# ============================================================
doc.add_page_break()
cap_p = doc.add_paragraph()
cap_r = cap_p.add_run('\u88682. ')
cap_r.bold = True
cap_r.font.size = Pt(10)
cap_p.add_run('Phase 2\uff1a\u6771\u5317\u6307\u6a19\u30fb\u4ea4\u7d61\u56e0\u5b50\u8abf\u6574\u3092\u542b\u3080\u56de\u5e30\u30e2\u30c7\u30eb\u3002').font.size = Pt(10)

table2 = doc.add_table(rows=7, cols=5)
set_table_borders(table2)
t2_headers = ['\u30e2\u30c7\u30eb', '\u5f93\u5c5e\u5909\u6570', '\u6771\u5317\u4fc2\u6570/\u52b9\u679c\u91cf', 'P\u5024', '\u6709\u610f\u6027']
for i, h in enumerate(t2_headers):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t2_data = [
    ['\u30e2\u30c7\u30eb1', '\u672a\u8abf\u6574t\u691c\u5b9a',
     f'd = {reg["model1_unadjusted"]["cohens_d"]:.3f}',
     f'{reg["model1_unadjusted"]["p_value"]:.2e}', '***'],
    ['\u30e2\u30c7\u30eb2', '\u795e\u7d4c\u969c\u5bb3\u6027\u75bc\u75db\u85ac/\u624b\u8853\uff08\u5b8c\u5168\u8abf\u6574\uff09',
     f'\u03b2 = {reg["model2_adjusted"]["tohoku_coef"]:.1f}',
     f'{reg["model2_adjusted"]["tohoku_p"]:.4f}', 'ns'],
    ['\u30e2\u30c7\u30eb3', '\u30b3\u30a2\u85ac\uff08PGB+MGB\uff09\uff08\u5b8c\u5168\u8abf\u6574\uff09',
     f'\u03b2 = {reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}',
     f'{reg["model3_core_neuropathic"]["tohoku_p"]:.4f}', 'ns'],
    ['\u30e2\u30c7\u30eb4', '\u795e\u7d4c\u30d6\u30ed\u30c3\u30af/\u624b\u8853\uff08\u5b8c\u5168\u8abf\u6574\uff09',
     f'\u03b2 = {reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}',
     f'{reg["model4_nerve_blocks"]["tohoku_p"]:.4f}', 'ns'],
    ['\u30e2\u30c7\u30eb5', '\u795e\u7d4c\u969c\u5bb3\u6027\u75bc\u75db\u85ac\uff08\u6025\u6027\u671f+\u4ea4\u7d61\u8abf\u6574\uff09',
     f'\u03b2 = {reg["model5_integrated"]["tohoku_coef"]:.1f}',
     f'{reg["model5_integrated"]["tohoku_p"]:.4f}', 'ns'],
    ['\u8abf\u6574\u6e08CPSP', '\u4ea4\u7d61\u9664\u53bb\u6b8b\u5dee',
     f'd = {reg["adjusted_cpsp_test"]["cohens_d"]:.3f}',
     f'{reg["adjusted_cpsp_test"]["p_value"]:.4f}', 'ns'],
]
for r_idx, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        table2.rows[r_idx + 1].cells[c].text = val
        for paragraph in table2.rows[r_idx + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_r = note_p.add_run(
    '*** P < 0.001; ns = \u6709\u610f\u3067\u306a\u3044\u3002'
    '\u30e2\u30c7\u30eb2\u301c5\uff1a\u6771\u5317\u6307\u6a19\uff08\u4e8c\u5024\uff09\u3068\u4ea4\u7d61\u75be\u60a3\u30d7\u30ed\u30ad\u30b7\u306b\u3088\u308b\u91cd\u56de\u5e30\u5206\u6790\u3002'
    '\u8abf\u6574\u6e08CPSP\uff1a\u795e\u7d4c\u969c\u5bb3\u6027\u75bc\u75db\u85ac\u51e6\u65b9\u30924\u4ea4\u7d61\u30d7\u30ed\u30ad\u30b7\u306b\u56de\u5e30\u3057\u305f\u6b8b\u5dee\u3002')
note_r.font.size = Pt(8)
note_r.font.italic = True

# ============================================================
# \u88683
# ============================================================
doc.add_page_break()
cap_p = doc.add_paragraph()
cap_r = cap_p.add_run('\u88683. ')
cap_r.bold = True
cap_r.font.size = Pt(10)
cap_p.add_run('\u4ea4\u7d61\u56e0\u5b50\u8abf\u6574\u306b\u3088\u308b\u6771\u5317\u5730\u57df\u6307\u6a19\u306e\u5909\u5316\u3002').font.size = Pt(10)

unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attn = (1 - adj_d / unadj_d) * 100

table3 = doc.add_table(rows=4, cols=5)
set_table_borders(table3)
t3_headers = ['\u6307\u6a19', '\u672a\u8abf\u6574', '\u4ea4\u7d61\u56e0\u5b50\u8abf\u6574\u5f8c', '\u5909\u5316', '\u89e3\u91c8']
for i, h in enumerate(t3_headers):
    table3.rows[0].cells[i].text = h
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t3_data = [
    ["Cohen's d\uff08\u6771\u5317 vs \u975e\u6771\u5317\uff09",
     f'{unadj_d:.3f} (P < 0.001)',
     f'{adj_d:.3f} (P = {reg["adjusted_cpsp_test"]["p_value"]:.3f})',
     f'{attn:.0f}%\u6e1b\u5f31',
     '\u5927 \u2192 \u5c0f\u52b9\u679c\u91cf'],
    ['\u6771\u5317\u5e73\u5747\u6307\u6a19',
     f'{reg["model1_unadjusted"]["tohoku_mean"]:.1f}',
     f'{reg["adjusted_cpsp_test"]["tohoku_mean"]:+.1f}\uff08\u6b8b\u5dee\uff09',
     '\u2014',
     '\u904e\u5270\u306e\u5927\u90e8\u5206\u306f\u4ea4\u7d61\u3067\u8aac\u660e'],
    ['\u975e\u6771\u5317\u5e73\u5747\u6307\u6a19',
     f'{reg["model1_unadjusted"]["non_tohoku_mean"]:.1f}',
     f'{reg["adjusted_cpsp_test"]["non_tohoku_mean"]:+.1f}\uff08\u6b8b\u5dee\uff09',
     '\u2014',
     '\u53c2\u7167\u7fa4'],
]
for r_idx, row_data in enumerate(t3_data):
    for c, val in enumerate(row_data):
        table3.rows[r_idx + 1].cells[c].text = val
        for paragraph in table3.rows[r_idx + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_r = note_p.add_run(
    '\u4ea4\u7d61\u56e0\u5b50\uff1a\u7d4c\u53e3\u8840\u7cd6\u964d\u4e0b\u85ac\uff08\u7cd6\u5c3f\u75c5\u30d7\u30ed\u30ad\u30b7\uff09\u3001\u5e2f\u72b6\u7652\u75b9\u6297\u30a6\u30a4\u30eb\u30b9\u85ac\u3001'
    '\u6297\u3046\u3064\u85ac\uff08\u30c7\u30e5\u30ed\u30ad\u30bb\u30c1\u30f3\u9664\u304f\uff09\u3001\u6297\u4e0d\u5b89\u85ac\u3002')
note_r.font.size = Pt(8)
note_r.font.italic = True

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(EJP_DIR, 'EJP_tables_JA.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
