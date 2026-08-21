#!/usr/bin/env python3
"""Create cover letter for JoPPP submission."""
import os
import json
import csv
import datetime
import numpy as np
from docx import Document
from docx.shared import Pt, Cm

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
JOPP_DIR = os.path.join(OUTPUT_DIR, 'joppp')
os.makedirs(JOPP_DIR, exist_ok=True)

# Load data
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except Exception:
                    pass
        rows.append(r)

PREF_EN = {
    '\u5317\u6d77\u9053': 'Hokkaido',
    '\u9752\u68ee\u770c': 'Aomori', '\u5ca9\u624b\u770c': 'Iwate', '\u5bae\u57ce\u770c': 'Miyagi',
    '\u79cb\u7530\u770c': 'Akita', '\u5c71\u5f62\u770c': 'Yamagata', '\u798f\u5cf6\u770c': 'Fukushima',
    '\u8328\u57ce\u770c': 'Ibaraki', '\u6803\u6728\u770c': 'Tochigi', '\u7fa4\u99ac\u770c': 'Gunma',
    '\u57fc\u7389\u770c': 'Saitama', '\u5343\u8449\u770c': 'Chiba', '\u6771\u4eac\u90fd': 'Tokyo',
    '\u795e\u5948\u5ddd\u770c': 'Kanagawa',
    '\u65b0\u6f5f\u770c': 'Niigata', '\u5bcc\u5c71\u770c': 'Toyama', '\u77f3\u5ddd\u770c': 'Ishikawa',
    '\u798f\u4e95\u770c': 'Fukui', '\u5c71\u68a8\u770c': 'Yamanashi', '\u9577\u91ce\u770c': 'Nagano',
    '\u5c90\u961c\u770c': 'Gifu', '\u9759\u5ca1\u770c': 'Shizuoka', '\u611b\u77e5\u770c': 'Aichi',
    '\u4e09\u91cd\u770c': 'Mie', '\u6ecb\u8cc0\u770c': 'Shiga', '\u4eac\u90fd\u5e9c': 'Kyoto',
    '\u5927\u962a\u5e9c': 'Osaka', '\u5175\u5eab\u770c': 'Hyogo', '\u5948\u826f\u770c': 'Nara',
    '\u548c\u6b4c\u5c71\u770c': 'Wakayama', '\u9ce5\u53d6\u770c': 'Tottori', '\u5cf6\u6839\u770c': 'Shimane',
    '\u5ca1\u5c71\u770c': 'Okayama', '\u5e83\u5cf6\u770c': 'Hiroshima', '\u5c71\u53e3\u770c': 'Yamaguchi',
    '\u5fb3\u5cf6\u770c': 'Tokushima', '\u9999\u5ddd\u770c': 'Kagawa', '\u611b\u5a9b\u770c': 'Ehime',
    '\u9ad8\u77e5\u770c': 'Kochi', '\u798f\u5ca1\u770c': 'Fukuoka', '\u4f50\u8cc0\u770c': 'Saga',
    '\u9577\u5d0e\u770c': 'Nagasaki', '\u718a\u672c\u770c': 'Kumamoto', '\u5927\u5206\u770c': 'Oita',
    '\u5bae\u5d0e\u770c': 'Miyazaki', '\u9e7f\u5150\u5cf6\u770c': 'Kagoshima', '\u6c96\u7e04\u770c': 'Okinawa',
}

acute_vals = [r['acute_analgesic_per_surgery'] for r in rows]
min_row = min(rows, key=lambda r: r['acute_analgesic_per_surgery'])
max_row = max(rows, key=lambda r: r['acute_analgesic_per_surgery'])
fold_ratio = max(acute_vals) / min(acute_vals)
min_pref = PREF_EN.get(min_row['pref_name'], min_row['pref_name'])
max_pref = PREF_EN.get(max_row['pref_name'], max_row['pref_name'])

tohoku_vals = [r['acute_analgesic_per_surgery'] for r in rows if r['is_tohoku'] == '1']
non_tohoku_vals = [r['acute_analgesic_per_surgery'] for r in rows if r['is_tohoku'] == '0']
pooled_sd = np.sqrt(((len(tohoku_vals) - 1) * np.var(tohoku_vals, ddof=1) +
                     (len(non_tohoku_vals) - 1) * np.var(non_tohoku_vals, ddof=1)) /
                    (len(tohoku_vals) + len(non_tohoku_vals) - 2))
tohoku_d = (np.mean(tohoku_vals) - np.mean(non_tohoku_vals)) / pooled_sd

unadj_d = reg['model1_unadjusted']['cohens_d']
adj_d = reg['adjusted_cpsp_test']['cohens_d']
attenuation = (1 - adj_d / unadj_d) * 100

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

today = datetime.date.today().strftime('%B %d, %Y')

doc.add_paragraph(today)
doc.add_paragraph()

doc.add_paragraph('Editors')
doc.add_paragraph('Journal of Pharmaceutical Policy and Practice')
doc.add_paragraph('Taylor & Francis')
doc.add_paragraph()

doc.add_paragraph('Dear Editors,')
doc.add_paragraph()

title = ("Regional variation in acute and neuropathic pain medication prescribing "
         "across Japan's 47 prefectures: a population-level ecological study")

doc.add_paragraph(
    f'We respectfully submit the enclosed manuscript entitled '
    f'\u201c{title}\u201d '
    f'for consideration as a Research Article in the Journal of Pharmaceutical Policy and Practice.'
)

doc.add_paragraph(
    'Cross-cultural studies have long characterized Japanese people as stoic toward pain, '
    'yet whether this stoicism varies within Japan has never been examined at the population level. '
    'Using Japan\u2019s National Database (NDB) Open Data\u2014capturing virtually all insurance-reimbursed '
    'healthcare for approximately 125 million people\u2014we mapped pain-related prescribing across all 47 prefectures '
    'and nine regional blocks.'
)

doc.add_paragraph('Our key findings are:')

doc.add_paragraph(
    f'(1) Acute perioperative analgesic prescribing varied {fold_ratio:.2f}-fold across prefectures '
    f'({min_pref} {min(acute_vals):.2f} to {max_pref} {max(acute_vals):.2f}), '
    f'demonstrating that Japan\u2019s pain culture is not monolithic.',
    style='List Bullet'
)

doc.add_paragraph(
    f'(2) Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014prescribed more, '
    f'not fewer, analgesics than the rest of the country (Cohen\u2019s d = {tohoku_d:.2f}).',
    style='List Bullet'
)

doc.add_paragraph(
    f'(3) The large regional variation in outpatient neuropathic pain prescribing (Cohen\u2019s d = {unadj_d:.2f}) '
    f'was attenuated by {attenuation:.0f}% after adjustment for confounding disease burden, '
    f'particularly diabetes, indicating that apparent cultural patterns are substantially shaped by comorbidity.',
    style='List Bullet'
)

doc.add_paragraph(
    'We believe this manuscript is well suited for the Journal of Pharmaceutical Policy and Practice. '
    'The topic sits at the intersection of pharmaceutical policy and clinical practice: Japan\u2019s universal, '
    'standardised-reimbursement system offers a natural laboratory for estimating policy-relevant variation '
    'under uniform pricing, and the findings challenge assumptions that can influence prescribing decisions. '
    'The ecological, population-complete design and transparent confounder-adjustment framework should be '
    'replicable in other countries with national claims databases.'
)

doc.add_paragraph(
    'This manuscript has not been published previously and is not under consideration elsewhere. '
    'All authors have approved the manuscript and agree with its submission to the Journal of Pharmaceutical Policy and Practice. '
    'The study used only publicly available aggregate data, so ethical approval was not required. '
    'The author reports no conflicts of interest to declare.'
)

doc.add_paragraph(
    'We confirm that this manuscript complies with the STROBE guidelines for reporting observational studies, '
    'and a completed STROBE checklist is included with this submission.'
)

doc.add_paragraph(
    'Thank you for considering our work. We look forward to your response.'
)

doc.add_paragraph()
doc.add_paragraph('Sincerely,')
doc.add_paragraph()
doc.add_paragraph('Tatsuki Onishi')
doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('[Address]')
doc.add_paragraph('E-mail: [email]')

outpath = os.path.join(JOPP_DIR, 'JoPPP_cover_letter.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
