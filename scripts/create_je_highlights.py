#!/usr/bin/env python3
"""Create Highlights file for Journal of Epidemiology submission.

JE requires:
- 3-5 bullet points
- Maximum 150 characters (including spaces) per bullet
- Separate editable file named 'Highlights'
"""

from docx import Document
from docx.shared import Pt, Cm
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JE_DIR = os.path.join(OUTPUT_DIR, 'je')
os.makedirs(JE_DIR, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

h = doc.add_heading('Highlights', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'

highlights = [
    'Perioperative analgesic prescribing varies 1.97-fold across Japan\u2019s 47 prefectures.',
    'Nearly twofold variation persists after age-sex standardisation (SCR range, 64\u2013148).',
    'Prescribing variation is dissociated from symptom burden (CSLC r=0.03), indicating unwarranted variation.',
    'Diabetes drug prescribing (r=0.87) is the dominant confounder of neuropathic pain patterns.',
    'NDB Open Data triangulated with CSLC enable demand\u2013supply analysis at no cost.',
]

for i, hl in enumerate(highlights, 1):
    char_count = len(hl)
    p = doc.add_paragraph(hl, style='List Bullet')
    status = 'OK' if char_count <= 150 else f'OVER by {char_count - 150}'
    print(f'  Highlight {i}: {char_count} chars ({status})')

outpath = os.path.join(JE_DIR, 'JE_highlights.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
