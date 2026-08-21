#!/usr/bin/env python3
"""Create Japanese manuscript for Journal of Epidemiology (JE).

JE accepts manuscripts in English only for publication, but this Japanese
version serves as reference material for the author and co-reviewers.
Structure mirrors the English version (create_je_docx_en.py).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re
import csv
import numpy as np
from collections import defaultdict

# ============================================================
# Paths
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
JE_DIR = os.path.join(OUTPUT_DIR, 'je')
os.makedirs(JE_DIR, exist_ok=True)

FIG_DIR = OUTPUT_DIR

# ============================================================
# Load data
# ============================================================
with open(os.path.join(OUTPUT_DIR, 'cpsp_regression_summary.json'), 'r') as f:
    reg = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'scr_summary.json'), 'r') as f:
    scr = json.load(f)

with open(os.path.join(OUTPUT_DIR, 'cslc_analysis.json'), 'r') as f:
    cslc = json.load(f)

rows = []
with open(os.path.join(OUTPUT_DIR, 'cpsp_integrated_results.csv'), 'r', encoding='utf-8') as f:
    for r in csv.DictReader(f):
        for k in r:
            if k not in ('pref_name', 'region', 'is_tohoku', 'pref_code'):
                try:
                    r[k] = float(r[k])
                except:
                    pass
        r['pref_code'] = int(r['pref_code'])
        r['is_tohoku'] = int(float(r['is_tohoku']))
        rows.append(r)

REGION_JA = {
    '北海道': '北海道', '東北': '東北', '関東': '関東',
    '北陸・甲信越': '北陸・甲信越', '東海': '東海', '近畿': '近畿',
    '中国': '中国', '四国': '四国', '九州・沖縄': '九州・沖縄',
}
REGION_ORDER = ['北海道', '東北', '関東', '北陸・甲信越', '東海', '近畿', '中国', '四国', '九州・沖縄']

# ============================================================
# REFERENCES
# ============================================================
ref_list = [
    # 1 Wennberg
    'Wennberg JE. Tracking Medicine: A Researcher\'s Quest to Understand Health Care. '
    'New York, NY: Oxford University Press; 2010.',
    # 2 Corallo
    'Corallo AN, Croxford R, Goodman DC, et al. A systematic review of medical practice '
    'variation in OECD countries. Health Policy. 2014;114:5\u201314.',
    # 3 MHLW NDB
    '厚生労働省. NDBオープンデータ 第10回. '
    'https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html. '
    '2025年1月15日アクセス.',
    # 4 Taira
    'Taira K, Mori T, Ishimaru M, et al. Regional inequality in dental care utilisation '
    'in Japan: an ecological study using the National Database of Health Insurance Claims. '
    'Lancet Reg Health West Pac. 2021;12:100170.',
    # 5 Wakaizumi
    'Wakaizumi K, Tanaka C, Shinohara Y, et al. Geographical variation in high-impact '
    'chronic pain and psychological associations at the regional level. '
    'Front Public Health. 2024;12:1482177.',
    # 6 Takahashi
    'Takahashi R, Miyashita M, Nakazawa Y, Wada S, Matsuoka Y. Population-based claims study '
    'of regional and hospital function differences in opioid prescribing for cancer patients '
    'who died in hospital in Japan. Jpn J Clin Oncol. 2025;55:1372\u20137.',
    # 7 CSLC
    '厚生労働省. 国民生活基礎調査（2022年）. '
    'https://www.mhlw.go.jp/toukei/saikin/hw/k-tyosa/k-tyosa22/. '
    '2025年3月1日アクセス.',
    # 8 von Elm STROBE
    'von Elm E, Altman DG, Egger M, et al. The STROBE statement: guidelines for reporting '
    'observational studies. Lancet. 2007;370:1453\u20137.',
    # 9 Benchimol RECORD
    'Benchimol EI, Smeeth L, Guttmann A, et al. The RECORD statement. '
    'PLoS Med. 2015;12:e1001885.',
    # 10 Kehlet
    'Kehlet H, Jensen TS, Woolf CJ. Persistent postsurgical pain: risk factors and prevention. '
    'Lancet. 2006;367:1618\u201325.',
    # 11 Anderson
    'Anderson KO, Green CR, Payne R. Racial and ethnic disparities in pain. '
    'J Pain. 2009;10:1187\u2013204.',
    # 12 Campbell
    'Campbell CM, Edwards RR. Ethnic differences in pain and pain management. '
    'Pain Manag. 2012;2:219\u201330.',
    # 13 Callister
    'Callister LC. Cultural influences on pain perceptions and behaviors. '
    'Home Health Care Manag Pract. 2003;15:207\u201311.',
    # 14 Hobara
    'Hobara M. Beliefs about appropriate pain behavior: cross-cultural and sex differences '
    'between Japanese and Euro-Americans. Eur J Pain. 2005;9:389\u201393.',
    # 15 Pfizer
    'ファイザー株式会社. 47都道府県 長く続く痛みに関する実態調査（2012年 vs 2017年比較）. '
    'https://www.pfizer.co.jp/pfizer/company/press/2017/2017_08_23.html. '
    '2022年5月28日アクセス. アーカイブ: '
    'https://web.archive.org/web/20220528073616/'
    'https://www.pfizer.co.jp/pfizer/company/press/2017/2017_08_23.html',
    # 16 Raja IASP
    'Raja SN, Carr DB, Cohen M, et al. The revised IASP definition of pain. '
    'Pain. 2020;161:1976\u201382.',
    # 17 Onishi
    'Onishi T, Onishi Y. Normalized pulse volume as a superior predictor of respiration recovery. '
    'F1000Research. 2024;13:233.',
]


def cite(*nums):
    return '{' + ','.join(str(n) for n in nums) + '}'


def add_ref_runs(p, text):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            p.add_run(part)


# ============================================================
# Document creation
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Yu Mincho'
font.size = Pt(10.5)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

# Page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Yu Gothic'
    return h


def wc(text):
    return len(re.sub(r'\{[^}]+\}', '', text).split())


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


# Map figure numbers to source PNG files (Japanese versions)
FIG_SOURCES = {
    1: os.path.join(FIG_DIR, 'fig1_neuropathic_unadjusted.png'),
    2: os.path.join(FIG_DIR, 'fig2_confounder_correlations.png'),
    3: os.path.join(FIG_DIR, 'fig4_region_unadj_vs_adj.png'),
    4: os.path.join(FIG_DIR, 'fig_cslc_demand_supply_ja.png'),
}


def add_inline_figure(caption_text, fig_num):
    fig_path = FIG_SOURCES.get(fig_num)
    if fig_path and os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        r0 = p.add_run(f'[図{fig_num}をここに挿入]')
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f'図{fig_num}. ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(caption_text)
    r2.font.size = Pt(10)


# ============================================================
# Computed values
# ============================================================
unadj_d = reg["model1_unadjusted"]["cohens_d"]
adj_d = reg["adjusted_cpsp_test"]["cohens_d"]
attenuation = (1 - adj_d / unadj_d) * 100

scr_analgesic_range = scr['analgesic_inpatient']['scr_range']
scr_analgesic_ratio = scr['analgesic_inpatient']['variation_ratio']
scr_neuro_range = scr['neuropathic_outpatient']['scr_range']
scr_neuro_ratio = scr['neuropathic_outpatient']['variation_ratio']

region_data = defaultdict(list)
for r in rows:
    region_data[r['region']].append(r['acute_analgesic_per_surgery'])

# ============================================================
# TITLE PAGE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    '日本47都道府県における疼痛関連処方の地域差：'
    'NDBオープンデータを用いた生態学的研究'
)
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('大西 龍樹')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('麻酔科, [所属施設], [住所], [都市], 〒[郵便番号], 日本')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('責任著者: ')
r.bold = True
p.add_run(
    '大西龍樹, 麻酔科, [所属施設], '
    '[住所], [都市], 〒[郵便番号], 日本. E-mail: [email]'
)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('論文種別: ')
r.bold = True
p.add_run('原著論文（Original Article）')

p = doc.add_paragraph()
r = p.add_run('キーワード: ')
r.bold = True
p.add_run('鎮痛薬; 生態学的研究; 医療格差; 日本; 処方')

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_text('抄録', level=1)

abstract_bg = (
    '単一の国民皆保険制度内においても疼痛関連処方に地域差が存在する可能性があるが、'
    'その規模と決定因子は日本では十分に解明されていない。'
    '処方差が症状負荷の差（warranted variation）を反映するのか、'
    '診療パターンの差（unwarranted variation）なのかは未検討である。'
)

abstract_methods = (
    '本研究は、NDBオープンデータ（第10回、2023年4月〜2024年3月）を用いて、'
    '47都道府県の県レベル集計データを分析する生態学的研究である。'
    '第1相は急性周術期鎮痛薬処方（手術あたり）を、'
    '第2相は外来神経障害性疼痛薬処方を慢性疼痛の代理指標として検討し、'
    '交絡疾患プロキシ（糖尿病、帯状疱疹、うつ病、不安症）による回帰調整を行った。'
    '間接年齢・性別標準化による標準化レセプト比（SCR）も算出した。'
)

abstract_results = (
    f'急性鎮痛薬処方は都道府県間で1.97倍の差を示した'
    f'（Kruskal–Wallis P<0.001）。'
    f'未調整の神経障害性疼痛薬処方は顕著な地域集積を示したが、'
    f'交絡疾患プロキシ（特に糖尿病薬、r=0.87）が都道府県間分散の'
    f'{reg["model2_adjusted"]["R2"]*100:.0f}%を説明した。'
    f'交絡調整後、見かけの地域間差は大幅に減弱し非有意となった。'
    f'CSLC有訴者率は{cslc["symptom_rate_stats"]["range_ratio"]:.2f}倍の差に留まり、'
    f'処方との相関は認められなかった'
    f'（r={cslc["correlations"]["symptom_vs_acute"]["pearson_r"]:.2f}）。'
)

abstract_conclusions = (
    '年齢・性別標準化後も約2倍の国内差が持続し、症状負荷と乖離していた。'
    'これはWennbergの不当な差異の基準を満たし、供給側要因が処方の不均一性を駆動していることを示唆する。'
    '神経障害性疼痛薬の生態学的研究では、糖尿病をはじめとする交絡疾患を考慮する必要がある。'
)

for label, text in [('背景', abstract_bg), ('方法', abstract_methods),
                     ('結果', abstract_results), ('結論', abstract_conclusions)]:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_text('緒言', level=1)

intro1 = (
    f'医療における地域差は、Wennbergの小地域差研究以来、'
    f'医療の質に関する重要な課題として認識されてきた。{cite(1)}'
    f'処方率が疾病負荷や人口構成で説明できる範囲を超えて地域間で異なる場合、'
    f'そのような不当な差異（unwarranted variation）は過剰治療または過少治療の可能性を示唆し、'
    f'体系的な調査が必要である。{cite(1,2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro1)

intro2 = (
    f'日本の国民皆保険制度、統一薬価制度、'
    f'およびNDB（National Database of Health Insurance Claims）'
    f'—約1億2,500万人の被保険者の医療利用をほぼ全数捕捉する{cite(3)}—は、'
    f'診療差異を研究するための独特な環境を提供する。'
    f'Tairaらは、NDB由来の標準化レセプト比を用いて歯科診療の地域格差を示した。{cite(4)}'
    f'慢性疼痛有病率の1.6倍の地域差{cite(5)}や、がんオピオイド処方の4倍差{cite(6)}'
    f'も報告されているが、NDBを用いて47都道府県の疼痛関連処方差を包括的に'
    f'マッピングした研究はない。'
)
p = doc.add_paragraph()
add_ref_runs(p, intro2)

intro_cslc = (
    f'Wennbergの枠組みは、疾病負荷の違いに起因する正当な差異（warranted variation）と、'
    f'医師の診療スタイルや地域規範等の供給側要因に起因する不当な差異（unwarranted variation）を'
    f'区別する。{cite(1,2)}'
    f'この枠組みを適用するには、独立した需要側ベンチマークが必要である。'
    f'国民生活基礎調査（CSLC）は、厚生労働省が実施する全国代表世帯調査であり、'
    f'都道府県別の自覚症状有訴者率を提供している。{cite(7)}'
)
p = doc.add_paragraph()
add_ref_runs(p, intro_cslc)

intro3 = (
    '疼痛管理は生態学的分析に特に適している。鎮痛薬処方は疾病負荷と診療パターンの'
    '両方を反映するためである。周術期設定は方法論的優位性を有する：'
    '本分析の全患者は手術のために入院しているため、地域在住型処方研究における'
    '主要な交絡因子である医療アクセスが設計上中和される。'
    '外来神経障害性疼痛薬は慢性疼痛負荷の集団レベルの代理指標となりうるが、'
    '複数の適応症（特に糖尿病性神経障害）を有するため、'
    '慎重な交絡調整が必要である。'
)
doc.add_paragraph(intro3)

intro4 = (
    '本研究の目的は4つである：(1)47都道府県における急性周術期鎮痛薬処方の地域差を'
    'マッピングすること、(2)交絡疾患を調整した上で、外来神経障害性疼痛薬処方を'
    '慢性疼痛の代理指標として検討すること、'
    '(3)見かけの地域パターンに対する交絡因子の寄与を定量化すること、'
    '(4)処方差異が症状負荷と乖離しているか（＝不当な差異の基準を満たすか）を評価すること。'
)
doc.add_paragraph(intro4)

# ============================================================
# METHODS
# ============================================================
add_heading_text('方法', level=1)

add_heading_text('研究デザインと報告', level=2)
m1 = (
    f'本研究は、NDBオープンデータの都道府県レベル集計データを分析した生態学的研究である。'
    f'STROBE声明{cite(8)}およびRECORD拡張{cite(9)}に従って報告する。'
    f'公開されている集計データのみを使用したため、「人を対象とする生命科学・医学系研究に'
    f'関する倫理指針」に基づき倫理審査は不要であった。'
)
p = doc.add_paragraph()
add_ref_runs(p, m1)

add_heading_text('データソース', level=2)
m2 = (
    f'NDBオープンデータ第10回（2023年4月〜2024年3月）を使用した。{cite(3)}'
    f'NDBは国民皆保険制度の全保険者からのレセプトを捕捉し、約1億2,500万人の'
    f'被保険者を包含する。処方および処置の集計データは都道府県レベル（n=47）で'
    f'公表され、10件未満のセルは秘匿される。'
    f'都道府県別人口推計（2023年10月、5歳階級別・男女別）は'
    f'総務省統計局から取得した。'
)
p = doc.add_paragraph()
add_ref_runs(p, m2)

add_heading_text('地域分類', level=2)
doc.add_paragraph(
    '都道府県は総務省統計局の分類に従い、9つの地域ブロックに分類した：'
    '北海道（1）、東北（6）、関東（7）、北陸・甲信越（6）、東海（4）、'
    '近畿（6）、中国（5）、四国（4）、九州・沖縄（8）。'
)

add_heading_text('第1相：急性周術期鎮痛薬処方', level=2)
doc.add_paragraph(
    '周術期疼痛管理に一般的に使用される3つの鎮痛薬クラスの入院処方データを抽出した：'
    '114類（解熱鎮痛薬、NSAIDsおよびアセトアミノフェンを含む）、'
    '811類（あへんアルカロイド系麻薬）、821類（合成麻薬、フェンタニルおよびペチジンを含む）。'
    '手術あたり鎮痛薬指数は、各都道府県について入院鎮痛薬処方単位数を'
    '入院手術件数（Kコード手術）で除して算出した。'
)

add_heading_text('第2相：外来神経障害性疼痛薬処方', level=2)
doc.add_paragraph(
    '外来経口神経障害性疼痛薬5種を抽出した：'
    'プレガバリン（78規格）、ミロガバリン（8規格）、'
    'デュロキセチン（33規格）、トラマドール（3規格）、ノイロトロピン（1規格）。'
    '手術あたり神経障害性疼痛薬処方指数を算出し、'
    '人口あたり処方率も追加算出した。'
)

add_heading_text('交絡疾患プロキシ', level=2)
doc.add_paragraph(
    '外来データから4つの交絡疾患プロキシを抽出した：'
    '経口血糖降下薬（261規格、糖尿病性神経障害の代理指標）、'
    '帯状疱疹抗ウイルス薬（47規格、帯状疱疹後神経痛の代理指標）、'
    'デュロキセチンを除く抗うつ薬（128規格、うつ病の代理指標）、'
    '抗不安薬（112規格、不安障害の代理指標）。'
    '各プロキシは一貫性のため手術あたりで表現した。'
    '外来神経ブロック処置件数（73コード）も追加の独立プロキシとした。'
)

add_heading_text('需要側ベンチマーク：国民生活基礎調査', level=2)
m_cslc = (
    f'都道府県別の自覚症状有訴者率を国民生活基礎調査（CSLC）2022年版から取得した。{cite(7)}'
    f'CSLCは約30万世帯を対象とする全国代表世帯調査であり、'
    f'有訴者率（人口千対）が47都道府県について公表されている。'
    f'自覚症状有訴者率は、全国的に筋骨格系疼痛（腰痛・肩こり）が最多であるため、'
    f'疼痛関連需要の生態学的プロキシとして使用した。'
    f'需要\u2013供給ミスマッチ指数は、z標準化した処方率とz標準化した有訴者率の差として算出した。'
)
p = doc.add_paragraph()
add_ref_runs(p, m_cslc)

add_heading_text('統計解析', level=2)
doc.add_paragraph(
    '第1相の地域差はKruskal–Wallis検定で評価し、'
    'Bonferroni補正付きMann–Whitney U検定で事後比較を行った。'
    '効果量はCohen\'s dで定量化した。'
    '第2相では、交絡因子を段階的に調整する5つの回帰モデルを構築した：'
    'モデル1（未調整、地域間比較）、'
    'モデル2（神経障害性疼痛 ~ 糖尿病+帯状疱疹+抗うつ薬+抗不安薬+地域）、'
    'モデル3（コア神経障害性薬のみ ~ 同交絡因子）、'
    'モデル4（神経ブロック ~ 同交絡因子）、'
    'モデル5（神経障害性疼痛 ~ 急性鎮痛薬指数+交絡因子）。'
    '調整CPSP指数は、神経障害性疼痛処方を4つの交絡プロキシに回帰した残差として導出した。'
)

m7b = (
    f'標準化レセプト比（SCR）は、Tairaら{cite(4)}に従い間接年齢・性別標準化により算出した。'
    f'全国の年齢・性別別処方率（5歳階級18区分×2性別）を各都道府県の人口構成に適用した。'
    f'全解析にPython 3.11（NumPy 1.24、SciPy 1.11）を使用した。'
)
p = doc.add_paragraph()
add_ref_runs(p, m7b)

# ============================================================
# RESULTS
# ============================================================
add_heading_text('結果', level=1)

add_heading_text('第1相：急性周術期鎮痛薬処方の地域差', level=2)
doc.add_paragraph(
    '2023年4月〜2024年3月の間に、NDBは47都道府県で7,903,515件の入院手術と'
    '274,579,851単位の鎮痛薬処方を記録した。'
    '全国平均手術あたり鎮痛薬指数は35.78（SD 5.56）であり、'
    '岐阜の25.20から鹿児島の49.75まで1.97倍の差があった'
    '（9地域間Kruskal–Wallis P<0.001; 表1）。'
)

# === TABLE 1 ===
p_cap = doc.add_paragraph()
p_cap.paragraph_format.space_before = Pt(14)
r_cap = p_cap.add_run('表1. ')
r_cap.bold = True
r_cap.font.size = Pt(10)
p_cap.add_run(
    '9地域ブロック別の入院鎮痛薬処方（手術あたり）。'
    '値は平均（SD）。Kruskal–Wallis P<0.001。'
).font.size = Pt(10)

t1 = doc.add_table(rows=1 + len(REGION_ORDER), cols=4, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t1.rows[0].cells
for i, h in enumerate(['地域', 'n', '平均（SD）', '範囲']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, reg_name in enumerate(REGION_ORDER):
    vals = region_data[reg_name]
    row = t1.rows[idx + 1].cells
    row[0].text = REGION_JA[reg_name]
    row[1].text = str(len(vals))
    row[2].text = f'{np.mean(vals):.2f}（{np.std(vals, ddof=1):.2f}）' if len(vals) > 1 else f'{np.mean(vals):.2f}'
    row[3].text = f'{min(vals):.2f}–{max(vals):.2f}'
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t1)
doc.add_paragraph()

doc.add_paragraph(
    '顕著な地域集積が認められた。東海と近畿が最も低い指数を示し、'
    '九州・沖縄と北海道が最も高かった。'
    '事後ペアワイズ比較（Bonferroni補正Mann–Whitney U検定）により、'
    '複数の地域ブロック間で有意な差が確認された。'
    'このパターンは3薬剤クラスすべてで一貫していた：'
    'NSAIDs、あへんアルカロイド、合成麻薬。'
)

add_heading_text('第2相：外来神経障害性疼痛薬処方', level=2)
r3_text = (
    f'外来神経障害性疼痛薬処方は合計2,289,549,163単位であり、'
    f'プレガバリン（40.2%）、ノイロトロピン（20.1%）、ミロガバリン（19.6%）、'
    f'デュロキセチン（15.3%）、トラマドール（4.9%）で構成された。'
    f'神経障害性疼痛薬処方指数にも顕著な地域間差が認められた'
    f'（Kruskal–Wallis P<0.001; 図1）。'
)
doc.add_paragraph(r3_text)

add_inline_figure(
    '都道府県別の外来神経障害性疼痛薬処方（手術あたり）。'
    '棒は地域ブロック別に色とハッチングで区別。'
    '破線は全国平均を示す。',
    1
)

add_heading_text('交絡分析と調整', level=2)
r4a_text = (
    f'神経障害性疼痛薬処方は交絡疾患プロキシと強い相関を示した。'
    f'糖尿病薬処方が最も強い相関（r=0.87, P<0.001）を示し、'
    f'次いで抗不安薬（r=0.75）、抗うつ薬（r=0.46）、'
    f'帯状疱疹抗ウイルス薬（r=0.19）であった。'
    f'4つの交絡因子は神経障害性疼痛薬処方の都道府県間分散の'
    f'{reg["model2_adjusted"]["R2"]*100:.1f}%を集合的に説明した'
    f'（R\u00b2={reg["model2_adjusted"]["R2"]:.3f}; 図2）。'
)
doc.add_paragraph(r4a_text)

add_inline_figure(
    '47都道府県における神経障害性疼痛薬処方と交絡疾患プロキシの相関。'
    '糖尿病薬が最も強い相関を示す（r=0.87）。',
    2
)

r4b_text = (
    f'交絡因子調整後、見かけの地域間クラスタリングは大幅に減弱した（表2）。'
    f'未調整モデルで有意であった地域指標変数は、'
    f'すべてのモデル仕様（モデル2〜5）で非有意となり、'
    f'交絡疾患が観察されたパターンの主因であることが示された。'
)
doc.add_paragraph(r4b_text)

# === TABLE 2 ===
p_cap2 = doc.add_paragraph()
p_cap2.paragraph_format.space_before = Pt(14)
r_cap2 = p_cap2.add_run('表2. ')
r_cap2.bold = True
r_cap2.font.size = Pt(10)
p_cap2.add_run(
    '外来神経障害性疼痛薬処方の回帰モデル。'
).font.size = Pt(10)

models = [
    ('モデル1（未調整）', '—',
     'Kruskal–Wallis H, P<0.001', '9地域オムニバス検定'),
    ('モデル2（全交絡因子）',
     f'R\u00b2={reg["model2_adjusted"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model2_adjusted"]["R2_adj"]:.3f}',
     f'地域β={reg["model2_adjusted"]["tohoku_coef"]:.1f}, '
     f'P={reg["model2_adjusted"]["tohoku_p"]:.3f}'),
    ('モデル3（コア神経障害薬）',
     f'R\u00b2={reg["model3_core_neuropathic"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model3_core_neuropathic"]["R2_adj"]:.3f}',
     f'地域β={reg["model3_core_neuropathic"]["tohoku_coef"]:.1f}, '
     f'P={reg["model3_core_neuropathic"]["tohoku_p"]:.3f}'),
    ('モデル4（神経ブロック）',
     f'R\u00b2={reg["model4_nerve_blocks"]["R2"]:.3f}',
     f'R\u00b2adj={reg["model4_nerve_blocks"]["R2_adj"]:.3f}',
     f'地域β={reg["model4_nerve_blocks"]["tohoku_coef"]:.2f}, '
     f'P={reg["model4_nerve_blocks"]["tohoku_p"]:.3f}'),
    ('モデル5（統合）',
     f'R\u00b2={reg["model5_integrated"]["R2"]:.3f}',
     f'急性β={reg["model5_integrated"]["acute_pain_coef"]:.2f}, P={reg["model5_integrated"]["acute_pain_p"]:.3f}',
     f'地域β={reg["model5_integrated"]["tohoku_coef"]:.1f}, '
     f'P={reg["model5_integrated"]["tohoku_p"]:.3f}'),
]

t2 = doc.add_table(rows=1 + len(models), cols=4, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['モデル', '交絡因子R\u00b2', '主要統計量', '備考']):
    t2.rows[0].cells[i].text = h
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

for idx, (name, r2_val, key_stat, note) in enumerate(models):
    row = t2.rows[idx + 1].cells
    row[0].text = name
    row[1].text = r2_val
    row[2].text = key_stat
    row[3].text = note
    for cell in row:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

set_table_borders(t2)
doc.add_paragraph()

r5a_text = (
    f'調整CPSP指数は、未調整データとは顕著に異なる地理的パターンを示した（図3）。'
    f'未調整分析で高い処方を示していた地域は、交絡調整後に有意な超過を示さなくなり'
    f'（Kruskal–Wallis P>0.05）、'
    f'元のクラスタリングが主として交絡疾患負荷の地域差によるものであったことが示された。'
)
doc.add_paragraph(r5a_text)

add_inline_figure(
    '神経障害性疼痛薬処方の地域比較：（a）未調整、（b）交絡調整後。'
    'エラーバーはSDを示す。',
    3
)

add_heading_text('需要\u2013供給の乖離', level=2)
r_cslc1 = (
    f'CSLC有訴者率は{cslc["symptom_rate_stats"]["min"]:.1f}（{cslc["symptom_rate_stats"]["min_pref"]}）から'
    f'{cslc["symptom_rate_stats"]["max"]:.1f}（{cslc["symptom_rate_stats"]["max_pref"]}）'
    f'（人口千対）の範囲であり、{cslc["symptom_rate_stats"]["range_ratio"]:.2f}倍の差であった'
    f'（平均{cslc["symptom_rate_stats"]["mean"]:.1f}、SD {cslc["symptom_rate_stats"]["sd"]:.1f}）。'
    f'一方、急性鎮痛薬処方は1.97倍、年齢・性別標準化後の神経障害性疼痛薬処方は'
    f'{scr_neuro_ratio:.1f}倍の差があり、処方の不均一性は症状負荷の差を大幅に上回っていた。'
)
doc.add_paragraph(r_cslc1)

r_cslc2 = (
    f'CSLC有訴者率は急性鎮痛薬処方と有意な相関を示さなかった'
    f'（r={cslc["correlations"]["symptom_vs_acute"]["pearson_r"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_acute"]["pearson_p"]:.3f}; '
    f'Spearman \u03c1={cslc["correlations"]["symptom_vs_acute"]["spearman_rho"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_acute"]["spearman_p"]:.3f}; 図4）。'
    f'人口あたり神経障害性疼痛薬処方との相関も認められなかった'
    f'（r={cslc["correlations"]["symptom_vs_neuro_percapita"]["pearson_r"]:.3f}, '
    f'P={cslc["correlations"]["symptom_vs_neuro_percapita"]["pearson_p"]:.3f}）。'
    f'相関の欠如は、処方差異が都道府県間の症状負荷の違いでは説明できないことを示す。'
)
doc.add_paragraph(r_cslc2)

add_inline_figure(
    '需要\u2013供給の乖離：CSLC有訴者率（人口千対）と急性鎮痛薬処方（手術あたり）。'
    '各点は1都道府県を示し、マーカー形状で地域ブロックを区別。'
    '相関の欠如（r=0.03）は処方差異が症状負荷と乖離していることを示す。',
    4
)

add_heading_text('急性・慢性処方の統合', level=2)
doc.add_paragraph(
    f'急性周術期処方は未調整神経障害性疼痛薬処方と正の相関を示した'
    f'（r=0.38, P=0.008）。交絡調整後、この相関は減弱した（r=0.29, P=0.052）。'
    f'モデル5では、急性疼痛指数は有意な予測因子であったが'
    f'（β={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}）、'
    f'地域指標は交絡調整後に非有意であった。'
)

add_heading_text('年齢・性別標準化レセプト比', level=2)
doc.add_paragraph(
    f'間接年齢・性別標準化後、入院鎮痛薬SCRは'
    f'{scr_analgesic_range[0]:.1f}から{scr_analgesic_range[1]:.1f}の範囲であり'
    f'（{scr_analgesic_ratio:.1f}倍差）、処方の不均一性が都道府県の'
    f'年齢・性別構成の違いに帰因しないことを確認した。'
    f'外来神経障害性疼痛薬SCRは{scr_neuro_range[0]:.1f}から'
    f'{scr_neuro_range[1]:.1f}（{scr_neuro_ratio:.1f}倍差）であった。'
)

# ============================================================
# DISCUSSION
# ============================================================
add_heading_text('考察', level=1)

doc.add_paragraph(
    '本研究は、NDBオープンデータを用いて、日本の47都道府県すべてにおける'
    '周術期および慢性疼痛関連処方をマッピングした初めての研究であり、'
    '独立した需要側指標との対照によりWennbergの不当な差異の基準を'
    '満たすか否かを正式に評価した初めての研究でもある。'
    '4つの主要な知見が得られた。'
)

add_heading_text('主要な知見', level=2)
d1 = (
    f'第一に、急性周術期鎮痛薬処方は都道府県間で1.97倍の差を示した。'
    f'これはTairaら{cite(4)}が報告した歯科診療利用の地域格差に匹敵し、'
    f'Wennberg{cite(1)}が記述した供給感応的差異パターンと一致する。'
    f'この差異は頑健であり、年齢・性別標準化後も持続し'
    f'（SCR範囲{scr_analgesic_range[0]:.0f}–{scr_analgesic_range[1]:.0f}）、'
    f'3つの鎮痛薬クラスすべてで一貫していた。'
)
p = doc.add_paragraph()
add_ref_runs(p, d1)

d2 = (
    f'第二に、方法論的に最も重要な知見として、神経障害性疼痛薬処方の顕著な地域間'
    f'クラスタリングは、交絡疾患プロキシ、'
    f'特に糖尿病薬処方（r=0.87）によって大部分が説明された。'
    f'調整後、見かけの地域間差は大幅に減弱し非有意となった。'
    f'これは重要な示唆を持つ：神経障害性疼痛薬を慢性疼痛の代理指標とする'
    f'生態学的研究では、糖尿病性神経障害を考慮しなければならない。'
)
doc.add_paragraph(d2)

d_cslc = (
    f'第三に、最も新規性の高い知見として、処方差異が症状負荷と乖離していた。'
    f'CSLC有訴者率は{cslc["symptom_rate_stats"]["range_ratio"]:.2f}倍の差に留まり、'
    f'鎮痛薬処方との相関は認められなかった（r=0.03, P=0.85）。'
    f'この乖離はWennbergの枠組みにおける不当な差異の基準を正式に満たす：{cite(1)}'
    f'供給側（処方）が需要側（症状負荷）と独立して変動しており、'
    f'疾病負荷ではなく診療パターンが不均一性を駆動していることを示す。'
    f'この知見はOECD諸国で報告されている供給感応的医療差異と一致する。{cite(2)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d_cslc)

d3 = (
    f'第四に、急性と交絡調整後慢性指数の正の相関（r=0.29, P=0.052）'
    f'およびモデル5における有意な急性疼痛予測因子'
    f'（β={reg["model5_integrated"]["acute_pain_coef"]:.2f}, '
    f'P={reg["model5_integrated"]["acute_pain_p"]:.3f}）は、'
    f'急性疼痛管理の地域パターンとその後の慢性疼痛関連処方との間に控えめな関連が'
    f'あることを示唆する。これは、急性術後疼痛強度がCPSPの危険因子であるという'
    f'個体レベルのエビデンスと一致する。{cite(10)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d3)

add_heading_text('臨床的示唆', level=2)
d4 = (
    f'1.97倍の国内差は、記述疫学を超えた広範な示唆を持つ。'
    f'文化的ステレオタイプが疼痛評価と処方行動に影響することは国際的に示されており、'
    f'Andersonら{cite(11)}は人種・民族による系統的格差を示し、'
    f'CampbellとEdwards{cite(12)}は文化的疼痛行動に関する臨床家の期待が'
    f'系統的な過少治療・過剰治療につながりうることを示した。'
    f'日本は文化的均質性が高く患者は疼痛に我慢強いと特徴づけられることが多いが、{cite(13,14)}'
    f'本データは、この均質とされる集団内でさえ処方パターンが約2倍異なることを示した。'
    f'ファイザー2017年調査でも、「治療せずに痛みを我慢している」慢性疼痛患者の'
    f'割合は都道府県間で48.7%（大阪）から81.6%（山梨）まで異なり、'
    f'秋田県（東北地方）は痛みを「我慢すべき」と回答した割合が全国最低（60.2%）であり、'
    f'文化的ステレオタイプとは矛盾する結果であった。{cite(15)}'
)
p = doc.add_paragraph()
add_ref_runs(p, d4)

d5 = (
    f'改定IASP定義は疼痛を本質的に主観的なものと記述している。{cite(16)}'
    f'地域ラベルや文化的一般化は個別化疼痛評価の代替にはならない。'
    f'客観的侵害受容モニタリングが周術期評価の標準化に貢献しうる。{cite(17)}'
    f'需要\u2013供給分析もこの点を補強する：同等の症状負荷を持つ都道府県でも'
    f'処方が大きく異なり、患者の需要ではなく臨床家の要因が治療強度を決定していることを示す。'
)
p = doc.add_paragraph()
add_ref_runs(p, d5)

add_heading_text('長所と限界', level=2)
doc.add_paragraph(
    '長所として、日本の全保険診療を網羅する全数データの使用、急性・慢性疼痛プロキシの'
    '新規統合、透明な交絡調整手法、医療アクセスを交絡因子として中和する'
    '周術期デザイン、およびNDB処方データと独立した世帯調査（CSLC）との'
    '需要\u2013供給分析のトライアンギュレーションが挙げられる。'
)

doc.add_paragraph(
    '限界は生態学的デザインに固有のものである。分析単位は都道府県であり個人ではなく、'
    '生態学的相関は個体レベルの関連を反映しない可能性がある（生態学的誤謬）。'
    'NDBオープンデータは診断コードを欠くため、神経障害性疼痛薬プロキシは'
    'CPSP以外の全適応を捕捉する。手術症例構成や医師密度などの'
    '未測定交絡因子が残存差異に寄与する可能性がある。'
    'CSLCの有訴者率は疼痛に特異的ではなく自覚症状全般の指標であるが、'
    '全国的に筋骨格系疼痛症状（腰痛・肩こり）が最多の訴えであり、'
    '疼痛関連需要の合理的な生態学的プロキシとして機能する。'
)

add_heading_text('結論', level=2)
doc.add_paragraph(
    '日本の47都道府県において、疼痛関連処方は年齢・性別標準化後も約2倍の差が持続し、'
    '症状負荷と乖離していた。このパターンはWennbergの不当な差異の基準を満たし、'
    '供給側要因が処方の不均一性を駆動していることを示唆する。'
    '交絡疾患—特に糖尿病—は神経障害性疼痛薬処方の見かけの地域パターンを実質的に変化させ、'
    '生態学的研究で考慮されなければならない。'
    '臨床家は地域的・文化的仮定ではなく、個別の評価に基づいて'
    '鎮痛薬の意思決定を行うべきである。'
)

# ============================================================
# Acknowledgements, Funding, etc.
# ============================================================
doc.add_paragraph()
add_heading_text('謝辞', level=1)
doc.add_paragraph(
    'NDBオープンデータを公開している厚生労働省に感謝する。'
)
p_ai = doc.add_paragraph()
p_ai.add_run('AIの使用: ').bold = True
p_ai.add_run(
    'データ処理および原稿作成の一部に生成AI（Claude, Anthropic）を使用した。'
    '原稿の正確性と内容について著者が全責任を負う。'
)
p_coi = doc.add_paragraph()
p_coi.add_run('利益相反: ').bold = True
p_coi.add_run('申告すべきものなし。')

doc.add_paragraph()
add_heading_text('資金', level=1)
doc.add_paragraph(
    '本研究は、公的・民間・非営利セクターのいずれの機関からも特定の助成を受けていない。')

doc.add_paragraph()
add_heading_text('著者貢献', level=1)
doc.add_paragraph(
    'TO: 研究の着想、データ管理、解析、方法論、ソフトウェア、可視化、'
    '原稿執筆（初稿・査読修正）。')

doc.add_paragraph()
add_heading_text('データ利用可能性', level=1)
doc.add_paragraph(
    '公開データに基づく。本研究で使用したNDBオープンデータは厚生労働省ウェブサイト'
    '（https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177221_00016.html）'
    'から一般に入手可能である。解析コードは'
    'https://github.com/bougtoir/ndb-pain-regional-variation-japan で公開している。')

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
add_heading_text('文献', level=1)

for i, ref_text in enumerate(ref_list, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.bold = True
    p.add_run(ref_text)

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(JE_DIR, 'JE_manuscript_JA.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
