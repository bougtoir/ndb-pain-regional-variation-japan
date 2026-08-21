#!/usr/bin/env python3
"""Create STROBE checklist for JoPPP submission.

STROBE Statement for cross-sectional ecological study.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
JOPP_DIR = os.path.join(OUTPUT_DIR, 'joppp')
os.makedirs(JOPP_DIR, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = 1  # landscape
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15

h = doc.add_heading('STROBE Statement\u2014Checklist of items that should be included in reports of cross-sectional studies', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


items = [
    ('', 'Item No', 'Recommendation', 'Reported on page / section'),
    ('Title and abstract', '1', '(a) Indicate the study\u2019s design with a commonly used term in the title or the abstract\n'
     '(b) Provide in the abstract an informative and balanced summary of what was done and what was found',
     '(a) Title: \u201cecological study\u201d in title\n(b) Structured abstract with Background, Methods, Results, Conclusion'),
    ('', '', '', ''),
    ('INTRODUCTION', '', '', ''),
    ('Background/rationale', '2', 'Explain the scientific background and rationale for the investigation being reported',
     'Background paragraphs 1\u20133'),
    ('Objectives', '3', 'State specific objectives, including any prespecified hypotheses',
     'Background paragraph 4: three objectives stated'),
    ('', '', '', ''),
    ('METHODS', '', '', ''),
    ('Study design', '4', 'Present key elements of study design early in the paper',
     'Methods: Study design and reporting'),
    ('Setting', '5', 'Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection',
     'Methods: Data source (NDB 10th ed., April 2023\u2013March 2024)'),
    ('Participants', '6', '(a) Give the eligibility criteria, and the sources and methods of selection of participants',
     '(a) Methods: all insured individuals in Japan (~125M)\nN/A for ecological study: unit is prefecture, not individual'),
    ('Variables', '7', 'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers',
     'Methods: Phase 1 (analgesic index), Phase 2 (neuropathic pain index), Confounders (4 proxies)'),
    ('Data sources/measurement', '8', 'For each variable of interest, give sources of data and details of methods of assessment',
     'Methods: NDB drug classes, formulation counts, procedure codes specified'),
    ('Bias', '9', 'Describe any efforts to address potential sources of bias',
     'Methods: confounder adjustment; Discussion: Strengths and limitations'),
    ('Study size', '10', 'Explain how the study size was arrived at',
     'Methods: all 47 prefectures (population-complete data)'),
    ('Quantitative variables', '11', 'Explain how quantitative variables were handled in the analyses',
     'Methods: per-surgery indices, regression models, residual derivation'),
    ('Statistical methods', '12', '(a) Describe all statistical methods\n'
     '(b) Describe any methods used to examine subgroups and interactions\n'
     '(c) Explain how missing data were addressed\n'
     '(d) If applicable, describe analytical methods taking account of sampling strategy\n'
     '(e) Describe any sensitivity analyses',
     '(a) Methods: Statistical analysis (Kruskal\u2013Wallis, regression, correlation)\n'
     '(b) Regional blocks, Tohoku indicator\n'
     '(c) NDB suppresses cells <10; no missing prefectures\n'
     '(d) N/A (population-complete)\n'
     '(e) Models 2\u20135 as sensitivity'),
    ('', '', '', ''),
    ('RESULTS', '', '', ''),
    ('Participants', '13', '(a) Report numbers of individuals at each stage of study\n'
     '(b) Give reasons for non-participation at each stage\n'
     '(c) Consider use of a flow diagram',
     '(a) Results: 7,903,515 surgeries, 274,579,851 analgesic units, 47 prefectures\n'
     '(b) N/A (ecological, all prefectures included)\n'
     '(c) N/A'),
    ('Descriptive data', '14', '(a) Give characteristics of study participants\n'
     '(b) Indicate number of participants with missing data for each variable',
     '(a) Table 1: regional summary\n'
     '(b) No missing prefectures; cell suppression noted'),
    ('Outcome data', '15', 'Report numbers of outcome events or summary measures',
     'Results: all indices reported with means, SDs, ranges'),
    ('Main results', '16', '(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates\n'
     '(b) Report category boundaries when continuous variables were categorized\n'
     '(c) If relevant, consider translating estimates into meaningful clinical measures',
     '(a) Table 2: unadjusted and adjusted results\n'
     '(b) Tohoku binary indicator defined\n'
     '(c) Cohen\u2019s d, fold-change, percentage attenuation'),
    ('Other analyses', '17', 'Report other analyses done\u2014e.g., analyses of subgroups and interactions, and sensitivity analyses',
     'Models 2\u20135: sensitivity analyses; Phase 1\u20132 integration; standardised claim ratios'),
    ('', '', '', ''),
    ('DISCUSSION', '', '', ''),
    ('Key results', '18', 'Summarise key results with reference to study objectives',
     'Discussion paragraph 1'),
    ('Limitations', '19', 'Discuss limitations of the study, taking into account sources of potential bias or imprecision',
     'Discussion: Strengths and limitations'),
    ('Interpretation', '20', 'Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence',
     'Discussion: all subsections'),
    ('Generalisability', '21', 'Discuss the generalisability (external validity) of the study results',
     'Discussion: Clinical implications; Implications and future directions'),
    ('', '', '', ''),
    ('OTHER INFORMATION', '', '', ''),
    ('Funding', '22', 'Give the source of funding and the role of the funders for the present study',
     'Funding sources section'),
    ('', '', '', ''),
    ('RECORD-specific items', '', '', ''),
    ('Data source', 'RECORD 6.1', 'The methods of study population selection (e.g., codes or algorithms used to identify subjects) should be listed in detail',
     'Methods: drug class codes, formulation counts, procedure codes specified'),
    ('Data cleaning', 'RECORD 6.2', 'Any validation studies of the codes or algorithms used to select the population should be referenced',
     'N/A: NDB Open Data is aggregate; validation not applicable to publicly released aggregate statistics'),
    ('Data linkage', 'RECORD 6.3', 'If the study involved linkage of databases, the methods and data quality checks should be described',
     'N/A: single data source (NDB Open Data)'),
    ('Ethics', 'RECORD 12.1', 'Authors should describe the extent to which investigators had access to the database and the steps taken to ensure data security',
     'Methods: publicly available aggregate data; no individual-level access'),
]

table = doc.add_table(rows=len(items), cols=4)
set_table_borders(table)

col_widths = [Cm(4.5), Cm(2.0), Cm(12.0), Cm(8.0)]
for row in table.rows:
    for idx, width in enumerate(col_widths):
        row.cells[idx].width = width

for r_idx, (section, item_no, rec, reported) in enumerate(items):
    row = table.rows[r_idx]
    row.cells[0].text = section
    row.cells[1].text = item_no
    row.cells[2].text = rec
    row.cells[3].text = reported

    for c_idx in range(4):
        for paragraph in row.cells[c_idx].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                if r_idx == 0:
                    run.bold = True
                    run.font.size = Pt(9)

    if section in ('INTRODUCTION', 'METHODS', 'RESULTS', 'DISCUSSION', 'OTHER INFORMATION', 'RECORD-specific items'):
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

outpath = os.path.join(JOPP_DIR, 'JoPPP_STROBE_checklist.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
