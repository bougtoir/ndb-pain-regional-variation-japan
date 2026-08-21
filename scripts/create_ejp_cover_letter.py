#!/usr/bin/env python3
"""Create cover letter for European Journal of Pain submission."""

from docx import Document
from docx.shared import Pt, Cm
import os
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
EJP_DIR = os.path.join(OUTPUT_DIR, 'ejp')
os.makedirs(EJP_DIR, exist_ok=True)

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

today = datetime.date.today().strftime('%B %d, %Y')

# Date
p = doc.add_paragraph(today)
doc.add_paragraph()

# Addressee
doc.add_paragraph('Professor Didier Bouhassira')
doc.add_paragraph('Editor-in-Chief')
doc.add_paragraph('European Journal of Pain')
doc.add_paragraph()

# Salutation
doc.add_paragraph('Dear Professor Bouhassira,')
doc.add_paragraph()

# Body
doc.add_paragraph(
    'We respectfully submit the enclosed manuscript entitled '
    '\u201cRegional Heterogeneity in Pain-Related Prescribing Across Japan\u2019s 47 Prefectures '
    'Challenges a Stoic Monolithic Patient Stereotype: an ecological study\u201d '
    'for consideration as an Original Article in the European Journal of Pain.'
)

doc.add_paragraph(
    'Cross-cultural studies have long characterized Japanese people as stoic toward pain, '
    'yet whether this endurance varies within Japan has never been examined at the population level. '
    'Using Japan\u2019s National Database (NDB) Open Data\u2014capturing virtually all insurance-reimbursed '
    'healthcare for 125 million people\u2014we mapped pain-related prescribing across all 47 prefectures '
    'and nine regional blocks.'
)

doc.add_paragraph(
    'Our key findings are:'
)
doc.add_paragraph(
    '(1) Acute perioperative analgesic prescribing varied 1.97-fold across prefectures '
    '(Gifu 25.20 to Kagoshima 49.75), demonstrating that Japan\u2019s pain culture is not monolithic.',
    style='List Bullet'
)
doc.add_paragraph(
    '(2) Tohoku\u2014traditionally considered Japan\u2019s most stoic region\u2014prescribed more, '
    'not fewer, analgesics than the national average (Cohen\u2019s d = 0.87).',
    style='List Bullet'
)
doc.add_paragraph(
    '(3) The large regional variation in outpatient neuropathic pain prescribing '
    '(used as a chronic postsurgical pain proxy) was largely explained by confounding disease prevalence, '
    'particularly diabetes, after systematic confounder adjustment.',
    style='List Bullet'
)

doc.add_paragraph(
    'We believe this manuscript is well suited for the European Journal of Pain for several reasons. '
    'First, the topic directly relates to cultural influences on pain, a theme central to EJP\u2019s scope. '
    'Second, one of our key references\u2014Hobara (2005), demonstrating cultural differences in pain beliefs '
    'between Japanese and Euro-Americans\u2014was published in your journal, establishing a direct lineage. '
    'Third, our finding that monolithic cultural labeling carries clinical risk is relevant to European '
    'readers who increasingly treat patients from diverse cultural backgrounds, including a growing '
    'Japanese expatriate population.'
)

doc.add_paragraph(
    'The study leverages freely available open data from Japan\u2019s universal healthcare system, '
    'offering a methodological framework that could be adapted to European countries with similar '
    'population-level databases. The transparent confounder-adjustment methodology demonstrated here '
    'addresses a critical gap in ecological pain research.'
)

doc.add_paragraph(
    'This manuscript has not been published previously and is not under consideration elsewhere. '
    'All authors have approved the manuscript and agree with its submission to the European Journal of Pain. '
    'The study used only publicly available aggregate data, so ethical approval was not required. '
    'The authors have no conflicts of interest to declare.'
)

doc.add_paragraph(
    'We confirm that this manuscript complies with the STROBE guidelines for reporting observational '
    'studies, and a completed STROBE checklist is included with this submission.'
)

doc.add_paragraph(
    'Thank you for considering our work. We look forward to your response.'
)

doc.add_paragraph()

doc.add_paragraph('Sincerely,')
doc.add_paragraph()
doc.add_paragraph('Tatsuki Onishi')
doc.add_paragraph('Department of Anesthesiology')
doc.add_paragraph('[Institution]')
doc.add_paragraph('[Address]')
doc.add_paragraph('E-mail: [email]')

outpath = os.path.join(EJP_DIR, 'EJP_cover_letter.docx')
doc.save(outpath)
print(f'Saved: {outpath}')
